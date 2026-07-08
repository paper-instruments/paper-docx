"""Tests for the ported verify_docx package-fact checks (harness/checks.py)."""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest

from docx import Document

from .harness import checks
from .harness.paths import FIXTURES_DIR, fixture_path, iter_fixture_docx_paths

MINIMAL = "generated/minimal-clean/minimal.docx"
BROKEN_REL = "generated/corrupt/broken-rel.docx"
MALFORMED = "generated/corrupt/malformed-xml.docx"
NUMBERING = "generated/feature-isolated/numbering-custom.docx"


def _clean_fixture_relpaths() -> list:
    """Every non-corrupt fixture in the corpus."""
    return [
        rel
        for rel in (
            p.relative_to(FIXTURES_DIR).as_posix() for p in iter_fixture_docx_paths()
        )
        if rel.split("/")[1] != "corrupt"
    ]


def _mutated_copy(src: Path, dst: Path, part_name: str, old: bytes, new: bytes) -> Path:
    """Copy the package at `src` to `dst` with one byte-level part substitution."""
    with zipfile.ZipFile(src) as zin:
        names = zin.namelist()
        parts = {name: zin.read(name) for name in names}
    assert old in parts[part_name], f"{old!r} not found in {part_name}"
    parts[part_name] = parts[part_name].replace(old, new)
    with zipfile.ZipFile(dst, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name in names:
            zout.writestr(name, parts[name])
    return dst


class DescribeFindMissingRequiredParts:
    def it_returns_empty_for_a_wellformed_package(self):
        assert checks.find_missing_required_parts(fixture_path(MINIMAL)) == []

    def it_reports_each_absent_required_part(self, tmp_path: Path):
        crippled = tmp_path / "no-document.docx"
        with zipfile.ZipFile(fixture_path(MINIMAL)) as zin:
            with zipfile.ZipFile(crippled, "w") as zout:
                for name in zin.namelist():
                    if name != "word/document.xml":
                        zout.writestr(name, zin.read(name))
        assert checks.find_missing_required_parts(crippled) == ["word/document.xml"]


class DescribeFindUnparseableXmlParts:
    def it_returns_empty_for_a_wellformed_package(self):
        assert checks.find_unparseable_xml_parts(fixture_path(MINIMAL)) == []

    def it_reports_the_truncated_document_part(self):
        failures = checks.find_unparseable_xml_parts(fixture_path(MALFORMED))
        assert [name for name, _ in failures] == ["word/document.xml"]


class DescribeFindBrokenRelationshipTargets:
    def it_returns_empty_for_a_wellformed_package(self):
        assert checks.find_broken_relationship_targets(fixture_path(MINIMAL)) == []

    def it_reports_the_dangling_target(self):
        findings = checks.find_broken_relationship_targets(fixture_path(BROKEN_REL))
        assert findings == [("word/_rels/document.xml.rels", "media/missing-image.png")]

    def and_it_reports_targets_that_escape_the_package_root(self, tmp_path: Path):
        """`../../...` resolves outside the package even when a same-named part
        exists inside it — no conforming OPC loader can resolve it."""
        mutated = _mutated_copy(
            fixture_path(MINIMAL),
            tmp_path / "root-escape.docx",
            "word/_rels/document.xml.rels",
            b'Target="styles.xml"',
            b'Target="../../word/styles.xml"',
        )
        findings = checks.find_broken_relationship_targets(mutated)
        assert findings == [("word/_rels/document.xml.rels", "../../word/styles.xml")]


class DescribeFindUndefinedStyleReferences:
    @pytest.mark.parametrize("relpath", _clean_fixture_relpaths())
    def it_returns_empty_for_every_clean_fixture(self, relpath: str):
        assert checks.find_undefined_style_references(fixture_path(relpath)) == []

    def it_reports_a_reference_to_an_undefined_style(self, tmp_path: Path):
        mutated = _mutated_copy(
            fixture_path(MINIMAL),
            tmp_path / "bad-style.docx",
            "word/document.xml",
            b'w:val="Heading1"',
            b'w:val="NoSuchStyle27"',
        )
        assert checks.find_undefined_style_references(mutated) == [
            ("word/document.xml", "NoSuchStyle27")
        ]

    def and_it_resolves_glossary_references_in_the_glossary_scope(self, tmp_path: Path):
        """word/glossary/ is its own story with its own styles table; its
        references must not be checked against word/styles.xml."""
        W = b"http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        glossary_doc = (
            b'<w:glossaryDocument xmlns:w="' + W + b'"><w:docParts/>'
            b'<w:p><w:pPr><w:pStyle w:val="GlossaryOnlyStyle"/></w:pPr>'
            b"<w:r><w:t>placeholder</w:t></w:r></w:p></w:glossaryDocument>"
        )
        glossary_styles = (
            b'<w:styles xmlns:w="' + W + b'">'
            b'<w:style w:type="paragraph" w:styleId="GlossaryOnlyStyle">'
            b'<w:name w:val="Glossary Only"/></w:style></w:styles>'
        )
        path = tmp_path / "with-glossary.docx"
        with zipfile.ZipFile(fixture_path(MINIMAL)) as zin:
            with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_DEFLATED) as zout:
                for name in zin.namelist():
                    zout.writestr(name, zin.read(name))
                zout.writestr("word/glossary/document.xml", glossary_doc)
                zout.writestr("word/glossary/styles.xml", glossary_styles)
        assert checks.find_undefined_style_references(path) == [], (
            "glossary style defined in word/glossary/styles.xml must resolve"
        )


class DescribeFindUndefinedNumberingReferences:
    @pytest.mark.parametrize("relpath", _clean_fixture_relpaths())
    def it_returns_empty_for_every_clean_fixture(self, relpath: str):
        assert checks.find_undefined_numbering_references(fixture_path(relpath)) == []

    def it_reports_a_dangling_num_id(self, tmp_path: Path):
        mutated = _mutated_copy(
            fixture_path(NUMBERING),
            tmp_path / "bad-num.docx",
            "word/document.xml",
            b'<w:numId w:val="42"/>',
            b'<w:numId w:val="999"/>',
        )
        findings = checks.find_undefined_numbering_references(mutated)
        assert findings == [("word/document.xml", "999")] * 3

    def it_reports_references_when_no_numbering_part_exists(self, tmp_path: Path):
        """With numbering.xml gone, EVERY numId reference dangles — including
        the ones the default template's styles.xml makes."""
        source = fixture_path(NUMBERING)
        stripped = tmp_path / "no-numbering-part.docx"
        with zipfile.ZipFile(source) as zin:
            with zipfile.ZipFile(stripped, "w") as zout:
                for name in zin.namelist():
                    if name != "word/numbering.xml":
                        zout.writestr(name, zin.read(name))
        findings = checks.find_undefined_numbering_references(stripped)
        document_findings = [f for f in findings if f[0] == "word/document.xml"]
        assert document_findings == [("word/document.xml", "42")] * 3
        assert ("word/styles.xml", "1") in findings, (
            "the default template's list styles also reference numbering and"
            " must be reported once the part is gone"
        )


class DescribeFindFakeBulletParagraphs:
    def it_flags_literal_bullet_paragraphs(self, tmp_path: Path):
        path = tmp_path / "fake-bullets.docx"
        document = Document()
        document.add_paragraph("A perfectly normal paragraph.")
        document.add_paragraph("• fake bullet item")
        document.add_paragraph("- another fake bullet")
        document.save(str(path))
        assert checks.find_fake_bullet_paragraphs(path) == [2, 3]

    def and_it_exempts_bullet_looking_text_that_has_real_numbering(
        self, tmp_path: Path
    ):
        """The numPr guard is the point of the check: text starting with a
        literal dash is fine when the paragraph carries real Word numbering."""
        from docx.oxml.ns import nsdecls
        from docx.oxml.parser import parse_xml

        path = tmp_path / "dash-with-numpr.docx"
        document = Document()
        p = document.add_paragraph("- dash-led text on a genuinely numbered paragraph")
        p._p.get_or_add_pPr().append(  # noqa: SLF001 - test authoring
            parse_xml(
                f'<w:numPr {nsdecls("w")}>'
                '<w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>'
            )
        )
        document.add_paragraph("- and one without numbering, which IS flagged")
        document.save(str(path))
        assert checks.find_fake_bullet_paragraphs(path) == [2]
