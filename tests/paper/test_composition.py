"""Cross-document composition."""

from __future__ import annotations

import io
from pathlib import Path

import pytest

import docx
from docx.composition import append_document, insert_blocks_from
from docx.errors import (
    DocumentProtectedError,
    UnsupportedStructureError,
)

from .harness import checks
from .harness.paths import fixture_path

MINIMAL = "generated/minimal-clean/minimal.docx"

#: a valid 1x1 red PNG (for image-copy tests)
TINY_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000000108020000009077"
    "53de0000000c4944415408d763f8cfc000000301010018dd8db00000000049"
    "454e44ae426082"
)


def _saved(document, path: Path) -> Path:
    document.save(str(path))
    checks.assert_package_facts_clean(path)
    return path


def _source_with_styles():
    source = docx.Document()
    source.add_heading("Clause Library", level=1)
    source.add_paragraph("Indemnity: the supplier shall hold harmless.")
    source.add_heading("Payment Terms", level=2)
    source.add_paragraph("Payment falls due after thirty days.")
    return source


def _source_with_named_blocks():
    source = docx.Document()
    for name in ("Block A", "Block B", "Block C", "Block D"):
        source.add_paragraph(name)
    return source


class DescribeInsertBlocksFrom:
    def it_copies_a_styled_range_adopting_the_house_look(self, tmp_path: Path):
        source = _source_with_styles()
        destination = docx.Document(str(fixture_path(MINIMAL)))
        report = insert_blocks_from(
            destination,
            source,
            "Payment Terms",
            count=2,
            anchor="Second body paragraph",
        )
        assert report.inserted_blocks == 2
        out = _saved(destination, tmp_path / "out.docx")
        reopened = docx.Document(str(out))
        texts = [p.text for p in reopened.paragraphs]
        assert "Payment Terms" in texts
        assert "Payment falls due after thirty days." in texts
        heading = next(p for p in reopened.paragraphs if p.text == "Payment Terms")
        assert heading.style.name == "Heading 2"
        # match_by_name: the destination's own Heading 2 definition wins
        assert "Heading2" not in report.imported_styles
        payload = report.to_dict()
        assert payload["schema"] == "paper_composition"
        assert "word/styles.xml" in payload["declared_parts"]

    def it_selects_the_range_by_start_and_end_anchor(self, tmp_path: Path):
        source = _source_with_styles()
        destination = docx.Document(str(fixture_path(MINIMAL)))
        report = insert_blocks_from(
            destination,
            source,
            "Clause Library",
            end_anchor="Payment falls due after thirty days.",
            anchor="Second body paragraph",
        )
        assert report.inserted_blocks == 4
        out = _saved(destination, tmp_path / "inclusive-defaults.docx")
        reopened = docx.Document(str(out))
        expected = [
            "Clause Library",
            "Indemnity: the supplier shall hold harmless.",
            "Payment Terms",
            "Payment falls due after thirty days.",
        ]
        assert [p.text for p in reopened.paragraphs if p.text in expected] == expected

    @pytest.mark.parametrize(
        ("include_start", "include_end", "expected"),
        [
            pytest.param(
                True,
                True,
                ["Block A", "Block B", "Block C", "Block D"],
                id="both",
            ),
            pytest.param(
                False,
                True,
                ["Block B", "Block C", "Block D"],
                id="exclude-start",
            ),
            pytest.param(
                True,
                False,
                ["Block A", "Block B", "Block C"],
                id="exclude-end",
            ),
            pytest.param(
                False,
                False,
                ["Block B", "Block C"],
                id="exclude-both",
            ),
        ],
    )
    def it_controls_endpoint_inclusion_and_ignores_count_with_an_end_anchor(
        self,
        include_start: bool,
        include_end: bool,
        expected: list[str],
        tmp_path: Path,
    ):
        source = _source_with_named_blocks()
        destination = docx.Document()
        destination.add_paragraph("Destination anchor")

        report = insert_blocks_from(
            destination,
            source,
            "Block A",
            end_anchor="Block D",
            count=99,
            include_start=include_start,
            include_end=include_end,
            anchor="Destination anchor",
        )

        assert report.inserted_blocks == len(expected)
        out = _saved(destination, tmp_path / "endpoint-range.docx")
        reopened = docx.Document(str(out))
        assert [
            paragraph.text
            for paragraph in reopened.paragraphs
            if paragraph.text.startswith("Block")
        ] == expected

    def it_excludes_the_count_based_start_without_consuming_the_count(
        self, tmp_path: Path
    ):
        source = _source_with_named_blocks()
        destination = docx.Document()
        destination.add_paragraph("Destination anchor")

        report = insert_blocks_from(
            destination,
            source,
            "Block A",
            count=2,
            include_start=False,
            anchor="Destination anchor",
        )

        assert report.inserted_blocks == 2
        out = _saved(destination, tmp_path / "count-range.docx")
        reopened = docx.Document(str(out))
        assert [
            paragraph.text
            for paragraph in reopened.paragraphs
            if paragraph.text.startswith("Block")
        ] == ["Block B", "Block C"]

    def it_requires_an_end_anchor_to_exclude_the_end(self):
        source = _source_with_named_blocks()
        destination = docx.Document()
        destination.add_paragraph("Destination anchor")
        before = destination.element.xml

        with pytest.raises(ValueError, match="requires end_anchor"):
            insert_blocks_from(
                destination,
                source,
                "Block A",
                include_end=False,
                anchor="Destination anchor",
            )

        assert destination.element.xml == before

    def it_imports_missing_style_definitions(self, tmp_path: Path):
        from docx.enum.style import WD_STYLE_TYPE

        source = docx.Document()
        style = source.styles.add_style("ClauseTerm", WD_STYLE_TYPE.PARAGRAPH)
        style.font.italic = True
        source.add_paragraph("A specially styled clause.", style="ClauseTerm")
        destination = docx.Document(str(fixture_path(MINIMAL)))
        report = insert_blocks_from(
            destination,
            source,
            "specially styled",
            anchor="Second body paragraph",
        )
        assert "ClauseTerm" in report.imported_styles
        out = _saved(destination, tmp_path / "out.docx")
        reopened = docx.Document(str(out))
        copied = next(
            p for p in reopened.paragraphs if "specially styled" in p.text
        )
        assert copied.style.name == "ClauseTerm"

    def it_renames_colliding_but_different_styles_on_request(
        self, tmp_path: Path
    ):
        from docx.enum.style import WD_STYLE_TYPE
        from docx.shared import Pt

        source = docx.Document()
        source_style = source.styles.add_style("HouseTerm", WD_STYLE_TYPE.PARAGRAPH)
        source_style.font.size = Pt(16)
        source.add_paragraph("Large-type clause.", style="HouseTerm")
        destination = docx.Document(str(fixture_path(MINIMAL)))
        destination_style = destination.styles.add_style(
            "HouseTerm", WD_STYLE_TYPE.PARAGRAPH
        )
        destination_style.font.size = Pt(9)
        report = insert_blocks_from(
            destination,
            source,
            "Large-type clause.",
            anchor="Second body paragraph",
            styles="import_renamed",
        )
        assert report.renamed_styles == {"HouseTerm": "HouseTerm (imported)"}
        out = _saved(destination, tmp_path / "out.docx")
        reopened = docx.Document(str(out))
        copied = next(p for p in reopened.paragraphs if "Large-type" in p.text)
        # content keeps its SOURCE look via the renamed clone
        assert copied.style.name == "HouseTerm (imported)"
        assert copied.style.font.size.pt == 16
        original = {s.name for s in reopened.styles}
        assert "HouseTerm" in original
        assert "HouseTerm (imported)" in original

    def it_remaps_numbering_to_fresh_restarted_definitions(self, tmp_path: Path):
        from docx.numbering import apply_numbering, ensure_decimal_definition, list_numbering

        source = docx.Document()
        source_num_id = ensure_decimal_definition(source)
        for text in ("First imported item", "Second imported item"):
            apply_numbering(source.add_paragraph(text), num_id=source_num_id)
        destination = docx.Document(str(fixture_path(MINIMAL)))
        destination_num_id = ensure_decimal_definition(destination)
        report = insert_blocks_from(
            destination,
            source,
            "First imported item",
            count=2,
            anchor="Second body paragraph",
        )
        new_id = report.numbering_map[source_num_id]
        assert new_id != destination_num_id
        out = _saved(destination, tmp_path / "out.docx")
        reopened = docx.Document(str(out))
        numbered = list_numbering(reopened).numbered_paragraphs
        copied = [p for p in numbered if "imported item" in p.text]
        assert [p.num_id for p in copied] == [new_id, new_id]

    def it_copies_images_as_new_parts_with_fresh_relationships(
        self, tmp_path: Path
    ):
        source = docx.Document()
        source.add_paragraph("Before the picture.")
        source.add_picture(io.BytesIO(TINY_PNG))
        destination = docx.Document(str(fixture_path(MINIMAL)))
        report = insert_blocks_from(
            destination,
            source,
            "Before the picture.",
            count=2,
            anchor="Second body paragraph",
        )
        assert len(report.media_copied) == 1
        out = _saved(destination, tmp_path / "out.docx")  # rels verified clean
        import zipfile

        with zipfile.ZipFile(out) as zf:
            assert any(name.startswith("word/media/") for name in zf.namelist())

    def it_recreates_external_hyperlinks(self, tmp_path: Path):
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml.ns import nsdecls
        from docx.oxml.parser import parse_xml

        source = docx.Document()
        paragraph = source.add_paragraph("Visit ")
        r_id = source.part.relate_to(
            "https://paper.example/terms", RT.HYPERLINK, is_external=True
        )
        paragraph._p.append(
            parse_xml(
                f'<w:hyperlink {nsdecls("w", "r")} r:id="{r_id}">'
                "<w:r><w:t>the terms page</w:t></w:r></w:hyperlink>"
            )
        )
        destination = docx.Document(str(fixture_path(MINIMAL)))
        insert_blocks_from(
            destination, source, "the terms page", anchor="Second body paragraph"
        )
        out = _saved(destination, tmp_path / "out.docx")
        reopened = docx.Document(str(out))
        external = [
            rel.target_ref
            for rel in reopened.part.rels.values()
            if rel.is_external and rel.reltype == RT.HYPERLINK
        ]
        assert "https://paper.example/terms" in external

    def it_renames_colliding_bookmarks_and_remaps_refs(self, tmp_path: Path):
        from docx.oxml.ns import nsdecls
        from docx.oxml.parser import parse_xml

        w = nsdecls("w")
        source = docx.Document()
        target = source.add_paragraph()
        target._p.append(parse_xml(f'<w:bookmarkStart {w} w:id="1" w:name="Term"/>'))
        target._p.append(parse_xml(f"<w:r {w}><w:t>the definition body</w:t></w:r>"))
        target._p.append(parse_xml(f'<w:bookmarkEnd {w} w:id="1"/>'))
        ref_paragraph = source.add_paragraph()
        ref_paragraph._p.append(
            parse_xml(
                f'<w:fldSimple {w} w:instr=" REF Term \\h ">'
                "<w:r><w:t>(cross-reference)</w:t></w:r></w:fldSimple>"
            )
        )
        destination = docx.Document(str(fixture_path(MINIMAL)))
        first = destination.paragraphs[1]
        first._p.append(parse_xml(f'<w:bookmarkStart {w} w:id="7" w:name="Term"/>'))
        first._p.append(parse_xml(f'<w:bookmarkEnd {w} w:id="7"/>'))
        report = insert_blocks_from(
            destination,
            source,
            "the definition body",
            count=2,
            anchor="Second body paragraph",
        )
        assert report.bookmarks_renamed == {"Term": "Term_imported"}
        xml = destination.element.xml
        assert 'w:name="Term_imported"' in xml
        assert " REF Term_imported " in xml

    def it_refuses_a_range_with_pending_revisions(self):
        from docx.search import find_one

        source = docx.Document(str(fixture_path(MINIMAL)))
        find_one(source, "perfectly ordinary").replace(
            "quite modern", tracked=True, author="Alice Editor"
        )
        destination = docx.Document(str(fixture_path(MINIMAL)))
        with pytest.raises(UnsupportedStructureError, match="accept or reject"):
            insert_blocks_from(
                destination,
                source,
                "First body paragraph",
                anchor="Second body paragraph",
            )

    def it_refuses_comment_anchors_and_note_references_in_the_range(self):
        destination = docx.Document(str(fixture_path(MINIMAL)))
        comments_source = docx.Document(
            str(fixture_path("generated/feature-isolated/comments.docx"))
        )
        with pytest.raises(UnsupportedStructureError, match="comment"):
            insert_blocks_from(
                destination,
                comments_source,
                "carries the first comment",
                anchor="Second body paragraph",
            )
        notes_source = docx.Document(
            str(fixture_path("generated/feature-isolated/footnotes-endnotes.docx"))
        )
        with pytest.raises(UnsupportedStructureError, match="footnote"):
            insert_blocks_from(
                destination,
                notes_source,
                "carries a footnote reference",
                anchor="Second body paragraph",
            )

    def it_refuses_embedded_objects(self):
        from docx.oxml.ns import nsdecls
        from docx.oxml.parser import parse_xml

        source = docx.Document()
        p = source.add_paragraph("Embedded thing: ")
        p._p.append(
            parse_xml(f"<w:r {nsdecls('w')}><w:object/></w:r>")
        )
        destination = docx.Document(str(fixture_path(MINIMAL)))
        with pytest.raises(UnsupportedStructureError, match="OLE"):
            insert_blocks_from(
                destination, source, "Embedded thing", anchor="Second body paragraph"
            )

    def it_refuses_a_protected_destination(self):
        source = _source_with_styles()
        destination = docx.Document(
            str(fixture_path("generated/feature-isolated/protected-readonly.docx"))
        )
        with pytest.raises(DocumentProtectedError):
            insert_blocks_from(
                destination, source, "Payment Terms", anchor="locked"
            )

    def it_validates_the_styles_mode(self):
        with pytest.raises(ValueError, match="styles"):
            insert_blocks_from(
                docx.Document(str(fixture_path(MINIMAL))),
                _source_with_styles(),
                "Payment Terms",
                anchor="Second body paragraph",
                styles="steal",
            )


class DescribeAppendDocument:
    def it_appends_with_a_page_break_by_default(self, tmp_path: Path):
        destination = docx.Document(str(fixture_path(MINIMAL)))
        report = append_document(destination, _source_with_styles())
        assert report.inserted_blocks == 4
        out = _saved(destination, tmp_path / "out.docx")
        reopened = docx.Document(str(out))
        texts = [p.text for p in reopened.paragraphs]
        assert texts.index("Second body paragraph, equally unremarkable.") < texts.index(
            "Clause Library"
        )
        assert 'w:type="page"' in reopened.element.xml

    def it_appends_flush_when_continuous(self, tmp_path: Path):
        destination = docx.Document(str(fixture_path(MINIMAL)))
        append_document(destination, _source_with_styles(), section="continuous")
        assert 'w:type="page"' not in destination.element.xml

    def it_keeps_destination_headers(self, tmp_path: Path):
        destination = docx.Document(
            str(fixture_path("generated/feature-isolated/header-footer-sections.docx"))
        )
        header_parts_before = {
            str(rel.target_part.partname)
            for rel in destination.part.rels.values()
            if not rel.is_external and "header" in str(rel.target_part.partname)
        }
        append_document(destination, _source_with_styles())
        header_parts_after = {
            str(rel.target_part.partname)
            for rel in destination.part.rels.values()
            if not rel.is_external and "header" in str(rel.target_part.partname)
        }
        assert header_parts_after == header_parts_before
        assert (
            destination.sections[0].header.paragraphs[0].text
            == "Header for section one"
        )

    def it_validates_the_section_argument(self):
        with pytest.raises(ValueError, match="section"):
            append_document(
                docx.Document(str(fixture_path(MINIMAL))),
                _source_with_styles(),
                section="fancy",
            )
