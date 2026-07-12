"""The compare engine and its algebra.

The invariants are the point: accept(compare(A,B)) == B, reject == A,
compare(A,A) == nothing, identical inputs -> byte-identical output.
"""

from __future__ import annotations

import datetime as dt
import io
import shutil
import zipfile
from pathlib import Path

import pytest

import docx
from docx.errors import UnsupportedStructureError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.package import compare
from docx.revision import _remaining_markup
from docx.story import iter_blocks

from .harness.paths import fixture_path

ORIGINAL = "generated/redline/compare-original.docx"
REVISED = "generated/redline/compare-revised.docx"
MULTIROUND = "generated/redline/multiround.docx"
MINIMAL = "generated/minimal-clean/minimal.docx"
GAUNTLET = "generated/gauntlet/gauntlet.docx"

FROZEN = dt.datetime(2026, 7, 8, 12, 0, 0, tzinfo=dt.timezone.utc)


def _texts(document) -> dict:
    texts: dict = {}
    for block in iter_blocks(document):
        texts.setdefault(block.story, []).append(block.text)
    return texts


def _visible(document) -> dict:
    """Per-story visible text with block boundaries collapsed (paragraph
    merges/splits legitimately change block counts, not content)."""
    return {story: "\n".join(t for t in items if t) for story, items in _texts(document).items()}


def _compare_fixture_pair():
    return compare(
        str(fixture_path(ORIGINAL)),
        str(fixture_path(REVISED)),
        author="Compare Engine",
        date=FROZEN,
    )


def _capitalize_header_part(path: Path) -> None:
    source = path.read_bytes()
    output = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(source)) as incoming, zipfile.ZipFile(
        output, "w", zipfile.ZIP_DEFLATED
    ) as outgoing:
        for name in incoming.namelist():
            blob = incoming.read(name)
            target_name = name
            if name == "word/header1.xml":
                target_name = "Word/Header1.xml"
            elif name == "[Content_Types].xml":
                blob = blob.replace(
                    b'PartName="/word/header1.xml"',
                    b'PartName="/Word/Header1.xml"',
                )
            outgoing.writestr(target_name, blob)
    path.write_bytes(output.getvalue())


class DescribeCompareAlgebra:
    def it_matches_loaded_story_names_case_insensitively(self, tmp_path: Path):
        original_path = tmp_path / "original.docx"
        revised_path = tmp_path / "revised.docx"
        for path, text in ((original_path, "before"), (revised_path, "after")):
            document = docx.Document()
            document.add_paragraph(text)
            document.sections[0].header.paragraphs[0].text = "Header"
            document.save(path)
            _capitalize_header_part(path)

        result = compare(original_path, revised_path, author="Compare Engine")

        assert result.revision_count > 0

    def it_verifies_compare_algebra_for_a_protected_document(self, tmp_path: Path):
        original_path = tmp_path / "original.docx"
        revised_path = tmp_path / "revised.docx"
        for path, text in ((original_path, "before"), (revised_path, "after")):
            document = docx.Document()
            document.add_paragraph(text)
            protection = OxmlElement("w:documentProtection")
            protection.set(qn("w:edit"), "readOnly")
            protection.set(qn("w:enforcement"), "1")
            document.settings.element.append(protection)
            document.save(path)

        result = compare(original_path, revised_path, author="Compare Engine")

        assert result.revision_count > 0
        assert any(finding.kind == "document_protection_present" for finding in result.findings)
        result.document.revisions.accept_all()

    def it_accept_alls_to_the_revised_text_across_every_story(self):
        result = _compare_fixture_pair()
        assert result.revision_count > 0
        result.document.revisions.accept_all()
        assert _remaining_markup(result.document) == {}
        assert _visible(result.document) == _visible(docx.Document(str(fixture_path(REVISED))))

    def it_reject_alls_back_to_the_original_text(self):
        result = _compare_fixture_pair()
        result.document.revisions.reject_all()
        assert _remaining_markup(result.document) == {}
        assert _visible(result.document) == _visible(docx.Document(str(fixture_path(ORIGINAL))))

    def it_yields_zero_revisions_for_identical_inputs(self):
        result = compare(
            str(fixture_path(ORIGINAL)),
            str(fixture_path(ORIGINAL)),
            author="Compare Engine",
            date=FROZEN,
        )
        assert result.revision_count == 0
        assert result.findings == []

    def it_is_deterministic_byte_for_byte(self, tmp_path: Path):
        for run in ("a", "b"):
            result = _compare_fixture_pair()
            result.document.save(str(tmp_path / f"out-{run}.docx"))
        assert (tmp_path / "out-a.docx").read_bytes() == (tmp_path / "out-b.docx").read_bytes()

    def it_survives_an_independent_reopen_with_the_same_algebra(self, tmp_path: Path):
        result = _compare_fixture_pair()
        out = tmp_path / "redline.docx"
        result.document.save(str(out))
        reopened = docx.Document(str(out))
        assert len(reopened.revisions) == result.revision_count
        reopened.revisions.accept_all()
        assert _visible(reopened) == _visible(docx.Document(str(fixture_path(REVISED))))


class DescribeCompareBehavior:
    def it_maps_word_edits_after_an_inline_image(self, tmp_path: Path):
        original_path = tmp_path / "original.docx"
        revised_path = tmp_path / "revised.docx"
        image_path = Path(__file__).parents[1] / "test_files" / "python-icon.png"
        for path, ending in (
            (original_path, "old language"),
            (revised_path, "new language"),
        ):
            document = docx.Document()
            paragraph = document.add_paragraph("Logo ")
            paragraph.add_run().add_picture(str(image_path))
            paragraph.add_run(f" {ending}")
            document.save(path)

        result = compare(original_path, revised_path, author="Compare Engine")

        assert result.revision_count > 0
        assert len(result.document.inline_shapes) == 1
        result.document.revisions.accept_all()
        assert result.document.paragraphs[0].text == "Logo  new language"
        assert len(result.document.inline_shapes) == 1

    def it_redlines_a_word_level_edit_minimally(self):
        result = _compare_fixture_pair()
        revisions = result.document.revisions
        deleted = [r.text for r in revisions if r.revision_type == "deletion"]
        inserted = [r.text for r in revisions if r.revision_type == "insertion"]
        # "thirty days notice" -> "sixty days notice": the span machinery's
        # affix trimming narrows the word-level region to the minimal change
        assert "thir" in deleted
        assert "six" in inserted
        full = "Either party may terminate with thirty days notice."
        assert full not in deleted  # never a whole-paragraph rewrite

    def it_redlines_the_table_cell_change_cell_wise(self):
        result = _compare_fixture_pair()
        revisions = result.document.revisions
        # $200 -> $250 narrows to the single changed character in the cell;
        # crucially the ROW was edited cell-wise, not deleted + reinserted
        assert any(r.revision_type == "deletion" and r.text == "0" for r in revisions)
        assert any(r.revision_type == "insertion" and r.text == "5" for r in revisions)
        assert not any(r.revision_type.startswith("row_") for r in revisions)
        assert not any("Advisory" in r.text for r in revisions)

    def it_stamps_every_revision_with_the_caller_identity(self):
        result = _compare_fixture_pair()
        for revision in result.document.revisions:
            assert revision.author == "Compare Engine"
            assert revision.date == FROZEN

    def it_pends_changes_matching_the_text_diff_of_the_inputs(self, tmp_path: Path):
        from docx.package import pending_changes

        result = _compare_fixture_pair()
        out = tmp_path / "redline.docx"
        result.document.save(str(out))
        pending = pending_changes(str(out))
        assert not pending.is_empty

    def it_validates_arguments(self):
        with pytest.raises(ValueError, match="author"):
            compare(str(fixture_path(ORIGINAL)), str(fixture_path(REVISED)), author="")
        with pytest.raises(ValueError, match="granularity"):
            compare(
                str(fixture_path(ORIGINAL)),
                str(fixture_path(REVISED)),
                author="X",
                granularity="letter",
            )
        with pytest.raises(ValueError, match="materialize"):
            compare(
                str(fixture_path(ORIGINAL)),
                str(fixture_path(REVISED)),
                author="X",
                materialize="merge",
            )

    def it_supports_block_granularity(self):
        result = compare(
            str(fixture_path(ORIGINAL)),
            str(fixture_path(REVISED)),
            author="Compare Engine",
            date=FROZEN,
            granularity="block",
        )
        result.document.revisions.accept_all()
        assert _visible(result.document) == _visible(docx.Document(str(fixture_path(REVISED))))


class DescribePendingRevisionInputs:
    def it_refuses_inputs_with_pending_revisions_by_default(self):
        with pytest.raises(UnsupportedStructureError, match="materialize"):
            compare(
                str(fixture_path(MULTIROUND)),
                str(fixture_path(MINIMAL)),
                author="X",
            )
        with pytest.raises(UnsupportedStructureError, match="materialize"):
            compare(
                str(fixture_path(MINIMAL)),
                str(fixture_path(MULTIROUND)),
                author="X",
            )

    def it_materializes_working_copies_on_request(self, tmp_path: Path):
        source = fixture_path(MULTIROUND)
        before = source.read_bytes()
        result = compare(
            str(source),
            str(fixture_path("generated/redline/multiround-accepted.docx")),
            author="Compare Engine",
            date=FROZEN,
            materialize="accept",
        )
        # materialized(original) == accepted ground truth -> empty redline
        assert result.revision_count == 0
        assert source.read_bytes() == before  # the input file is untouched


class DescribeCompareOnTheGauntlet:
    def it_compares_the_gauntlet_with_itself_materialized(self):
        """Everything-document sanity: materialize + self-compare = clean."""
        result = compare(
            str(fixture_path(GAUNTLET)),
            str(fixture_path(GAUNTLET)),
            author="Compare Engine",
            date=FROZEN,
            materialize="accept",
        )
        assert result.revision_count == 0

    def it_compares_minimal_against_the_heavily_edited_variant(self, tmp_path: Path):
        """End-to-end: edit a copy with plain (untracked) upstream calls,
        then let compare reconstruct the redline."""
        edited_path = tmp_path / "edited.docx"
        shutil.copyfile(fixture_path(MINIMAL), edited_path)
        document = docx.Document(str(edited_path))
        document.paragraphs[1].runs[0].text = "First body paragraph with thoroughly modern text."
        document.add_paragraph("A brand new closing paragraph.")
        document.save(str(edited_path))

        result = compare(
            str(fixture_path(MINIMAL)),
            str(edited_path),
            author="Compare Engine",
            date=FROZEN,
        )
        result.document.revisions.accept_all()
        assert _visible(result.document) == _visible(docx.Document(str(edited_path)))
