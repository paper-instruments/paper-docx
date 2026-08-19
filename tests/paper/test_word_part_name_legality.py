"""Part-name legality, pinned to measured Microsoft Word verdicts.

Word was shown seven spellings of the same renamed media part. One rule accounts
for every verdict: a part name may contain only ASCII characters that are legal
in a URI path, and a percent escape is acceptable only where it encodes such a
character. ``[Content_Types].xml`` is a reserved package item rather than an OPC
part name and is exempt — it carries ``[`` and ``]``, neither of which is legal.

Ledger: ``verifying-against-word/WORD-VERDICTS.md``, "Member names".
"""

from __future__ import annotations

import unicodedata
from pathlib import Path

import pytest

import docx
from docx import _paperpkg
from docx._zipguard import _validate_member_name
from docx.errors import PackageLimitError

from .test_audit_zip_preflight import _write_archive

_NFC_NAME = "word/media/" + unicodedata.normalize("NFC", "imagé1.png")
_NFD_NAME = "word/media/" + unicodedata.normalize("NFD", "imagé1.png")


class DescribeWordVerdictMemberNames:
    """One test per row of the ledger's seven-row member-name table."""

    def it_accepts_a_plain_ascii_renamed_media_part(self):
        # Word: OPENS (the rename mechanism control)
        _validate_member_name("word/media/renamed1.png")

    def it_accepts_a_percent_escaped_ascii_space(self):
        # Word: OPENS. This is the row a literal-character pass placed ahead of
        # the escape walk breaks, because "%" is not itself a legal literal.
        _validate_member_name("word/media/my%20image.png")

    def it_refuses_a_literal_space_naming_the_character(self):
        # Word: REFUSES
        with pytest.raises(PackageLimitError) as exc:
            _validate_member_name("word/media/my image.png")
        message = str(exc.value)
        assert "' '" in message
        assert "not legal in an OPC part name" in message

    def it_refuses_a_raw_non_ascii_name_in_composed_form(self):
        # Word: REFUSES
        with pytest.raises(PackageLimitError) as exc:
            _validate_member_name(_NFC_NAME)
        assert "non-ASCII character" in str(exc.value)

    def it_refuses_a_raw_non_ascii_name_in_decomposed_form(self):
        # Word: REFUSES. Refused for being non-ASCII, not for its normal form.
        with pytest.raises(PackageLimitError) as exc:
            _validate_member_name(_NFD_NAME)
        message = str(exc.value)
        assert "non-ASCII character" in message
        assert "Unicode-normalized" not in message

    def it_refuses_a_percent_escaped_composed_non_ascii_name(self):
        # Word: REFUSES
        with pytest.raises(PackageLimitError) as exc:
            _validate_member_name("word/media/imag%C3%A91.png")
        message = str(exc.value)
        assert "percent-escapes the non-ASCII byte 0xC3" in message
        assert "may escape only ASCII characters" in message

    def it_refuses_a_percent_escaped_decomposed_non_ascii_name(self):
        # Word: REFUSES
        with pytest.raises(PackageLimitError) as exc:
            _validate_member_name("word/media/image%CC%811.png")
        assert "percent-escapes the non-ASCII byte 0xCC" in str(exc.value)


class DescribeContentTypesExemption:
    def it_accepts_the_content_types_stream_despite_its_brackets(self):
        _validate_member_name("[Content_Types].xml")

    def it_still_reopens_a_document_it_just_saved(self, tmp_path: Path):
        path = tmp_path / "roundtrip.docx"
        document = docx.Document()
        document.add_paragraph("exempt")
        document.save(path)

        reopened = docx.Document(path)

        assert [paragraph.text for paragraph in reopened.paragraphs] == ["exempt"]

    def it_exempts_a_case_variant_spelling_of_the_same_package_item(self):
        # A case-variant content-types stream is the same package item, and must
        # keep reaching its own content-types diagnosis rather than be renamed a
        # character defect by this rule.
        _validate_member_name("[CONTENT_TYPES].XML")

    def it_does_not_exempt_brackets_in_any_other_member(self):
        with pytest.raises(PackageLimitError) as exc:
            _validate_member_name("word/[Content_Types].xml")
        assert "not legal in an OPC part name" in str(exc.value)


class DescribeMemberNameMessagesAreNotShadowed:
    def it_still_names_a_control_character_rather_than_an_illegal_literal(self):
        with pytest.raises(PackageLimitError) as exc:
            _validate_member_name("word/med\x01ia.png")
        assert "contains a control character" in str(exc.value)

    def it_still_names_a_malformed_escape_rather_than_an_illegal_literal(self):
        with pytest.raises(PackageLimitError) as exc:
            _validate_member_name("word/media/x%ZZ.png")
        assert "noncanonical percent escape" in str(exc.value)

    def it_still_names_an_unsafe_escape_of_an_unreserved_character(self):
        with pytest.raises(PackageLimitError) as exc:
            _validate_member_name("word/media/x%41.png")
        assert "unsafe percent escape" in str(exc.value)


class DescribePartNameLegalityOnRead:
    def it_refuses_a_package_carrying_a_literal_space_in_a_member_name(
        self, tmp_path: Path
    ):
        path = tmp_path / "spaced.docx"
        _write_archive(path, (("word/media/my image.png", b"payload"),))

        with pytest.raises(PackageLimitError, match="not legal in an OPC part name"):
            _paperpkg._read_zip(path)

    def it_reads_a_package_whose_member_name_escapes_an_ascii_space(
        self, tmp_path: Path
    ):
        path = tmp_path / "escaped.docx"
        _write_archive(path, (("word/media/my%20image.png", b"payload"),))

        parts, _ = _paperpkg._read_zip(path)

        assert parts == {"word/media/my%20image.png": b"payload"}
