"""Regressions for ZIP preflight structural checks."""

from __future__ import annotations

import struct
import zipfile
from pathlib import Path

import pytest

import docx
from docx import _paperpkg
from docx.errors import PackageLimitError
from docx.opc import phys_pkg

_END_RECORD = struct.Struct("<4s4H2LH")
_ZIP64_END_RECORD = struct.Struct("<4sQ2H2L4Q")
_ZIP64_LOCATOR = struct.Struct("<4sLQL")


def _write_archive(path: Path, members: tuple[tuple[str, bytes], ...]) -> None:
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_STORED) as archive:
        for name, data in members:
            archive.writestr(name, data)


def _zip64_count_only_archive(member_count: int) -> bytes:
    zip64_end = _ZIP64_END_RECORD.pack(
        b"PK\x06\x06",
        44,
        45,
        45,
        0,
        0,
        member_count,
        member_count,
        0,
        0,
    )
    locator = _ZIP64_LOCATOR.pack(b"PK\x06\x07", 0, 0, 1)
    legacy_end = _END_RECORD.pack(
        b"PK\x05\x06",
        0,
        0,
        0xFFFF,
        0xFFFF,
        0xFFFFFFFF,
        0xFFFFFFFF,
        0,
    )
    return zip64_end + locator + legacy_end


class DescribeCentralDirectoryPreflight:
    def it_opens_package_kernel_paths_once_before_preflight(
        self,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ):
        path = tmp_path / "selected.zip"
        replacement = tmp_path / "replacement.zip"
        _write_archive(path, (("selected.bin", b"selected"),))
        _write_archive(replacement, (("replacement.bin", b"replacement"),))
        original_preflight = _paperpkg.preflight_zip

        def preflight_then_swap(source: object) -> None:
            original_preflight(source)
            replacement.replace(path)

        monkeypatch.setattr(_paperpkg, "preflight_zip", preflight_then_swap)

        parts, _ = _paperpkg._read_zip(path)

        assert parts == {"selected.bin": b"selected"}

    def it_opens_document_paths_once_before_preflight(
        self,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ):
        path = tmp_path / "selected.docx"
        replacement = tmp_path / "replacement.docx"
        selected = docx.Document()
        selected.add_paragraph("selected")
        selected.save(path)
        other = docx.Document()
        other.add_paragraph("replacement")
        other.save(replacement)
        original_preflight = phys_pkg.preflight_zip

        def preflight_then_swap(source: object) -> None:
            original_preflight(source)
            replacement.replace(path)

        monkeypatch.setattr(phys_pkg, "preflight_zip", preflight_then_swap)

        reopened = docx.Document(path)

        assert [paragraph.text for paragraph in reopened.paragraphs] == ["selected"]

    def it_refuses_a_zip64_count_that_cannot_fit_in_the_central_directory(
        self,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ):
        path = tmp_path / "too-many.docx"
        path.write_bytes(_zip64_count_only_archive(100))
        constructions = 0

        class UnexpectedZipFile:
            def __init__(self, *_args, **_kwargs):
                nonlocal constructions
                constructions += 1
                raise AssertionError("ZipFile must not be constructed before count refusal")

        monkeypatch.setattr(_paperpkg.zipfile, "ZipFile", UnexpectedZipFile)

        with pytest.raises(PackageLimitError, match="too small for its member count"):
            _paperpkg._read_zip(path)

        assert constructions == 0

    def it_preflights_an_ordinary_document_open_before_constructing_zipfile(
        self,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ):
        path = tmp_path / "too-many-on-open.docx"
        path.write_bytes(_zip64_count_only_archive(100))
        constructions = 0

        class UnexpectedZipFile:
            def __init__(self, *_args, **_kwargs):
                nonlocal constructions
                constructions += 1
                raise AssertionError("ZipFile must not be constructed before count refusal")

        monkeypatch.setattr(phys_pkg, "ZipFile", UnexpectedZipFile)

        with pytest.raises(PackageLimitError, match="too small for its member count"):
            docx.Document(path)

        assert constructions == 0

    def it_scans_records_instead_of_trusting_a_forged_low_count(
        self,
        tmp_path: Path,
        monkeypatch: pytest.MonkeyPatch,
    ):
        path = tmp_path / "forged-count.docx"
        _write_archive(path, (("one.bin", b"1"), ("two.bin", b"2")))
        data = bytearray(path.read_bytes())
        end_offset = data.rfind(b"PK\x05\x06")
        struct.pack_into("<HH", data, end_offset + 8, 1, 1)
        path.write_bytes(data)
        constructions = 0

        class UnexpectedZipFile:
            def __init__(self, *_args, **_kwargs):
                nonlocal constructions
                constructions += 1
                raise AssertionError("ZipFile must not be constructed before count refusal")

        monkeypatch.setattr(_paperpkg.zipfile, "ZipFile", UnexpectedZipFile)

        with pytest.raises(PackageLimitError, match="member count"):
            _paperpkg._read_zip(path)

        assert constructions == 0


_END_RECORD = struct.Struct("<4s4H2LH")
_CENTRAL_HEADER = struct.Struct("<4s6H3L5H2L")


def _prefixed_archive(clean: bytes, prefix: bytes) -> bytes:
    """Return `clean` with `prefix` in front and every offset rebased.

    The result is a fully self-consistent ZIP: the end record points at the real
    central-directory position and each central record points at its real local header. So it
    passes every structural check in the preflight and is refused only for not beginning at
    byte 0 -- which is what makes the verdict attributable.
    """
    body = bytearray(prefix + clean)
    end_offset = clean.rfind(b"PK\x05\x06")
    fields = list(_END_RECORD.unpack_from(clean, end_offset))
    central_size, central_offset = fields[5], fields[6]
    cursor = central_offset + len(prefix)
    limit = cursor + central_size
    while cursor < limit:
        record = _CENTRAL_HEADER.unpack_from(bytes(body), cursor)
        struct.pack_into("<L", body, cursor + 42, record[16] + len(prefix))
        cursor += _CENTRAL_HEADER.size + record[10] + record[11] + record[12]
    fields[6] = central_offset + len(prefix)
    _END_RECORD.pack_into(body, end_offset + len(prefix), *fields)
    return bytes(body)


class DescribeArchiveMustBeginAtByteZero:
    """Word refuses a package with bytes in front of it (Word for Mac, 2026-08-18).

    Every reader below Word accepts such a file and reports the correct content, so nothing but
    a guard here can tell the caller the document is unusable.
    """

    @staticmethod
    def _clean(tmp_path: Path) -> bytes:
        source = tmp_path / "clean.docx"
        docx.Document().save(str(source))
        return source.read_bytes()

    def it_refuses_a_self_consistent_prefixed_archive(self, tmp_path: Path):
        target = tmp_path / "prefixed.docx"
        target.write_bytes(_prefixed_archive(self._clean(tmp_path), b"# self-extracting stub\n"))

        with pytest.raises(PackageLimitError, match="does not begin with a ZIP local file header"):
            docx.Document(str(target))

    def it_opens_the_same_document_without_the_prefix(self, tmp_path: Path):
        """The control that makes the refusal above attributable to the prefix alone."""
        target = tmp_path / "clean-control.docx"
        target.write_bytes(self._clean(tmp_path))

        assert isinstance(docx.Document(str(target)), docx.document.Document)

    def it_still_opens_an_archive_carrying_a_declared_comment(self, tmp_path: Path):
        """Bytes AFTER the archive, declared correctly, stay legal.

        Separates this rule from the end-of-central-directory rule: one is about where the
        archive starts, the other about what follows it.
        """
        target = tmp_path / "commented.docx"
        target.write_bytes(self._clean(tmp_path))
        with zipfile.ZipFile(target, "a") as archive:
            archive.comment = b"produced by an internal build pipeline"

        assert isinstance(docx.Document(str(target)), docx.document.Document)

    def it_refuses_a_prefixed_original_through_patch_save(self, tmp_path: Path):
        """`patch_save` reads the original through the same preflight.

        This is the coupling that makes a verbatim-copy gate unnecessary: the no-op byte-copy
        path used to reproduce a prefixed original verbatim, and it can no longer be reached
        with one.
        """
        original = tmp_path / "prefixed.docx"
        original.write_bytes(_prefixed_archive(self._clean(tmp_path), b"#" * 64))

        with pytest.raises(PackageLimitError, match="does not begin with a ZIP local file header"):
            _paperpkg.patch_save(original, docx.Document(), tmp_path / "out.docx")
