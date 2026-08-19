"""Word's content-type rule is about core parts, not media parts.

Measured in Word (see ``verifying-against-word/WORD-VERDICTS.md``, "Content types"):

- ``image_octet_stream`` — a displayed image declared ``application/octet-stream`` —
  **OPENS**.
- ``styles_generic_ct`` — ``word/styles.xml`` declared ``application/xml`` — **REFUSES**.

The pair is what documents the split. Either test alone invites re-tightening the media
side or loosening the core side.
"""

from __future__ import annotations

import io
import zipfile

import pytest

import docx
from docx.errors import PackageLimitError
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from .test_composition import TINY_PNG
from .test_practical_opc_hardening import _rewrite_package

_STYLES_CT = CT.WML_STYLES.encode()


def _content_types(data: bytes) -> bytes:
    with zipfile.ZipFile(io.BytesIO(data)) as package:
        return package.read("[Content_Types].xml")


def _document_with_image_bytes() -> bytes:
    document = docx.Document()
    document.add_paragraph("before the image")
    document.add_picture(io.BytesIO(TINY_PNG))
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _image_blobs(document) -> list[bytes]:
    return [
        rel.target_part.blob
        for rel in document.part.rels.values()
        if not rel.is_external and rel.reltype == RT.IMAGE
    ]


def it_opens_a_displayed_image_declared_application_octet_stream():
    def transform(name: str, blob: bytes):
        if name == "[Content_Types].xml":
            blob = blob.replace(
                b'ContentType="image/png"', b'ContentType="application/octet-stream"'
            )
        return name, blob

    data = _rewrite_package(_document_with_image_bytes(), transform)
    assert b'Extension="png" ContentType="application/octet-stream"' in _content_types(data)

    document = docx.Document(io.BytesIO(data))

    assert document.paragraphs[0].text == "before the image"
    assert len(document.inline_shapes) == 1
    assert _image_blobs(document) == [TINY_PNG]


def it_still_refuses_a_core_part_declared_with_a_generic_content_type():
    def transform(name: str, blob: bytes):
        if name == "[Content_Types].xml":
            blob = blob.replace(_STYLES_CT, b"application/xml")
        return name, blob

    data = _rewrite_package(_document_with_image_bytes(), transform)
    assert b'PartName="/word/styles.xml" ContentType="application/xml"' in _content_types(data)

    with pytest.raises(PackageLimitError, match="invalid content type"):
        docx.Document(io.BytesIO(data))
