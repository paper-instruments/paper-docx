"""Regressions for comment ownership and anchor-id preflights."""

from __future__ import annotations

import pytest

import docx
from docx.commentops import (
    COMMENTS_EXTENDED_RELATIONSHIP_TYPE,
    anchored_text,
    comment_thread,
    is_resolved,
    parent_of,
    reply,
    resolve,
)
from docx.errors import PaperRefusal, UnsupportedStructureError
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import qn
from docx.search import find_one
from docx.shared import Inches

from .harness.contract import assert_refusal_atomic

_W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"
_W15_PARA_ID = f"{{{_W15_NS}}}paraId"
_W15_PARA_ID_PARENT = f"{{{_W15_NS}}}paraIdParent"


def _commented_document():
    document = docx.Document()
    document.add_paragraph("comment target")
    comment = find_one(document, "comment target").comment(
        "Parent", author="Reviewer"
    )
    return document, comment


def _assert_parent_reply_topology(document, parent_id: int, child_id: int):
    comments = {comment.comment_id: comment for comment in document.comments}
    assert set(comments) == {parent_id, child_id}
    assert anchored_text(document, comments[parent_id]) == "comment target"
    assert anchored_text(document, comments[child_id]) == "comment target"

    marker_order = [
        (marker.tag.rsplit("}", 1)[-1], int(marker.get(qn("w:id"))))
        for marker in document.element.body.iter(
            qn("w:commentRangeStart"),
            qn("w:commentRangeEnd"),
            qn("w:commentReference"),
        )
    ]
    assert marker_order == [
        ("commentRangeStart", parent_id),
        ("commentRangeStart", child_id),
        ("commentRangeEnd", child_id),
        ("commentRangeEnd", parent_id),
        ("commentReference", parent_id),
        ("commentReference", child_id),
    ]

    comments_root = document.part.part_related_by(RT.COMMENTS)._element  # noqa: SLF001
    comment_elements = {
        int(element.get(qn("w:id"))): element for element in comments_root
    }
    para_ids = {
        comment_id: list(comment_elements[comment_id].iter(qn("w:p")))[-1].get(
            qn("w14:paraId")
        )
        for comment_id in (parent_id, child_id)
    }
    assert all(para_ids.values())

    extended_root = document.part.part_related_by(
        COMMENTS_EXTENDED_RELATIONSHIP_TYPE
    )._element  # noqa: SLF001
    entries = {entry.get(_W15_PARA_ID): entry for entry in extended_root}
    assert entries[para_ids[parent_id]].get(_W15_PARA_ID_PARENT) is None
    assert (
        entries[para_ids[child_id]].get(_W15_PARA_ID_PARENT)
        == para_ids[parent_id]
    )
    return para_ids[parent_id], para_ids[child_id]


class DescribeCommentRelationshipOwnership:
    @pytest.mark.parametrize("operation", ["resolve", "thread"])
    def it_refuses_multiple_comments_relationships_atomically(
        self, operation: str
    ):
        document, comment = _commented_document()
        package = document.part.package
        assert package is not None
        duplicate = Part(
            PackURI("/word/comments-duplicate.xml"),
            CT.WML_COMMENTS,
            b"<comments/>",
            package,
        )
        document.part.relate_to(duplicate, RT.COMMENTS)

        error = assert_refusal_atomic(
            document,
            lambda candidate: (
                resolve(candidate, comment)
                if operation == "resolve"
                else comment_thread(candidate)
            ),
            PaperRefusal,
        )

        assert isinstance(error, UnsupportedStructureError)
        assert "multiple comments relationships" in str(error)

    def it_refuses_multiple_main_document_relationships_before_mutation(self):
        document, comment = _commented_document()
        package = document.part.package
        assert package is not None
        package.rels.add_relationship(
            RT.OFFICE_DOCUMENT,
            document.part,
            "rIdDuplicateMain",
        )
        before = comment._comment_elm.xml  # noqa: SLF001

        with pytest.raises(UnsupportedStructureError, match="multiple main-document"):
            comment.add_paragraph("must not be added")

        assert comment._comment_elm.xml == before  # noqa: SLF001


class DescribeCommentContentMutation:
    def it_preserves_parent_resolution_and_child_link_when_appending_text(self):
        document, parent = _commented_document()
        child = reply(
            document,
            parent,
            "Reply",
            author="Second Reviewer",
        )
        resolve(document, parent)

        parent.add_paragraph("Additional context")

        assert is_resolved(document, parent)
        assert parent_of(document, child) == parent.comment_id

    def it_preserves_a_reply_link_when_appending_a_table(self):
        document, parent = _commented_document()
        child = reply(
            document,
            parent,
            "Reply",
            author="Second Reviewer",
        )

        child.add_table(1, 1, Inches(1))

        assert parent_of(document, child) == parent.comment_id


class DescribeCommentInspection:
    def it_does_not_create_a_comments_part_for_an_uncommented_document(self):
        document = docx.Document()
        package = document.part.package
        assert package is not None
        relationships_before = tuple(document.part.rels)
        parts_before = tuple(part.partname for part in package.iter_parts())

        assert comment_thread(document) == ()

        assert tuple(document.part.rels) == relationships_before
        assert tuple(part.partname for part in package.iter_parts()) == parts_before


class DescribeCommentAnchorIds:
    def it_retains_conformant_parent_and_reply_anchors_after_round_trip(
        self, tmp_path
    ):
        document, parent = _commented_document()
        child = reply(document, parent, "Reply", author="Second Reviewer")
        assert parent.comment_id != child.comment_id

        # Microsoft's pinned SDK test gives each threaded comment its own
        # anchor triple and links the child through CommentEx.ParaIdParent:
        # https://github.com/dotnet/Open-XML-SDK/blob/431ab05cf160248cc3885a4a766026d4f8243792/test/DocumentFormat.OpenXml.Tests/ConformanceTest/CommentEx/TestEntities.cs#L89-L115
        para_ids = _assert_parent_reply_topology(
            document, parent.comment_id, child.comment_id
        )

        output = tmp_path / "threaded-comments.docx"
        document.save(str(output))
        reopened = docx.Document(str(output))
        assert (
            _assert_parent_reply_topology(
                reopened, parent.comment_id, child.comment_id
            )
            == para_ids
        )

    def it_compares_anchor_marker_ids_numerically(self):
        document, comment = _commented_document()
        comment._comment_elm.set(qn("w:id"), "1")  # noqa: SLF001
        for marker in document.element.body.iter(
            qn("w:commentRangeStart"),
            qn("w:commentRangeEnd"),
            qn("w:commentReference"),
        ):
            marker.set(qn("w:id"), "01")

        assert anchored_text(document, comment) == "comment target"
        child = reply(document, comment, "Reply", author="Second Reviewer")
        assert parent_of(document, child) == 1

    @pytest.mark.parametrize(
        "marker_name",
        ["commentRangeStart", "commentRangeEnd", "commentReference"],
    )
    @pytest.mark.parametrize("bad_id", [None, "broken"])
    def it_refuses_invalid_anchor_marker_ids_atomically(
        self, marker_name: str, bad_id: str | None
    ):
        document, comment = _commented_document()
        marker = next(document.element.body.iter(qn(f"w:{marker_name}")))
        if bad_id is None:
            del marker.attrib[qn("w:id")]
        else:
            marker.set(qn("w:id"), bad_id)

        error = assert_refusal_atomic(
            document,
            lambda candidate: reply(
                candidate, comment, "Reply", author="Second Reviewer"
            ),
            PaperRefusal,
        )

        assert isinstance(error, UnsupportedStructureError)
        assert f"malformed {marker_name} comment marker" in str(error)
