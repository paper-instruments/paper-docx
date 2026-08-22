"""Tests for docx.search: find_text/Span, run-preserving replace, and
tracked replace over the revision vocabulary.

Span mapping is tested BY USE (perform a replace, assert the outcome) —
never by inspecting private offsets.
"""

from __future__ import annotations

import datetime as dt
import shutil
from pathlib import Path

import pytest
from lxml import etree

import docx
import docx._clock
from docx.errors import (
    AmbiguousTargetError,
    BoundaryViolationError,
    TargetNotFoundError,
    UnsupportedStructureError,
)
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.search import find_one, find_text, normalize_text
from docx.story import iter_blocks

from .harness.contract import assert_changed_parts, assert_refusal_atomic, save_and_reopen
from .harness.paths import fixture_path

FRAGMENTED = "generated/feature-isolated/fragmented-runs.docx"
TRACKED = "generated/feature-isolated/tracked-ins-del.docx"
CONTROLS = "generated/feature-isolated/content-control.docx"
GAUNTLET = "generated/gauntlet/gauntlet.docx"
MINIMAL = "generated/minimal-clean/minimal.docx"

RATE_TEXT = "$75–100/hr on a “full-service” basis"
FROZEN = dt.datetime(2026, 7, 7, 12, 0, 0, tzinfo=dt.timezone.utc)
W = nsdecls("w")


def _doc(relpath: str):
    return docx.Document(str(fixture_path(relpath)))


class DescribeNormalizeText:
    @pytest.mark.parametrize(
        ("raw", "expected"),
        [
            ("‘quoted’", "'quoted'"),
            ("“quoted”", '"quoted"'),
            ("75–100", "75-100"),
            ("em—dash", "em-dash"),
            ("minus−sign", "minus-sign"),
            ("no break", "no break"),  # NBSP
            ("figure space", "figure space"),
            ("narrow break", "narrow break"),
            ("thin space", "thin space"),
            ("soft­hyphen", "softhyphen"),
            ("tab\there", "tab here"),
            ("cr\rhere", "cr here"),
            ("many   spaces\n\n here", "many spaces here"),
            ("CaseFOLD", "casefold"),
        ],
    )
    def it_applies_each_pinned_normalization_rule(self, raw: str, expected: str):
        assert normalize_text(raw) == expected


class DescribeFindText:
    def it_matches_ascii_needles_against_smart_characters_across_runs(self):
        span = find_one(_doc(FRAGMENTED), '$75-100/hr on a "full-service" basis')
        assert span.text == RATE_TEXT  # raw text captured verbatim, 8 runs deep
        assert not span.crosses_paragraphs

    def it_matches_through_no_break_spaces(self):
        span = find_one(_doc(FRAGMENTED), "Net 30 payment terms")
        assert "Net 30" in span.text  # captured text preserves the raw NBSP

    def it_matches_across_a_paragraph_boundary(self):
        spans = find_text(_doc(MINIMAL), "ordinary text. Second body paragraph")
        assert len(spans) == 1
        assert spans[0].crosses_paragraphs

    def it_returns_matches_in_document_order_with_nth_selection(self):
        document = _doc(TRACKED)
        matches = find_text(document, "Paragraph")
        assert len(matches) == 2
        assert find_text(document, "Paragraph", nth=2)[0].text == matches[1].text
        assert find_text(document, "Paragraph", nth=3) == []

    def it_scopes_to_a_story_part(self):
        document = _doc(GAUNTLET)
        everywhere = find_text(document, "Gauntlet header, section one")
        assert {span.story for span in everywhere} == {"word/header1.xml"}
        assert find_text(document, "Gauntlet header, section one",
                         story="word/document.xml") == []

    def it_ranks_by_proximity_to_the_near_text(self):
        document = _doc(GAUNTLET)
        span = find_text(document, "Gauntlet numbered item", near="item two", nth=1)[0]
        follow_on = span.text  # nearest match to "item two"
        assert follow_on == "Gauntlet numbered item"
        # prove it selected the second occurrence: replace it and look
        span.replace("Gauntlet renumbered item")
        texts = [b.text for b in iter_blocks(document)]
        assert "Gauntlet numbered item one" in texts
        assert "Gauntlet renumbered item two" in texts

    def it_honors_the_view_parameter(self):
        document = _doc(TRACKED)
        assert find_text(document, "forty-two") == []  # deleted text, current view
        assert len(find_text(document, "forty-two", view="original")) == 1
        assert len(find_text(document, "forty-two", view="all")) == 1
        assert find_text(document, "forty-seven", view="original") == []

    def it_finds_text_inside_text_boxes_and_controls(self):
        document = _doc(GAUNTLET)
        boxed = find_one(document, "Text living inside the text box.")
        assert boxed.in_text_box
        controlled = find_one(document, "controlled text")
        assert controlled.in_content_control


class DescribeFindOne:
    def it_refuses_zero_matches(self):
        with pytest.raises(TargetNotFoundError, match="no match"):
            find_one(_doc(MINIMAL), "text that does not exist anywhere")

    def it_refuses_ambiguity_without_disambiguators(self):
        with pytest.raises(AmbiguousTargetError, match="disambiguate"):
            find_one(_doc(TRACKED), "Paragraph")

    def it_resolves_ambiguity_with_nth(self):
        span = find_one(_doc(TRACKED), "Paragraph", nth=1)
        assert span.text == "Paragraph"


class DescribePlainReplace:
    def it_preserves_untouched_run_formatting(self, tmp_path: Path):
        document = _doc(FRAGMENTED)
        find_one(document, "$75-100/hr").replace("$85–110/hr")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        paragraph = reopened.paragraphs[0]
        assert paragraph.text == (
            "Consulting rate: $85–110/hr on a “full-service” basis"
            " — travel time billed at $37.50/hr."
        )
        italic_runs = [r.text for r in paragraph.runs if r.italic]
        assert "".join(italic_runs) == "“full-service”"  # untouched formatting island
        bold_text = "".join(r.text for r in paragraph.runs if r.bold)
        assert bold_text.startswith("$85")

    def it_survives_a_bold_to_italic_formatting_transition(self, tmp_path: Path):
        document = _doc(FRAGMENTED)
        find_one(document, '100/hr on a "full-').replace("90/hr on any “full-")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert "90/hr on any “full-service”" in reopened.paragraphs[0].text

    def it_restores_text_and_formatting_when_inverting_a_uniform_span(
        self, tmp_path: Path
    ):
        """Invariant: replace(x->y) then (y->x) restores text and formatting.

        Holds fully for spans of uniform formatting (here: the italic
        island, split across two runs)."""
        document = _doc(FRAGMENTED)
        find_one(document, "full-service").replace("bespoke")
        find_one(document, "bespoke").replace("full-service")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        paragraph = reopened.paragraphs[0]
        assert paragraph.text == (
            "Consulting rate: $75–100/hr on a “full-service” basis"
            " — travel time billed at $37.50/hr."
        )
        italic_text = "".join(r.text for r in paragraph.runs if r.italic)
        assert italic_text == "“full-service”"
        bold_text = "".join(r.text for r in paragraph.runs if r.bold)
        assert bold_text == "$75–100/hr"

    def it_restores_text_and_outside_formatting_for_mixed_spans(
        self, tmp_path: Path
    ):
        """A span covering a formatting transition collapses ITS OWN interior
        formatting into the start run when replaced — that information is
        destroyed by any replacement. The
        inverse still restores the visible text exactly and never disturbs
        formatting outside the span."""
        document = _doc(FRAGMENTED)
        find_one(document, RATE_TEXT).replace("something else entirely")
        find_one(document, "something else entirely").replace(RATE_TEXT)
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        paragraph = reopened.paragraphs[0]
        assert paragraph.text == (
            "Consulting rate: $75–100/hr on a “full-service” basis"
            " — travel time billed at $37.50/hr."
        )
        # outside the span, formatting is untouched
        plain_runs = [r.text for r in paragraph.runs if not r.bold and not r.italic]
        assert "Consulting rate: " in plain_runs[0]
        assert any("travel time billed" in text for text in plain_runs)

    def it_keeps_the_changed_part_budget_to_the_document_part(self, tmp_path: Path):
        source = fixture_path(FRAGMENTED)
        working = tmp_path / "work.docx"
        shutil.copyfile(source, working)
        document = docx.Document(str(working))
        find_one(document, "$75-100/hr").replace("$95–120/hr")
        out = tmp_path / "out.docx"
        docx.package.patch_save(working, document, out)
        assert_changed_parts(working, out, {"word/document.xml"})


class DescribePreservationPolicies:
    def _insertion_document(self, *, revision_id: str = "41"):
        document = docx.Document()
        insertion = parse_xml(
            f'<w:ins {W} w:id="{revision_id}" w:author="Alice"'
            ' w:date="2026-07-07T12:00:00Z">'
            "<w:r><w:t>pending wording</w:t></w:r></w:ins>"
        )
        document.add_paragraph()._p.append(insertion)
        return document, insertion

    def it_preserves_an_existing_insertion_and_its_projections(self, tmp_path: Path):
        document, insertion = self._insertion_document()
        attributes = dict(insertion.attrib)
        original = [b.text for b in iter_blocks(document, view="original")]
        result = find_one(document, "pending").replace(
            "revised", preserve_revision=True
        )
        assert result.preserved_revision_ids == (41,)
        assert not result.preserved_structure
        assert dict(insertion.attrib) == attributes
        assert [b.text for b in iter_blocks(document, view="original")] == original
        path = tmp_path / "preserved-insertion.docx"
        document.save(path)
        accepted = docx.Document(path)
        accepted.revisions.accept_all()
        assert "revised wording" in [b.text for b in iter_blocks(accepted)]
        rejected = docx.Document(path)
        rejected.revisions.reject_all()
        assert "revised wording" not in [b.text for b in iter_blocks(rejected)]

    def it_keeps_safe_insertion_refusal_as_the_default(self):
        document, _ = self._insertion_document()
        with pytest.raises(UnsupportedStructureError, match="pending tracked insertion"):
            find_one(document, "pending").replace("revised")

    def it_refuses_a_base_text_and_insertion_crossing(self):
        document, insertion = self._insertion_document()
        insertion.addprevious(parse_xml(f'<w:r {W}><w:t>base </w:t></w:r>'))
        with pytest.raises(UnsupportedStructureError, match="mixes base text"):
            find_one(document, "base pending").replace(
                "combined", preserve_revision=True
            )

    def it_refuses_a_revision_nested_inside_the_preserved_insertion(self):
        document = docx.Document()
        document.add_paragraph()._p.append(
            parse_xml(
                f'<w:ins {W} w:id="41" w:author="Alice"'
                ' w:date="2026-07-07T12:00:00Z">'
                '<w:r><w:t xml:space="preserve">pending </w:t></w:r>'
                '<w:del w:id="42" w:author="Bob" w:date="2026-07-07T12:00:00Z">'
                '<w:r><w:delText xml:space="preserve">dropped </w:delText>'
                "</w:r></w:del>"
                "<w:r><w:t>wording</w:t></w:r></w:ins>"
            )
        )
        # the deletion is hidden from the current view but sits BETWEEN the
        # matched atoms; an in-place edit would leave it stranded
        with pytest.raises(UnsupportedStructureError, match="nested or mixed"):
            find_one(document, "pending wording").replace(
                "revised", preserve_revision=True
            )

    def it_refuses_a_span_crossing_two_insertions(self):
        document, insertion = self._insertion_document()
        insertion.addnext(
            parse_xml(
                f'<w:ins {W} w:id="55" w:author="Bob"'
                ' w:date="2026-07-07T12:00:00Z">'
                "<w:r><w:t> and more</w:t></w:r></w:ins>"
            )
        )
        with pytest.raises(UnsupportedStructureError, match="crosses multiple"):
            find_one(document, "wording and more").replace(
                "revised", preserve_revision=True
            )

    def it_refuses_a_tracked_move_destination(self):
        document = docx.Document()
        document.add_paragraph()._p.append(
            parse_xml(
                f'<w:moveTo {W} w:id="9" w:author="Alice"'
                ' w:date="2026-07-07T12:00:00Z">'
                "<w:r><w:t>moved wording</w:t></w:r></w:moveTo>"
            )
        )
        with pytest.raises(UnsupportedStructureError, match="tracked moves"):
            find_one(document, "moved wording").replace(
                "revised", preserve_revision=True
            )

    def it_refuses_preservation_outside_the_current_view(self):
        document, _ = self._insertion_document()
        span = find_text(document, "pending wording", view="all")[0]
        with pytest.raises(UnsupportedStructureError, match="view='current'"):
            span.replace("revised", preserve_revision=True)

    def it_preserves_the_exact_text_element_graph_and_distribution(self, tmp_path: Path):
        document = docx.Document()
        paragraph = document.add_paragraph()
        for text in ("alpha", " pay", "ment", " terms"):
            paragraph.add_run(text)
        paragraph.runs[1].bold = True
        proofing_marker = parse_xml(f'<w:proofErr {W} w:type="spellStart"/>')
        paragraph.runs[1]._r.addnext(proofing_marker)
        elements = tuple(paragraph._p.iter(qn("w:t")))
        runs = tuple(paragraph._p.iter(qn("w:r")))
        graph = tuple((e, e.tag, tuple(e.attrib.items()), tuple(e)) for e in elements)
        run_properties = tuple(
            etree.tostring(run.find(qn("w:rPr")))
            if run.find(qn("w:rPr")) is not None else None
            for run in runs
        )
        result = find_one(document, "alpha payment terms").replace(
            "alpha settlement terms", preserve_structure=True
        )
        assert result.preserved_structure
        assert [e.text for e in elements] == ["alpha", " set", "tlem", "ent terms"]
        assert tuple((e, e.tag, tuple(e.attrib.items()), tuple(e)) for e in elements) == graph
        assert tuple(paragraph._p.iter(qn("w:r"))) == runs
        assert proofing_marker.getparent() is paragraph._p
        assert tuple(
            etree.tostring(run.find(qn("w:rPr")))
            if run.find(qn("w:rPr")) is not None else None
            for run in runs
        ) == run_properties
        reopened = save_and_reopen(document, tmp_path / "exact.docx")
        assert reopened.paragraphs[-1].text == "alpha settlement terms"

    def it_keeps_empty_nodes_and_consumes_a_mutated_exact_span(self):
        document = docx.Document()
        paragraph = document.add_paragraph()
        for text in ("ab", "cd", "ef"):
            paragraph.add_run(text)
        elements = tuple(paragraph._p.iter(qn("w:t")))
        span = find_one(document, "abcdef")
        span.replace("x", preserve_structure=True)
        assert [e.text for e in elements] == ["x", "", ""]
        with pytest.raises(TargetNotFoundError, match="structure-preserving"):
            span.replace("again")
        assert find_one(document, "x").text == "x"

    def it_refuses_exact_edge_whitespace_without_changing_xml_space(self):
        document = docx.Document()
        paragraph = document.add_paragraph("plain")
        element = paragraph._p.find(".//" + qn("w:t"))
        assert element is not None
        assert element.get(qn("xml:space")) is None
        before = etree.tostring(paragraph._p)
        with pytest.raises(UnsupportedStructureError, match="edge whitespace"):
            find_one(document, "plain").replace(" plain", preserve_structure=True)
        assert etree.tostring(paragraph._p) == before

    def it_refuses_an_exact_plan_that_would_hollow_a_bookmark(self):
        document = docx.Document()
        paragraph = document.add_paragraph()
        first = paragraph.add_run("outside")
        start = parse_xml(f'<w:bookmarkStart {W} w:id="9" w:name="target"/>')
        first._r.addnext(start)
        inside = paragraph.add_run("inside")
        end = parse_xml(f'<w:bookmarkEnd {W} w:id="9"/>')
        inside._r.addnext(end)
        before = etree.tostring(paragraph._p)
        with pytest.raises(UnsupportedStructureError, match="hollow"):
            find_one(document, "outsideinside").replace(
                "x", preserve_structure=True
            )
        assert etree.tostring(paragraph._p) == before

    def it_refuses_hollowing_a_bookmark_whose_markers_span_paragraphs(self):
        document = docx.Document()
        first = document.add_paragraph()._p
        first.append(
            parse_xml(f'<w:bookmarkStart {W} w:id="10" w:name="target"/>')
        )
        document.add_paragraph("inside")
        last = document.add_paragraph()._p
        last.append(parse_xml(f'<w:bookmarkEnd {W} w:id="10"/>'))
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match="hollow"):
            find_one(document, "inside").replace("", preserve_structure=True)

        assert document.element.xml == before

    def it_keeps_a_revision_preserving_noop_reusable_across_a_bookmark(self):
        document = docx.Document()
        insertion = parse_xml(
            f'<w:ins {W} w:id="42" w:author="Alice">'
            '<w:r><w:t>outside</w:t></w:r>'
            '<w:bookmarkStart w:id="11" w:name="target"/>'
            '<w:r><w:t>inside</w:t></w:r>'
            '<w:bookmarkEnd w:id="11"/>'
            "</w:ins>"
        )
        document.add_paragraph()._p.append(insertion)
        span = find_one(document, "outsideinside")
        before = document.element.xml

        result = span.replace("outsideinside", preserve_revision=True)

        assert result.preserved_revision_ids == (42,)
        assert document.element.xml == before
        span.replace("outsideinside", preserve_revision=True)

    def it_leaves_an_exact_noop_reusable(self):
        document = docx.Document()
        document.add_paragraph("same")
        span = find_one(document, "same")
        before = document.element.xml
        assert span.replace("same", preserve_structure=True).preserved_structure
        assert document.element.xml == before
        span.replace("same", preserve_structure=True)

    def it_combines_revision_and_structure_preservation(self):
        document, insertion = self._insertion_document()
        text_element = next(insertion.iter(qn("w:t")))
        result = find_one(document, "pending").replace(
            "current", preserve_revision=True, preserve_structure=True
        )
        assert result.preserved_structure
        assert result.preserved_revision_ids == (41,)
        assert text_element.getparent() is not None

    def it_keeps_an_exact_patch_save_to_the_changed_story(self, tmp_path: Path):
        source = fixture_path(FRAGMENTED)
        working = tmp_path / "work.docx"
        shutil.copyfile(source, working)
        document = docx.Document(str(working))
        find_one(document, "$75-100/hr").replace(
            "$85–110/hr", preserve_structure=True
        )
        out = tmp_path / "out.docx"
        docx.package.patch_save(working, document, out)
        assert_changed_parts(working, out, {"word/document.xml"})

    @pytest.mark.parametrize("revision_id", ["bad", ""])
    def it_refuses_unreportable_insertion_ids(self, revision_id: str):
        document, insertion = self._insertion_document(revision_id=revision_id)
        if not revision_id:
            del insertion.attrib[qn("w:id")]
        with pytest.raises(UnsupportedStructureError, match="w:id"):
            find_one(document, "pending").replace("current", preserve_revision=True)

    @pytest.mark.parametrize("policy", ["preserve_structure", "preserve_revision"])
    def it_refuses_tracked_preservation_combinations(self, policy: str):
        document = docx.Document()
        document.add_paragraph("target")
        with pytest.raises(ValueError, match=policy):
            find_one(document, "target").replace(
                "changed", tracked=True, author="Editor", **{policy: True}
            )


class DescribeReplaceRefusals:
    def it_refuses_spans_over_deleted_text(self):
        document = _doc(TRACKED)
        span = find_one(document, "forty-two", view="all")
        with pytest.raises(UnsupportedStructureError, match="tracked-deleted"):
            span.replace("anything")

    def it_refuses_cross_paragraph_spans(self):
        document = _doc(MINIMAL)
        span = find_one(document, "ordinary text. Second body paragraph")
        with pytest.raises(BoundaryViolationError, match="paragraph boundary"):
            span.replace("anything")

    def it_refuses_spans_crossing_a_content_control_boundary(self):
        document = _doc(CONTROLS)
        span = find_one(document, "follows: controlled")
        with pytest.raises(BoundaryViolationError, match="content-control"):
            span.replace("anything")

    def it_refuses_stale_spans(self):
        document = _doc(MINIMAL)
        span = find_one(document, "perfectly ordinary text")
        find_one(document, "perfectly ordinary").replace("thoroughly mundane")
        with pytest.raises(TargetNotFoundError, match="stale"):
            span.replace("anything")

    def it_refuses_atomically(self):
        """A refused replace leaves no trace, in memory or on disk."""
        document = _doc(CONTROLS)
        span = find_one(document, "follows: controlled")
        assert_refusal_atomic(
            document,
            lambda doc: span.replace("anything"),
            BoundaryViolationError,
            on_disk=(fixture_path(CONTROLS),),
        )


class DescribeTrackedReplace:
    def it_marks_only_the_minimal_changed_span(self, tmp_path: Path):
        """The redline marks `75-10 -> 85-11`, not the sentence (pinned)."""
        document = _doc(FRAGMENTED)
        result = find_one(document, "$75-100/hr").replace(
            "$85–110/hr", tracked=True, author="Carol QA", date=FROZEN
        )
        assert result.deleted_text == "75–10"
        assert result.inserted_text == "85–11"
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        blocks = list(iter_blocks(reopened))
        assert "$85–110/hr" in blocks[0].text  # current view: change applied
        original = list(iter_blocks(reopened, view="original"))
        assert "$75–100/hr" in original[0].text  # original view: change absent

    def it_keeps_deleted_text_in_delText_never_live_wt(self, tmp_path: Path):
        document = _doc(FRAGMENTED)
        find_one(document, "$75-100/hr").replace(
            "$85–110/hr", tracked=True, author="Carol QA", date=FROZEN
        )
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        body = reopened.element.body
        assert body.xpath("//w:del//w:delText"), "deleted text must be in w:delText"
        assert not body.xpath("//w:del//w:t"), "live w:t inside w:del is corrupt"

    def it_allocates_unique_increasing_revision_ids(self):
        document = _doc(TRACKED)  # fixture already holds ids 11, 12, 21
        first = find_one(document, "Paragraph before").replace(
            "Paragraph just before", tracked=True, author="Carol QA", date=FROZEN
        )
        second = find_one(document, "Paragraph after").replace(
            "Paragraph right after", tracked=True, author="Carol QA", date=FROZEN
        )
        all_ids = [int(v) for v in document.element.body.xpath(
            "//w:ins/@w:id | //w:del/@w:id"
        )]
        assert len(all_ids) == len(set(all_ids)), "revision ids must be unique"
        assert min(first.revision_ids) > 21
        assert min(second.revision_ids) > max(first.revision_ids)

    def it_stamps_dates_from_the_injectable_clock(
        self, monkeypatch: pytest.MonkeyPatch
    ):
        monkeypatch.setattr(docx._clock, "now", lambda: FROZEN)
        document = _doc(MINIMAL)
        find_one(document, "perfectly ordinary").replace(
            "thoroughly mundane", tracked=True, author="Carol QA"
        )
        (ins,) = document.element.body.xpath("//w:ins")
        assert ins.get(qn("w:date")) == "2026-07-07T12:00:00Z"

    def it_preserves_run_formatting_on_both_sides(self, tmp_path: Path):
        document = _doc(FRAGMENTED)
        find_one(document, "$75-100/hr").replace(
            "$85–110/hr", tracked=True, author="Carol QA", date=FROZEN
        )
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        body = reopened.element.body
        # deleted text keeps each SOURCE run's rPr (one w:r per source run)
        del_rprs = body.xpath("//w:del/w:r/w:rPr")
        assert del_rprs, "deletion lost its runs"
        assert all(r.find(qn("w:b")) is not None for r in del_rprs), (
            "deleted side lost bold"
        )
        (ins_rpr,) = body.xpath("//w:ins/w:r/w:rPr")
        assert ins_rpr.find(qn("w:b")) is not None, "inserted side lost bold"

    def it_requires_an_author(self):
        document = _doc(MINIMAL)
        span = find_one(document, "perfectly ordinary")
        with pytest.raises(ValueError, match="author"):
            span.replace("x", tracked=True)

    def it_refuses_a_replacement_equal_to_the_existing_text(self):
        document = _doc(MINIMAL)
        span = find_one(document, "perfectly ordinary")
        with pytest.raises(TargetNotFoundError, match="nothing to change"):
            span.replace("perfectly ordinary", tracked=True, author="Carol QA")

    def it_refuses_cross_paragraph_tracked_targets(self):
        document = _doc(MINIMAL)
        span = find_one(document, "ordinary text. Second body paragraph")
        with pytest.raises(BoundaryViolationError):
            span.replace("anything", tracked=True, author="Carol QA")

    def it_keeps_the_package_clean_and_budgeted(self, tmp_path: Path):
        from .harness import checks

        source = fixture_path(FRAGMENTED)
        working = tmp_path / "work.docx"
        shutil.copyfile(source, working)
        document = docx.Document(str(working))
        find_one(document, "$75-100/hr").replace(
            "$85–110/hr", tracked=True, author="Carol QA", date=FROZEN
        )
        out = tmp_path / "out.docx"
        docx.package.patch_save(working, document, out)
        assert_changed_parts(working, out, {"word/document.xml"})
        checks.assert_package_facts_clean(out)

    @pytest.mark.lo_smoke
    def it_produces_output_libreoffice_can_open(self, tmp_path: Path):
        from .harness.lo import assert_libreoffice_opens

        document = _doc(FRAGMENTED)
        find_one(document, "$75-100/hr").replace(
            "$85–110/hr", tracked=True, author="Carol QA", date=FROZEN
        )
        out = tmp_path / "out.docx"
        document.save(str(out))
        assert_libreoffice_opens(out)
