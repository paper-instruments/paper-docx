"""Finalize, scrub, and protection-aware mutations."""

from __future__ import annotations

import zipfile
from pathlib import Path

import pytest

import docx
from docx.errors import DocumentProtectedError, UnsupportedStructureError
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.protection import acknowledge_protection, protection_status
from docx.revision import _remaining_markup

from .harness.contract import save_and_reopen
from .harness.paths import fixture_path

MULTIROUND = "generated/redline/multiround.docx"
GAUNTLET = "generated/gauntlet/gauntlet.docx"
MINIMAL = "generated/minimal-clean/minimal.docx"
PROTECTED_FORMS = "generated/feature-isolated/protected-forms.docx"
PROTECTED_READONLY = "generated/feature-isolated/protected-readonly.docx"
PROTECTED_TRACKED = "generated/feature-isolated/protected-tracked.docx"


def _doc(relpath: str):
    return docx.Document(str(fixture_path(relpath)))


class DescribeFinalize:
    def it_accepts_everything_and_certifies_zero_markup(self, tmp_path: Path):
        document = _doc(MULTIROUND)
        assert document.finalize() == 20
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _remaining_markup(reopened) == {}
        assert len(reopened.revisions) == 0

    def it_rejects_on_request(self, tmp_path: Path):
        document = _doc(MULTIROUND)
        assert document.finalize(revisions="reject") == 20
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _remaining_markup(reopened) == {}

    def it_validates_the_revisions_argument(self):
        with pytest.raises(ValueError, match="accept.*reject|reject.*accept"):
            _doc(MULTIROUND).finalize(revisions="resolve")

    def it_refuses_typed_naming_what_blocked_it(self):
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        document = _doc(MULTIROUND)
        tbl = document.tables[0]._tbl
        tbl.tblPr.append(
            parse_xml(
                f'<w:tblPrChange {nsdecls("w")} w:id="950" w:author="A"'
                ' w:date="2026-06-01T09:30:00Z"><w:tblPr/></w:tblPrChange>'
            )
        )
        before = document.element.xml
        with pytest.raises(UnsupportedStructureError, match="table_property_change"):
            document.finalize()
        assert document.element.xml == before  # atomic: nothing half-final


class DescribeScrub:
    def it_removes_package_root_comment_relationships(self, tmp_path: Path):
        document = docx.Document()
        document.comments.add_comment("review")
        comments_part = document.part.part_related_by(RT.COMMENTS)
        document.part.package.relate_to(comments_part, RT.COMMENTS)

        report = document.scrub(
            metadata=False,
            track_changes_setting=False,
        )
        output = tmp_path / "out.docx"
        document.save(output)

        assert "word/comments.xml" in report.removed_parts
        assert all(
            rel.reltype != RT.COMMENTS
            for rel in document.part.package.rels.values()
        )
        with zipfile.ZipFile(output) as archive:
            assert "word/comments.xml" not in archive.namelist()

    def it_scrubs_the_finalized_gauntlet_clean(self, tmp_path: Path):
        document = _doc(GAUNTLET)
        document.finalize()
        report = document.scrub(rsids=True)
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _remaining_markup(reopened) == {}
        import zipfile

        with zipfile.ZipFile(tmp_path / "out.docx") as zf:
            names = zf.namelist()
            assert not any("comments" in n for n in names)
            document_xml = zf.read("word/document.xml")
            settings_xml = zf.read("word/settings.xml")
        assert b"commentRangeStart" not in document_xml
        assert b"commentReference" not in document_xml
        assert b"w:rsid" not in document_xml
        assert b"<w:rsids>" not in settings_xml
        core = reopened.core_properties
        assert core.author == "" and core.last_modified_by == ""
        assert "word/comments.xml" in {str(p) for p in report.removed_parts}
        assert report.comment_anchors_removed > 0
        assert report.rsid_attributes_removed > 0
        payload = report.to_dict()
        assert payload["schema"] == "paper_scrub_report"

    def it_matches_the_package_diff_to_the_report(self, tmp_path: Path):
        """Report-matches-diff: every removed part is in the report; changed
        parts are exactly the ones scrubbing legally touches."""
        from docx.package import diff_package

        source = fixture_path(GAUNTLET)
        document = docx.Document(str(source))
        document.finalize()
        report = document.scrub(rsids=True)
        out = tmp_path / "scrubbed.docx"
        document.save(str(out))
        diff = diff_package(str(source), str(out))
        assert set(diff.removed) == set(report.removed_parts)
        allowed_changes = {
            "word/document.xml",
            "word/settings.xml",
            "word/_rels/document.xml.rels",
            "[Content_Types].xml",
            "docProps/core.xml",
            "docProps/app.xml",
            "_rels/.rels",
            "word/footnotes.xml",
            "word/endnotes.xml",
        }
        changed = set(diff.semantic_changed_parts())
        header_footer_ok = {
            name
            for name in changed
            if name.startswith(("word/header", "word/footer"))
        }
        assert changed - header_footer_ok <= allowed_changes, changed
        assert not diff.added

    def it_refuses_metadata_scrub_with_pending_revisions(self):
        document = _doc(MULTIROUND)
        before = document.element.xml
        with pytest.raises(UnsupportedStructureError, match="finalize"):
            document.scrub()
        assert document.element.xml == before

    def and_it_scrubs_non_metadata_targets_despite_pending_revisions(self):
        document = _doc(MULTIROUND)
        report = document.scrub(metadata=False)
        assert "word/comments.xml" in {str(p) for p in report.removed_parts}

    def it_keeps_comments_when_asked(self, tmp_path: Path):
        document = _doc(GAUNTLET)
        document.finalize()
        report = document.scrub(comments=False)
        assert report.removed_parts == [] or all(
            "comments" not in str(p) for p in report.removed_parts
        )
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert len(reopened.comments) > 0

    def it_keeps_hidden_text_in_a_strict_comments_part_when_asked(self):
        document = docx.Document()
        comment = document.comments.add_comment(
            "private review note", author="Reviewer"
        )
        hidden_run = comment.paragraphs[0].runs[0]
        hidden_run.font.hidden = True
        relationship = next(
            rel for rel in document.part.rels.values() if rel.reltype == RT.COMMENTS
        )
        relationship._reltype = (  # noqa: SLF001 - strict-package fixture
            "http://purl.oclc.org/ooxml/officeDocument/relationships/comments"
        )

        report = document.scrub(
            comments=False,
            metadata=False,
            track_changes_setting=False,
            hidden_text=True,
        )

        assert report.hidden_runs_removed == 0
        assert hidden_run._r.getparent() is not None  # noqa: SLF001
        assert comment.text == "private review note"

    def it_removes_hidden_text_only_on_explicit_request(self, tmp_path: Path):
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        for hidden_text in (False, True):
            document = _doc(MINIMAL)
            p = document.add_paragraph()
            p._p.append(
                parse_xml(
                    f"<w:r {nsdecls('w')}><w:rPr><w:vanish/></w:rPr>"
                    "<w:t>invisible ink</w:t></w:r>"
                )
            )
            report = document.scrub(hidden_text=hidden_text)
            reopened = save_and_reopen(
                document, tmp_path / f"out-{hidden_text}.docx"
            )
            body_xml = reopened.element.xml
            if hidden_text:
                assert report.hidden_runs_removed == 1
                assert "invisible ink" not in body_xml
            else:
                assert report.hidden_runs_removed == 0
                assert "invisible ink" in body_xml

    def it_reports_protection_and_never_removes_it(self, tmp_path: Path):
        document = _doc(PROTECTED_FORMS)
        with pytest.raises(DocumentProtectedError):
            document.scrub()
        acknowledge_protection(document)
        report = document.scrub()
        assert report.document_protection == {
            "edit": "forms",
            "enforced": True,
            "note": "reported, never removed (docx.protection)",
        }
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert protection_status(reopened).enforced is True  # never stripped


class DescribeProtectionAwareness:
    """Paper mutating APIs refuse on enforced protection; upstream untouched."""

    @pytest.mark.parametrize(
        "relpath, edit",
        [
            (PROTECTED_FORMS, "forms"),
            (PROTECTED_READONLY, "readOnly"),
            (PROTECTED_TRACKED, "trackedChanges"),
        ],
    )
    def it_reads_protection_status(self, relpath: str, edit: str):
        status = protection_status(_doc(relpath))
        assert status.edit == edit and status.enforced is True
        assert status.blocks_paper_edits is True

    def it_reports_no_protection_on_ordinary_documents(self):
        status = protection_status(_doc(MINIMAL))
        assert status.edit is None and status.enforced is False

    def it_refuses_span_replacement_atomically(self):
        from docx.search import find_one

        document = _doc(PROTECTED_READONLY)
        before = document.element.xml
        with pytest.raises(DocumentProtectedError, match="readOnly"):
            find_one(document, "locked").replace("unlocked")
        assert document.element.xml == before

    def it_names_the_override_path_in_the_refusal(self):
        from docx.search import find_one

        document = _doc(PROTECTED_FORMS)
        with pytest.raises(DocumentProtectedError, match="acknowledge_protection"):
            find_one(document, "locked").replace("unlocked")

    def it_refuses_the_other_paper_organs_too(self):
        from docx.blocks import RichParagraph, TextRun, insert_blocks_after
        from docx.controls import set_control_value
        from docx.numbering import apply_list_style
        from docx.search import replace_all
        from docx.tableops import update_cell

        document = _doc(PROTECTED_READONLY)
        with pytest.raises(DocumentProtectedError):
            replace_all(document, "locked", "unlocked")
        with pytest.raises(DocumentProtectedError):
            insert_blocks_after(
                document,
                "Paragraph after the form control.",
                blocks=[RichParagraph(runs=[TextRun("added")])],
            )
        with pytest.raises(DocumentProtectedError):
            apply_list_style(document.paragraphs[0], "Heading 1")
        with pytest.raises(DocumentProtectedError):
            set_control_value(document, "value", tag="inline-field-1")
        with pytest.raises(DocumentProtectedError):
            document.finalize()
        # upstream table added via upstream API (untouched by protection)...
        table = document.add_table(rows=1, cols=1)
        # ...but the paper-docx verb on it still refuses
        with pytest.raises(DocumentProtectedError):
            update_cell(table, 0, 0, "x")

    def it_allows_edits_after_one_document_level_acknowledgment(
        self, tmp_path: Path
    ):
        from docx.search import find_one

        document = _doc(PROTECTED_FORMS)
        status = acknowledge_protection(document)
        assert status.edit == "forms"
        find_one(document, "locked").replace("reviewed and edited")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        # the acknowledgment is in-memory only: a fresh open is protected
        # again, and the protection setting itself was never touched
        fresh_status = protection_status(reopened)
        assert fresh_status.enforced is True and fresh_status.acknowledged is False
        with pytest.raises(DocumentProtectedError):
            find_one(reopened, "reviewed and edited").replace("again")

    def it_leaves_upstream_apis_untouched(self, tmp_path: Path):
        """Strict superset: upstream editing works exactly as before."""
        document = _doc(PROTECTED_READONLY)
        document.add_paragraph("upstream additions are not policed")
        document.paragraphs[0].add_run(" upstream run")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert "upstream additions are not policed" in [
            p.text for p in reopened.paragraphs
        ]
