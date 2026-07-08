"""Tests for the CONVENTIONS §7 package kernel (docx.package / docx._paperpkg).

Every pinned §7 requirement is covered here: the meaningful-whitespace trap,
the no-op byte-identity round trip, zip determinism, single-part change
isolation, and mid-write failure injection.
"""

from __future__ import annotations

import io
import json
import os
import zipfile
from pathlib import Path

import pytest
from lxml import etree

import docx
from docx.package import PackageDiff, diff_package, patch_save, xml_equivalent

from .harness.paths import FIXTURES_DIR, fixture_path, iter_fixture_docx_paths

MINIMAL = "generated/minimal-clean/minimal.docx"
LO_TRACKED = "libreoffice/feature-isolated/tracked-ins-del.docx"

W_DECL = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'


def _noop_roundtrip_relpaths() -> list:
    """Every non-corrupt fixture — patch_save must be byte-stable on all."""
    return [
        rel
        for rel in (
            p.relative_to(FIXTURES_DIR).as_posix() for p in iter_fixture_docx_paths()
        )
        if rel.split("/")[1] not in ("corrupt", "large")
    ]


class DescribeXmlEquivalent:
    def it_accepts_identical_documents(self):
        doc = f'<w:document {W_DECL}><w:body><w:p/></w:body></w:document>'
        assert xml_equivalent(doc.encode(), doc.encode())

    def it_ignores_namespace_prefix_choices(self):
        a = f'<w:document {W_DECL}><w:body/></w:document>'
        b = (
            '<x:document xmlns:x="http://schemas.openxmlformats.org/'
            'wordprocessingml/2006/main"><x:body/></x:document>'
        )
        assert xml_equivalent(a.encode(), b.encode())

    def it_ignores_attribute_order(self):
        a = f'<w:p {W_DECL}><w:r><w:t w:a="1" w:b="2">x</w:t></w:r></w:p>'
        b = f'<w:p {W_DECL}><w:r><w:t w:b="2" w:a="1">x</w:t></w:r></w:p>'
        assert xml_equivalent(a.encode(), b.encode())

    def it_respects_child_order(self):
        a = f'<w:p {W_DECL}><w:r><w:t>one</w:t></w:r><w:r><w:t>two</w:t></w:r></w:p>'
        b = f'<w:p {W_DECL}><w:r><w:t>two</w:t></w:r><w:r><w:t>one</w:t></w:r></w:p>'
        assert not xml_equivalent(a.encode(), b.encode())

    def it_treats_a_preserved_trailing_space_as_meaningful(self):
        """THE §7 trap test: two documents differing only by a trailing space
        inside a text node are NOT equivalent."""
        a = (
            f'<w:p {W_DECL}><w:r>'
            '<w:t xml:space="preserve">rate is </w:t></w:r></w:p>'
        )
        b = (
            f'<w:p {W_DECL}><w:r>'
            '<w:t xml:space="preserve">rate is</w:t></w:r></w:p>'
        )
        assert not xml_equivalent(a.encode(), b.encode())

    def it_treats_whitespace_only_text_nodes_as_content(self):
        a = f'<w:p {W_DECL}><w:r><w:t xml:space="preserve"> </w:t></w:r></w:p>'
        b = f'<w:p {W_DECL}><w:r><w:t xml:space="preserve"></w:t></w:r></w:p>'
        assert not xml_equivalent(a.encode(), b.encode())

    def it_raises_loudly_on_malformed_xml(self):
        with pytest.raises(etree.XMLSyntaxError):
            xml_equivalent(b"<w:p>not closed", b"<w:p/>")


class DescribeDiffPackage:
    def it_reports_no_semantic_change_for_identical_packages(self):
        path = fixture_path(MINIMAL)
        diff = diff_package(path, path)
        assert diff.is_semantically_empty
        assert diff.added == () and diff.removed == () and diff.changed == ()

    def it_reports_added_and_removed_parts(self, tmp_path: Path):
        source = fixture_path(MINIMAL)
        modified = tmp_path / "extra-part.docx"
        with zipfile.ZipFile(source) as zin:
            with zipfile.ZipFile(modified, "w") as zout:
                for name in zin.namelist():
                    if name != "word/fontTable.xml":
                        zout.writestr(name, zin.read(name))
                zout.writestr("word/extra.xml", b"<extra/>")
        diff = diff_package(source, modified)
        assert diff.added == ("word/extra.xml",)
        assert diff.removed == ("word/fontTable.xml",)

    def it_classifies_binary_part_changes_by_bytes(self, tmp_path: Path):
        source = fixture_path("generated/feature-isolated/comments.docx")
        modified = tmp_path / "thumb.docx"
        with zipfile.ZipFile(source) as zin:
            with zipfile.ZipFile(modified, "w") as zout:
                for name in zin.namelist():
                    blob = zin.read(name)
                    if name == "docProps/thumbnail.jpeg":
                        blob = blob + b"\x00"
                    zout.writestr(name, blob)
        diff = diff_package(source, modified)
        (item,) = diff.changed
        assert item.part == "docProps/thumbnail.jpeg"
        assert item.kind == "binary" and item.semantic_change

    def it_produces_deterministic_json(self, tmp_path: Path):
        source = fixture_path(LO_TRACKED)
        resaved = tmp_path / "resaved.docx"
        docx.Document(str(source)).save(str(resaved))
        payload_1 = json.dumps(diff_package(source, resaved).to_dict())
        payload_2 = json.dumps(diff_package(source, resaved).to_dict())
        assert payload_1 == payload_2
        parsed = json.loads(payload_1)
        assert parsed["schema"] == "paper_package_diff" and parsed["version"] == 1

    def it_sees_upstream_resave_of_foreign_bytes_as_semantically_empty(
        self, tmp_path: Path
    ):
        source = fixture_path(LO_TRACKED)
        resaved = tmp_path / "resaved.docx"
        docx.Document(str(source)).save(str(resaved))
        diff = diff_package(source, resaved)
        assert diff.is_semantically_empty
        assert diff.changed, "expected bytewise reserialization churn"


class DescribePatchSave:
    @pytest.mark.parametrize("relpath", _noop_roundtrip_relpaths())
    def it_makes_a_noop_roundtrip_byte_identical(self, relpath: str, tmp_path: Path):
        """THE §7 no-op invariant, on every clean fixture in the corpus."""
        source = fixture_path(relpath)
        out = tmp_path / "out.docx"
        result = patch_save(source, docx.Document(str(source)), out)
        assert result.verbatim_copy
        assert out.read_bytes() == source.read_bytes()

    def it_isolates_a_single_part_edit(self, tmp_path: Path):
        """§7: single-part edit -> exactly that part differs, bytewise."""
        source = fixture_path(LO_TRACKED)
        out = tmp_path / "out.docx"
        document = docx.Document(str(source))
        document.add_paragraph("One new paragraph.")
        result = patch_save(source, document, out)
        assert result.changed_parts == ("word/document.xml",)
        assert not result.verbatim_copy
        with zipfile.ZipFile(source) as za, zipfile.ZipFile(out) as zb:
            # entry ORDER follows the candidate serialization (pinned F2);
            # the part SET and all unchanged bytes must match the original
            assert sorted(za.namelist()) == sorted(zb.namelist())
            differing = [n for n in sorted(za.namelist()) if za.read(n) != zb.read(n)]
        assert differing == ["word/document.xml"]

    def it_writes_deterministic_zip_output(self, tmp_path: Path):
        source = fixture_path(MINIMAL)

        def build(out: Path) -> None:
            document = docx.Document(str(source))
            document.add_paragraph("Deterministic addition.")
            patch_save(source, document, out)

        out_1, out_2 = tmp_path / "one.docx", tmp_path / "two.docx"
        build(out_1)
        build(out_2)
        assert out_1.read_bytes() == out_2.read_bytes()

    def it_allows_saving_over_the_original(self, tmp_path: Path):
        import shutil

        working = tmp_path / "doc.docx"
        shutil.copyfile(fixture_path(MINIMAL), working)
        document = docx.Document(str(working))
        document.add_paragraph("Edited in place.")
        result = patch_save(working, document, working)
        assert result.changed_parts == ("word/document.xml",)
        reopened = docx.Document(str(working))
        assert reopened.paragraphs[-1].text == "Edited in place."

    def it_survives_a_midwrite_failure_with_the_original_intact(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ):
        """§7 failure injection: a crash between temp-write and rename leaves
        the existing output file untouched and no temp litter behind."""
        source = fixture_path(MINIMAL)
        out = tmp_path / "out.docx"
        sentinel = b"pre-existing bytes that must survive"
        out.write_bytes(sentinel)

        document = docx.Document(str(source))
        document.add_paragraph("Never lands.")

        real_replace = os.replace

        def exploding_replace(src, dst):  # noqa: ANN001
            if Path(dst) == out:
                raise OSError("injected mid-write failure")
            return real_replace(src, dst)

        monkeypatch.setattr(os, "replace", exploding_replace)
        with pytest.raises(OSError, match="injected"):
            patch_save(source, document, out)
        assert out.read_bytes() == sentinel
        assert list(tmp_path.iterdir()) == [out], "temp file litter left behind"

    def it_restores_foreign_authored_bytes_for_untouched_parts(self, tmp_path: Path):
        """The point of the kernel: LibreOffice-authored bytes survive an
        open/edit/save cycle untouched for every part we didn't edit."""
        source = fixture_path(LO_TRACKED)
        out = tmp_path / "out.docx"
        document = docx.Document(str(source))
        document.add_paragraph("Narrow edit.")
        result = patch_save(source, document, out)
        assert "word/styles.xml" in result.restored_parts
        with zipfile.ZipFile(source) as za, zipfile.ZipFile(out) as zb:
            assert za.read("word/styles.xml") == zb.read("word/styles.xml")

    def it_round_trips_through_the_contract_harness(self, tmp_path: Path):
        """patch_save output must satisfy the changed-part budget assertion
        the rest of the suite uses."""
        from .harness.contract import assert_changed_parts

        source = fixture_path(MINIMAL)
        out = tmp_path / "out.docx"
        document = docx.Document(str(source))
        document.add_paragraph("Budgeted change.")
        patch_save(source, document, out)
        assert_changed_parts(source, out, {"word/document.xml"})


class DescribePublicSurface:
    def it_reexports_the_kernel_from_docx_package(self):
        import docx.package as package_module

        assert package_module.xml_equivalent is xml_equivalent
        assert package_module.PackageDiff is PackageDiff

    def it_leaves_the_upstream_package_surface_intact(self):
        from docx.package import ImageParts, Package

        assert Package is not None and ImageParts is not None
