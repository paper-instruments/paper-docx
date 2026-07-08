"""Self-tests for the contract-harness utilities (CONVENTIONS §4).

The five contract assertions ship in Phase 1, before any organ exists, so
these tests exercise them through upstream APIs (and synthetic operations for
the refusal-atomicity helper). Organ phases then reuse the same utilities
against their own APIs.
"""

from __future__ import annotations

import datetime as dt
import shutil
from pathlib import Path

import pytest

import docx
from docx.document import Document

from .harness import contract
from .harness.clock import PAPER_TEST_INSTANT, FrozenClock
from .harness.paths import fixture_path
from .harness.pkgdiff import diff_parts

MINIMAL = "generated/minimal-clean/minimal.docx"
GAUNTLET = "generated/gauntlet/gauntlet.docx"
TRACKED = "generated/feature-isolated/tracked-ins-del.docx"
HDRFTR = "generated/feature-isolated/header-footer-sections.docx"

NOOP_ROUNDTRIP_FIXTURES = (
    MINIMAL,
    TRACKED,
    HDRFTR,
    "generated/feature-isolated/comments.docx",
    "generated/feature-isolated/footnotes-endnotes.docx",
    "generated/feature-isolated/fragmented-runs.docx",
    GAUNTLET,
    "libreoffice/feature-isolated/tracked-ins-del.docx",
    "libreoffice/feature-isolated/textbox.docx",
)


class DescribeSaveAndReopen:
    def it_returns_a_document_freshly_loaded_from_disk(self, tmp_path: Path):
        document = docx.Document(str(fixture_path(MINIMAL)))
        document.add_paragraph("A sentence added in memory.")
        reopened = contract.save_and_reopen(document, tmp_path / "out.docx")
        assert reopened is not document
        assert isinstance(reopened, Document)
        assert reopened.paragraphs[-1].text == "A sentence added in memory."


class DescribeChangedPartBudget:
    @pytest.mark.parametrize("relpath", NOOP_ROUNDTRIP_FIXTURES)
    def it_sees_a_noop_open_save_roundtrip_as_semantically_clean(
        self, relpath: str, tmp_path: Path
    ):
        """Open -> save changes nothing semantically, for our files AND
        LibreOffice-authored files. (Byte identity is Phase 2's patch_save
        invariant; upstream save() reserializes freely.)"""
        source = fixture_path(relpath)
        resaved = tmp_path / "resaved.docx"
        docx.Document(str(source)).save(str(resaved))
        contract.assert_no_op_roundtrip_is_semantically_clean(source, resaved)

    def it_passes_when_exactly_the_expected_part_changed(self, tmp_path: Path):
        source = fixture_path(MINIMAL)
        modified = tmp_path / "modified.docx"
        document = docx.Document(str(source))
        document.add_paragraph("A brand-new paragraph.")
        document.save(str(modified))
        diff = contract.assert_changed_parts(source, modified, {"word/document.xml"})
        assert diff.semantic_changed == ("word/document.xml",)

    def it_fails_when_an_unexpected_part_changes(self, tmp_path: Path):
        source = fixture_path(HDRFTR)
        modified = tmp_path / "modified.docx"
        document = docx.Document(str(source))
        document.sections[0].header.paragraphs[0].text = "Rewritten header"
        document.add_paragraph("Body change too.")
        document.save(str(modified))
        with pytest.raises(AssertionError, match="changed parts"):
            contract.assert_changed_parts(source, modified, {"word/document.xml"})

    def it_fails_when_an_expected_change_is_absent(self, tmp_path: Path):
        source = fixture_path(MINIMAL)
        resaved = tmp_path / "resaved.docx"
        docx.Document(str(source)).save(str(resaved))
        with pytest.raises(AssertionError, match="changed parts"):
            contract.assert_changed_parts(source, resaved, {"word/document.xml"})

    def it_classifies_a_changed_malformed_part_as_semantic_conservatively(
        self, tmp_path: Path
    ):
        """A byte-changed part with unparseable (even non-UTF-8) content can
        never be proven equivalent — it must count as semantic change, not
        crash the diff and not silently pass."""
        import zipfile

        source = fixture_path(MINIMAL)
        mangled = tmp_path / "mangled.docx"
        with zipfile.ZipFile(source) as zin:
            names = zin.namelist()
            parts = {name: zin.read(name) for name in names}
        parts["word/document.xml"] = b"\xff\xfe garbage that is not XML at all"
        with zipfile.ZipFile(mangled, "w") as zout:
            for name in names:
                zout.writestr(name, parts[name])
        diff = diff_parts(source, mangled)
        assert "word/document.xml" in diff.semantic_changed

    def it_reports_byte_churn_separately_from_semantic_change(self, tmp_path: Path):
        """A foreign-authored package (LibreOffice bytes) resaved by upstream
        reserializes bytewise without changing meaning — the churn patch_save
        (Phase 2) exists to eliminate. Self-generated fixtures don't show this
        (lxml is byte-stable over its own serialization), so a LibreOffice
        fixture is the source here."""
        source = fixture_path("libreoffice/feature-isolated/tracked-ins-del.docx")
        resaved = tmp_path / "resaved.docx"
        docx.Document(str(source)).save(str(resaved))
        diff = diff_parts(source, resaved)
        assert diff.is_semantically_empty
        assert "word/document.xml" in diff.byte_changed, (
            "expected upstream save() to reserialize LibreOffice-authored parts"
            " bytewise; if this ever becomes byte-stable, patch_save's no-op test"
            " can tighten"
        )


class DescribeRefusalAtomicity:
    """CONVENTIONS §1.3: refuse-then-nothing-changed, in memory and on disk."""

    def it_passes_when_the_operation_refuses_and_mutates_nothing(self, tmp_path: Path):
        path = tmp_path / "doc.docx"
        shutil.copyfile(fixture_path(MINIMAL), path)
        document = docx.Document(str(path))

        def refusing_operation(doc: Document) -> None:
            assert doc.paragraphs  # reads are fine
            raise ValueError("refused: synthetic refusal for harness self-test")

        raised = contract.assert_refusal_atomic(
            document, refusing_operation, ValueError, on_disk=(path,)
        )
        assert "synthetic refusal" in str(raised)

    def it_fails_when_no_refusal_is_raised(self):
        document = docx.Document(str(fixture_path(MINIMAL)))
        with pytest.raises(AssertionError, match="must refuse"):
            contract.assert_refusal_atomic(document, lambda doc: None, ValueError)

    def it_fails_when_the_tree_mutates_before_raising(self):
        """Mutate-then-validate is exactly the bug class this helper catches."""
        document = docx.Document(str(fixture_path(MINIMAL)))

        def leaky_operation(doc: Document) -> None:
            doc.add_paragraph("partial mutation that should not survive")
            raise ValueError("refused too late")

        with pytest.raises(AssertionError, match="in-memory package"):
            contract.assert_refusal_atomic(document, leaky_operation, ValueError)

    def and_it_fails_when_package_level_relationships_mutate(self):
        """_rels/.rels lives on the package object, not on any part — a
        mutation there must not slip past the snapshot."""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        document = docx.Document(str(fixture_path(MINIMAL)))

        def rels_leaking_operation(doc: Document) -> None:
            package = doc.part.package
            assert package is not None
            package.rels.get_or_add_ext_rel(RT.HYPERLINK, "https://example.com/leak")
            raise ValueError("refused after touching package rels")

        with pytest.raises(AssertionError, match="in-memory package"):
            contract.assert_refusal_atomic(document, rels_leaking_operation, ValueError)

    def and_it_fails_when_disk_bytes_change(self, tmp_path: Path):
        path = tmp_path / "doc.docx"
        shutil.copyfile(fixture_path(MINIMAL), path)
        document = docx.Document(str(path))

        def disk_writing_operation(doc: Document) -> None:
            path.write_bytes(b"PK\x03\x04 clobbered")
            raise ValueError("refused after writing")

        with pytest.raises(AssertionError, match="on disk"):
            contract.assert_refusal_atomic(
                document, disk_writing_operation, ValueError, on_disk=(path,)
            )


class DescribeFrozenClock:
    def it_always_returns_the_same_instant(self):
        clock = FrozenClock()
        assert clock.now() is clock.now()
        assert clock.now() == PAPER_TEST_INSTANT

    def it_is_timezone_aware(self, frozen_clock: FrozenClock):
        assert frozen_clock.now().tzinfo is not None

    def and_it_accepts_a_custom_instant(self):
        instant = dt.datetime(2030, 12, 31, 23, 59, 59, tzinfo=dt.timezone.utc)
        assert FrozenClock(instant).now() == instant
