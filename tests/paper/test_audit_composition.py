"""Adversarial ownership and composition audit cases."""

from __future__ import annotations

import pytest

import docx
from docx.blocks import insert_section_after
from docx.bookmarks import (
    _NAME_RE,
    _iter_field_instructions,
    create_bookmark,
    delete_bookmark,
    list_bookmarks,
)
from docx.commentops import anchored_text, resolve
from docx.composition import append_document, insert_blocks_from
from docx.enum.style import WD_STYLE_TYPE
from docx.errors import BoundaryViolationError, UnsupportedStructureError
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.search import find_one

_SECT_PR = qn("w:sectPr")
_SDT = qn("w:sdt")


def _package_state(document) -> tuple:
    parts = []
    for part in document.part.package.iter_parts():
        relationships = tuple(
            sorted(
                (
                    rel.rId,
                    rel.reltype,
                    str(rel.target_ref),
                    rel.is_external,
                )
                for rel in part.rels.values()
            )
        )
        parts.append((str(part.partname), part.blob, relationships))
    return tuple(sorted(parts, key=lambda item: item[0]))


def _clear_body(document) -> None:
    body = document.element.body
    for child in list(body):
        if child.tag != _SECT_PR:
            body.remove(child)


def _insert_before_sect(document, element) -> None:
    body = document.element.body
    sect_pr = body.find(_SECT_PR)
    if sect_pr is None:
        body.append(element)
    else:
        sect_pr.addprevious(element)


def _protect(document) -> None:
    protection = OxmlElement("w:documentProtection")
    protection.set(qn("w:edit"), "readOnly")
    protection.set(qn("w:enforcement"), "1")
    document.settings.element.append(protection)


def _add_bookmark(paragraph, name: str, bookmark_id: str, text: str) -> None:
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    run = OxmlElement("w:r")
    run.add_t(text)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.extend((start, run, end))


def _append_complex_reference(document, fragments: list[str]) -> None:
    paragraph = document.add_paragraph()

    def field_char(kind: str):
        run = OxmlElement("w:r")
        char = OxmlElement("w:fldChar")
        char.set(qn("w:fldCharType"), kind)
        run.append(char)
        paragraph._p.append(run)

    field_char("begin")
    for fragment in fragments:
        run = OxmlElement("w:r")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = fragment
        run.append(instruction)
        paragraph._p.append(run)
    field_char("separate")
    result = OxmlElement("w:r")
    result.add_t("(reference)")
    paragraph._p.append(result)
    field_char("end")


def _top_level_control(
    *, text: str, data_bound: bool = False, control_id: int = 7
):
    control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    id_element = OxmlElement("w:id")
    id_element.set(qn("w:val"), str(control_id))
    properties.append(id_element)
    if data_bound:
        binding = OxmlElement("w:dataBinding")
        binding.set(qn("w:xpath"), "/root/value")
        binding.set(qn("w:storeItemID"), "{AUDIT-STORE}")
        properties.append(binding)
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    run.add_t(text)
    paragraph.append(run)
    content.append(paragraph)
    control.extend((properties, content))
    return control


def _unsupported_top_level(kind: str):
    if kind == "customXml":
        wrapper = OxmlElement("w:customXml")
        paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        run.add_t("custom XML payload")
        paragraph.append(run)
        wrapper.append(paragraph)
        return wrapper
    chunk = OxmlElement("w:altChunk")
    chunk.set(qn("r:id"), "rIdMissingChunk")
    return chunk


def it_refuses_foreign_span_anchors_before_destination_protection() -> None:
    source = docx.Document()
    source.add_paragraph("Shared anchor text")
    destination = docx.Document()
    destination.add_paragraph("Shared anchor text")
    _protect(destination)
    foreign_span = find_one(source, "Shared anchor")
    source_before = _package_state(source)
    destination_before = _package_state(destination)

    with pytest.raises(BoundaryViolationError, match="different document"):
        insert_section_after(
            destination, foreign_span, heading="Unsafe", paragraphs=[]
        )
    with pytest.raises(BoundaryViolationError, match="different document"):
        create_bookmark(destination, foreign_span, "UnsafeBookmark")

    assert _package_state(source) == source_before
    assert _package_state(destination) == destination_before


def it_refuses_a_foreign_comment_proxy_even_with_a_colliding_id() -> None:
    source = docx.Document()
    source_run = source.add_paragraph("Source anchor").runs[0]
    foreign_comment = source.add_comment(
        source_run, text="source comment", author="Source"
    )
    destination = docx.Document()
    destination_run = destination.add_paragraph("Destination anchor").runs[0]
    destination.add_comment(
        destination_run, text="destination comment", author="Destination"
    )
    source_before = _package_state(source)
    destination_before = _package_state(destination)

    with pytest.raises(BoundaryViolationError, match="different document"):
        anchored_text(destination, foreign_comment)
    with pytest.raises(BoundaryViolationError, match="different document"):
        resolve(destination, foreign_comment)

    assert _package_state(source) == source_before
    assert _package_state(destination) == destination_before


@pytest.mark.parametrize("mode", ["append", "range"])
def it_composes_top_level_content_controls_as_whole_blocks(mode: str) -> None:
    source = docx.Document()
    _clear_body(source)
    source_control = _top_level_control(text="Controlled source block")
    _insert_before_sect(source, source_control)
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")

    if mode == "append":
        report = append_document(destination, source, section="continuous")
    else:
        report = insert_blocks_from(
            destination,
            source,
            "Controlled source block",
            anchor="Destination anchor",
        )

    copied = [child for child in destination.element.body if child.tag == _SDT]
    assert report.inserted_blocks == 1
    assert len(copied) == 1
    assert copied[0] is not source_control
    assert "".join(node.text or "" for node in copied[0].iter(qn("w:t"))) == (
        "Controlled source block"
    )


def it_refuses_a_data_bound_control_before_composition_mutates() -> None:
    source = docx.Document()
    _clear_body(source)
    _insert_before_sect(
        source, _top_level_control(text="Bound value", data_bound=True)
    )
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    before = _package_state(destination)

    with pytest.raises(UnsupportedStructureError, match="data-bound"):
        append_document(destination, source, section="continuous")

    assert _package_state(destination) == before


def it_keeps_all_content_control_ids_unique_when_appending_after_a_control() -> None:
    source = docx.Document()
    _clear_body(source)
    _insert_before_sect(
        source, _top_level_control(text="First imported control", control_id=2)
    )
    _insert_before_sect(
        source, _top_level_control(text="Second imported control", control_id=1)
    )
    destination = docx.Document()
    _clear_body(destination)
    _insert_before_sect(
        destination, _top_level_control(text="Existing control", control_id=1)
    )

    report = append_document(destination, source, section="continuous")

    control_ids = []
    for control in (child for child in destination.element.body if child.tag == _SDT):
        properties = control.find(qn("w:sdtPr"))
        id_element = properties.find(qn("w:id"))
        control_ids.append(int(id_element.get(qn("w:val"))))
    assert report.inserted_blocks == 2
    assert control_ids == [1, 2, 3]


@pytest.mark.parametrize("kind", ["customXml", "altChunk"])
def it_refuses_unsupported_top_level_content_instead_of_omitting_it(
    kind: str,
) -> None:
    source = docx.Document()
    _clear_body(source)
    source.add_paragraph("Safe-looking source block")
    _insert_before_sect(source, _unsupported_top_level(kind))
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    before = _package_state(destination)

    with pytest.raises(UnsupportedStructureError, match=kind):
        append_document(destination, source, section="continuous")

    assert _package_state(destination) == before


def it_refuses_unsupported_content_between_range_endpoints_atomically() -> None:
    source = docx.Document()
    _clear_body(source)
    first = source.add_paragraph("First selected block")
    second = source.add_paragraph("Second selected block")
    first._p.addnext(_unsupported_top_level("customXml"))
    assert second._p.getprevious().tag == qn("w:customXml")
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    before = _package_state(destination)

    with pytest.raises(UnsupportedStructureError, match="customXml"):
        insert_blocks_from(
            destination,
            source,
            "First selected block",
            count=2,
            anchor="Destination anchor",
        )

    assert _package_state(destination) == before


def it_authors_a_legal_fresh_bookmark_and_rewrites_a_split_lowercase_ref() -> None:
    name = "A" * 40
    source = docx.Document()
    target = source.add_paragraph()
    _add_bookmark(target, name, "1", "Source bookmark target")
    _append_complex_reference(
        source,
        [" r", f"ef {name[:20]}", f"{name[20:]} \\h "],
    )
    destination = docx.Document()
    anchor = destination.add_paragraph("Destination anchor")
    _add_bookmark(anchor, name, "8", "")

    report = insert_blocks_from(
        destination,
        source,
        "Source bookmark target",
        count=2,
        anchor="Destination anchor",
    )

    fresh = report.bookmarks_renamed[name]
    assert len(fresh) <= 40
    assert _NAME_RE.fullmatch(fresh)
    instructions = [
        instruction
        for instruction, _nodes in _iter_field_instructions(destination.element)
    ]
    assert any(f"ref {fresh}".casefold() in value.casefold() for value in instructions)


def it_sees_a_case_insensitive_split_quoted_reference_before_bookmark_delete() -> None:
    document = docx.Document()
    target = document.add_paragraph()
    _add_bookmark(target, "KeyPhrase", "4", "Referenced target")
    _append_complex_reference(
        document, [" pa", 'geref "key', 'phrase" \\h ']
    )
    before = _package_state(document)

    with pytest.raises(UnsupportedStructureError, match="referenced by"):
        delete_bookmark(document, "KeyPhrase")

    assert _package_state(document) == before


def it_uses_full_run_content_for_bookmark_and_comment_anchor_text() -> None:
    document = docx.Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("A")
    run.add_tab()
    run.add_text("B")
    run.add_break()
    run.add_text("C")
    run._r.append(OxmlElement("w:noBreakHyphen"))
    run.add_text("D")
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), "12")
    start.set(qn("w:name"), "RichAnchor")
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), "12")
    run._r.addprevious(start)
    run._r.addnext(end)
    comment = document.add_comment(run, text="note", author="Reviewer")

    bookmarks = {item.name: item for item in list_bookmarks(document)}
    assert bookmarks["RichAnchor"].text == "A\tB\nC-D"
    assert anchored_text(document, comment) == "A\tB\nC-D"


def it_refuses_malformed_ids_before_style_import() -> None:
    source = docx.Document()
    style = source.styles.add_style("AuditImportedStyle", WD_STYLE_TYPE.PARAGRAPH)
    paragraph = source.add_paragraph(style=style)
    _add_bookmark(paragraph, "Malformed", "not-a-number", "Malformed target")
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    before = _package_state(destination)

    with pytest.raises(UnsupportedStructureError, match="non-numeric"):
        insert_blocks_from(
            destination,
            source,
            "Malformed target",
            anchor="Destination anchor",
        )

    assert _package_state(destination) == before
    assert "AuditImportedStyle" not in {style.name for style in destination.styles}
