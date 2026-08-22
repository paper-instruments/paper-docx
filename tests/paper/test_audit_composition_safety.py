"""Atomicity regressions for cross-document composition."""

from __future__ import annotations

import pytest

import docx
from docx.composition import insert_blocks_from
from docx.enum.style import WD_STYLE_TYPE
from docx.errors import TargetNotFoundError, UnsupportedStructureError
from docx.numbering import apply_numbering, ensure_decimal_definition
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement


def _package_state(document) -> tuple:
    parts = []
    for part in document.part.package.iter_parts():
        relationships = tuple(
            sorted(
                (
                    rel.rId,
                    rel.reltype,
                    str(rel.target_ref),
                    rel.is_external,
                )
                for rel in part.rels.values()
            )
        )
        parts.append((str(part.partname), part.blob, relationships))
    return tuple(sorted(parts, key=lambda item: item[0]))


def _append_chart_reference(document, paragraph, r_id: str = "rId99") -> None:
    chart_part = Part(
        PackURI("/word/charts/chart1.xml"),
        CT.DML_CHART,
        b'<c:chartSpace xmlns:c="http://schemas.openxmlformats.org/drawingml/2006/chart"/>',
        document.part.package,
    )
    document.part.load_rel(RT.CHART, chart_part, r_id)
    drawing = OxmlElement("w:drawing")
    chart = OxmlElement("c:chart")
    chart.set(qn("r:id"), r_id)
    drawing.append(chart)
    paragraph.add_run()._r.append(drawing)


def _append_style_reference(paragraph, tag: str, style_id: str) -> None:
    properties = (
        paragraph._p.get_or_add_pPr()
        if tag == "w:pStyle"
        else paragraph.runs[0]._r.get_or_add_rPr()
    )
    reference = OxmlElement(tag)
    reference.set(qn("w:val"), style_id)
    properties.append(reference)


def _clear_body(document) -> None:
    body = document.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _insert_before_section_properties(document, element) -> None:
    section_properties = document.element.body.find(qn("w:sectPr"))
    if section_properties is None:
        document.element.body.append(element)
    else:
        section_properties.addprevious(element)


def _top_level_control(text: str):
    control = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    run = OxmlElement("w:r")
    run.add_t(text)
    paragraph.append(run)
    content.append(paragraph)
    control.append(content)
    return control


@pytest.mark.parametrize(
    ("start", "end", "include_start", "include_end", "count"),
    [
        pytest.param("Block A", "Block A", False, True, 1, id="same-exclude-start"),
        pytest.param("Block A", "Block A", True, False, 1, id="same-exclude-end"),
        pytest.param("Block A", "Block B", False, False, 1, id="adjacent-exclude-both"),
        pytest.param("Block B", "Block A", True, True, 1, id="reversed"),
        pytest.param("Block C", None, False, True, 1, id="adjusted-out-of-bounds"),
    ],
)
def it_refuses_empty_reversed_and_out_of_bounds_adjusted_ranges_atomically(
    start: str,
    end: str | None,
    include_start: bool,
    include_end: bool,
    count: int,
) -> None:
    source = docx.Document()
    for text in ("Block A", "Block B", "Block C"):
        source.add_paragraph(text)
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    before = _package_state(destination)

    with pytest.raises(TargetNotFoundError):
        insert_blocks_from(
            destination,
            source,
            start,
            end_anchor=end,
            count=count,
            include_start=include_start,
            include_end=include_end,
            anchor="Destination anchor",
        )

    assert _package_state(destination) == before


@pytest.mark.parametrize("include_control", [True, False])
def it_treats_a_top_level_control_as_one_includable_endpoint(
    include_control: bool,
) -> None:
    source = docx.Document()
    _clear_body(source)
    source.add_paragraph("Selected paragraph")
    _insert_before_section_properties(source, _top_level_control("Control boundary"))
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")

    report = insert_blocks_from(
        destination,
        source,
        "Selected paragraph",
        end_anchor="Control boundary",
        include_end=include_control,
        anchor="Destination anchor",
    )

    copied_controls = destination.element.body.findall(qn("w:sdt"))
    assert report.inserted_blocks == (2 if include_control else 1)
    assert len(copied_controls) == int(include_control)


def it_preserves_unsupported_physical_children_inside_an_adjusted_slice() -> None:
    source = docx.Document()
    _clear_body(source)
    source.add_paragraph("Start boundary")
    first = source.add_paragraph("First selected block")
    source.add_paragraph("Second selected block")
    source.add_paragraph("End boundary")
    unsupported = OxmlElement("w:customXml")
    unsupported.append(OxmlElement("w:p"))
    first._p.addnext(unsupported)
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    before = _package_state(destination)

    with pytest.raises(UnsupportedStructureError, match="customXml"):
        insert_blocks_from(
            destination,
            source,
            "Start boundary",
            end_anchor="End boundary",
            include_start=False,
            include_end=False,
            anchor="Destination anchor",
        )

    assert _package_state(destination) == before


def it_refuses_a_chart_relationship_before_copying_stale_rids() -> None:
    source = docx.Document()
    paragraph = source.add_paragraph("Source chart")
    _append_chart_reference(source, paragraph)
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    before = _package_state(destination)

    with pytest.raises(UnsupportedStructureError, match="chart relationship"):
        insert_blocks_from(
            destination,
            source,
            "Source chart",
            anchor="Destination anchor",
        )

    assert _package_state(destination) == before
    assert "rId99" not in destination.element.xml


def it_refuses_a_source_num_whose_abstract_definition_is_missing() -> None:
    source = docx.Document()
    num_id = ensure_decimal_definition(source)
    paragraph = source.add_paragraph("Source numbered paragraph")
    apply_numbering(paragraph, num_id=num_id)
    source_numbering = source.part.numbering_part.element
    source_num = next(
        num
        for num in source_numbering.findall(qn("w:num"))
        if num.get(qn("w:numId")) == str(num_id)
    )
    abstract_id = source_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    source_abstract = next(
        abstract
        for abstract in source_numbering.findall(qn("w:abstractNum"))
        if abstract.get(qn("w:abstractNumId")) == abstract_id
    )
    source_numbering.remove(source_abstract)
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    before = _package_state(destination)

    with pytest.raises(UnsupportedStructureError, match="missing abstract"):
        insert_blocks_from(
            destination,
            source,
            "Source numbered paragraph",
            anchor="Destination anchor",
        )

    assert _package_state(destination) == before


def it_refuses_malformed_destination_numbering_before_importing_a_style() -> None:
    source = docx.Document()
    style = source.styles.add_style("AuditCompositionStyle", WD_STYLE_TYPE.PARAGRAPH)
    num_id = ensure_decimal_definition(source)
    paragraph = source.add_paragraph("Styled numbered source", style=style)
    apply_numbering(paragraph, num_id=num_id)
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    ensure_decimal_definition(destination)
    destination_num = destination.part.numbering_part.element.find(qn("w:num"))
    destination_num.set(qn("w:numId"), "not-a-number")
    before = _package_state(destination)

    with pytest.raises(UnsupportedStructureError, match="destination numbering.*non-numeric"):
        insert_blocks_from(
            destination,
            source,
            "Styled numbered source",
            anchor="Destination anchor",
        )

    assert _package_state(destination) == before
    assert "AuditCompositionStyle" not in {style.name for style in destination.styles}


@pytest.mark.parametrize(
    ("reference_tag", "destination_style_type"),
    [
        pytest.param("w:pStyle", WD_STYLE_TYPE.PARAGRAPH, id="paragraph-style"),
        pytest.param("w:rStyle", WD_STYLE_TYPE.CHARACTER, id="character-style"),
    ],
)
def it_refuses_undefined_source_style_references_before_destination_binding(
    reference_tag: str,
    destination_style_type: WD_STYLE_TYPE,
) -> None:
    source = docx.Document()
    paragraph = source.add_paragraph("Source with an undefined style")
    _append_style_reference(paragraph, reference_tag, "SharedStyleId")
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    destination_style = destination.styles.add_style(
        "Unrelated destination style", destination_style_type
    )
    destination_style._element.set(qn("w:styleId"), "SharedStyleId")
    before = _package_state(destination)

    with pytest.raises(
        UnsupportedStructureError, match="undefined source style 'SharedStyleId'"
    ):
        insert_blocks_from(
            destination,
            source,
            "Source with an undefined style",
            anchor="Destination anchor",
        )

    assert _package_state(destination) == before
    assert all(
        paragraph.text != "Source with an undefined style"
        for paragraph in destination.paragraphs
    )


@pytest.mark.parametrize("chain_tag", ["w:basedOn", "w:link", "w:next"])
def it_refuses_undefined_source_style_chain_dependencies_atomically(chain_tag: str) -> None:
    source = docx.Document()
    source_style = source.styles.add_style("Source chain leaf", WD_STYLE_TYPE.PARAGRAPH)
    dependency = OxmlElement(chain_tag)
    dependency.set(qn("w:val"), "MissingDependency")
    source_style._element.append(dependency)
    source.add_paragraph("Source with a broken style chain", style=source_style)
    destination = docx.Document()
    destination.add_paragraph("Destination anchor")
    destination_style = destination.styles.add_style(
        "Unrelated chain destination", WD_STYLE_TYPE.PARAGRAPH
    )
    destination_style._element.set(qn("w:styleId"), "MissingDependency")
    before = _package_state(destination)

    with pytest.raises(
        UnsupportedStructureError, match="undefined source style 'MissingDependency'"
    ):
        insert_blocks_from(
            destination,
            source,
            "Source with a broken style chain",
            anchor="Destination anchor",
        )

    assert _package_state(destination) == before
    assert "Source chain leaf" not in {style.name for style in destination.styles}
