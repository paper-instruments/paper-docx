"""Regression coverage for structural editing refusal atomicity."""

from __future__ import annotations

import pytest
from lxml import etree

from docx import Document
from docx.controls import get_control
from docx.errors import UnsupportedStructureError
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.tableops import insert_row_after


def _empty_control(tag: str, *, with_content: bool):
    sdt = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    tag_element = OxmlElement("w:tag")
    tag_element.set(qn("w:val"), tag)
    properties.append(tag_element)
    properties.append(OxmlElement("w:text"))
    sdt.append(properties)
    if with_content:
        sdt.append(OxmlElement("w:sdtContent"))
    return sdt


def _append_revision_marker(table, container_tag: str, marker_tag: str) -> None:
    row = table.rows[0]._tr
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    if container_tag == "w:tr":
        container = row
    elif container_tag == "w:trPr":
        container = row.get_or_add_trPr()
    elif container_tag == "w:tblPrEx":
        container = OxmlElement(container_tag)
        row.insert(0, container)
    elif container_tag == "w:tcPr":
        container = cell._tc.get_or_add_tcPr()
    elif container_tag == "w:pPr":
        container = paragraph._p.get_or_add_pPr()
    elif container_tag == "w:rPr":
        container = paragraph.runs[0]._r.get_or_add_rPr()
    else:
        assert container_tag == "w:numPr"
        container = OxmlElement(container_tag)
        paragraph._p.get_or_add_pPr().append(container)

    marker = OxmlElement(marker_tag)
    marker.set(qn("w:id"), "77")
    marker.set(qn("w:author"), "Reviewer")
    snapshot_tags = {
        "w:trPrChange": "w:trPr",
        "w:tblPrExChange": "w:tblPrEx",
        "w:tcPrChange": "w:tcPr",
        "w:pPrChange": "w:pPr",
        "w:rPrChange": "w:rPr",
    }
    snapshot_tag = snapshot_tags.get(marker_tag)
    if snapshot_tag is not None:
        marker.append(OxmlElement(snapshot_tag))
    container.append(marker)


class DescribeEmptyContentControls:
    @pytest.mark.parametrize("with_content", [False, True])
    def it_authors_paragraph_content_for_a_block_control(self, with_content: bool):
        document = Document()
        sdt = _empty_control("block", with_content=with_content)
        body = document.element.body
        body.insert(body.index(body.sectPr), sdt)

        get_control(document, tag="block").set_value("replacement")

        content = sdt.find(qn("w:sdtContent"))
        assert content is not None
        assert [child.tag for child in content] == [qn("w:p")]
        paragraph = content.find(qn("w:p"))
        assert paragraph.find(qn("w:r")).find(qn("w:t")).text == "replacement"

    @pytest.mark.parametrize("with_content", [False, True])
    def it_keeps_inline_content_inline(self, with_content: bool):
        document = Document()
        paragraph = document.add_paragraph()
        sdt = _empty_control("inline", with_content=with_content)
        paragraph._p.append(sdt)

        get_control(document, tag="inline").set_value("replacement")

        content = sdt.find(qn("w:sdtContent"))
        assert content is not None
        assert [child.tag for child in content] == [qn("w:r")]
        assert content.find(qn("w:r")).find(qn("w:t")).text == "replacement"

    def it_refuses_an_empty_control_with_an_ambiguous_parent(self):
        document = Document()
        wrapper = OxmlElement("w:customXml")
        sdt = _empty_control("ambiguous", with_content=True)
        wrapper.append(sdt)
        body = document.element.body
        body.insert(body.index(body.sectPr), wrapper)
        before = etree.tostring(sdt)

        with pytest.raises(UnsupportedStructureError, match="ambiguous block/inline"):
            get_control(document, tag="ambiguous").set_value("replacement")

        assert etree.tostring(sdt) == before


class DescribeDetachedRowPopulation:
    def it_refuses_a_control_wrapped_template_atomically(self):
        document = Document()
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "plain template"
        cell = table.cell(0, 1)
        for child in list(cell._tc):
            if child.tag != qn("w:tcPr"):
                cell._tc.remove(child)

        sdt = _empty_control("wrapped", with_content=True)
        content = sdt.find(qn("w:sdtContent"))
        paragraph = OxmlElement("w:p")
        run = OxmlElement("w:r")
        run.add_t("template")
        paragraph.append(run)
        content.append(paragraph)
        cell._tc.append(sdt)
        before = table._tbl.xml

        with pytest.raises(UnsupportedStructureError, match="content control"):
            insert_row_after(table, 0, ["first was populated", "replacement"])

        assert table._tbl.xml == before
        assert len(table.rows) == 1


class DescribeTrackedRowTemplateRefusal:
    @pytest.mark.parametrize("marker_tag", ["w:ins", "w:del"])
    def it_does_not_duplicate_tracked_row_ids(self, marker_tag: str):
        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "template"
        _append_revision_marker(table, "w:trPr", marker_tag)
        before = table._tbl.xml

        with pytest.raises(UnsupportedStructureError, match="tracked revision metadata"):
            insert_row_after(table, 0, ["new row"])

        assert table._tbl.xml == before
        assert len(table.rows) == 1
        assert [node.get(qn("w:id")) for node in table._tbl.iter(qn(marker_tag))] == [
            "77"
        ]

    @pytest.mark.parametrize(
        ("container_tag", "marker_tag"),
        [
            pytest.param("w:tr", "w:moveFromRangeStart", id="move-range"),
            pytest.param("w:trPr", "w:trPrChange", id="row-properties"),
            pytest.param("w:tblPrEx", "w:tblPrExChange", id="table-property-exceptions"),
            pytest.param("w:tcPr", "w:cellIns", id="cell-insertion"),
            pytest.param("w:tcPr", "w:cellDel", id="cell-deletion"),
            pytest.param("w:tcPr", "w:cellMerge", id="cell-merge"),
            pytest.param("w:tcPr", "w:tcPrChange", id="cell-properties"),
            pytest.param("w:pPr", "w:pPrChange", id="paragraph-properties"),
            pytest.param("w:rPr", "w:rPrChange", id="run-properties"),
            pytest.param("w:numPr", "w:numberingChange", id="numbering-properties"),
        ],
    )
    def it_refuses_other_revision_metadata_that_the_row_clone_would_retain(
        self, container_tag: str, marker_tag: str
    ):
        document = Document()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "template"
        _append_revision_marker(table, container_tag, marker_tag)
        before = table._tbl.xml

        with pytest.raises(UnsupportedStructureError, match="tracked revision metadata"):
            insert_row_after(table, 0, ["new row"])

        assert table._tbl.xml == before
        assert len(table.rows) == 1
        assert [node.get(qn("w:id")) for node in table._tbl.iter(qn(marker_tag))] == [
            "77"
        ]


class DescribeTrackedBlockIdentityPreflight:
    """List blocks add numbering definitions while building nodes, so a
    malformed identity must refuse before that side-part mutation."""

    def it_leaves_numbering_untouched_when_tracked_author_is_malformed(self):
        import io

        from docx.blocks import ListBlock, insert_blocks_after

        document = Document()
        document.add_paragraph("anchor paragraph")
        before = io.BytesIO()
        document.save(before)

        with pytest.raises(ValueError, match="author"):
            insert_blocks_after(
                document,
                "anchor paragraph",
                blocks=[ListBlock(items=["item one"], kind="bullet")],
                tracked=True,
                author="bad\x00author",
            )

        after = io.BytesIO()
        document.save(after)
        assert after.getvalue() == before.getvalue()

    def and_a_string_date_refuses_before_any_mutation(self):
        import io

        from docx.blocks import tracked_delete_paragraphs

        document = Document()
        document.add_paragraph("anchor paragraph")
        document.add_paragraph("paragraph to delete")
        before = io.BytesIO()
        document.save(before)

        with pytest.raises(TypeError, match="datetime"):
            tracked_delete_paragraphs(
                document, "paragraph to delete", author="A", date="2026-01-01"
            )

        after = io.BytesIO()
        document.save(after)
        assert after.getvalue() == before.getvalue()
