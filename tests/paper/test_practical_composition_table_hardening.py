"""Focused practical composition, table, and numbering hardening."""

from __future__ import annotations

import pytest

import docx
from docx.bookmarks import create_bookmark
from docx.composition import insert_blocks_from
from docx.enum.style import WD_STYLE_TYPE
from docx.errors import TargetNotFoundError, UnsupportedStructureError
from docx.numbering import apply_numbering, ensure_decimal_definition, list_numbering
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.search import find_one
from docx.tableops import update_cell


def _field_char(kind: str):
    run = OxmlElement("w:r")
    marker = OxmlElement("w:fldChar")
    marker.set(qn("w:fldCharType"), kind)
    run.append(marker)
    return run


def it_arms_copied_fields_for_update() -> None:
    source = docx.Document()
    paragraph = source.add_paragraph("Field source")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), " PAGE ")
    paragraph._p.append(field)
    destination = docx.Document()
    destination.add_paragraph("Destination")

    insert_blocks_from(destination, source, "Field source", anchor="Destination")

    update = destination.settings.element.find(qn("w:updateFields"))
    assert update is not None
    assert update.get(qn("w:val")) == "true"


def it_reports_a_numbered_paragraphs_visible_layout_grid_column() -> None:
    document = docx.Document()
    num_id = ensure_decimal_definition(document)
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1))
    paragraph = table.cell(0, 2).paragraphs[0]
    paragraph.text = "third grid column"
    apply_numbering(paragraph, num_id=num_id)

    item = next(
        item
        for item in list_numbering(document).numbered_paragraphs
        if item.text == "third grid column"
    )

    assert item.table_cell == (0, 2, 0)


def it_refuses_updating_a_cell_with_a_cross_paragraph_field() -> None:
    document = docx.Document()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.paragraphs[0]._p.append(_field_char("begin"))
    cell.add_paragraph("result")
    cell.add_paragraph()._p.append(_field_char("end"))

    with pytest.raises(UnsupportedStructureError, match="Word field"):
        update_cell(document.tables[0], 0, 0, "replacement")

    assert cell.paragraphs[1].text == "result"


def it_refuses_filling_an_empty_cell_around_a_typed_control() -> None:
    document = docx.Document()
    cell = document.add_table(rows=1, cols=1).cell(0, 0)
    control = OxmlElement("w:sdt")
    properties = OxmlElement("w:sdtPr")
    properties.append(OxmlElement("w14:checkbox"))
    control.append(properties)
    control.append(OxmlElement("w:sdtContent"))
    cell.paragraphs[0]._p.append(control)

    with pytest.raises(UnsupportedStructureError, match="Control.set_value"):
        update_cell(document.tables[0], 0, 0, "YES")

    assert cell.text == ""


def it_updates_a_cell_by_visible_grid_column_after_grid_before() -> None:
    document = docx.Document()
    table = document.add_table(rows=1, cols=3)
    row = table.rows[0]
    row._tr.remove(row._tr.tc_lst[0])
    grid_before = OxmlElement("w:gridBefore")
    grid_before.set(qn("w:val"), "1")
    row._tr.get_or_add_trPr().append(grid_before)
    row.cells[0].text = "visible column one"

    update_cell(table, 0, 1, "changed")

    assert row.cells[0].text == "changed"
    with pytest.raises(TargetNotFoundError, match="has no cell"):
        update_cell(table, 0, 0, "wrong")


def it_reports_a_numbering_level_defined_by_an_instance_override() -> None:
    document = docx.Document()
    num_id = ensure_decimal_definition(document)
    numbering = document.part.numbering_part.element
    num = next(
        node for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) == str(num_id)
    )
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "4")
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "4")
    num_format = OxmlElement("w:numFmt")
    num_format.set(qn("w:val"), "lowerLetter")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%5.")
    level.extend((num_format, level_text))
    override.append(level)
    num.append(override)
    paragraph = document.add_paragraph("override level")
    apply_numbering(paragraph, num_id=num_id, level=4)

    report = list_numbering(document)

    assert report.numbered_paragraphs[-1].level == 4
    assert any(level.level == 4 for level in report.definitions[-1].levels)


def it_inserts_composed_numbering_before_the_cleanup_marker() -> None:
    source = docx.Document()
    source_id = ensure_decimal_definition(source)
    source_paragraph = source.add_paragraph("Numbered source")
    apply_numbering(source_paragraph, num_id=source_id)
    destination = docx.Document()
    destination.add_paragraph("Destination")
    ensure_decimal_definition(destination)
    cleanup = OxmlElement("w:numIdMacAtCleanup")
    cleanup.set(qn("w:val"), "1")
    destination.part.numbering_part.element.append(cleanup)

    insert_blocks_from(destination, source, "Numbered source", anchor="Destination")

    assert destination.part.numbering_part.element[-1].tag == qn("w:numIdMacAtCleanup")


def it_quotes_a_renamed_styleref_and_carries_its_numbering_dependency() -> None:
    source = docx.Document()
    source_num_id = ensure_decimal_definition(source)
    source_style = source.styles.add_style("ClauseRef", WD_STYLE_TYPE.PARAGRAPH)
    paragraph_properties = source_style._element.get_or_add_pPr()
    numbering_properties = OxmlElement("w:numPr")
    numbering_id = OxmlElement("w:numId")
    numbering_id.set(qn("w:val"), str(source_num_id))
    numbering_properties.append(numbering_id)
    paragraph_properties.append(numbering_properties)
    source_paragraph = source.add_paragraph("Field source")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), " STYLEREF ClauseRef ")
    source_paragraph._p.append(field)

    destination = docx.Document()
    destination.add_paragraph("Destination")
    destination.styles.add_style("ClauseRef", WD_STYLE_TYPE.PARAGRAPH)
    ensure_decimal_definition(destination)

    report = insert_blocks_from(
        destination,
        source,
        "Field source",
        anchor="Destination",
        styles="import_renamed",
    )

    imported_field = next(destination.element.iter(qn("w:fldSimple")))
    assert imported_field.get(qn("w:instr")) == ' STYLEREF "ClauseRef (imported)" '
    assert source_num_id in report.numbering_map


def it_remaps_bookmark_identifiers_inside_word_formulas() -> None:
    source = docx.Document()
    source.add_paragraph("100")
    create_bookmark(source, find_one(source, "100"), "GrossIncome")
    formula = OxmlElement("w:fldSimple")
    formula.set(qn("w:instr"), " =ROUND(GrossIncome,0) ")
    source.paragraphs[0]._p.append(formula)

    destination = docx.Document()
    destination.add_paragraph("existing")
    create_bookmark(destination, find_one(destination, "existing"), "GrossIncome")

    report = insert_blocks_from(
        destination, source, "100", anchor="existing"
    )

    imported_formula = next(
        field
        for field in destination.element.iter(qn("w:fldSimple"))
        if "ROUND" in (field.get(qn("w:instr")) or "")
    )
    assert report.bookmarks_renamed["GrossIncome"] == "GrossIncome_imported"
    assert "GrossIncome_imported" in imported_formula.get(qn("w:instr"))


def it_renames_bookmarks_that_collide_only_by_case() -> None:
    source = docx.Document()
    source.add_paragraph("source term")
    create_bookmark(source, find_one(source, "source term"), "term")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), " REF term ")
    source.paragraphs[0]._p.append(field)
    destination = docx.Document()
    destination.add_paragraph("destination term")
    create_bookmark(destination, find_one(destination, "destination term"), "Term")

    report = insert_blocks_from(
        destination, source, "source term", anchor="destination term"
    )

    assert report.bookmarks_renamed["term"] == "term_imported"
    imported_field = next(
        node
        for node in destination.element.iter(qn("w:fldSimple"))
        if "term_imported" in (node.get(qn("w:instr")) or "")
    )
    assert imported_field is not None


def it_refuses_a_bookmark_reference_whose_definition_is_outside_the_range() -> None:
    source = docx.Document()
    source.add_paragraph("defined term")
    create_bookmark(source, find_one(source, "defined term"), "Term")
    reference_paragraph = source.add_paragraph("Reference block")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), " REF Term ")
    reference_paragraph._p.append(field)
    destination = docx.Document()
    destination.add_paragraph("Destination")
    before = destination.element.xml

    with pytest.raises(UnsupportedStructureError, match="outside the copied range"):
        insert_blocks_from(
            destination, source, "Reference block", anchor="Destination"
        )

    assert destination.element.xml == before


def it_refuses_numbering_that_depends_on_a_linked_style() -> None:
    source = docx.Document()
    num_id = ensure_decimal_definition(source)
    numbering = source.part.numbering_part.element
    num = next(
        node
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) == str(num_id)
    )
    abstract_id = num.find(qn("w:abstractNumId")).get(qn("w:val"))
    abstract = next(
        node
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) == abstract_id
    )
    style_link = OxmlElement("w:numStyleLink")
    style_link.set(qn("w:val"), "NumLinkStyle")
    abstract.append(style_link)
    paragraph = source.add_paragraph("Linked numbering")
    apply_numbering(paragraph, num_id=num_id)
    destination = docx.Document()
    destination.add_paragraph("Destination")
    ensure_decimal_definition(destination)
    before = destination.element.xml

    with pytest.raises(UnsupportedStructureError, match="linked numbering style"):
        insert_blocks_from(
            destination, source, "Linked numbering", anchor="Destination"
        )

    assert destination.element.xml == before
