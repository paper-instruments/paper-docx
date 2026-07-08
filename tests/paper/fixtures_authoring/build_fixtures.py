#!/usr/bin/env python3
"""Authoring script for the *generated* provenance bucket of the fixture corpus.

Run manually from the repo root:

    uv run --no-sync python tests/paper/fixtures_authoring/build_fixtures.py

This script is development tooling, NOT code under test and NOT run by tests.
Fixtures it writes are frozen by ``tests/paper/fixtures/MANIFEST.sha256``; the
manifest test fails on any hash drift, so regenerating a fixture is always a
deliberate, reviewed act (CONVENTIONS §4).

Determinism: every date stamped into a fixture is a fixed constant below, and
generated packages are rewritten with fixed zip entry timestamps so that
re-running the script on unchanged inputs yields byte-identical fixtures.
LibreOffice-exported variants (``--lo``) are kept verbatim as LibreOffice
wrote them — their provenance is the point — so only their *part content* is
checked, not their zip metadata.

The exotic WordprocessingML vocabulary (tracked changes, content controls,
text boxes, footnotes/endnotes, numbering definitions) is hand-written XML
below, mined from ECMA-376 shapes; python-docx only assembles the base
document. Provenance for this bucket is honestly labeled "generated".
"""

from __future__ import annotations

import datetime as dt
import io
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from pathlib import Path
from typing import Callable, Dict, Iterable, List, Tuple

REPO_ROOT = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(REPO_ROOT))

import docx  # noqa: E402
from docx import Document  # noqa: E402
from docx.oxml.ns import nsdecls, qn  # noqa: E402
from docx.oxml.parser import parse_xml  # noqa: E402

from tests.paper.harness import checks  # noqa: E402
from tests.paper.harness.lo import libreoffice_converts, soffice_binary  # noqa: E402
from tests.paper.harness.paths import FIXTURES_DIR  # noqa: E402

# ---------------------------------------------------------------------------
# fixed constants — all dates in fixtures come from here
# ---------------------------------------------------------------------------

AUTHOR_A = "Alice Editor"
AUTHOR_B = "Bob Reviewer"
DATE_A = "2026-06-01T09:30:00Z"
DATE_B = "2026-06-02T14:45:00Z"
COMMENT_INSTANT = dt.datetime(2026, 6, 3, 10, 15, 0, tzinfo=dt.timezone.utc)

#: Fixed zip entry timestamp for generated fixtures (zip epoch).
FIXED_ZIP_DATE = (1980, 1, 1, 0, 0, 0)

GENERATED_DIR = FIXTURES_DIR / "generated"
LIBREOFFICE_DIR = FIXTURES_DIR / "libreoffice"

W = nsdecls("w")

# ---------------------------------------------------------------------------
# zip plumbing
# ---------------------------------------------------------------------------


def read_parts(data: bytes) -> Tuple[Dict[str, bytes], List[str]]:
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        order = zf.namelist()
        return {name: zf.read(name) for name in order}, order


def write_parts(parts: Dict[str, bytes], order: Iterable[str]) -> bytes:
    """Zip `parts` in `order` with fixed entry timestamps (deterministic)."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for name in order:
            info = zipfile.ZipInfo(name, date_time=FIXED_ZIP_DATE)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o644 << 16
            zf.writestr(info, parts[name])
    return buf.getvalue()


def save_frozen(
    document: "docx.document.Document",
    out_path: Path,
    *,
    post: Callable[[Dict[str, bytes], List[str]], None] | None = None,
) -> None:
    """Save `document`, optionally transform its parts, and freeze the zip."""
    buf = io.BytesIO()
    document.save(buf)
    parts, order = read_parts(buf.getvalue())
    if post is not None:
        post(parts, order)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_bytes(write_parts(parts, order))


# ---------------------------------------------------------------------------
# part-transform helpers (used via save_frozen(post=...))
# ---------------------------------------------------------------------------


def add_content_type_overrides(parts: Dict[str, bytes], overrides: Dict[str, str]) -> None:
    """Insert Override elements for {part-name: content-type} into [Content_Types].xml."""
    blob = parts["[Content_Types].xml"].decode("utf-8")
    inserts = "".join(
        f'<Override PartName="/{name}" ContentType="{ctype}"/>' for name, ctype in overrides.items()
    )
    assert "</Types>" in blob
    parts["[Content_Types].xml"] = blob.replace("</Types>", inserts + "</Types>").encode("utf-8")


def add_document_relationships(parts: Dict[str, bytes], rels: Dict[str, str]) -> None:
    """Append {reltype: target} relationships to word/_rels/document.xml.rels."""
    name = "word/_rels/document.xml.rels"
    blob = parts[name].decode("utf-8")
    next_rid = 1 + max(
        (int(m.group(1)) for m in re.finditer(r'Id="rId(\d+)"', blob)), default=0
    )
    inserts = []
    for reltype, target in rels.items():
        inserts.append(f'<Relationship Id="rId{next_rid}" Type="{reltype}" Target="{target}"/>')
        next_rid += 1
    assert "</Relationships>" in blob
    parts[name] = blob.replace("</Relationships>", "".join(inserts) + "</Relationships>").encode(
        "utf-8"
    )


def add_part(parts: Dict[str, bytes], order: List[str], name: str, blob: bytes) -> None:
    assert name not in parts
    parts[name] = blob
    order.append(name)


NUMBERING_ABSTRACT_XML = (
    '<w:abstractNum w:abstractNumId="90">'
    '<w:multiLevelType w:val="hybridMultilevel"/>'
    '<w:lvl w:ilvl="0"><w:start w:val="1"/><w:numFmt w:val="decimal"/>'
    '<w:lvlText w:val="%1."/><w:lvlJc w:val="left"/>'
    '<w:pPr><w:ind w:left="720" w:hanging="360"/></w:pPr></w:lvl>'
    '<w:lvl w:ilvl="1"><w:start w:val="1"/><w:numFmt w:val="lowerLetter"/>'
    '<w:lvlText w:val="%2)"/><w:lvlJc w:val="left"/>'
    '<w:pPr><w:ind w:left="1440" w:hanging="360"/></w:pPr></w:lvl>'
    "</w:abstractNum>"
)
NUMBERING_NUM_XML = '<w:num w:numId="42"><w:abstractNumId w:val="90"/></w:num>'


def add_custom_numbering_definition(parts: Dict[str, bytes], order: List[str]) -> None:
    """Add abstractNum 90 / num 42 to word/numbering.xml, respecting child order.

    ``w:numbering`` requires all ``w:abstractNum`` elements before all
    ``w:num`` elements, so the abstract definition is inserted before the
    first existing ``w:num`` and the ``w:num`` goes at the end.
    """
    del order  # unused; numbering.xml already exists in the default template
    blob = parts["word/numbering.xml"].decode("utf-8")
    first_num = blob.find("<w:num ")
    assert first_num != -1, "default template numbering.xml lost its w:num elements?"
    blob = blob[:first_num] + NUMBERING_ABSTRACT_XML + blob[first_num:]
    assert "</w:numbering>" in blob
    blob = blob.replace("</w:numbering>", NUMBERING_NUM_XML + "</w:numbering>")
    parts["word/numbering.xml"] = blob.encode("utf-8")


FOOTNOTES_XML = (
    f"<w:footnotes {W}>"
    '<w:footnote w:type="separator" w:id="-1">'
    "<w:p><w:r><w:separator/></w:r></w:p></w:footnote>"
    '<w:footnote w:type="continuationSeparator" w:id="0">'
    "<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>"
    '<w:footnote w:id="1"><w:p><w:r><w:footnoteRef/></w:r>'
    '<w:r><w:t xml:space="preserve"> The footnote body text, hand-authored.</w:t></w:r>'
    "</w:p></w:footnote>"
    "</w:footnotes>"
).encode("utf-8")

ENDNOTES_XML = (
    f"<w:endnotes {W}>"
    '<w:endnote w:type="separator" w:id="-1">'
    "<w:p><w:r><w:separator/></w:r></w:p></w:endnote>"
    '<w:endnote w:type="continuationSeparator" w:id="0">'
    "<w:p><w:r><w:continuationSeparator/></w:r></w:p></w:endnote>"
    '<w:endnote w:id="1"><w:p><w:r><w:endnoteRef/></w:r>'
    '<w:r><w:t xml:space="preserve"> The endnote body text, hand-authored.</w:t></w:r>'
    "</w:p></w:endnote>"
    "</w:endnotes>"
).encode("utf-8")

FOOTNOTES_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"
ENDNOTES_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml"
FOOTNOTES_RT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes"
ENDNOTES_RT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes"


def add_footnotes_and_endnotes_parts(parts: Dict[str, bytes], order: List[str]) -> None:
    add_part(parts, order, "word/footnotes.xml", FOOTNOTES_XML)
    add_part(parts, order, "word/endnotes.xml", ENDNOTES_XML)
    add_content_type_overrides(
        parts, {"word/footnotes.xml": FOOTNOTES_CT, "word/endnotes.xml": ENDNOTES_CT}
    )
    add_document_relationships(
        parts, {FOOTNOTES_RT: "footnotes.xml", ENDNOTES_RT: "endnotes.xml"}
    )


# ---------------------------------------------------------------------------
# live-tree building blocks (hand-written WordprocessingML fragments)
# ---------------------------------------------------------------------------


def append_block(document: "docx.document.Document", element) -> None:
    """Append a block element to the document body, before the trailing sectPr."""
    body = document.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(element)
    else:
        body.append(element)


def tracked_change_paragraph() -> "docx.oxml.xmlchemy.BaseOxmlElement":
    """`w:p` with a deletion by Bob and an insertion by Alice, plain runs around."""
    return parse_xml(
        f"<w:p {W}>"
        '<w:r><w:t xml:space="preserve">The quarterly total was </w:t></w:r>'
        f'<w:del w:id="11" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
        '<w:r><w:delText xml:space="preserve">forty-two</w:delText></w:r>'
        "</w:del>"
        f'<w:ins w:id="12" w:author="{AUTHOR_A}" w:date="{DATE_A}">'
        "<w:r><w:t>forty-seven</w:t></w:r>"
        "</w:ins>"
        '<w:r><w:t xml:space="preserve"> units.</w:t></w:r>'
        "</w:p>"
    )


def tracked_insertion_paragraph() -> "docx.oxml.xmlchemy.BaseOxmlElement":
    """`w:p` whose only content is an insertion by Bob (second distinct author)."""
    return parse_xml(
        f"<w:p {W}>"
        f'<w:ins w:id="21" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
        "<w:r><w:t>This whole sentence was inserted with tracking on.</w:t></w:r>"
        "</w:ins>"
        "</w:p>"
    )


def block_content_control() -> "docx.oxml.xmlchemy.BaseOxmlElement":
    return parse_xml(
        f"<w:sdt {W}>"
        "<w:sdtPr>"
        '<w:alias w:val="ClauseBlock"/><w:tag w:val="clause-block-1"/><w:id w:val="1001"/>'
        "</w:sdtPr>"
        "<w:sdtContent>"
        "<w:p><w:r><w:t>Text inside a block-level content control.</w:t></w:r></w:p>"
        "</w:sdtContent>"
        "</w:sdt>"
    )


def inline_content_control_paragraph() -> "docx.oxml.xmlchemy.BaseOxmlElement":
    return parse_xml(
        f"<w:p {W}>"
        '<w:r><w:t xml:space="preserve">Inline control follows: </w:t></w:r>'
        "<w:sdt>"
        "<w:sdtPr>"
        '<w:alias w:val="InlineField"/><w:tag w:val="inline-field-1"/><w:id w:val="1002"/>'
        "</w:sdtPr>"
        "<w:sdtContent><w:r><w:t>controlled text</w:t></w:r></w:sdtContent>"
        "</w:sdt>"
        '<w:r><w:t xml:space="preserve"> and text after it.</w:t></w:r>'
        "</w:p>"
    )


#: A wps (wordprocessingShape) inline text box. Deliberately the plain
#: `w:drawing` form without the `mc:AlternateContent` + VML fallback wrapper
#: desktop Word emits — the realistic double-content shape is requested from a
#: human in FIXTURE-REQUESTS.md.
TEXTBOX_PARAGRAPH_XML = (
    "<w:p"
    ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
    ' xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"'
    ' xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"'
    ' xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
    "<w:r>"
    "<w:drawing>"
    '<wp:inline distT="0" distB="0" distL="0" distR="0">'
    '<wp:extent cx="2286000" cy="571500"/>'
    '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
    '<wp:docPr id="7" name="Text Box 7"/>'
    "<wp:cNvGraphicFramePr/>"
    "<a:graphic>"
    '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
    "<wps:wsp>"
    '<wps:cNvSpPr txBox="1"/>'
    "<wps:spPr>"
    '<a:xfrm><a:off x="0" y="0"/><a:ext cx="2286000" cy="571500"/></a:xfrm>'
    '<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
    '<a:solidFill><a:srgbClr val="FFFFFF"/></a:solidFill>'
    '<a:ln w="9525"><a:solidFill><a:srgbClr val="000000"/></a:solidFill></a:ln>'
    "</wps:spPr>"
    "<wps:txbx>"
    "<w:txbxContent>"
    "<w:p><w:r><w:t>Text living inside the text box.</w:t></w:r></w:p>"
    "</w:txbxContent>"
    "</wps:txbx>"
    '<wps:bodyPr rot="0" vert="horz" wrap="square" anchor="t" anchorCtr="0"/>'
    "</wps:wsp>"
    "</a:graphicData>"
    "</a:graphic>"
    "</wp:inline>"
    "</w:drawing>"
    "</w:r>"
    "</w:p>"
)


def textbox_paragraph() -> "docx.oxml.xmlchemy.BaseOxmlElement":
    return parse_xml(TEXTBOX_PARAGRAPH_XML)


def note_reference_paragraphs() -> List["docx.oxml.xmlchemy.BaseOxmlElement"]:
    """Two paragraphs referencing footnote 1 and endnote 1 respectively."""
    footnote_p = parse_xml(
        f"<w:p {W}>"
        "<w:r><w:t>This sentence carries a footnote reference.</w:t></w:r>"
        '<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
        '<w:footnoteReference w:id="1"/></w:r>'
        "</w:p>"
    )
    endnote_p = parse_xml(
        f"<w:p {W}>"
        "<w:r><w:t>This sentence carries an endnote reference.</w:t></w:r>"
        '<w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr>'
        '<w:endnoteReference w:id="1"/></w:r>'
        "</w:p>"
    )
    return [footnote_p, endnote_p]


def apply_custom_numbering(paragraph, ilvl: int) -> None:
    """Give `paragraph` a numPr referencing the hand-authored numId 42."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_pr.append(
        parse_xml(
            f'<w:numPr {W}><w:ilvl w:val="{ilvl}"/><w:numId w:val="42"/></w:numPr>'
        )
    )


def add_fragmented_rate_paragraph(document: "docx.document.Document") -> None:
    """The classic fragmented-run sentence: `$75–100/hr` with smart quotes.

    Run boundaries fall mid-token and mid-quoted-phrase; formatting changes
    (bold, italic) occur inside what reads as one phrase. Trailing/leading
    spaces ride on their own runs so `xml:space="preserve"` semantics are
    exercised.
    """
    p = document.add_paragraph()
    p.add_run("Consulting rate: ")
    p.add_run("$75").bold = True
    p.add_run("–").bold = True  # en dash, its own run
    p.add_run("100/hr").bold = True
    p.add_run(" on a ")
    p.add_run("“full-").italic = True  # left smart quote, split mid-word
    p.add_run("service”").italic = True
    p.add_run(" basis — travel time billed at $37.50/hr.")


def add_nbsp_paragraph(document: "docx.document.Document") -> None:
    document.add_paragraph(
        "Net 30 payment terms apply, with a 10 % late surcharge."
    )


def freeze_comment_dates(document: "docx.document.Document") -> None:
    """Overwrite the wall-clock date upstream `add_comment` stamps (determinism)."""
    comments_part = document.part._comments_part  # noqa: SLF001 - authoring tool
    for comment_elm in comments_part._element.comment_lst:  # noqa: SLF001
        comment_elm.date = COMMENT_INSTANT


#: Comment styles Word defines when the first comment is inserted. Upstream
#: add_comment references them (`CommentText` pStyle, `CommentReference`
#: rStyle) but the default template doesn't define them; a real Word package
#: defines both, so fixtures do too.
COMMENT_STYLES_XML = (
    '<w:style w:type="character" w:styleId="CommentReference">'
    '<w:name w:val="annotation reference"/>'
    '<w:rPr><w:sz w:val="16"/><w:szCs w:val="16"/></w:rPr></w:style>'
    '<w:style w:type="paragraph" w:styleId="CommentText">'
    '<w:name w:val="annotation text"/><w:basedOn w:val="Normal"/>'
    '<w:rPr><w:sz w:val="20"/><w:szCs w:val="20"/></w:rPr></w:style>'
)


def add_comment_style_definitions(parts: Dict[str, bytes], order: List[str]) -> None:
    del order
    blob = parts["word/styles.xml"].decode("utf-8")
    assert 'w:styleId="CommentReference"' not in blob
    assert "</w:styles>" in blob
    parts["word/styles.xml"] = blob.replace(
        "</w:styles>", COMMENT_STYLES_XML + "</w:styles>"
    ).encode("utf-8")


# ---------------------------------------------------------------------------
# fixture builders
# ---------------------------------------------------------------------------


def build_minimal(out: Path) -> None:
    doc = Document()
    doc.add_heading("Minimal Clean Document", level=1)
    doc.add_paragraph("First body paragraph with perfectly ordinary text.")
    doc.add_paragraph("Second body paragraph, equally unremarkable.")
    save_frozen(doc, out)


def build_tracked_ins_del(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Paragraph before any tracked edits.")
    append_block(doc, tracked_change_paragraph())
    append_block(doc, tracked_insertion_paragraph())
    doc.add_paragraph("Paragraph after the tracked edits.")
    save_frozen(doc, out)


def build_comments(out: Path) -> None:
    doc = Document()
    p1 = doc.add_paragraph("This paragraph carries the first comment.")
    p2 = doc.add_paragraph("A second paragraph with another comment target.")
    doc.add_comment(p1.runs, text="Please double-check this figure.", author=AUTHOR_A, initials="AE")
    doc.add_comment(p2.runs, text="Approved as written.", author=AUTHOR_B, initials="BR")
    freeze_comment_dates(doc)
    save_frozen(doc, out, post=add_comment_style_definitions)


def build_content_control(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Paragraph before the content controls.")
    append_block(doc, block_content_control())
    append_block(doc, inline_content_control_paragraph())
    doc.add_paragraph("Paragraph after the content controls.")
    save_frozen(doc, out)


def build_textbox(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Body text before the text box.")
    append_block(doc, textbox_paragraph())
    doc.add_paragraph("Body text after the text box.")
    save_frozen(doc, out)


def build_footnotes_endnotes(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Opening paragraph without any notes.")
    for p in note_reference_paragraphs():
        append_block(doc, p)
    save_frozen(doc, out, post=add_footnotes_and_endnotes_parts)


def build_header_footer_sections(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Section one body paragraph.")
    s1 = doc.sections[0]
    s1.header.is_linked_to_previous = False
    s1.header.paragraphs[0].text = "Header for section one"
    s1.footer.is_linked_to_previous = False
    s1.footer.paragraphs[0].text = "Footer for section one"
    s1.different_first_page_header_footer = True
    s1.first_page_header.is_linked_to_previous = False
    s1.first_page_header.paragraphs[0].text = "First-page header for section one"

    s2 = doc.add_section()
    doc.add_paragraph("Section two body paragraph.")
    s2.header.is_linked_to_previous = False
    s2.header.paragraphs[0].text = "Header for section two"
    s2.footer.is_linked_to_previous = False
    s2.footer.paragraphs[0].text = "Footer for section two"
    save_frozen(doc, out)


def build_table_merged_nested(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Paragraph before the merged and nested tables.")
    table = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            table.cell(r, c).text = f"R{r}C{c}"
    table.cell(0, 0).merge(table.cell(0, 1))  # horizontal: gridSpan
    table.cell(1, 2).merge(table.cell(2, 2))  # vertical: vMerge
    inner = table.cell(2, 0).add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            inner.cell(r, c).text = f"N{r}{c}"
    doc.add_paragraph("Paragraph after the tables.")
    save_frozen(doc, out)


def build_numbering_custom(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Introduction to the hand-numbered list.")
    apply_custom_numbering(doc.add_paragraph("First numbered item"), ilvl=0)
    apply_custom_numbering(doc.add_paragraph("Second numbered item"), ilvl=0)
    apply_custom_numbering(doc.add_paragraph("Nested lettered item"), ilvl=1)
    save_frozen(doc, out, post=add_custom_numbering_definition)


def build_fragmented_runs(out: Path) -> None:
    doc = Document()
    add_fragmented_rate_paragraph(doc)
    add_nbsp_paragraph(doc)
    save_frozen(doc, out)


def build_gauntlet(out: Path) -> None:
    """Everything ugly combined in one document."""
    doc = Document()
    doc.add_heading("Gauntlet Document", level=1)

    s1 = doc.sections[0]
    s1.header.is_linked_to_previous = False
    s1.header.paragraphs[0].text = "Gauntlet header, section one"
    s1.footer.is_linked_to_previous = False
    s1.footer.paragraphs[0].text = "Gauntlet footer, section one"
    s1.different_first_page_header_footer = True
    s1.first_page_header.is_linked_to_previous = False
    s1.first_page_header.paragraphs[0].text = "Gauntlet first-page header"

    add_fragmented_rate_paragraph(doc)
    add_nbsp_paragraph(doc)
    append_block(doc, tracked_change_paragraph())
    append_block(doc, tracked_insertion_paragraph())

    commented = doc.add_paragraph("This gauntlet paragraph carries a comment.")
    doc.add_comment(commented.runs, text="Gauntlet comment one.", author=AUTHOR_A, initials="AE")

    append_block(doc, block_content_control())
    append_block(doc, inline_content_control_paragraph())
    append_block(doc, textbox_paragraph())
    for p in tracked_move_paragraphs():
        append_block(doc, p)
    for p in format_change_paragraphs():
        append_block(doc, p)
    for p in field_paragraphs():
        append_block(doc, p)
    append_block(doc, placeholder_control_paragraph())
    append_block(doc, bookmark_paragraph())
    for p in toc_field_paragraphs():
        append_block(doc, p)
    noisy = doc.add_paragraph("Gauntlet paragrah with proofing noise.")
    noisy._p.insert(  # noqa: SLF001 - authoring tool
        list(noisy._p).index(noisy.runs[0]._r),
        parse_xml(f'<w:proofErr {W} w:type="spellStart"/>'),
    )
    noisy._p.append(parse_xml(f'<w:proofErr {W} w:type="spellEnd"/>'))
    for p in note_reference_paragraphs():
        append_block(doc, p)

    table = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            table.cell(r, c).text = f"R{r}C{c}"
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(1, 2).merge(table.cell(2, 2))
    inner = table.cell(2, 0).add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            inner.cell(r, c).text = f"N{r}{c}"

    apply_custom_numbering(doc.add_paragraph("Gauntlet numbered item one"), ilvl=0)
    apply_custom_numbering(doc.add_paragraph("Gauntlet numbered item two"), ilvl=0)

    # v0.11: row revisions, rich format changes, tracked paragraph merge/split
    append_block(doc, row_revision_table(base_id=81))
    for p in rich_format_change_paragraphs():
        append_block(doc, p)
    for p in paragraph_merge_paragraphs():
        append_block(doc, p)

    s2 = doc.add_section()
    doc.add_paragraph("Gauntlet section two body paragraph.")
    s2.header.is_linked_to_previous = False
    s2.header.paragraphs[0].text = "Gauntlet header, section two"

    freeze_comment_dates(doc)

    def post(parts: Dict[str, bytes], order: List[str]) -> None:
        add_footnotes_and_endnotes_parts(parts, order)
        add_custom_numbering_definition(parts, order)
        add_comment_style_definitions(parts, order)
        add_placeholder_style_definition(parts, order)

    save_frozen(doc, out, post=post)


def build_large(out: Path) -> None:
    doc = Document()
    for i in range(1, 5001):
        if i % 250 == 1:
            doc.add_heading(f"Chapter {i // 250 + 1}", level=2)
        doc.add_paragraph(f"Paragraph {i}: steady-state filler text for the perf smoke fixture.")
    save_frozen(doc, out)


def tracked_move_paragraphs() -> "List[docx.oxml.xmlchemy.BaseOxmlElement]":
    """Three paragraphs: a moveFrom source, an untouched middle, a moveTo
    destination — the shape Word emits for drag/cut-paste with tracking on."""
    moved_text = "The indemnity clause relocated by tracked move."
    source = parse_xml(
        f"<w:p {W}>"
        f'<w:moveFromRangeStart w:id="61" w:author="{AUTHOR_A}" w:date="{DATE_A}"'
        ' w:name="move1"/>'
        f'<w:moveFrom w:id="62" w:author="{AUTHOR_A}" w:date="{DATE_A}">'
        f"<w:r><w:t>{moved_text}</w:t></w:r>"
        "</w:moveFrom>"
        '<w:moveFromRangeEnd w:id="61"/>'
        "</w:p>"
    )
    middle = parse_xml(f"<w:p {W}><w:r><w:t>Paragraph between the move ends.</w:t></w:r></w:p>")
    destination = parse_xml(
        f"<w:p {W}>"
        f'<w:moveToRangeStart w:id="63" w:author="{AUTHOR_A}" w:date="{DATE_A}"'
        ' w:name="move1"/>'
        f'<w:moveTo w:id="64" w:author="{AUTHOR_A}" w:date="{DATE_A}">'
        f"<w:r><w:t>{moved_text}</w:t></w:r>"
        "</w:moveTo>"
        '<w:moveToRangeEnd w:id="63"/>'
        "</w:p>"
    )
    return [source, middle, destination]


def build_tracked_moves(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Paragraph before the tracked move.")
    for p in tracked_move_paragraphs():
        append_block(doc, p)
    doc.add_paragraph("Paragraph after the tracked move.")
    save_frozen(doc, out)


def format_change_paragraphs() -> "List[docx.oxml.xmlchemy.BaseOxmlElement]":
    """A run made bold with tracking on (w:rPrChange) and a paragraph
    re-centered with tracking on (w:pPrChange)."""
    run_change = parse_xml(
        f"<w:p {W}><w:r><w:rPr><w:b/>"
        f'<w:rPrChange w:id="71" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
        "<w:rPr/></w:rPrChange>"
        "</w:rPr><w:t>This text was bolded with tracking on.</w:t></w:r></w:p>"
    )
    par_change = parse_xml(
        f"<w:p {W}><w:pPr>"
        '<w:jc w:val="center"/>'
        f'<w:pPrChange w:id="72" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
        "<w:pPr/></w:pPrChange>"
        "</w:pPr><w:r><w:t>This paragraph was centered with tracking on.</w:t></w:r></w:p>"
    )
    return [run_change, par_change]


def build_format_changes(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Paragraph before the tracked formatting changes.")
    for p in format_change_paragraphs():
        append_block(doc, p)
    save_frozen(doc, out)


def field_paragraphs() -> "List[docx.oxml.xmlchemy.BaseOxmlElement]":
    """Both field shapes: a w:fldSimple DATE and a complex PAGEREF field
    (fldChar begin / instrText / separate / cached result / end)."""
    simple = parse_xml(
        f"<w:p {W}>"
        '<w:r><w:t xml:space="preserve">Dated: </w:t></w:r>'
        '<w:fldSimple w:instr=" DATE \\@ &quot;MMMM d, yyyy&quot; ">'
        "<w:r><w:t>June 1, 2026</w:t></w:r>"
        "</w:fldSimple>"
        "</w:p>"
    )
    complex_field = parse_xml(
        f"<w:p {W}>"
        '<w:r><w:t xml:space="preserve">See page </w:t></w:r>'
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> PAGEREF _RefAnchor \\h </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        "<w:r><w:t>14</w:t></w:r>"
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        '<w:r><w:t xml:space="preserve"> for the details.</w:t></w:r>'
        "</w:p>"
    )
    return [simple, complex_field]


def build_fields(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Paragraph before the fields.")
    for p in field_paragraphs():
        append_block(doc, p)
    save_frozen(doc, out)


PLACEHOLDER_TEXT = "Click or tap here to enter text."

PLACEHOLDER_STYLE_XML = (
    '<w:style w:type="character" w:styleId="PlaceholderText">'
    '<w:name w:val="Placeholder Text"/>'
    '<w:rPr><w:color w:val="808080"/></w:rPr></w:style>'
)


def add_placeholder_style_definition(parts: Dict[str, bytes], order: List[str]) -> None:
    del order
    blob = parts["word/styles.xml"].decode("utf-8")
    assert 'w:styleId="PlaceholderText"' not in blob
    parts["word/styles.xml"] = blob.replace(
        "</w:styles>", PLACEHOLDER_STYLE_XML + "</w:styles>"
    ).encode("utf-8")


def placeholder_control_paragraph() -> "docx.oxml.xmlchemy.BaseOxmlElement":
    """An unfilled form control exactly as Word templates ship them:
    w:showingPlcHdr set, PlaceholderText-styled prompt run."""
    return parse_xml(
        f"<w:p {W}>"
        '<w:r><w:t xml:space="preserve">Client name: </w:t></w:r>'
        "<w:sdt><w:sdtPr>"
        '<w:alias w:val="ClientName"/><w:tag w:val="client-name"/>'
        '<w:id w:val="2001"/><w:showingPlcHdr/>'
        "</w:sdtPr><w:sdtContent>"
        '<w:r><w:rPr><w:rStyle w:val="PlaceholderText"/></w:rPr>'
        f"<w:t>{PLACEHOLDER_TEXT}</w:t></w:r>"
        "</w:sdtContent></w:sdt>"
        "</w:p>"
    )


def build_placeholder_control(out: Path) -> None:
    doc = Document()
    append_block(doc, placeholder_control_paragraph())
    doc.add_paragraph("Paragraph after the form control.")
    save_frozen(doc, out, post=add_placeholder_style_definition)


def bookmark_paragraph() -> "docx.oxml.xmlchemy.BaseOxmlElement":
    """A named non-empty bookmark (cross-reference target) plus Word's
    ubiquitous point bookmark `_GoBack`."""
    return parse_xml(
        f"<w:p {W}>"
        '<w:r><w:t xml:space="preserve">See </w:t></w:r>'
        '<w:bookmarkStart w:id="1" w:name="DefinedTerm"/>'
        "<w:r><w:t>the Master Agreement</w:t></w:r>"
        '<w:bookmarkEnd w:id="1"/>'
        '<w:r><w:t xml:space="preserve"> for definitions.</w:t></w:r>'
        '<w:bookmarkStart w:id="2" w:name="_GoBack"/>'
        '<w:bookmarkEnd w:id="2"/>'
        "</w:p>"
    )


def build_bookmarks(out: Path) -> None:
    doc = Document()
    append_block(doc, bookmark_paragraph())
    doc.add_paragraph("Paragraph after the bookmarks.")
    save_frozen(doc, out)


def build_noisy_markup(out: Path) -> None:
    """proofErr spell-check noise + comment anchors + _GoBack around ordinary
    paragraphs — the markup Word scatters through virtually every saved doc.
    (Hand-built; a real-Word capture is requested in FIXTURE-REQUESTS.md.)"""
    doc = Document()
    p1 = doc.add_paragraph("Paragrah with a spelling issue.")
    p1._p.insert(  # noqa: SLF001 - authoring tool
        list(p1._p).index(p1.runs[0]._r),
        parse_xml(f'<w:proofErr {W} w:type="spellStart"/>'),
    )
    p1._p.append(parse_xml(f'<w:proofErr {W} w:type="spellEnd"/>'))
    p1._p.append(parse_xml(f'<w:bookmarkStart {W} w:id="9" w:name="_GoBack"/>'))
    p1._p.append(parse_xml(f'<w:bookmarkEnd {W} w:id="9"/>'))
    p2 = doc.add_paragraph("This clause carries a reviewer comment.")
    doc.add_comment(p2.runs, text="Noise-tolerance target.", author=AUTHOR_A, initials="AE")
    doc.add_paragraph("Paragraph after the noisy ones.")
    freeze_comment_dates(doc)
    save_frozen(doc, out, post=add_comment_style_definitions)


def toc_field_paragraphs() -> "List[docx.oxml.xmlchemy.BaseOxmlElement]":
    """A TOC-shaped multi-paragraph complex field: begin+instr+separate in the
    first paragraph, an entry paragraph in the middle, end in the last —
    the canonical shape whose cross-block state the v0.1 review caught."""
    first = parse_xml(
        f"<w:p {W}>"
        '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
        '<w:r><w:instrText xml:space="preserve"> TOC \\o &quot;1-3&quot; \\h </w:instrText></w:r>'
        '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
        "<w:r><w:t>Chapter One entry	4</w:t></w:r>"
        "</w:p>"
    )
    middle = parse_xml(
        f"<w:p {W}><w:r><w:t>Chapter Two entry	9</w:t></w:r></w:p>"
    )
    last = parse_xml(
        f"<w:p {W}>"
        "<w:r><w:t>Chapter Three entry	14</w:t></w:r>"
        '<w:r><w:fldChar w:fldCharType="end"/></w:r>'
        "</w:p>"
    )
    return [first, middle, last]


def build_toc_field(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Contents heading before the TOC field.")
    for p in toc_field_paragraphs():
        append_block(doc, p)
    doc.add_paragraph("Body paragraph after the TOC field.")
    save_frozen(doc, out)


# ---------------------------------------------------------------------------
# v0.11 redline-pipeline fixtures (PLAN-v0.11 Phase 0)
# ---------------------------------------------------------------------------


def rich_format_change_paragraphs() -> "List[docx.oxml.xmlchemy.BaseOxmlElement]":
    """Format-change revisions whose stored previous properties are NON-empty
    (reject must restore them, not just drop the change element), plus a
    paragraph-MARK formatting change (`w:pPr/w:rPr/w:rPrChange`)."""
    run_change = parse_xml(
        f"<w:p {W}><w:r><w:rPr><w:b/>"
        f'<w:rPrChange w:id="91" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
        '<w:rPr><w:i/><w:sz w:val="28"/></w:rPr></w:rPrChange>'
        "</w:rPr><w:t>Delivery follows the schedule in Exhibit A.</w:t></w:r></w:p>"
    )
    par_change = parse_xml(
        f"<w:p {W}><w:pPr>"
        '<w:jc w:val="center"/>'
        f'<w:pPrChange w:id="92" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
        '<w:pPr><w:jc w:val="right"/></w:pPr></w:pPrChange>'
        "</w:pPr><w:r><w:t>This paragraph was right-aligned before re-centering.</w:t></w:r></w:p>"
    )
    mark_change = parse_xml(
        f"<w:p {W}><w:pPr><w:rPr><w:b/>"
        f'<w:rPrChange w:id="93" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
        "<w:rPr/></w:rPrChange>"
        "</w:rPr></w:pPr><w:r><w:t>The paragraph mark itself was re-formatted.</w:t></w:r></w:p>"
    )
    return [run_change, par_change, mark_change]


def build_format_changes_rich(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Paragraph before the rich formatting changes.")
    for p in rich_format_change_paragraphs():
        append_block(doc, p)
    doc.add_paragraph("Paragraph after the rich formatting changes.")
    save_frozen(doc, out)


def paragraph_merge_paragraphs() -> "List[docx.oxml.xmlchemy.BaseOxmlElement]":
    """Two tracked paragraph-mark revisions: a deleted mark (Word's deleted
    pilcrow — accept merges the two paragraphs) and an inserted mark (a
    tracked split — reject merges them back)."""
    return [
        parse_xml(
            f"<w:p {W}><w:pPr><w:rPr>"
            f'<w:del w:id="94" w:author="{AUTHOR_B}" w:date="{DATE_B}"/>'
            "</w:rPr></w:pPr>"
            '<w:r><w:t xml:space="preserve">This sentence continues </w:t></w:r></w:p>'
        ),
        parse_xml(f"<w:p {W}><w:r><w:t>onto the following line.</w:t></w:r></w:p>"),
        parse_xml(
            f"<w:p {W}><w:pPr><w:rPr>"
            f'<w:ins w:id="95" w:author="{AUTHOR_A}" w:date="{DATE_A}"/>'
            "</w:rPr></w:pPr>"
            '<w:r><w:t xml:space="preserve">A tracked split divides </w:t></w:r></w:p>'
        ),
        parse_xml(f"<w:p {W}><w:r><w:t>this once-single sentence.</w:t></w:r></w:p>"),
    ]


def build_paragraph_merge(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Paragraph before the tracked paragraph-mark revisions.")
    for p in paragraph_merge_paragraphs():
        append_block(doc, p)
    doc.add_paragraph("Paragraph after the tracked paragraph-mark revisions.")
    save_frozen(doc, out)


def row_revision_table(base_id: int) -> "docx.oxml.xmlchemy.BaseOxmlElement":
    """A table with a plain header row, a row inserted with tracking on
    (`w:trPr/w:ins`, ins-wrapped cell content, ins-stamped cell paragraph
    marks) and a row deleted with tracking on (`w:trPr/w:del`, `w:delText`
    cell content, del-stamped marks) — Word's row-revision shapes."""
    ids = list(range(base_id, base_id + 10))
    tc_w = '<w:tcPr><w:tcW w:w="4675" w:type="dxa"/></w:tcPr>'

    def inserted_cell(text: str, mark_id: int, content_id: int) -> str:
        return (
            f"<w:tc>{tc_w}<w:p>"
            f'<w:pPr><w:rPr><w:ins w:id="{mark_id}" w:author="{AUTHOR_A}"'
            f' w:date="{DATE_A}"/></w:rPr></w:pPr>'
            f'<w:ins w:id="{content_id}" w:author="{AUTHOR_A}" w:date="{DATE_A}">'
            f"<w:r><w:t>{text}</w:t></w:r></w:ins>"
            "</w:p></w:tc>"
        )

    def deleted_cell(text: str, mark_id: int, content_id: int) -> str:
        return (
            f"<w:tc>{tc_w}<w:p>"
            f'<w:pPr><w:rPr><w:del w:id="{mark_id}" w:author="{AUTHOR_B}"'
            f' w:date="{DATE_B}"/></w:rPr></w:pPr>'
            f'<w:del w:id="{content_id}" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
            f"<w:r><w:delText>{text}</w:delText></w:r></w:del>"
            "</w:p></w:tc>"
        )

    header = (
        f"<w:tr><w:tc>{tc_w}<w:p><w:r><w:t>Item</w:t></w:r></w:p></w:tc>"
        f"<w:tc>{tc_w}<w:p><w:r><w:t>Amount</w:t></w:r></w:p></w:tc></w:tr>"
    )
    inserted_row = (
        f'<w:tr><w:trPr><w:ins w:id="{ids[0]}" w:author="{AUTHOR_A}"'
        f' w:date="{DATE_A}"/></w:trPr>'
        + inserted_cell("Filing fee", ids[1], ids[2])
        + inserted_cell("$100", ids[3], ids[4])
        + "</w:tr>"
    )
    deleted_row = (
        f'<w:tr><w:trPr><w:del w:id="{ids[5]}" w:author="{AUTHOR_B}"'
        f' w:date="{DATE_B}"/></w:trPr>'
        + deleted_cell("Old charge", ids[6], ids[7])
        + deleted_cell("$50", ids[8], ids[9])
        + "</w:tr>"
    )
    return parse_xml(
        f"<w:tbl {W}>"
        '<w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr>'
        '<w:tblGrid><w:gridCol w:w="4675"/><w:gridCol w:w="4675"/></w:tblGrid>'
        + header
        + inserted_row
        + deleted_row
        + "</w:tbl>"
    )


def plain_result_table() -> "docx.oxml.xmlchemy.BaseOxmlElement":
    """`row_revision_table` with all row revisions accepted (hand-computed)."""
    tc_w = '<w:tcPr><w:tcW w:w="4675" w:type="dxa"/></w:tcPr>'
    return parse_xml(
        f"<w:tbl {W}>"
        '<w:tblPr><w:tblW w:w="0" w:type="auto"/></w:tblPr>'
        '<w:tblGrid><w:gridCol w:w="4675"/><w:gridCol w:w="4675"/></w:tblGrid>'
        f"<w:tr><w:tc>{tc_w}<w:p><w:r><w:t>Item</w:t></w:r></w:p></w:tc>"
        f"<w:tc>{tc_w}<w:p><w:r><w:t>Amount</w:t></w:r></w:p></w:tc></w:tr>"
        f"<w:tr><w:tc>{tc_w}<w:p><w:r><w:t>Filing fee</w:t></w:r></w:p></w:tc>"
        f"<w:tc>{tc_w}<w:p><w:r><w:t>$100</w:t></w:r></w:p></w:tc></w:tr>"
        "</w:tbl>"
    )


def build_row_revisions(out: Path) -> None:
    doc = Document()
    doc.add_paragraph("Paragraph before the row-revision table.")
    append_block(doc, row_revision_table(base_id=81))
    doc.add_paragraph("Paragraph after the row-revision table.")
    save_frozen(doc, out)


COMMENTS_CT = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
COMMENTS_RT = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"

MULTIROUND_COMMENTS_XML = (
    f"<w:comments {W}>"
    f'<w:comment w:id="0" w:author="{AUTHOR_B}" w:date="{DATE_B}" w:initials="BR">'
    '<w:p><w:pPr><w:pStyle w:val="CommentText"/></w:pPr>'
    '<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr><w:annotationRef/></w:r>'
    '<w:r><w:t xml:space="preserve">Confirm the execution venue.</w:t></w:r>'
    "</w:p></w:comment>"
    "</w:comments>"
).encode("utf-8")


def add_multiround_comment_part(parts: Dict[str, bytes], order: List[str]) -> None:
    add_part(parts, order, "word/comments.xml", MULTIROUND_COMMENTS_XML)
    add_content_type_overrides(parts, {"word/comments.xml": COMMENTS_CT})
    add_document_relationships(parts, {COMMENTS_RT: "comments.xml"})
    add_comment_style_definitions(parts, order)


MULTIROUND_MOVED_TEXT = "The indemnity clause shall survive termination of this agreement."


def multiround_paragraphs() -> "List[docx.oxml.xmlchemy.BaseOxmlElement]":
    """The two-author, two-round redline body (FIXTURE-REQUESTS §15 bootstrap).

    Every revision type v0.11 resolves, in one document: rPrChange (empty and
    rich stored properties), pPrChange, a mark-stamped whole-paragraph move,
    tracked row insert/delete, tracked paragraph merge + split, and a comment
    anchored inside a tracked insertion. Hand-built; the real-Word capture
    requested in FIXTURE-REQUESTS.md supersedes this bootstrap on arrival.
    """
    blocks: "List[docx.oxml.xmlchemy.BaseOxmlElement]" = [
        parse_xml(
            f"<w:p {W}><w:r><w:t>Engagement letter, revised across two rounds.</w:t></w:r></w:p>"
        ),
        parse_xml(
            f"<w:p {W}>"
            '<w:r><w:t xml:space="preserve">This </w:t></w:r>'
            "<w:r><w:rPr><w:b/>"
            f'<w:rPrChange w:id="101" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
            "<w:rPr/></w:rPrChange>"
            "</w:rPr><w:t>agreement</w:t></w:r>"
            '<w:r><w:t xml:space="preserve"> is made between the parties.</w:t></w:r>'
            "</w:p>"
        ),
        # move source: content and the paragraph mark both move away
        parse_xml(
            f"<w:p {W}>"
            f'<w:pPr><w:rPr><w:moveFrom w:id="102" w:author="{AUTHOR_A}"'
            f' w:date="{DATE_A}"/></w:rPr></w:pPr>'
            f'<w:moveFromRangeStart w:id="103" w:author="{AUTHOR_A}"'
            f' w:date="{DATE_A}" w:name="mrMove1"/>'
            f'<w:moveFrom w:id="104" w:author="{AUTHOR_A}" w:date="{DATE_A}">'
            f"<w:r><w:t>{MULTIROUND_MOVED_TEXT}</w:t></w:r>"
            "</w:moveFrom>"
            '<w:moveFromRangeEnd w:id="103"/>'
            "</w:p>"
        ),
        parse_xml(
            f"<w:p {W}><w:r><w:t>Middle paragraph between the move sites.</w:t></w:r></w:p>"
        ),
        # move destination
        parse_xml(
            f"<w:p {W}>"
            f'<w:pPr><w:rPr><w:moveTo w:id="105" w:author="{AUTHOR_A}"'
            f' w:date="{DATE_A}"/></w:rPr></w:pPr>'
            f'<w:moveToRangeStart w:id="106" w:author="{AUTHOR_A}"'
            f' w:date="{DATE_A}" w:name="mrMove1"/>'
            f'<w:moveTo w:id="107" w:author="{AUTHOR_A}" w:date="{DATE_A}">'
            f"<w:r><w:t>{MULTIROUND_MOVED_TEXT}</w:t></w:r>"
            "</w:moveTo>"
            '<w:moveToRangeEnd w:id="106"/>'
            "</w:p>"
        ),
        parse_xml(
            f"<w:p {W}><w:pPr>"
            '<w:jc w:val="center"/>'
            f'<w:pPrChange w:id="108" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
            '<w:pPr><w:jc w:val="right"/></w:pPr></w:pPrChange>'
            "</w:pPr><w:r><w:t>Payment is due within thirty days of invoice.</w:t></w:r></w:p>"
        ),
        parse_xml(
            f"<w:p {W}><w:r><w:rPr><w:b/>"
            f'<w:rPrChange w:id="109" w:author="{AUTHOR_B}" w:date="{DATE_B}">'
            '<w:rPr><w:i/><w:sz w:val="28"/></w:rPr></w:rPrChange>'
            "</w:rPr><w:t>Delivery follows the schedule in Exhibit A.</w:t></w:r></w:p>"
        ),
        row_revision_table(base_id=110),
    ]
    blocks.extend(paragraph_merge_paragraphs())
    blocks.append(
        parse_xml(
            f"<w:p {W}>"
            '<w:r><w:t xml:space="preserve">Signed at </w:t></w:r>'
            f'<w:ins w:id="122" w:author="{AUTHOR_A}" w:date="{DATE_A}">'
            '<w:commentRangeStart w:id="0"/>'
            "<w:r><w:t>the offices of the Client</w:t></w:r>"
            '<w:commentRangeEnd w:id="0"/>'
            '<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
            '<w:commentReference w:id="0"/></w:r>'
            "</w:ins>"
            "<w:r><w:t>.</w:t></w:r>"
            "</w:p>"
        )
    )
    blocks.append(
        parse_xml(
            f"<w:p {W}><w:r><w:t>Closing paragraph after all tracked activity.</w:t></w:r></w:p>"
        )
    )
    return blocks


def build_redline_multiround(out: Path) -> None:
    doc = Document()
    for block in multiround_paragraphs():
        append_block(doc, block)
    save_frozen(doc, out, post=add_multiround_comment_part)


def build_redline_multiround_accepted(out: Path) -> None:
    """`multiround.docx` with every revision applied, hand-computed to Word's
    Accept All semantics — the resolution ground truth until the real-Word
    capture (FIXTURE-REQUESTS §16) supersedes it."""
    doc = Document()
    for block in [
        parse_xml(
            f"<w:p {W}><w:r><w:t>Engagement letter, revised across two rounds.</w:t></w:r></w:p>"
        ),
        parse_xml(
            f"<w:p {W}>"
            '<w:r><w:t xml:space="preserve">This </w:t></w:r>'
            "<w:r><w:rPr><w:b/></w:rPr><w:t>agreement</w:t></w:r>"
            '<w:r><w:t xml:space="preserve"> is made between the parties.</w:t></w:r>'
            "</w:p>"
        ),
        # accepted move: source paragraph gone, destination plain
        parse_xml(
            f"<w:p {W}><w:r><w:t>Middle paragraph between the move sites.</w:t></w:r></w:p>"
        ),
        parse_xml(f"<w:p {W}><w:r><w:t>{MULTIROUND_MOVED_TEXT}</w:t></w:r></w:p>"),
        parse_xml(
            f"<w:p {W}><w:pPr><w:jc w:val=\"center\"/></w:pPr>"
            "<w:r><w:t>Payment is due within thirty days of invoice.</w:t></w:r></w:p>"
        ),
        parse_xml(
            f"<w:p {W}><w:r><w:rPr><w:b/></w:rPr>"
            "<w:t>Delivery follows the schedule in Exhibit A.</w:t></w:r></w:p>"
        ),
        plain_result_table(),
        # merged pair (deleted mark applied), split pair (inserted mark kept)
        parse_xml(
            f"<w:p {W}><w:r>"
            '<w:t xml:space="preserve">This sentence continues </w:t></w:r>'
            "<w:r><w:t>onto the following line.</w:t></w:r></w:p>"
        ),
        parse_xml(
            f'<w:p {W}><w:r><w:t xml:space="preserve">A tracked split divides </w:t></w:r></w:p>'
        ),
        parse_xml(f"<w:p {W}><w:r><w:t>this once-single sentence.</w:t></w:r></w:p>"),
        parse_xml(
            f"<w:p {W}>"
            '<w:r><w:t xml:space="preserve">Signed at </w:t></w:r>'
            '<w:commentRangeStart w:id="0"/>'
            "<w:r><w:t>the offices of the Client</w:t></w:r>"
            '<w:commentRangeEnd w:id="0"/>'
            '<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
            '<w:commentReference w:id="0"/></w:r>'
            "<w:r><w:t>.</w:t></w:r>"
            "</w:p>"
        ),
        parse_xml(
            f"<w:p {W}><w:r><w:t>Closing paragraph after all tracked activity.</w:t></w:r></w:p>"
        ),
    ]:
        append_block(doc, block)
    save_frozen(doc, out, post=add_multiround_comment_part)
    blob = part_bytes(out, "word/document.xml")
    for marker in (b"<w:ins ", b"<w:del ", b"<w:moveFrom", b"<w:moveTo", b"PrChange"):
        assert marker not in blob, f"accepted ground truth still carries {marker!r}"


def build_compare_original(out: Path) -> None:
    doc = Document()
    doc.add_heading("Service Agreement", level=1)
    doc.add_paragraph("This agreement governs the provision of consulting services.")
    doc.add_paragraph("The consultant will deliver monthly progress reports.")
    doc.add_paragraph("Either party may terminate with thirty days notice.")
    doc.add_heading("Fees", level=2)
    doc.add_paragraph("Fees are payable within thirty days of invoice.")
    doc.add_paragraph("Late payments accrue interest at one percent monthly.")
    table = doc.add_table(rows=3, cols=2)
    for row, (service, rate) in zip(
        table.rows, [("Service", "Rate"), ("Advisory", "$200"), ("Drafting", "$150")]
    ):
        row.cells[0].text = service
        row.cells[1].text = rate
    doc.add_heading("Confidentiality", level=2)
    doc.add_paragraph("Each party shall protect the other's confidential information.")
    save_frozen(doc, out)


def build_compare_revised(out: Path) -> None:
    """`compare-original.docx` edited with tracking OFF: a word-level rewording,
    a deleted paragraph, an added paragraph, a changed table cell, and the
    Confidentiality section moved ahead of Fees."""
    doc = Document()
    doc.add_heading("Service Agreement", level=1)
    doc.add_paragraph(
        "This agreement governs the provision of consulting and advisory services."
    )
    doc.add_paragraph("Either party may terminate with sixty days notice.")
    doc.add_heading("Confidentiality", level=2)
    doc.add_paragraph("Each party shall protect the other's confidential information.")
    doc.add_heading("Fees", level=2)
    doc.add_paragraph("Fees are payable within thirty days of invoice.")
    doc.add_paragraph("Late payments accrue interest at one and a half percent monthly.")
    doc.add_paragraph("A retainer of one thousand dollars is due on signing.")
    table = doc.add_table(rows=3, cols=2)
    for row, (service, rate) in zip(
        table.rows, [("Service", "Rate"), ("Advisory", "$250"), ("Drafting", "$150")]
    ):
        row.cells[0].text = service
        row.cells[1].text = rate
    save_frozen(doc, out)


def add_document_protection(parts: Dict[str, bytes], edit: str) -> None:
    """Insert `w:documentProtection` into word/settings.xml at its schema
    position (before `w:defaultTabStop`; CT_Settings child sequence)."""
    blob = parts["word/settings.xml"].decode("utf-8")
    marker = "<w:defaultTabStop"
    assert marker in blob, "default template settings.xml lost defaultTabStop?"
    protection = f'<w:documentProtection w:edit="{edit}" w:enforcement="1"/>'
    parts["word/settings.xml"] = blob.replace(marker, protection + marker, 1).encode("utf-8")


def build_protected(out: Path, edit: str) -> None:
    """A document Word treats as protected (Restrict Editing, no password):
    `w:documentProtection w:edit="<mode>" w:enforcement="1"`."""
    doc = Document()
    doc.add_paragraph("This template is locked; only sanctioned edits apply.")
    append_block(doc, inline_content_control_paragraph())
    doc.add_paragraph("Paragraph after the form control.")
    save_frozen(doc, out, post=lambda parts, order: add_document_protection(parts, edit))


def build_protected_forms(out: Path) -> None:
    build_protected(out, "forms")


def build_protected_readonly(out: Path) -> None:
    build_protected(out, "readOnly")


def build_protected_tracked(out: Path) -> None:
    build_protected(out, "trackedChanges")


def build_corrupt_broken_rel(minimal: Path, out: Path) -> None:
    """Valid zip, valid XML, but document.xml.rels points at a missing part."""
    parts, order = read_parts(minimal.read_bytes())

    def post(parts: Dict[str, bytes], order: List[str]) -> None:
        del order
        name = "word/_rels/document.xml.rels"
        blob = parts[name].decode("utf-8")
        dangling = (
            '<Relationship Id="rId99"'
            ' Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"'
            ' Target="media/missing-image.png"/>'
        )
        parts[name] = blob.replace("</Relationships>", dangling + "</Relationships>").encode(
            "utf-8"
        )

    post(parts, order)
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_bytes(write_parts(parts, order))


def build_corrupt_malformed_xml(minimal: Path, out: Path) -> None:
    """Valid zip whose word/document.xml is truncated mid-element."""
    parts, order = read_parts(minimal.read_bytes())
    parts["word/document.xml"] = parts["word/document.xml"][:-40]
    out.parent.mkdir(parents=True, exist_ok=True)
    out.write_bytes(write_parts(parts, order))


# ---------------------------------------------------------------------------
# validation battery
# ---------------------------------------------------------------------------


def part_bytes(path: Path, name: str) -> bytes:
    with zipfile.ZipFile(path) as zf:
        return zf.read(name)


def part_names(path: Path) -> set:
    with zipfile.ZipFile(path) as zf:
        return set(zf.namelist())


#: fixture-stem -> list of (part-name, required-byte-pattern) feature probes.
FEATURE_PROBES: Dict[str, List[Tuple[str, bytes]]] = {
    "tracked-ins-del": [
        ("word/document.xml", b"<w:ins "),
        ("word/document.xml", b"<w:del "),
        ("word/document.xml", b"<w:delText"),
        ("word/document.xml", AUTHOR_A.encode()),
        ("word/document.xml", AUTHOR_B.encode()),
    ],
    "comments": [
        ("word/comments.xml", b"<w:comment "),
        ("word/comments.xml", AUTHOR_A.encode()),
        ("word/document.xml", b"commentReference"),
    ],
    "content-control": [
        ("word/document.xml", b"<w:sdt>"),
        ("word/document.xml", b"clause-block-1"),
        ("word/document.xml", b"inline-field-1"),
    ],
    "textbox": [
        ("word/document.xml", b"txbxContent"),
        ("word/document.xml", "Text living inside the text box.".encode()),
    ],
    "footnotes-endnotes": [
        ("word/footnotes.xml", b"<w:footnote "),
        ("word/endnotes.xml", b"<w:endnote "),
        ("word/document.xml", b"footnoteReference"),
        ("word/document.xml", b"endnoteReference"),
    ],
    "header-footer-sections": [
        ("word/document.xml", b"titlePg"),
    ],
    "table-merged-nested": [
        ("word/document.xml", b"gridSpan"),
        ("word/document.xml", b"vMerge"),
    ],
    "numbering-custom": [
        ("word/numbering.xml", b'w:abstractNumId="90"'),
        ("word/numbering.xml", b'<w:num w:numId="42"'),
        ("word/document.xml", b'<w:numId w:val="42"/>'),
    ],
    "fragmented-runs": [
        ("word/document.xml", "–".encode("utf-8")),
        ("word/document.xml", "“full-".encode("utf-8")),
        ("word/document.xml", " 30".encode("utf-8")),
    ],
    "tracked-moves": [
        ("word/document.xml", b"<w:moveFrom "),
        ("word/document.xml", b"<w:moveTo "),
        ("word/document.xml", b'w:name="move1"'),
    ],
    "format-changes": [
        ("word/document.xml", b"<w:rPrChange "),
        ("word/document.xml", b"<w:pPrChange "),
    ],
    "fields": [
        ("word/document.xml", b"<w:fldSimple "),
        ("word/document.xml", b'w:fldCharType="begin"'),
        ("word/document.xml", b"PAGEREF"),
    ],
    "placeholder-control": [
        ("word/document.xml", b"<w:showingPlcHdr/>"),
        ("word/document.xml", b"Click or tap here to enter text."),
        ("word/styles.xml", b'w:styleId="PlaceholderText"'),
    ],
    "bookmarks": [
        ("word/document.xml", b'w:name="DefinedTerm"'),
        ("word/document.xml", b'w:name="_GoBack"'),
    ],
    "noisy-markup": [
        ("word/document.xml", b"<w:proofErr "),
        ("word/document.xml", b"commentRangeStart"),
        ("word/document.xml", b'w:name="_GoBack"'),
    ],
    "toc-field": [
        ("word/document.xml", b" TOC "),
        ("word/document.xml", b'w:fldCharType="separate"'),
        ("word/document.xml", b"Chapter Two entry"),
    ],
    "row-revisions": [
        ("word/document.xml", b"<w:trPr><w:ins "),
        ("word/document.xml", b"<w:trPr><w:del "),
        ("word/document.xml", b"Filing fee"),
        ("word/document.xml", b"Old charge"),
    ],
    "format-changes-rich": [
        ("word/document.xml", b"Delivery follows the schedule in Exhibit A."),
        ("word/document.xml", b'<w:jc w:val="right"/>'),
        ("word/document.xml", b"The paragraph mark itself was re-formatted."),
    ],
    "paragraph-merge": [
        ("word/document.xml", b"This sentence continues "),
        ("word/document.xml", b"A tracked split divides "),
    ],
}

GAUNTLET_PROBES: List[Tuple[str, bytes]] = [
    probe for probes in FEATURE_PROBES.values() for probe in probes
]

#: probes for fixtures that deliberately do NOT fold into the gauntlet
#: (protection would poison every other gauntlet test; the redline bucket
#: fixtures are document-pair ground truth, not isolated features).
STANDALONE_PROBES: Dict[str, List[Tuple[str, bytes]]] = {
    "protected-forms": [
        ("word/settings.xml", b'w:edit="forms"'),
        ("word/settings.xml", b'w:enforcement="1"'),
    ],
    "protected-readonly": [("word/settings.xml", b'w:edit="readOnly"')],
    "protected-tracked": [("word/settings.xml", b'w:edit="trackedChanges"')],
    "multiround": [
        ("word/document.xml", b'w:name="mrMove1"'),
        ("word/document.xml", b"<w:rPrChange "),
        ("word/document.xml", b"<w:pPrChange "),
        ("word/document.xml", b"<w:trPr><w:ins "),
        ("word/document.xml", b"<w:trPr><w:del "),
        ("word/document.xml", b"<w:pPr><w:rPr><w:del "),
        ("word/document.xml", b"<w:pPr><w:rPr><w:moveFrom "),
        ("word/comments.xml", b"Confirm the execution venue."),
    ],
    "multiround-accepted": [
        ("word/document.xml", b"Filing fee"),
        ("word/document.xml", MULTIROUND_MOVED_TEXT.encode()),
        ("word/comments.xml", b"Confirm the execution venue."),
    ],
    "compare-original": [
        ("word/document.xml", b"monthly progress reports"),
        ("word/document.xml", b"$200"),
    ],
    "compare-revised": [
        ("word/document.xml", b"consulting and advisory services"),
        ("word/document.xml", b"$250"),
        ("word/document.xml", b"A retainer of one thousand dollars"),
    ],
}


def multi_section_probe(path: Path) -> None:
    document = part_bytes(path, "word/document.xml")
    assert document.count(b"<w:sectPr") >= 2, "expected at least two sections"
    header_parts = [n for n in part_names(path) if n.startswith("word/header")]
    assert len(header_parts) >= 2, f"expected >=2 header parts, found {header_parts}"


def validate_generated_fixture(
    path: Path,
    probes: List[Tuple[str, bytes]],
    *,
    run_lo: bool,
) -> None:
    checks.assert_package_facts_clean(path)
    reopened = docx.Document(str(path))
    assert reopened.paragraphs is not None
    for part_name, pattern in probes:
        blob = part_bytes(path, part_name)
        assert pattern in blob, f"{path.name}: {pattern!r} not found in {part_name}"
    if run_lo:
        ok, diagnostic = libreoffice_converts(path)
        assert ok, f"{path.name}: LibreOffice smoke failed: {diagnostic}"


def validate_corrupt_fixtures(broken_rel: Path, malformed: Path) -> None:
    findings = checks.find_broken_relationship_targets(broken_rel)
    assert findings, "broken-rel fixture is not detectably broken"
    assert not checks.find_unparseable_xml_parts(broken_rel)

    failures = checks.find_unparseable_xml_parts(malformed)
    assert any(name == "word/document.xml" for name, _ in failures), (
        "malformed-xml fixture did not break word/document.xml parseability"
    )


# ---------------------------------------------------------------------------
# LibreOffice-exported variants
# ---------------------------------------------------------------------------

#: fixture-stem -> probes an LO round-trip must still satisfy for the variant
#: to be kept. Weaker than FEATURE_PROBES where LibreOffice legitimately
#: rewrites markup (ids, run boundaries) — presence of the *feature* is what
#: matters, not the exact bytes.
LO_SURVIVAL_PROBES: Dict[str, List[Tuple[str, bytes]]] = {
    "tracked-ins-del": [
        ("word/document.xml", b"<w:ins "),
        ("word/document.xml", b"<w:del "),
        ("word/document.xml", b"<w:delText"),
    ],
    "comments": [
        ("word/comments.xml", b"<w:comment "),
    ],
    "content-control": [
        ("word/document.xml", b"<w:sdt"),
    ],
    "textbox": [
        ("word/document.xml", b"txbxContent"),
    ],
    "footnotes-endnotes": [
        ("word/footnotes.xml", b"<w:footnote "),
        ("word/endnotes.xml", b"<w:endnote "),
    ],
    "header-footer-sections": [],  # checked structurally by multi_section_probe
    "table-merged-nested": [
        ("word/document.xml", b"gridSpan"),
        ("word/document.xml", b"vMerge"),
    ],
    "numbering-custom": [
        ("word/document.xml", b"<w:numId"),
        ("word/numbering.xml", b"<w:num "),
    ],
    "fragmented-runs": [
        ("word/document.xml", "–".encode("utf-8")),
        ("word/document.xml", "“".encode("utf-8")),
    ],
}


def libreoffice_version() -> str:
    binary = soffice_binary()
    assert binary is not None
    out = subprocess.run([binary, "--version"], capture_output=True, text=True, check=True)
    return out.stdout.strip()


def export_libreoffice_variant(source: Path, out_dir: Path) -> Path:
    binary = soffice_binary()
    assert binary is not None
    with tempfile.TemporaryDirectory() as tmp:
        profile_dir = Path(tmp) / "lo-profile"
        completed = subprocess.run(
            [
                binary,
                f"-env:UserInstallation={profile_dir.as_uri()}",
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                tmp,
                str(source),
            ],
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
        produced = Path(tmp) / source.name
        if completed.returncode != 0 or not produced.is_file():
            raise RuntimeError(
                f"LibreOffice conversion of {source.name} failed:"
                f" rc={completed.returncode} stderr={completed.stderr.strip()}"
            )
        out_dir.mkdir(parents=True, exist_ok=True)
        target = out_dir / source.name
        shutil.copyfile(produced, target)
    return target


def build_libreoffice_bucket(report: List[str]) -> None:
    feature_dir = GENERATED_DIR / "feature-isolated"
    out_dir = LIBREOFFICE_DIR / "feature-isolated"
    for stem, probes in sorted(LO_SURVIVAL_PROBES.items()):
        source = feature_dir / f"{stem}.docx"
        target = export_libreoffice_variant(source, out_dir)
        try:
            for part_name, pattern in probes:
                blob = part_bytes(target, part_name)
                assert pattern in blob, f"{pattern!r} missing from {part_name}"
            if stem == "header-footer-sections":
                multi_section_probe(target)
            if stem == "fragmented-runs":
                document = part_bytes(target, "word/document.xml")
                assert document.count(b"<w:r>") + document.count(b"<w:r ") >= 4
            checks.assert_package_facts_clean(target)
            docx.Document(str(target))
        except AssertionError as exc:
            target.unlink()
            report.append(f"DROPPED libreoffice/{stem}.docx — feature did not survive: {exc}")
        else:
            report.append(f"kept libreoffice/feature-isolated/{stem}.docx")


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------


def main(argv: "List[str] | None" = None) -> int:
    """Build + validate the generated bucket; LO export only on explicit opt-in.

    Flags:
      --no-lo      skip the (read-only) LibreOffice smoke validation
      --lo-export  ALSO regenerate tests/paper/fixtures/libreoffice/ — this
                   overwrites the frozen LibreOffice bucket with fresh,
                   NON-deterministic LibreOffice output and therefore always
                   requires re-freezing the manifest and re-verifying the LO
                   sidecars' ground truth. Never a default side effect.
    """
    args = argv if argv is not None else sys.argv[1:]
    run_lo = "--no-lo" not in args
    export_lo = "--lo-export" in args
    report: List[str] = []

    feature_dir = GENERATED_DIR / "feature-isolated"
    builders: Dict[Path, Callable[[Path], None]] = {
        GENERATED_DIR / "minimal-clean" / "minimal.docx": build_minimal,
        feature_dir / "tracked-ins-del.docx": build_tracked_ins_del,
        feature_dir / "comments.docx": build_comments,
        feature_dir / "content-control.docx": build_content_control,
        feature_dir / "textbox.docx": build_textbox,
        feature_dir / "footnotes-endnotes.docx": build_footnotes_endnotes,
        feature_dir / "header-footer-sections.docx": build_header_footer_sections,
        feature_dir / "table-merged-nested.docx": build_table_merged_nested,
        feature_dir / "numbering-custom.docx": build_numbering_custom,
        feature_dir / "fragmented-runs.docx": build_fragmented_runs,
        feature_dir / "tracked-moves.docx": build_tracked_moves,
        feature_dir / "format-changes.docx": build_format_changes,
        feature_dir / "fields.docx": build_fields,
        feature_dir / "placeholder-control.docx": build_placeholder_control,
        feature_dir / "bookmarks.docx": build_bookmarks,
        feature_dir / "noisy-markup.docx": build_noisy_markup,
        feature_dir / "toc-field.docx": build_toc_field,
        feature_dir / "row-revisions.docx": build_row_revisions,
        feature_dir / "format-changes-rich.docx": build_format_changes_rich,
        feature_dir / "paragraph-merge.docx": build_paragraph_merge,
        feature_dir / "protected-forms.docx": build_protected_forms,
        feature_dir / "protected-readonly.docx": build_protected_readonly,
        feature_dir / "protected-tracked.docx": build_protected_tracked,
        GENERATED_DIR / "redline" / "multiround.docx": build_redline_multiround,
        GENERATED_DIR / "redline" / "multiround-accepted.docx": build_redline_multiround_accepted,
        GENERATED_DIR / "redline" / "compare-original.docx": build_compare_original,
        GENERATED_DIR / "redline" / "compare-revised.docx": build_compare_revised,
        GENERATED_DIR / "gauntlet" / "gauntlet.docx": build_gauntlet,
        GENERATED_DIR / "large" / "large-5000-paragraphs.docx": build_large,
    }
    for out_path, builder in builders.items():
        builder(out_path)
        report.append(f"built {out_path.relative_to(FIXTURES_DIR)}")

    minimal = GENERATED_DIR / "minimal-clean" / "minimal.docx"
    broken_rel = GENERATED_DIR / "corrupt" / "broken-rel.docx"
    malformed = GENERATED_DIR / "corrupt" / "malformed-xml.docx"
    build_corrupt_broken_rel(minimal, broken_rel)
    build_corrupt_malformed_xml(minimal, malformed)
    report.append("built corrupt fixtures: broken-rel.docx, malformed-xml.docx")

    for out_path in builders:
        stem = out_path.stem
        if stem == "gauntlet":
            probes = GAUNTLET_PROBES
        else:
            probes = FEATURE_PROBES.get(stem, []) + STANDALONE_PROBES.get(stem, [])
        validate_generated_fixture(out_path, probes, run_lo=run_lo and stem != "large-5000-paragraphs")
        if stem in ("header-footer-sections", "gauntlet"):
            multi_section_probe(out_path)
    validate_corrupt_fixtures(broken_rel, malformed)
    report.append("validated all generated fixtures (package facts, reopen, probes, LO smoke)")

    if export_lo:
        build_libreoffice_bucket(report)
        report.append(f"libreoffice version: {libreoffice_version()}")
        report.append(
            "NOTE: libreoffice/ bucket rewritten with fresh LibreOffice output —"
            " re-verify LO sidecar ground truth and re-run freeze_manifest.py"
        )

    report.append(f"python-docx base: {docx.__version__} / paper {docx.__paper_version__}")
    print("\n".join(report))
    return 0


if __name__ == "__main__":
    sys.exit(main())
