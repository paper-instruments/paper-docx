"""Regression tests for atomic revision resolution and discard cleanup."""

from __future__ import annotations

import struct
from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import pytest

import docx
from docx.commentops import (
    COMMENTS_EXTENDED_RELATIONSHIP_TYPE,
    COMMENTS_IDS_RELATIONSHIP_TYPE,
    reply,
    resolve,
)
from docx.errors import UnsupportedStructureError
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from .harness.paths import fixture_path

TRACKED_MOVES = "generated/feature-isolated/tracked-moves.docx"
NOTES = "generated/feature-isolated/footnotes-endnotes.docx"
_W15_NS = "http://schemas.microsoft.com/office/word/2012/wordml"


def _revision(tag: str, inner_xml: str, revision_id: int):
    return parse_xml(
        f'<w:{tag} {nsdecls("w")} w:id="{revision_id}" w:author="Audit">'
        f"{inner_xml}</w:{tag}>"
    )


def _wrap_in_insertion(element, revision_id: int) -> None:
    wrapper = _revision("ins", "", revision_id)
    element.addprevious(wrapper)
    wrapper.append(element)


def _bmp(red: int, green: int, blue: int) -> bytes:
    """Return a valid one-pixel 24-bit BMP without an image dependency."""
    return (
        struct.pack("<2sIHHI", b"BM", 58, 0, 0, 54)
        + struct.pack("<IiiHHIIiiII", 40, 1, 1, 1, 24, 0, 4, 2835, 2835, 0, 0)
        + bytes((blue, green, red, 0))
    )


def _image_r_id(run) -> str:
    blip = next(run._r.iter(qn("a:blip")))
    return blip.get(qn("r:embed"))


def _note_ids(document, reltype: str, note_tag: str) -> set[str]:
    root = document.part.part_related_by(reltype)._element  # noqa: SLF001
    return {
        note.get(qn("w:id"))
        for note in root
        if note.tag == qn(note_tag) and note.get(qn("w:id")) is not None
    }


class DescribeAtomicRefusal:
    @pytest.mark.parametrize(
        ("method", "first_tag", "first_inner", "bad_tag"),
        [
            (
                "accept_all",
                "ins",
                "<w:r><w:t>kept</w:t></w:r>",
                "del",
            ),
            (
                "reject_all",
                "del",
                "<w:r><w:delText>restored</w:delText></w:r>",
                "ins",
            ),
        ],
    )
    def it_prevalidates_malformed_comment_reference_ids(
        self, method: str, first_tag: str, first_inner: str, bad_tag: str
    ):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        paragraph.append(_revision(first_tag, first_inner, 1))
        paragraph.append(
            _revision(
                bad_tag,
                '<w:r><w:commentReference w:id="not-a-number"/></w:r>',
                2,
            )
        )
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match=r"malformed.*w:id"):
            getattr(document.revisions, method)()

        assert document.element.xml == before
        assert len(document.revisions) == 2

    def it_prevalidates_malformed_comment_range_ids(self):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        paragraph.append(
            parse_xml(
                f'<w:commentRangeStart {nsdecls("w")} w:id="broken"/>'
            )
        )
        paragraph.append(
            _revision(
                "ins", '<w:r><w:commentReference w:id="7"/></w:r>', 3
            )
        )
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match=r"malformed.*w:id"):
            document.revisions.reject_all()

        assert document.element.xml == before

    def it_refuses_stale_revision_and_revisions_snapshots(self):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        paragraph.append(
            _revision("ins", "<w:r><w:t>first</w:t></w:r>", 10)
        )
        paragraph.append(
            _revision("ins", "<w:r><w:t>second</w:t></w:r>", 11)
        )
        snapshot = document.revisions
        first, second = snapshot
        first.accept()
        before = document.element.xml

        for stale_action in (
            first.accept,
            second.accept,
            snapshot.accept_all,
            snapshot.reject_all,
        ):
            with pytest.raises(UnsupportedStructureError, match="stale"):
                stale_action()
            assert document.element.xml == before

        assert document.revisions.accept_all() == 1

    def it_refuses_a_mixed_author_move_unit_before_mutation(self):
        document = docx.Document(str(fixture_path(TRACKED_MOVES)))
        moves = [
            revision
            for revision in document.revisions
            if revision.revision_type in ("move_from", "move_to")
        ]
        for revision in moves:
            revision._element.set(qn("w:author"), "Alice")  # noqa: SLF001
        next(
            revision for revision in moves if revision.revision_type == "move_to"
        )._element.set(qn("w:author"), "Bob")  # noqa: SLF001
        snapshot = document.revisions
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match="authors differ"):
            snapshot.accept_all(author="Alice")

        assert document.element.xml == before
        assert len(document.revisions) == len(snapshot)

    def it_deduplicates_a_valid_author_filtered_move_unit(self):
        document = docx.Document(str(fixture_path(TRACKED_MOVES)))
        snapshot = document.revisions
        expected = len(snapshot)

        assert snapshot.accept_all(author="Alice Editor") == expected
        assert len(document.revisions) == 0


class DescribeDiscardCleanup:
    def it_drops_unreferenced_relationships_and_payloads_but_keeps_shared_ones(
        self, tmp_path: Path
    ):
        document = docx.Document()
        paragraph = document.add_paragraph()

        shared_keep = paragraph.add_run()
        shared_keep.add_picture(BytesIO(_bmp(255, 0, 0)))
        shared_drop = paragraph.add_run()
        shared_drop.add_picture(BytesIO(_bmp(255, 0, 0)))
        unique_drop = paragraph.add_run()
        unique_drop.add_picture(BytesIO(_bmp(0, 0, 255)))

        shared_r_id = _image_r_id(shared_keep)
        assert _image_r_id(shared_drop) == shared_r_id
        unique_r_id = _image_r_id(unique_drop)

        hyperlink_r_id = document.part.relate_to(
            "https://audit.example.invalid/", RT.HYPERLINK, is_external=True
        )
        hyperlink = parse_xml(
            f'<w:hyperlink {nsdecls("w", "r")} r:id="{hyperlink_r_id}">'
            "<w:r><w:t>discarded link</w:t></w:r></w:hyperlink>"
        )

        package = document.part.package
        assert package is not None
        object_part = Part(
            PackURI("/word/embeddings/audit-object.bin"),
            CT.OFC_OLE_OBJECT,
            b"discarded object payload",
            package,
        )
        object_r_id = document.part.relate_to(object_part, RT.OLE_OBJECT)
        object_run = parse_xml(
            f'<w:r {nsdecls("w", "r")} '
            'xmlns:o="urn:schemas-microsoft-com:office:office">'
            f'<w:object><o:OLEObject r:id="{object_r_id}"/></w:object></w:r>'
        )

        wrapper = _revision("ins", "", 20)
        shared_drop._r.addprevious(wrapper)
        wrapper.append(shared_drop._r)
        wrapper.append(unique_drop._r)
        wrapper.append(hyperlink)
        wrapper.append(object_run)

        shared_partname = str(
            document.part.rels[shared_r_id].target_part.partname
        ).lstrip("/")
        unique_partname = str(
            document.part.rels[unique_r_id].target_part.partname
        ).lstrip("/")
        object_partname = str(object_part.partname).lstrip("/")

        assert document.revisions.reject_all() == 1
        assert shared_r_id in document.part.rels
        assert unique_r_id not in document.part.rels
        assert hyperlink_r_id not in document.part.rels
        assert object_r_id not in document.part.rels

        output = tmp_path / "relationships.docx"
        document.save(str(output))
        with ZipFile(output) as package_zip:
            names = set(package_zip.namelist())
        assert shared_partname in names
        assert unique_partname not in names
        assert object_partname not in names

    def it_removes_unreferenced_footnote_and_endnote_bodies(
        self, tmp_path: Path
    ):
        document = docx.Document(str(fixture_path(NOTES)))
        footnote_ref = next(document.element.iter(qn("w:footnoteReference")))
        endnote_ref = next(document.element.iter(qn("w:endnoteReference")))
        footnote_id = footnote_ref.get(qn("w:id"))
        endnote_id = endnote_ref.get(qn("w:id"))
        _wrap_in_insertion(footnote_ref.getparent(), 30)
        _wrap_in_insertion(endnote_ref.getparent(), 31)

        assert document.revisions.reject_all() == 2
        assert footnote_id not in _note_ids(document, RT.FOOTNOTES, "w:footnote")
        assert endnote_id not in _note_ids(document, RT.ENDNOTES, "w:endnote")

        output = tmp_path / "notes.docx"
        document.save(str(output))
        reopened = docx.Document(str(output))
        assert footnote_id not in _note_ids(reopened, RT.FOOTNOTES, "w:footnote")
        assert endnote_id not in _note_ids(reopened, RT.ENDNOTES, "w:endnote")

    def it_preserves_a_note_body_with_a_surviving_shared_reference(self):
        document = docx.Document(str(fixture_path(NOTES)))
        footnote_ref = next(document.element.iter(qn("w:footnoteReference")))
        footnote_id = footnote_ref.get(qn("w:id"))
        _wrap_in_insertion(footnote_ref.getparent(), 32)
        surviving_run = document.add_paragraph().add_run()._r
        surviving_run.append(
            parse_xml(
                f'<w:footnoteReference {nsdecls("w")} w:id="{footnote_id}"/>'
            )
        )

        assert document.revisions.reject_all() == 1
        assert footnote_id in _note_ids(document, RT.FOOTNOTES, "w:footnote")

    def it_removes_only_matching_comments_extended_entries(self, tmp_path: Path):
        document = docx.Document()
        first_p = document.add_paragraph("first comment")
        second_p = document.add_paragraph("second comment")
        first = document.add_comment(
            first_p.runs, text="discard me", author="Audit"
        )
        second = document.add_comment(
            second_p.runs, text="keep me", author="Audit"
        )
        resolve(document, first)
        resolve(document, second)

        comments_root = document.part.part_related_by(RT.COMMENTS)._element  # noqa: SLF001
        comments = {int(comment.get(qn("w:id"))): comment for comment in comments_root}
        first_para_id = comments[first.comment_id].findall(qn("w:p"))[-1].get(
            qn("w14:paraId")
        )
        second_para_id = comments[second.comment_id].findall(qn("w:p"))[-1].get(
            qn("w14:paraId")
        )
        first_reference = next(
            reference
            for reference in document.element.iter(qn("w:commentReference"))
            if reference.get(qn("w:id")) == str(first.comment_id)
        )
        _wrap_in_insertion(first_reference.getparent(), 40)

        assert document.revisions.reject_all() == 1
        assert [comment.comment_id for comment in document.comments] == [
            second.comment_id
        ]
        extended_root = document.part.part_related_by(
            COMMENTS_EXTENDED_RELATIONSHIP_TYPE
        )._element  # noqa: SLF001
        remaining_para_ids = {
            entry.get(f"{{{_W15_NS}}}paraId") for entry in extended_root
        }
        assert first_para_id not in remaining_para_ids
        assert second_para_id in remaining_para_ids
        ids_root = document.part.part_related_by(
            COMMENTS_IDS_RELATIONSHIP_TYPE
        )._element  # noqa: SLF001
        remaining_identity = {
            entry.get(
                "{http://schemas.microsoft.com/office/word/2016/wordml/cid}paraId"
            )
            for entry in ids_root
        }
        assert first_para_id not in remaining_identity
        assert second_para_id in remaining_identity

        output = tmp_path / "comments.docx"
        document.save(str(output))
        reopened = docx.Document(str(output))
        reopened_extended = reopened.part.part_related_by(
            COMMENTS_EXTENDED_RELATIONSHIP_TYPE
        )._element  # noqa: SLF001
        assert first_para_id not in {
            entry.get(f"{{{_W15_NS}}}paraId") for entry in reopened_extended
        }

    def it_removes_an_entire_discarded_comment_thread(self):
        from docx.blocks import tracked_delete_paragraphs

        document = docx.Document()
        paragraph = document.add_paragraph("threaded comment")
        parent = document.add_comment(
            paragraph.runs, text="parent", author="Audit"
        )
        reply(document, parent, "reply", author="Audit")
        assert len(document.comments) == 2

        tracked_delete_paragraphs(
            document,
            "threaded comment",
            count=1,
            author="Audit",
        )
        assert document.revisions.accept_all() > 0

        assert len(document.comments) == 0
        extended_root = document.part.part_related_by(
            COMMENTS_EXTENDED_RELATIONSHIP_TYPE
        )._element  # noqa: SLF001
        assert len(extended_root) == 0
