"""Regressions for compare's raw-package and matcher preflights."""

from __future__ import annotations

import datetime as dt
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest

import docx
import docx._compare as compare_impl
from docx.errors import UnsupportedStructureError
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.package import compare

FROZEN = dt.datetime(2026, 7, 10, 12, 0, tzinfo=dt.timezone.utc)


def _save_pair(tmp_path: Path, build_original, build_revised):
    paths = []
    for name, build in (("original", build_original), ("revised", build_revised)):
        document = docx.Document()
        build(document)
        path = tmp_path / f"{name}.docx"
        document.save(path)
        paths.append(path)
    return paths


def _add_orphan_custom_xml(path: Path) -> None:
    with ZipFile(path, "r") as source:
        entries = [(info, source.read(info.filename)) for info in source.infolist()]

    marker = b"</Types>"
    override = (
        b'<Override PartName="/customXml/orphan.xml" '
        b'ContentType="application/xml"/>'
    )
    rewritten = []
    for info, data in entries:
        if info.filename == "[Content_Types].xml":
            assert marker in data
            data = data.replace(marker, override + marker, 1)
        rewritten.append((info, data))

    replacement = path.with_suffix(".replacement.docx")
    with ZipFile(replacement, "w", compression=ZIP_DEFLATED) as destination:
        for info, data in rewritten:
            destination.writestr(info, data)
        destination.writestr("customXml/orphan.xml", b"<orphan/>")
    replacement.replace(path)


def _guard_large_sequence_match(monkeypatch: pytest.MonkeyPatch) -> None:
    real_matcher = compare_impl.SequenceMatcher

    def guarded_matcher(isjunk=None, a="", b="", autojunk=True):
        if not isinstance(a, str) and len(a) * len(b) > 4:
            raise AssertionError("SequenceMatcher ran before its resource preflight")
        return real_matcher(isjunk, a, b, autojunk=autojunk)

    monkeypatch.setattr(compare_impl, "SequenceMatcher", guarded_matcher)
    monkeypatch.setattr(compare_impl, "_MAX_SEQUENCE_CELLS", 4)


class DescribeRawPackagePreflight:
    def it_treats_marker_only_move_markup_as_a_pending_revision(self, tmp_path: Path):
        def build(document):
            paragraph = document.add_paragraph("Move markers only")._p
            for tag, name in (
                ("w:moveFromRangeStart", "Move1"),
                ("w:moveFromRangeEnd", None),
                ("w:moveToRangeStart", "Move1"),
                ("w:moveToRangeEnd", None),
            ):
                marker = OxmlElement(tag)
                marker.set(qn("w:id"), "17")
                if name is not None:
                    marker.set(qn("w:name"), name)
                paragraph.append(marker)

        original, revised = _save_pair(tmp_path, build, build)
        with pytest.raises(UnsupportedStructureError, match="pending tracked revisions"):
            compare(original, revised, author="Reviewer", date=FROZEN)

    def it_refuses_an_orphan_before_loading_and_preserves_both_inputs(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ):
        def build(document):
            document.add_paragraph("Same visible document")

        original, revised = _save_pair(tmp_path, build, build)
        _add_orphan_custom_xml(revised)
        before = (original.read_bytes(), revised.read_bytes())

        def unexpected_document_load(*args, **kwargs):
            raise AssertionError("raw package preflight must run before Document()")

        monkeypatch.setattr(docx, "Document", unexpected_document_load)
        with pytest.raises(UnsupportedStructureError, match="unreachable package part"):
            compare(original, revised, author="Reviewer", date=FROZEN)

        assert (original.read_bytes(), revised.read_bytes()) == before


class DescribeSequenceMatcherBudget:
    def it_refuses_story_matching_before_quadratic_work(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ):
        def build(prefix):
            def _build(document):
                for index in range(3):
                    document.add_paragraph(f"{prefix} paragraph {index}")

            return _build

        original, revised = _save_pair(tmp_path, build("Old"), build("New"))
        _guard_large_sequence_match(monkeypatch)

        with pytest.raises(UnsupportedStructureError, match="sequence-matching budget"):
            compare(original, revised, author="Reviewer", date=FROZEN)

    def it_refuses_table_matching_before_quadratic_work(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ):
        def build(middle):
            def _build(document):
                table = document.add_table(rows=3, cols=1)
                for index, row in enumerate(table.rows):
                    row.cells[0].text = middle if index == 1 else f"Stable {index}"

            return _build

        original, revised = _save_pair(tmp_path, build("Old"), build("New"))
        _guard_large_sequence_match(monkeypatch)

        with pytest.raises(
            UnsupportedStructureError, match="table.*sequence-matching budget"
        ):
            compare(original, revised, author="Reviewer", date=FROZEN)

    def it_refuses_changed_text_before_quadratic_similarity_work(
        self, tmp_path: Path, monkeypatch: pytest.MonkeyPatch
    ):
        def build(text):
            def _build(document):
                document.add_paragraph(text)

            return _build

        original, revised = _save_pair(tmp_path, build("A" * 8), build("B" * 8))
        real_matcher = compare_impl.SequenceMatcher

        def guarded_matcher(isjunk=None, a="", b="", autojunk=True):
            if isinstance(a, str) and len(a) * len(b) > 4:
                raise AssertionError(
                    "text SequenceMatcher ran before its resource preflight"
                )
            return real_matcher(isjunk, a, b, autojunk=autojunk)

        monkeypatch.setattr(compare_impl, "SequenceMatcher", guarded_matcher)
        monkeypatch.setattr(compare_impl, "_MAX_TEXT_SEQUENCE_CELLS", 4)

        with pytest.raises(
            UnsupportedStructureError, match="changed-region text.*budget"
        ):
            compare(original, revised, author="Reviewer", date=FROZEN)

    def it_refuses_token_matching_before_quadratic_work(
        self, monkeypatch: pytest.MonkeyPatch
    ):
        monkeypatch.setattr(compare_impl, "_MAX_TEXT_SEQUENCE_CELLS", 4)

        with pytest.raises(
            UnsupportedStructureError, match="paragraph token diff.*budget"
        ):
            compare_impl._token_regions("one two three", "four five six")
