"""Tests for the everyday document shapes: noise-tolerant block ops,
break-tolerant replace, hyperlink-interior edits, and batch replace.

(Cell-wise table guards are covered in test_tableops_numbering.py.)
"""

from __future__ import annotations

import datetime as dt
from contextlib import contextmanager

import pytest

import docx
import docx.search as search_module
from docx.errors import BoundaryViolationError, UnsupportedStructureError
from docx.oxml.ns import nsdecls
from docx.oxml.parser import parse_xml
from docx.search import find_one, replace_all
from docx.story import iter_blocks

from .harness.paths import fixture_path

FROZEN = dt.datetime(2026, 7, 8, 9, 0, 0, tzinfo=dt.timezone.utc)
MINIMAL = "generated/minimal-clean/minimal.docx"
NOISY = "generated/feature-isolated/noisy-markup.docx"
FIELDS = "generated/feature-isolated/fields.docx"
W = nsdecls("w")


def _doc(relpath: str = MINIMAL):
    return docx.Document(str(fixture_path(relpath)))


class DescribeNoiseTolerantBlockOps:
    """proofErr/_GoBack/comment anchors are Word's ubiquitous noise —
    tracked block ops treat them as transparent instead of refusing."""

    def it_tracked_deletes_a_paragraph_with_proofing_noise(self):
        from docx.blocks import tracked_delete_paragraphs

        document = _doc(NOISY)
        result = tracked_delete_paragraphs(
            document, "Paragrah with a spelling issue.", count=1,
            author="Carol QA", date=FROZEN,
        )
        assert result.deleted_blocks == 1
        body_xml = document.element.body
        assert not body_xml.xpath("//w:del//w:proofErr"), (
            "proofErr must be dropped, not wrapped in the deletion"
        )
        assert body_xml.xpath('//w:bookmarkStart[@w:name="_GoBack"]'), (
            "point bookmark must survive in place"
        )

    def it_tracked_deletes_a_commented_paragraph_keeping_the_anchor(self):
        from docx.blocks import tracked_delete_paragraphs

        document = _doc(NOISY)
        tracked_delete_paragraphs(
            document, "This clause carries a reviewer comment.", count=1,
            author="Carol QA", date=FROZEN,
        )
        body_xml = document.element.body
        assert body_xml.xpath("//w:commentRangeStart"), "comment anchor lost"
        assert len(document.comments) == 1, "the comment itself must survive"

    def it_still_round_trips_reject_to_the_original_text(self):
        from docx.blocks import tracked_delete_paragraphs

        document = _doc(NOISY)
        pristine = [b.text for b in iter_blocks(document)]
        tracked_delete_paragraphs(
            document, "Paragrah with a spelling issue.", count=1,
            author="Carol QA", date=FROZEN,
        )
        document.revisions.reject_all()
        assert [b.text for b in iter_blocks(document)] == pristine


class DescribeBreakTolerantReplace:
    """Matching across a tab is fine; only a CHANGE crossing it refuses."""

    def _tabbed_doc(self):
        document = _doc()
        paragraph = document.add_paragraph()
        paragraph.add_run("Section 3.")
        paragraph.add_run().add_tab()
        paragraph.add_run("Termination")
        return document, paragraph

    def it_edits_one_segment_of_a_tab_crossing_span(self):
        document, paragraph = self._tabbed_doc()
        span = find_one(document, "Section 3. Termination")
        span.replace("Section 4. Termination")
        assert paragraph.text == "Section 4.\tTermination"

    def it_edits_the_trailing_segment_too(self):
        document, paragraph = self._tabbed_doc()
        find_one(document, "Section 3. Termination").replace("Section 3. Renewal")
        assert paragraph.text == "Section 3.\tRenewal"

    def it_narrows_tracked_replaces_the_same_way(self):
        document, paragraph = self._tabbed_doc()
        find_one(document, "Section 3. Termination").replace(
            "Section 4. Termination", tracked=True, author="Carol QA", date=FROZEN
        )
        document.revisions.accept_all()
        assert paragraph.text == "Section 4.\tTermination"
        document2, paragraph2 = self._tabbed_doc()
        find_one(document2, "Section 3. Termination").replace(
            "Section 4. Termination", tracked=True, author="Carol QA", date=FROZEN
        )
        document2.revisions.reject_all()
        assert paragraph2.text == "Section 3.\tTermination"

    def it_edits_both_segments_keeping_the_tab(self):
        """Whitespace in the replacement aligns with the existing tab: the
        document keeps its tab, both text segments change."""
        document, paragraph = self._tabbed_doc()
        find_one(document, "Section 3. Termination").replace(
            "Chapter Three - Termination"
        )
        assert paragraph.text == "Chapter Three -\tTermination"

    def it_still_refuses_changes_that_would_swallow_the_break(self):
        document, _ = self._tabbed_doc()
        span = find_one(document, "Section 3. Termination")
        with pytest.raises(UnsupportedStructureError, match="tab or line break"):
            span.replace("Section3Termination")  # no whitespace for the tab


class DescribeHyperlinkInteriorEdits:
    """Text inside one hyperlink is redlinable; crossing its boundary is not."""

    def _linked_doc(self):
        document = _doc()
        paragraph_xml = parse_xml(
            f"<w:p {W}>"
            '<w:r><w:t xml:space="preserve">As defined in </w:t></w:r>'
            '<w:hyperlink w:anchor="DefSection">'
            "<w:r><w:t>Section 3.2 (Payment Terms)</w:t></w:r>"
            "</w:hyperlink>"
            '<w:r><w:t xml:space="preserve"> of this Agreement.</w:t></w:r>'
            "</w:p>"
        )
        body = document.element.body
        body.insert(len(body) - 1, paragraph_xml)
        return document

    def it_tracked_replaces_text_inside_a_hyperlink(self):
        document = self._linked_doc()
        find_one(document, "Section 3.2 (Payment Terms)").replace(
            "Section 4.1 (Payment Terms)", tracked=True,
            author="Carol QA", date=FROZEN,
        )
        (ins,) = document.element.body.xpath("//w:hyperlink/w:ins")
        assert "4.1" in "".join(ins.itertext())
        document.revisions.accept_all()
        texts = [b.text for b in iter_blocks(document)]
        assert "As defined in Section 4.1 (Payment Terms) of this Agreement." in texts

    def it_rejects_back_to_the_original_link_text(self):
        document = self._linked_doc()
        pristine = [b.text for b in iter_blocks(document)]
        find_one(document, "Section 3.2").replace(
            "Section 9.9", tracked=True, author="Carol QA", date=FROZEN
        )
        document.revisions.reject_all()
        assert [b.text for b in iter_blocks(document)] == pristine

    def it_refuses_spans_crossing_the_hyperlink_boundary(self):
        document = self._linked_doc()
        span = find_one(document, "defined in Section 3.2")
        with pytest.raises(BoundaryViolationError, match="hyperlink boundary"):
            span.replace("defined by Section 3.2", tracked=True,
                         author="Carol QA", date=FROZEN)


class DescribeReplaceAll:
    """Single scan, reverse-document-order application, loud refusals."""

    def it_replaces_matches_sharing_one_run(self):
        document = _doc()
        document.add_paragraph("{{x}} then {{x}} then {{x}} in one run.")
        result = replace_all(document, "{{x}}", "VALUE")
        assert result.replaced_count == 3
        assert not result.refused
        texts = [b.text for b in iter_blocks(document)]
        assert "VALUE then VALUE then VALUE in one run." in texts

    def it_reports_refused_matches_instead_of_skipping_silently(self):
        document = _doc(FIELDS)
        # "June 1, 2026" also exists as editable base text nowhere; the only
        # match is the field result -> one refusal, zero replacements
        result = replace_all(document, "June 1, 2026", "July 9, 2026")
        assert result.replaced_count == 0
        assert len(result.refused) == 1
        assert result.refused[0]["error"] == "UnsupportedStructureError"

    def it_supports_tracked_batch_replacement(self):
        document = _doc()
        document.add_paragraph("Fee: {{fee}}. Late fee: {{fee}}.")
        result = replace_all(
            document, "{{fee}}", "$100", tracked=True, author="Carol QA", date=FROZEN
        )
        assert result.replaced_count == 2
        assert len({rid for r in result.results for rid in r.revision_ids}) >= 2
        document.revisions.accept_all()
        texts = [b.text for b in iter_blocks(document)]
        assert "Fee: $100. Late fee: $100." in texts

    def it_serializes_the_outcome(self):
        document = _doc()
        document.add_paragraph("{{y}}")
        payload = replace_all(document, "{{y}}", "z").to_dict()
        assert payload["schema"] == "paper_replace_all"
        assert payload["version"] == 1
        assert payload["replaced_count"] == 1
        nested = payload["results"][0]
        assert nested["schema"] == "paper_replace"
        assert nested["version"] == 1
        assert nested["preserved_structure"] is False
        assert nested["preserved_revision_ids"] == []

    def it_preserves_revision_identity_only_where_needed(self):
        document = _doc()
        document.add_paragraph("base term")
        document.add_paragraph()._p.append(
            parse_xml(
                f'<w:ins {W} w:id="801" w:author="Alice">'
                "<w:r><w:t>inserted term</w:t></w:r></w:ins>"
            )
        )
        result = replace_all(
            document, "term", "phrase", preserve_revision=True
        )
        assert result.replaced_count == 2
        assert sorted(item.preserved_revision_ids for item in result.results) == [
            (),
            (801,),
        ]
        assert document.element.body.xpath('//w:ins[@w:id="801"]')

    def it_reports_exact_structure_evidence_for_each_successful_match(self):
        document = _doc()
        document.add_paragraph("token token")
        result = replace_all(
            document, "token", "value", preserve_structure=True
        )
        assert result.replaced_count == 2
        assert all(item.preserved_structure for item in result.results)
        assert "value value" in [block.text for block in iter_blocks(document)]

    @pytest.mark.parametrize("preserve_structure", [False, True])
    def it_uses_only_the_outer_batch_transaction(
        self, monkeypatch, preserve_structure: bool
    ):
        document = _doc()
        document.add_paragraph("token token")
        transaction_count = 0
        original = search_module.rollback_on_error

        @contextmanager
        def counted_transaction(*args, **kwargs):
            nonlocal transaction_count
            transaction_count += 1
            with original(*args, **kwargs):
                yield

        monkeypatch.setattr(
            search_module, "rollback_on_error", counted_transaction
        )
        replace_all(
            document,
            "token",
            "value",
            preserve_structure=preserve_structure,
        )
        assert transaction_count == 1

    def it_locally_restores_a_late_per_match_refusal(self, monkeypatch):
        document = _doc()
        document.add_paragraph("token token")
        original = search_module._apply_text_assignments
        calls = 0

        def refuse_second(assignments):
            nonlocal calls
            calls += 1
            if calls == 2:
                assignments[0].element.text = "corrupt"
                raise UnsupportedStructureError("forced late refusal")
            original(assignments)

        monkeypatch.setattr(
            search_module, "_apply_text_assignments", refuse_second
        )
        result = replace_all(
            document, "token", "value", preserve_structure=True
        )
        assert result.replaced_count == 1
        assert len(result.refused) == 1
        assert "token value" in [block.text for block in iter_blocks(document)]

    def it_rolls_back_the_batch_after_an_unexpected_exact_failure(self, monkeypatch):
        document = _doc()
        paragraph = document.add_paragraph("token token")
        original = search_module._apply_text_assignments
        calls = 0

        def fail_second(assignments):
            nonlocal calls
            calls += 1
            if calls == 2:
                assignments[0].element.text = "corrupt"
                raise RuntimeError("forced unexpected failure")
            original(assignments)

        monkeypatch.setattr(search_module, "_apply_text_assignments", fail_second)
        with pytest.raises(RuntimeError, match="forced unexpected"):
            replace_all(document, "token", "value", preserve_structure=True)
        assert paragraph.text == "token token"
