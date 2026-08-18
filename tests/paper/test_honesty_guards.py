"""Tests for the write guards and package triage."""

from __future__ import annotations

import datetime as dt
import shutil
import struct
import zipfile
from pathlib import Path

import pytest

import docx
from docx.errors import TargetNotFoundError, UnsupportedStructureError
from docx.package import diagnose
from docx.search import find_one

from .harness.contract import assert_refusal_atomic
from .harness.paths import fixture_path

FROZEN = dt.datetime(2026, 7, 8, 9, 0, 0, tzinfo=dt.timezone.utc)
MINIMAL = "generated/minimal-clean/minimal.docx"
BOOKMARKS = "generated/feature-isolated/bookmarks.docx"
NOISY = "generated/feature-isolated/noisy-markup.docx"


def _doc(relpath: str = MINIMAL):
    return docx.Document(str(fixture_path(relpath)))


class DescribeControlCharacterRefusal:
    """Raw \\n/\\t in written text would be verified-but-false structure."""

    @pytest.mark.parametrize("bad", ["Line one\nLine two", "col\tcol", "a\rb"])
    def it_refuses_control_characters_in_replacement_text(self, bad: str):
        span = find_one(_doc(), "perfectly ordinary")
        with pytest.raises(ValueError, match="control character"):
            span.replace(bad)

    def it_refuses_control_characters_in_inserted_paragraphs(self):
        from docx.blocks import insert_section_after

        with pytest.raises(ValueError, match="control character"):
            insert_section_after(
                _doc(), "First body paragraph",
                heading="OK", paragraphs=["fine", "not\nfine"],
            )

    def it_refuses_control_characters_in_tracked_replacements(self):
        from docx.blocks import tracked_replace_paragraphs

        with pytest.raises(ValueError, match="control character"):
            tracked_replace_paragraphs(
                _doc(), "Second body paragraph", ["multi\nline"],
                author="Carol QA", date=FROZEN,
            )


class DescribeFakeBulletRefusal:
    """A list style whose numbering does not resolve must refuse."""

    def it_applies_a_list_style_whose_numbering_resolves(self):
        from docx.numbering import apply_list_style

        document = _doc()
        paragraph = document.add_paragraph("bulleted")
        apply_list_style(paragraph, "List Bullet")
        assert paragraph.style.name == "List Bullet"

    def it_refuses_a_list_style_with_dangling_numbering(self, tmp_path: Path):
        from docx.numbering import apply_list_style

        import re

        source = fixture_path(MINIMAL)
        stripped = tmp_path / "no-numbering.docx"
        with zipfile.ZipFile(source) as zin:
            with zipfile.ZipFile(stripped, "w") as zout:
                for name in zin.namelist():
                    if name == "word/numbering.xml":
                        continue
                    blob = zin.read(name)
                    if name == "word/_rels/document.xml.rels":
                        # drop the numbering relationship entirely (a dangling
                        # rel would break upstream's loader before our check)
                        blob = re.sub(rb'<Relationship [^>]*numbering\.xml"/>', b"", blob)
                    zout.writestr(name, blob)
        document = docx.Document(str(stripped))
        paragraph = document.add_paragraph("wants a bullet")
        with pytest.raises(TargetNotFoundError, match="fake"):
            apply_list_style(paragraph, "List Bullet")

    def it_still_applies_styles_without_numbering_bindings(self):
        from docx.numbering import apply_list_style

        document = _doc()
        paragraph = document.add_paragraph("just a heading")
        apply_list_style(paragraph, "Heading 3")
        assert paragraph.style.name == "Heading 3"


class DescribeUntrackedEditInsideInsertions:
    """Untracked edits must not rewrite text attributed to an author."""

    def it_refuses_untracked_replace_inside_a_pending_insertion(self, tmp_path: Path):
        path = tmp_path / "doc.docx"
        shutil.copyfile(fixture_path(MINIMAL), path)
        document = docx.Document(str(path))
        find_one(document, "perfectly ordinary").replace(
            "written by Alice", tracked=True, author="Alice Editor", date=FROZEN
        )

        def untracked_rewrite(doc):
            find_one(doc, "written by Alice").replace("forged words")

        assert_refusal_atomic(
            document, untracked_rewrite, UnsupportedStructureError, on_disk=(path,)
        )


class DescribeBookmarkHollowing:
    """Replaces must not silently empty cross-reference targets."""

    def it_refuses_a_replace_that_hollows_a_named_bookmark(self):
        document = _doc(BOOKMARKS)
        span = find_one(document, "See the Master Agreement for definitions.")
        with pytest.raises(UnsupportedStructureError, match="DefinedTerm"):
            span.replace("See the Purchase Agreement for definitions.")

    def it_allows_replacing_text_inside_the_bookmark(self):
        """Markers outside the span: the new text stays bookmarked."""
        document = _doc(BOOKMARKS)
        find_one(document, "the Master Agreement").replace("the Purchase Agreement")
        assert "the Purchase Agreement" in find_one(
            document, "See the Purchase Agreement for definitions."
        ).text

    def it_treats_point_bookmarks_as_transparent(self):
        """_GoBack noise must never block an edit."""
        document = _doc(NOISY)
        find_one(document, "Paragrah with a spelling issue.").replace(
            "Paragraph without a spelling issue."
        )


class DescribePackageDiagnosis:
    """Typed triage instead of raw KeyError/ValueError dead ends."""

    def it_diagnoses_a_healthy_docx(self):
        report = diagnose(fixture_path(MINIMAL))
        assert report.readable and report.kind == "docx" and not report.problems
        payload = report.to_dict()
        assert payload["schema"] == "paper_diagnosis" and payload["version"] == 1

    def it_diagnoses_missing_files(self, tmp_path: Path):
        report = diagnose(tmp_path / "nope.docx")
        assert not report.readable and report.kind == "missing"

    def it_diagnoses_encrypted_or_legacy_binaries(self, tmp_path: Path):
        cfb = tmp_path / "locked.docx"
        cfb.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 64)
        report = diagnose(cfb)
        assert not report.readable
        assert report.kind == "encrypted-or-legacy-binary"
        assert "password" in report.problems[0]

    def it_diagnoses_non_zip_content(self, tmp_path: Path):
        text = tmp_path / "notes.docx"
        text.write_bytes(b"just some text pretending to be a docx")
        assert diagnose(text).kind == "not-a-zip"

    @staticmethod
    def _prefixed(source: Path, prefix: bytes) -> bytes:
        """`source` with `prefix` in front and every ZIP offset rebased."""
        end_record = struct.Struct("<4s4H2LH")
        central = struct.Struct("<4s6H3L5H2L")
        clean = source.read_bytes()
        body = bytearray(prefix + clean)
        end_offset = clean.rfind(b"PK\x05\x06")
        fields = list(end_record.unpack_from(clean, end_offset))
        central_size, central_offset = fields[5], fields[6]
        cursor = central_offset + len(prefix)
        limit = cursor + central_size
        while cursor < limit:
            record = central.unpack_from(bytes(body), cursor)
            struct.pack_into("<L", body, cursor + 42, record[16] + len(prefix))
            cursor += central.size + record[10] + record[11] + record[12]
        fields[6] = central_offset + len(prefix)
        end_record.pack_into(body, end_offset + len(prefix), *fields)
        return bytes(body)

    @pytest.mark.parametrize(
        "prefix",
        [
            pytest.param(b"# self-extracting stub\n" + b"#" * 40, id="prefix-not-starting-PK"),
            pytest.param(b"PKSTUB" + b"#" * 58, id="prefix-starting-PK"),
        ],
    )
    def it_diagnoses_an_archive_that_does_not_begin_at_byte_zero(
        self, tmp_path: Path, prefix: bytes
    ):
        """Word refuses a prefixed package; `diagnose` must not call it healthy or non-ZIP.

        Before this guard the two shapes reported differently and both wrongly: a prefix
        starting with `PK` came back `readable=True, kind="docx"` -- the triage API vouching
        for a file Word refuses -- and any other prefix came back `not-a-zip`, which is false
        because the archive is present. Both now land on `unsafe-archive`; no new `kind` value
        was introduced.
        """
        target = tmp_path / "prefixed.docx"
        target.write_bytes(self._prefixed(fixture_path(MINIMAL), prefix))

        report = diagnose(target)

        assert not report.readable
        assert report.kind == "unsafe-archive"
        assert "does not begin" in report.problems[0]

    def it_diagnoses_macro_enabled_documents(self, tmp_path: Path):
        source = fixture_path(MINIMAL)
        docm = tmp_path / "macros.docm"
        with zipfile.ZipFile(source) as zin:
            with zipfile.ZipFile(docm, "w") as zout:
                for name in zin.namelist():
                    blob = zin.read(name)
                    if name == "[Content_Types].xml":
                        blob = blob.replace(
                            b"application/vnd.openxmlformats-officedocument"
                            b".wordprocessingml.document.main+xml",
                            b"application/vnd.ms-word.document.macroEnabled.main+xml",
                        )
                    zout.writestr(name, blob)
        report = diagnose(docm)
        assert report.kind == "docm" and not report.readable
        assert "macro" in report.problems[0]

    def it_diagnoses_the_corrupt_fixtures(self):
        report = diagnose(fixture_path("generated/corrupt/malformed-xml.docx"))
        # zip is fine; the damage is inside a part — readable at OPC level
        assert report.kind == "docx"

    def it_diagnoses_other_opc_packages(self, tmp_path: Path):
        fake_xlsx = tmp_path / "book.xlsx"
        with zipfile.ZipFile(fake_xlsx, "w") as zf:
            zf.writestr("[Content_Types].xml", "<Types xmlns='http://schemas.openxmlformats.org/package/2006/content-types'/>")
            zf.writestr("_rels/.rels", "<Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'/>")
            zf.writestr("xl/workbook.xml", "<workbook/>")
        report = diagnose(fake_xlsx)
        assert report.kind == "xlsx" and not report.readable
