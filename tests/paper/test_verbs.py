"""Tests for the v0.1 Phase 2 verbs (V1-V5)."""

from __future__ import annotations

import datetime as dt
import shutil
from pathlib import Path

import pytest

import docx
from docx.errors import (
    AmbiguousTargetError,
    TargetNotFoundError,
    UnsupportedStructureError,
)
from docx.oxml.ns import nsdecls
from docx.oxml.parser import parse_xml
from docx.search import find_one
from docx.story import iter_blocks

from .harness.contract import assert_changed_parts, save_and_reopen
from .harness.paths import fixture_path

FROZEN = dt.datetime(2026, 7, 8, 9, 0, 0, tzinfo=dt.timezone.utc)
MINIMAL = "generated/minimal-clean/minimal.docx"
PLACEHOLDER = "generated/feature-isolated/placeholder-control.docx"
W14 = nsdecls("w") + ' xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"'


def _doc(relpath: str = MINIMAL):
    return docx.Document(str(fixture_path(relpath)))


def _copy(relpath: str, tmp_path: Path) -> Path:
    target = tmp_path / Path(relpath).name
    shutil.copyfile(fixture_path(relpath), target)
    return target


def _add_sdt(document, xml: str) -> None:
    body = document.element.body
    body.insert(len(body) - 1, parse_xml(xml))


class DescribeContentControls:
    def it_enumerates_controls_with_identity_and_state(self):
        from docx.controls import list_controls

        (info,) = list_controls(_doc(PLACEHOLDER))
        assert info.tag == "client-name" and info.alias == "ClientName"
        assert info.showing_placeholder and not info.is_data_bound
        payload = info.to_dict()
        assert payload["value"] == "Click or tap here to enter text."

    def it_sets_text_values_clearing_placeholder_state(self, tmp_path: Path):
        from docx.controls import get_control, set_control_value

        path = _copy(PLACEHOLDER, tmp_path)
        document = docx.Document(str(path))
        set_control_value(document, "Paper Instruments LLC", tag="client-name")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        control = get_control(reopened, tag="client-name")
        assert control.value == "Paper Instruments LLC"
        assert not control.showing_placeholder

    def it_sets_checkboxes_with_state_and_glyph(self):
        from docx.controls import get_control, set_control_value

        document = _doc()
        _add_sdt(
            document,
            f'<w:p {W14}><w:sdt><w:sdtPr><w:tag w:val="agree"/>'
            '<w14:checkbox><w14:checked w14:val="0"/></w14:checkbox></w:sdtPr>'
            "<w:sdtContent><w:r><w:t>☐</w:t></w:r></w:sdtContent></w:sdt></w:p>",
        )
        set_control_value(document, True, tag="agree")
        control = get_control(document, tag="agree")
        assert control.value is True
        assert "☒" in "".join(document.element.body.itertext())
        with pytest.raises(ValueError, match="bool"):
            set_control_value(document, "yes", tag="agree")

    def it_validates_dropdown_choices(self):
        from docx.controls import get_control, set_control_value

        document = _doc()
        _add_sdt(
            document,
            f'<w:p {W14}><w:sdt><w:sdtPr><w:tag w:val="state"/>'
            '<w:dropDownList><w:listItem w:displayText="California" w:value="CA"/>'
            '<w:listItem w:displayText="New York" w:value="NY"/></w:dropDownList>'
            "</w:sdtPr><w:sdtContent><w:r><w:t>Choose.</w:t></w:r></w:sdtContent>"
            "</w:sdt></w:p>",
        )
        assert get_control(document, tag="state").info().choices == (
            "California", "New York",
        )
        set_control_value(document, "New York", tag="state")
        with pytest.raises(TargetNotFoundError, match="choices"):
            set_control_value(document, "Texas", tag="state")

    def it_sets_dates_stamping_full_date(self):
        from docx.controls import get_control, set_control_value

        document = _doc()
        _add_sdt(
            document,
            f'<w:p {W14}><w:sdt><w:sdtPr><w:tag w:val="signed"/>'
            "<w:date/></w:sdtPr>"
            "<w:sdtContent><w:r><w:t>pick a date</w:t></w:r></w:sdtContent>"
            "</w:sdt></w:p>",
        )
        set_control_value(document, dt.date(2026, 7, 8), tag="signed")
        assert get_control(document, tag="signed").value == "2026-07-08"
        (date_pr,) = document.element.body.xpath("//w:sdtPr/w:date")
        assert date_pr.get(docx.oxml.ns.qn("w:fullDate")) == "2026-07-08T00:00:00Z"

    def it_refuses_data_bound_controls(self):
        from docx.controls import set_control_value

        document = _doc()
        _add_sdt(
            document,
            f'<w:p {W14}><w:sdt><w:sdtPr><w:tag w:val="bound"/>'
            '<w:dataBinding w:xpath="/x"/></w:sdtPr>'
            "<w:sdtContent><w:r><w:t>synced</w:t></w:r></w:sdtContent></w:sdt></w:p>",
        )
        with pytest.raises(UnsupportedStructureError, match="data-bound"):
            set_control_value(document, "x", tag="bound")

    def it_refuses_ambiguous_and_missing_tags(self):
        from docx.controls import get_control

        document = _doc(PLACEHOLDER)
        with pytest.raises(TargetNotFoundError):
            get_control(document, tag="nope")
        with pytest.raises(ValueError, match="tag"):
            get_control(document)


class DescribeNumberingAuthoring:
    def it_creates_a_bullet_definition_in_a_definitionless_document(
        self, tmp_path: Path
    ):
        import re
        import zipfile

        from docx.numbering import apply_numbering, ensure_bullet_definition

        source = fixture_path(MINIMAL)
        stripped = tmp_path / "no-numbering.docx"
        with zipfile.ZipFile(source) as zin:
            with zipfile.ZipFile(stripped, "w") as zout:
                for name in zin.namelist():
                    if name == "word/numbering.xml":
                        continue
                    blob = zin.read(name)
                    if name == "word/_rels/document.xml.rels":
                        blob = re.sub(rb'<Relationship [^>]*numbering\.xml"/>', b"", blob)
                    zout.writestr(name, blob)
        document = docx.Document(str(stripped))
        num_id = ensure_bullet_definition(document)
        apply_numbering(document.add_paragraph("a real bullet"), num_id=num_id)
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        with zipfile.ZipFile(tmp_path / "out.docx") as zf:
            numbering = zf.read("word/numbering.xml")
        assert b'w:val="bullet"' in numbering
        assert reopened is not None

    def it_is_idempotent(self):
        from docx.numbering import ensure_bullet_definition, ensure_decimal_definition

        document = _doc()
        assert ensure_bullet_definition(document) == ensure_bullet_definition(document)
        assert ensure_decimal_definition(document) != ensure_bullet_definition(document)

    def it_restarts_numbering_with_a_fresh_definition(self):
        from docx.numbering import (
            apply_numbering,
            ensure_decimal_definition,
            restart_numbering,
        )

        document = _doc()
        num_id = ensure_decimal_definition(document)
        apply_numbering(document.add_paragraph("one"), num_id=num_id)
        restarted = restart_numbering(document, num_id=num_id)
        assert restarted != num_id
        apply_numbering(document.add_paragraph("one again"), num_id=restarted)
        (override,) = document.part.part_related_by(
            docx.opc.constants.RELATIONSHIP_TYPE.NUMBERING
        )._element.xpath(f'//w:num[@w:numId="{restarted}"]/w:lvlOverride/w:startOverride')
        assert override.get(docx.oxml.ns.qn("w:val")) == "1"

    def it_refuses_restarting_a_missing_definition(self):
        from docx.numbering import restart_numbering

        with pytest.raises(TargetNotFoundError):
            restart_numbering(_doc(), num_id=999)


class DescribeRichBlockInsertion:
    def it_inserts_paragraphs_lists_and_tables(self, tmp_path: Path):
        from docx.blocks import (
            ListBlock,
            RichParagraph,
            TableBlock,
            TextRun,
            insert_blocks_after,
        )
        from docx.package import patch_save

        source = _copy(MINIMAL, tmp_path)
        document = docx.Document(str(source))
        result = insert_blocks_after(
            document, "First body paragraph",
            blocks=[
                RichParagraph(runs=[TextRun("Key findings", bold=True),
                                    TextRun(" (draft)", italic=True)]),
                ListBlock(items=["Alpha", "Beta"], kind="decimal"),
                TableBlock(rows=[["Metric", "Value"], ["Uptime", "99.9%"]]),
            ],
        )
        assert result.inserted_blocks == 4  # paragraph + 2 list items + table
        out = tmp_path / "out.docx"
        patch_save(source, document, out)
        assert_changed_parts(
            source, out, {"word/document.xml", "word/numbering.xml"},
            expected_added=(), expected_removed=(),
        )
        reopened = docx.Document(str(out))
        blocks = list(iter_blocks(reopened))
        bold_para = next(b for b in blocks if b.text == "Key findings (draft)")
        assert bold_para.kind == "paragraph"
        table_block = next(b for b in blocks if b.kind == "table")
        assert "Uptime" in table_block.text

    def it_tracked_inserts_paragraphs_and_lists(self):
        from docx.blocks import ListBlock, RichParagraph, TextRun, insert_blocks_after

        document = _doc()
        pristine = [b.text for b in iter_blocks(document)]
        insert_blocks_after(
            document, "First body paragraph",
            blocks=[
                RichParagraph(runs=[TextRun("Tracked lead", bold=True)]),
                ListBlock(items=["one"]),
            ],
            tracked=True, author="Carol QA", date=FROZEN,
        )
        assert {r.author for r in document.revisions} == {"Carol QA"}
        document.revisions.reject_all()
        assert [b.text for b in iter_blocks(document)] == pristine

    def it_refuses_tracked_table_insertion(self):
        from docx.blocks import TableBlock, insert_blocks_after

        with pytest.raises(UnsupportedStructureError, match="tracked table"):
            insert_blocks_after(
                _doc(), "First body paragraph",
                blocks=[TableBlock(rows=[["x"]])],
                tracked=True, author="Carol QA", date=FROZEN,
            )

    def it_validates_block_shapes_loudly(self):
        from docx.blocks import ListBlock, TableBlock, insert_blocks_after

        with pytest.raises(ValueError, match="rectangular"):
            insert_blocks_after(
                _doc(), "First body paragraph",
                blocks=[TableBlock(rows=[["a", "b"], ["c"]])],
            )
        with pytest.raises(ValueError, match="kind"):
            insert_blocks_after(
                _doc(), "First body paragraph",
                blocks=[ListBlock(items=["x"], kind="fancy")],
            )


class DescribeSpanComments:
    def it_anchors_a_comment_to_exactly_the_span_text(self, tmp_path: Path):
        from docx.commentops import anchored_text

        path = _copy(MINIMAL, tmp_path)
        document = docx.Document(str(path))
        span = find_one(document, "perfectly ordinary")
        comment = span.comment("Check this.", author="Alice Editor", date=FROZEN)
        assert anchored_text(document, comment) == "perfectly ordinary"
        # anchoring splits are semantically neutral
        texts = [b.text for b in iter_blocks(document)]
        assert "First body paragraph with perfectly ordinary text." in texts

    def it_threads_replies_and_resolution_through_word_machinery(
        self, tmp_path: Path
    ):
        from docx.commentops import comment_thread, reply, resolve

        path = _copy(MINIMAL, tmp_path)
        document = docx.Document(str(path))
        comment = find_one(document, "perfectly ordinary").comment(
            "Check this.", author="Alice Editor", date=FROZEN
        )
        reply(document, comment, "Checked.", author="Bob Reviewer", date=FROZEN)
        resolve(document, comment)
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        thread = comment_thread(reopened)
        assert thread[0]["resolved"] is True
        assert thread[1]["parent_id"] == thread[0]["comment_id"]
        assert thread[1]["anchored_text"] == "perfectly ordinary"

    def it_refuses_comments_outside_the_main_story(self):
        document = _doc("generated/feature-isolated/header-footer-sections.docx")
        span = find_one(document, "Header for section one", story="word/header1.xml")
        with pytest.raises(UnsupportedStructureError, match="main document story"):
            span.comment("nope", author="Carol QA", date=FROZEN)


class DescribeTextDiff:
    def it_says_what_changed_not_just_which_part(self, tmp_path: Path):
        from docx.package import text_diff

        source = _copy(MINIMAL, tmp_path)
        document = docx.Document(str(source))
        find_one(document, "equally unremarkable").replace("quite peculiar")
        out = tmp_path / "out.docx"
        document.save(str(out))
        diff = text_diff(source, out)
        joined = "\n".join(line for s in diff.stories for line in s.diff_lines)
        assert "-[paragraph] Second body paragraph, equally unremarkable." in joined
        assert "+[paragraph] Second body paragraph, quite peculiar." in joined
        payload = diff.to_dict()
        assert payload["schema"] == "paper_text_diff" and payload["version"] == 1

    def it_summarizes_pending_revisions(self):
        from docx.package import pending_changes

        diff = pending_changes(
            fixture_path("generated/feature-isolated/tracked-ins-del.docx")
        )
        joined = "\n".join(line for s in diff.stories for line in s.diff_lines)
        assert "forty-two" in joined and "forty-seven" in joined

    def it_reports_empty_for_identical_documents(self):
        from docx.package import text_diff

        path = fixture_path(MINIMAL)
        assert text_diff(path, path).is_empty
