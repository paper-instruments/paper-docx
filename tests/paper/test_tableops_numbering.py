"""Tests for docx.tableops and docx.numbering (Phase 9)."""

from __future__ import annotations

import datetime as dt
import zipfile
from pathlib import Path

import pytest

import docx
from docx.errors import (
    AmbiguousTargetError,
    TargetNotFoundError,
    UnsupportedStructureError,
)
from docx.numbering import apply_list_style, apply_numbering, list_numbering
from docx.tableops import delete_row, find_table, insert_row_after, update_cell

from .harness.contract import assert_refusal_atomic, save_and_reopen
from .harness.paths import fixture_path

MINIMAL = "generated/minimal-clean/minimal.docx"
COMPLEX_TABLE = "generated/feature-isolated/table-merged-nested.docx"
NUMBERING = "generated/feature-isolated/numbering-custom.docx"
NUMBERING_LO = "libreoffice/feature-isolated/numbering-custom.docx"

FROZEN = dt.datetime(2026, 7, 7, 12, 0, 0, tzinfo=dt.timezone.utc)


def _doc(relpath: str):
    return docx.Document(str(fixture_path(relpath)))


def _doc_with_simple_table():
    document = _doc(MINIMAL)
    table = document.add_table(rows=2, cols=2)
    for r in range(2):
        for c in range(2):
            table.cell(r, c).text = f"cell {r}{c}"
    return document, table


class DescribeFindTable:
    def it_finds_the_table_by_normalized_cell_text(self):
        document, table = _doc_with_simple_table()
        found = find_table(document, near_text="CELL 10")  # casefolded matching
        assert found._tbl is table._tbl

    def it_refuses_when_no_table_matches(self):
        document, _ = _doc_with_simple_table()
        with pytest.raises(TargetNotFoundError, match="no table"):
            find_table(document, near_text="nothing like this")

    def it_refuses_ambiguity(self):
        document, _ = _doc_with_simple_table()
        second = document.add_table(rows=1, cols=1)
        second.cell(0, 0).text = "cell 10 duplicate"
        with pytest.raises(AmbiguousTargetError):
            find_table(document, near_text="cell 10")


class DescribeUpdateCell:
    def it_replaces_cell_text_through_the_span_machinery(self, tmp_path: Path):
        document, table = _doc_with_simple_table()
        result = update_cell(table, 0, 1, "updated value")
        assert result.deleted_text == "cell 01"
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert reopened.tables[0].cell(0, 1).text == "updated value"

    def it_fills_an_empty_cell(self, tmp_path: Path):
        document, table = _doc_with_simple_table()
        table.cell(1, 1).paragraphs[0].clear()
        update_cell(table, 1, 1, "was empty")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert reopened.tables[0].cell(1, 1).text == "was empty"

    def it_supports_tracked_updates(self, tmp_path: Path):
        document, table = _doc_with_simple_table()
        result = update_cell(
            table, 0, 0, "cell 99", tracked=True, author="Carol QA", date=FROZEN
        )
        assert result.tracked and result.revision_ids
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert reopened.tables[0].cell(0, 0).text.startswith("cell")
        revisions = reopened.revisions
        assert {r.author for r in revisions} == {"Carol QA"}
        reopened.revisions.accept_all()
        assert reopened.tables[0].cell(0, 0).text == "cell 99"

    def it_refuses_complex_tables_atomically(self):
        document = _doc(COMPLEX_TABLE)
        assert_refusal_atomic(
            document,
            lambda doc: update_cell(doc.tables[0], 0, 0, "x"),
            UnsupportedStructureError,
        )

    def it_refuses_out_of_range_addresses(self):
        _, table = _doc_with_simple_table()
        with pytest.raises(TargetNotFoundError, match="row 7"):
            update_cell(table, 7, 0, "x")
        with pytest.raises(TargetNotFoundError, match="column 9"):
            update_cell(table, 0, 9, "x")

    def it_requires_an_author_when_tracked(self):
        _, table = _doc_with_simple_table()
        with pytest.raises(ValueError, match="author"):
            update_cell(table, 0, 0, "x", tracked=True)


class DescribeRowOperations:
    def it_inserts_a_row_copying_neighbor_formatting(self, tmp_path: Path):
        document, table = _doc_with_simple_table()
        insert_row_after(table, 0, ["new a", "new b"])
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        grid = [[c.text for c in row.cells] for row in reopened.tables[0].rows]
        assert grid == [["cell 00", "cell 01"], ["new a", "new b"], ["cell 10", "cell 11"]]

    def it_pads_missing_values_with_empty_cells(self):
        _, table = _doc_with_simple_table()
        insert_row_after(table, 1, ["only one"])
        assert [c.text for c in table.rows[2].cells] == ["only one", ""]

    def it_rejects_too_many_values(self):
        _, table = _doc_with_simple_table()
        with pytest.raises(ValueError, match="values"):
            insert_row_after(table, 0, ["a", "b", "c"])

    def it_deletes_a_row(self, tmp_path: Path):
        document, table = _doc_with_simple_table()
        delete_row(table, 0)
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        grid = [[c.text for c in row.cells] for row in reopened.tables[0].rows]
        assert grid == [["cell 10", "cell 11"]]

    def it_refuses_deleting_the_last_row(self):
        document, table = _doc_with_simple_table()
        delete_row(table, 0)
        assert_refusal_atomic(
            document,
            lambda doc: delete_row(table, 0),
            UnsupportedStructureError,
        )

    def it_refuses_structural_ops_on_complex_tables(self):
        document = _doc(COMPLEX_TABLE)
        with pytest.raises(UnsupportedStructureError):
            insert_row_after(document.tables[0], 0, ["x"])
        with pytest.raises(UnsupportedStructureError):
            delete_row(document.tables[0], 0)


class DescribeListNumbering:
    def it_reports_definitions_with_levels(self):
        report = list_numbering(_doc(NUMBERING))
        custom = next(d for d in report.definitions if d.num_id == 42)
        assert custom.abstract_num_id == 90
        assert [(lvl.level, lvl.num_fmt) for lvl in custom.levels] == [
            (0, "decimal"),
            (1, "lowerLetter"),
        ]

    def it_reports_numbered_paragraphs_with_anchors_and_levels(self):
        report = list_numbering(_doc(NUMBERING))
        assert [(p.num_id, p.level, p.text) for p in report.numbered_paragraphs] == [
            (42, 0, "First numbered item"),
            (42, 0, "Second numbered item"),
            (42, 1, "Nested lettered item"),
        ]

    def it_reads_the_libreoffice_remapped_definition(self):
        """LO round-trip remapped numId 42 -> 7 (frozen in the sidecar)."""
        report = list_numbering(_doc(NUMBERING_LO))
        assert [(p.num_id, p.text) for p in report.numbered_paragraphs] == [
            (7, "First numbered item"),
            (7, "Second numbered item"),
            (7, "Nested lettered item"),
        ]

    def it_serializes_deterministically(self):
        import json

        payload = list_numbering(_doc(NUMBERING)).to_dict()
        assert payload["schema"] == "paper_numbering" and payload["version"] == 1
        assert json.dumps(payload) == json.dumps(list_numbering(_doc(NUMBERING)).to_dict())


class DescribeApplyNumbering:
    def it_applies_an_existing_definition(self, tmp_path: Path):
        document = _doc(NUMBERING)
        paragraph = document.add_paragraph("Newly numbered item")
        apply_numbering(paragraph, num_id=42, level=1)
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        report = list_numbering(reopened)
        assert (42, 1, "Newly numbered item") in [
            (p.num_id, p.level, p.text) for p in report.numbered_paragraphs
        ]

    def it_refuses_undefined_definitions(self):
        document = _doc(NUMBERING)
        paragraph = document.add_paragraph("target")
        assert_refusal_atomic(
            document,
            lambda doc: apply_numbering(paragraph, num_id=999),
            TargetNotFoundError,
        )

    def it_refuses_undefined_levels(self):
        document = _doc(NUMBERING)
        paragraph = document.add_paragraph("target")
        with pytest.raises(TargetNotFoundError, match="level 7"):
            apply_numbering(paragraph, num_id=42, level=7)

    def it_refuses_when_no_numbering_part_exists(self, tmp_path: Path):
        import re

        stripped = tmp_path / "no-numbering.docx"
        with zipfile.ZipFile(fixture_path(MINIMAL)) as zin:
            with zipfile.ZipFile(stripped, "w") as zout:
                for name in zin.namelist():
                    if "numbering" in name:
                        continue
                    blob = zin.read(name)
                    if name == "word/_rels/document.xml.rels":
                        blob = re.sub(rb"<Relationship [^>]*numbering[^>]*/>", b"", blob)
                    zout.writestr(name, blob)
        document = docx.Document(str(stripped))
        paragraph = document.add_paragraph("target")
        with pytest.raises(TargetNotFoundError, match="does not exist"):
            apply_numbering(paragraph, num_id=1)


class DescribeApplyListStyle:
    def it_applies_an_existing_style(self):
        document = _doc(MINIMAL)
        paragraph = document.add_paragraph("styled")
        apply_list_style(paragraph, "Heading 2")
        assert paragraph.style.name == "Heading 2"

    def it_refuses_undefined_styles(self):
        document = _doc(MINIMAL)
        paragraph = document.add_paragraph("styled")
        assert_refusal_atomic(
            document,
            lambda doc: apply_list_style(paragraph, "No Such List Style"),
            TargetNotFoundError,
        )
