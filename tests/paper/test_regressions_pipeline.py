"""Regression pins for confirmed defects.

Every test here reproduces a confirmed silent-corruption / false-state /
atomicity finding and pins its fix.
"""

from __future__ import annotations

import datetime as dt
import io
from pathlib import Path

import pytest

import docx
from docx.errors import UnsupportedStructureError
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.revision import _fallback_markup, _remaining_markup
from docx.story import iter_blocks

from .harness.paths import fixture_path
from .test_composition import TINY_PNG

W = nsdecls("w")
FROZEN = dt.datetime(2026, 7, 8, 15, 0, 0, tzinfo=dt.timezone.utc)
MINIMAL = "generated/minimal-clean/minimal.docx"


def _body_texts(document) -> list:
    return [
        b.text for b in iter_blocks(document) if b.story == "word/document.xml"
    ]


def _visible(document) -> dict:
    texts: dict = {}
    for block in iter_blocks(document):
        texts.setdefault(block.story, []).append(block.text)
    return {
        story: "\n".join(t for t in items if t) for story, items in texts.items()
    }


def _compare_pair(build_a, build_b, tmp_path, **kwargs):
    from docx.package import compare

    a, b = tmp_path / "a.docx", tmp_path / "b.docx"
    doc_a, doc_b = docx.Document(), docx.Document()
    build_a(doc_a)
    build_b(doc_b)
    doc_a.save(str(a))
    doc_b.save(str(b))
    kwargs.setdefault("author", "Sweep")
    kwargs.setdefault("date", FROZEN)
    return compare(str(a), str(b), **kwargs), str(a), str(b)


class DescribeResolutionRegressions:
    def it_fully_resolves_a_move_whose_source_sits_in_a_deleted_row(self):
        """Moves resolve FIRST in a batch — a row removal must never take a
        move site with it and then refuse claiming nothing changed."""
        document = docx.Document()
        body = document.element.body
        table = parse_xml(
            f"<w:tbl {W}><w:tblPr><w:tblW w:w='0' w:type='auto'/></w:tblPr>"
            "<w:tblGrid><w:gridCol w:w='3000'/></w:tblGrid>"
            "<w:tr><w:tc><w:tcPr/><w:p><w:r><w:t>surviving row</w:t></w:r></w:p></w:tc></w:tr>"
            "<w:tr><w:trPr><w:del w:id='10' w:author='A' w:date='2026-01-01T00:00:00Z'/></w:trPr>"
            "<w:tc><w:tcPr/><w:p>"
            "<w:moveFromRangeStart w:id='20' w:name='mv1'/>"
            "<w:moveFrom w:id='21' w:author='A' w:date='2026-01-01T00:00:00Z'>"
            "<w:r><w:t>moved sentence.</w:t></w:r></w:moveFrom>"
            "<w:moveFromRangeEnd w:id='20'/>"
            "</w:p></w:tc></w:tr></w:tbl>"
        )
        body.insert(0, table)
        destination = parse_xml(
            f"<w:p {W}><w:moveToRangeStart w:id='22' w:name='mv1'/>"
            "<w:moveTo w:id='23' w:author='A' w:date='2026-01-01T00:00:00Z'>"
            "<w:r><w:t>moved sentence.</w:t></w:r></w:moveTo>"
            "<w:moveToRangeEnd w:id='22'/></w:p>"
        )
        table.addnext(destination)
        assert document.revisions.accept_all() > 0
        assert _remaining_markup(document) == {}
        assert len(table.findall(qn("w:tr"))) == 1
        assert "moved sentence." in _body_texts(document)

    def it_never_merges_a_deleted_mark_past_a_block_level_content_control(self):
        document = docx.Document()
        body = document.element.body
        first = parse_xml(
            f"<w:p {W}><w:pPr><w:rPr>"
            "<w:del w:id='70' w:author='A' w:date='2026-01-01T00:00:00Z'/>"
            "</w:rPr></w:pPr><w:r><w:t>ALPHA </w:t></w:r></w:p>"
        )
        sdt = parse_xml(
            f"<w:sdt {W}><w:sdtPr><w:id w:val='1'/></w:sdtPr>"
            "<w:sdtContent><w:p><w:r><w:t>BETA </w:t></w:r></w:p></w:sdtContent></w:sdt>"
        )
        last = parse_xml(f"<w:p {W}><w:r><w:t>GAMMA</w:t></w:r></w:p>")
        body.insert(0, last)
        body.insert(0, sdt)
        body.insert(0, first)
        document.revisions.accept_all()
        texts = _body_texts(document)
        # ALPHA must not hop over the control; declining the merge keeps order
        assert texts.index("ALPHA ") < texts.index("BETA ")
        assert texts.index("BETA ") < texts.index("GAMMA")

    def it_refuses_overlapping_move_ranges_atomically(self):
        document = docx.Document()
        paragraph = parse_xml(
            f"<w:p {W}>"
            "<w:moveFromRangeStart w:id='30' w:name='mvA'/>"
            "<w:moveFromRangeStart w:id='31' w:name='mvB'/>"
            "<w:moveFrom w:id='32' w:author='A' w:date='2026-01-01T00:00:00Z'>"
            "<w:r><w:t>BRAVO</w:t></w:r></w:moveFrom>"
            "<w:moveFromRangeEnd w:id='30'/>"
            "<w:moveFromRangeEnd w:id='31'/>"
            "</w:p>"
        )
        document.element.body.insert(0, paragraph)
        before = document.element.xml
        with pytest.raises(
            UnsupportedStructureError, match="overlapping|pair|move"
        ):
            document.revisions.accept_all()
        assert document.element.xml == before

    def it_refuses_resolution_when_markup_hides_in_mc_fallback(self):
        document = docx.Document(str(fixture_path(MINIMAL)))
        alternate = parse_xml(
            f'<w:p {W} xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
            "<mc:AlternateContent>"
            "<mc:Choice Requires='wps'><w:r>"
            "<w:ins w:id='90' w:author='A' w:date='2026-01-01T00:00:00Z'>"
            "<w:r><w:t>choice text</w:t></w:r></w:ins></w:r></mc:Choice>"
            "<mc:Fallback><w:r>"
            "<w:ins w:id='91' w:author='A' w:date='2026-01-01T00:00:00Z'>"
            "<w:r><w:t>fallback text</w:t></w:r></w:ins></w:r></mc:Fallback>"
            "</mc:AlternateContent></w:p>"
        )
        document.element.body.insert(0, alternate)
        assert _fallback_markup(document) == {"ins": 1}
        before = document.element.xml
        with pytest.raises(UnsupportedStructureError, match="[Ff]allback"):
            document.revisions.accept_all()
        assert document.element.xml == before
        with pytest.raises(UnsupportedStructureError, match="[Ff]allback"):
            document.finalize()
        with pytest.raises(UnsupportedStructureError):
            document.scrub()  # metadata guard sees fallback markup too

    def it_enumerates_a_body_level_section_property_change(self):
        document = docx.Document(str(fixture_path(MINIMAL)))
        sect_pr = document.element.body.find(qn("w:sectPr"))
        sect_pr.append(
            parse_xml(
                f"<w:sectPrChange {W} w:id='95' w:author='A'"
                " w:date='2026-01-01T00:00:00Z'><w:sectPr/></w:sectPrChange>"
            )
        )
        revisions = document.revisions
        assert revisions.remaining_unsupported() == {"section_property_change": 1}
        before = document.element.xml
        with pytest.raises(UnsupportedStructureError):
            revisions.accept_all()
        assert document.element.xml == before

    def it_refuses_when_unpaired_move_markup_would_survive_resolution(self):
        document = docx.Document(str(fixture_path(MINIMAL)))
        # a stray range bracket pair with no wrappers: nothing enumerates it,
        # so a "successful" accept_all would leave it behind
        paragraph = document.paragraphs[1]._p
        paragraph.append(parse_xml(f"<w:moveToRangeStart {W} w:id='40' w:name='ghost'/>"))
        paragraph.append(parse_xml(f"<w:moveToRangeEnd {W} w:id='40'/>"))
        with pytest.raises(UnsupportedStructureError, match="pair"):
            document.revisions.accept_all()


class DescribeScrubRegressions:
    def it_never_fabricates_missing_metadata_or_settings_parts(self, tmp_path):
        import zipfile

        source = fixture_path(MINIMAL)
        stripped = tmp_path / "bare.docx"
        with zipfile.ZipFile(source) as zin, zipfile.ZipFile(stripped, "w") as zout:
            for name in zin.namelist():
                if name in ("docProps/core.xml", "word/settings.xml"):
                    continue
                blob = zin.read(name)
                if name == "_rels/.rels":
                    import re

                    blob = re.sub(rb"<Relationship [^>]*core-properties[^>]*/>", b"", blob)
                if name == "word/_rels/document.xml.rels":
                    import re

                    blob = re.sub(rb"<Relationship [^>]*settings[^>]*/>", b"", blob)
                zout.writestr(name, blob)
        document = docx.Document(str(stripped))
        report = document.scrub(rsids=True)
        assert report.metadata_fields_cleared == []
        assert report.track_changes_setting_removed is False
        out = tmp_path / "scrubbed.docx"
        document.save(str(out))
        with zipfile.ZipFile(out) as zf:
            names = set(zf.namelist())
        assert "docProps/core.xml" not in names  # nothing fabricated
        assert "word/settings.xml" not in names

    def it_reports_cascade_removed_comment_side_parts(self, tmp_path):
        from docx.commentops import reply

        document = docx.Document(
            str(fixture_path("generated/feature-isolated/comments.docx"))
        )
        comment = list(document.comments)[0]
        reply(document, comment, "threading part appears", author="Sweep")
        source = tmp_path / "with-thread.docx"
        document.save(str(source))

        from docx.package import diff_package

        reopened = docx.Document(str(source))
        reopened.finalize()
        report = reopened.scrub()
        out = tmp_path / "scrubbed.docx"
        reopened.save(str(out))
        diff = diff_package(str(source), str(out))
        assert set(diff.removed) == set(report.removed_parts)
        assert any("commentsExtended" in p for p in report.removed_parts)


class DescribeCompareRegressions:
    def it_redlines_nested_table_differences(self, tmp_path: Path):
        def build(inner_text):
            def _build(document):
                table = document.add_table(rows=1, cols=1)
                cell = table.cell(0, 0)
                cell.text = "outer cell"
                inner = cell.add_table(rows=1, cols=1)
                inner.cell(0, 0).text = inner_text

            return _build

        result, _a, b = _compare_pair(
            build("inner OLD"), build("inner NEW"), tmp_path
        )
        assert result.revision_count > 0  # never silently equal
        result.document.revisions.accept_all()
        assert _visible(result.document) == _visible(docx.Document(b))

    def it_keeps_the_algebra_when_a_cell_contains_a_tab(self, tmp_path: Path):
        def build(text):
            def _build(document):
                table = document.add_table(rows=1, cols=1)
                paragraph = table.cell(0, 0).paragraphs[0]
                paragraph._p.append(
                    parse_xml(
                        f"<w:r {W}><w:t>a</w:t><w:tab/><w:t>{text}</w:t></w:r>"
                    )
                )

            return _build

        result, a, b = _compare_pair(build("b OLD"), build("b NEW"), tmp_path)
        result.document.revisions.accept_all()
        assert _visible(result.document) == _visible(docx.Document(b))

    def it_orders_extra_inserted_rows_after_the_replaced_pair(
        self, tmp_path: Path
    ):
        def build(rows):
            def _build(document):
                table = document.add_table(rows=len(rows), cols=1)
                for index, value in enumerate(rows):
                    table.cell(index, 0).text = value

            return _build

        result, _a, b = _compare_pair(
            build(["row one"]),
            build(["row 1", "row two", "row three"]),
            tmp_path,
            granularity="block",
        )
        result.document.revisions.accept_all()
        grid = [
            [cell.text for cell in row.cells]
            for row in result.document.tables[0].rows
        ]
        assert grid == [["row 1"], ["row two"], ["row three"]]

    def it_emits_unique_revision_ids(self, tmp_path: Path):
        from collections import Counter

        result, _a, _b = _compare_pair(
            lambda d: [d.add_paragraph(t) for t in ("one", "two same", "three")],
            lambda d: [d.add_paragraph(t) for t in ("one CHANGED", "two same", "four", "five")],
            tmp_path,
        )
        ids = [
            node.get(qn("w:id"))
            for node in result.document.element.body.iter(
                qn("w:ins"), qn("w:del")
            )
        ]
        duplicates = {k: v for k, v in Counter(ids).items() if v > 1}
        assert duplicates == {}

    def it_falls_back_on_the_pristine_paragraph_never_nesting_revisions(
        self, tmp_path: Path
    ):
        def build(text):
            def _build(document):
                paragraph = document.add_paragraph()
                paragraph._p.append(
                    parse_xml(f"<w:r {W}><w:t>a b</w:t><w:tab/><w:t>{text}</w:t></w:r>")
                )

            return _build

        # region 1 crosses the tab (refuses); region 2 after it succeeds
        # first in reverse order — the fallback must start from pristine
        result, a, b = _compare_pair(build("c d"), build("c D"), tmp_path)
        body = result.document.element.body
        for deletion in body.iter(qn("w:del")):
            assert deletion.find(f".//{qn('w:del')}") is None
            assert deletion.find(f".//{qn('w:ins')}") is None
        for insertion in body.iter(qn("w:ins")):
            assert insertion.find(f".//{qn('w:delText')}") is None
        result.document.revisions.accept_all()
        assert _visible(result.document) == _visible(docx.Document(b))

    def it_refuses_field_result_changes_instead_of_flattening(self, tmp_path: Path):
        def build(result_text):
            def _build(document):
                paragraph = document.add_paragraph()
                paragraph._p.append(
                    parse_xml(
                        f"<w:p {W}></w:p>"
                    )[0] if False else parse_xml(
                        f"<w:r {W}><w:fldChar w:fldCharType='begin'/></w:r>"
                    )
                )
                paragraph._p.append(
                    parse_xml(
                        f"<w:r {W}><w:instrText xml:space='preserve'> DATE </w:instrText></w:r>"
                    )
                )
                paragraph._p.append(
                    parse_xml(f"<w:r {W}><w:fldChar w:fldCharType='separate'/></w:r>")
                )
                paragraph._p.append(
                    parse_xml(f"<w:r {W}><w:t>{result_text}</w:t></w:r>")
                )
                paragraph._p.append(
                    parse_xml(f"<w:r {W}><w:fldChar w:fldCharType='end'/></w:r>")
                )

            return _build

        with pytest.raises(UnsupportedStructureError, match="cannot safely redline"):
            _compare_pair(
                build("June 1, 2026"), build("July 9, 2026"), tmp_path
            )

    def it_stamps_one_edit_with_one_timestamp(self, tmp_path: Path):
        naive = dt.datetime(2026, 1, 2, 3, 4, 5)  # NAIVE on purpose

        def build(rows):
            def _build(document):
                table = document.add_table(rows=len(rows), cols=1)
                for index, value in enumerate(rows):
                    table.cell(index, 0).text = value

            return _build

        result, _a, _b = _compare_pair(
            build(["kept", "dropped"]), build(["kept"]), tmp_path, date=naive
        )
        dates = {
            node.get(qn("w:date"))
            for node in result.document.element.body.iter(
                qn("w:ins"), qn("w:del")
            )
            if node.get(qn("w:date"))
        }
        assert len(dates) == 1, dates


class DescribeCompositionRegressions:
    def it_drops_half_bookmarks_at_range_edges_with_a_finding(self, tmp_path):
        from docx.composition import insert_blocks_from

        source = docx.Document()
        first = source.add_paragraph("range starts here")
        second = source.add_paragraph("copied tail")
        first._p.append(parse_xml(f"<w:bookmarkStart {W} w:id='5' w:name='Wide'/>"))
        second._p.append(parse_xml(f"<w:bookmarkEnd {W} w:id='5'/>"))
        destination = docx.Document(str(fixture_path(MINIMAL)))
        target = destination.paragraphs[1]._p
        target.append(parse_xml(f"<w:bookmarkStart {W} w:id='9' w:name='Existing'/>"))
        destination.paragraphs[2]._p.append(
            parse_xml(f"<w:bookmarkEnd {W} w:id='9'/>")
        )
        report = insert_blocks_from(
            destination,
            source,
            "copied tail",  # only the paragraph with the dangling END
            anchor="First body paragraph",
        )
        assert any(
            f.kind == "bookmark_partially_in_range" for f in report.findings
        )
        # the copied end must not adopt id 9 and truncate 'Existing'
        from docx.bookmarks import list_bookmarks

        existing = next(
            b for b in list_bookmarks(destination) if b.name == "Existing"
        )
        assert "Second body paragraph" in existing.text

    def it_remaps_numbering_referenced_only_by_an_imported_style(self, tmp_path):
        from docx.composition import insert_blocks_from
        from docx.numbering import ensure_decimal_definition

        source = docx.Document()
        source_num = ensure_decimal_definition(source)
        style = parse_xml(
            f"<w:style {W} w:type='paragraph' w:styleId='NumberedClause'>"
            "<w:name w:val='NumberedClause'/>"
            f"<w:pPr><w:numPr><w:ilvl w:val='0'/><w:numId w:val='{source_num}'/>"
            "</w:numPr></w:pPr></w:style>"
        )
        source.styles.element.append(style)
        source.add_paragraph("clause body text", style="NumberedClause")
        destination = docx.Document(str(fixture_path(MINIMAL)))
        destination_num = ensure_decimal_definition(destination)
        report = insert_blocks_from(
            destination,
            source,
            "clause body text",
            anchor="Second body paragraph",
        )
        assert "NumberedClause" in report.imported_styles
        new_id = report.numbering_map[source_num]
        assert new_id != destination_num
        imported = next(
            s
            for s in destination.styles.element.findall(qn("w:style"))
            if s.get(qn("w:styleId")) == "NumberedClause"
        )
        num_ref = imported.find(f".//{qn('w:numId')}")
        assert num_ref.get(qn("w:val")) == str(new_id)

    def it_never_duplicates_style_ids_within_one_import_batch(self, tmp_path):
        from docx.composition import insert_blocks_from

        source = docx.Document()
        for style_id, name in (("X", "SourceX"), ("XImported", "Other")):
            source.styles.element.append(
                parse_xml(
                    f"<w:style {W} w:type='paragraph' w:styleId='{style_id}'>"
                    f"<w:name w:val='{name}'/><w:rPr><w:b/></w:rPr></w:style>"
                )
            )
        p1 = source.add_paragraph("styled one")
        p1._p.get_or_add_pPr().append(parse_xml(f"<w:pStyle {W} w:val='X'/>"))
        p2 = source.add_paragraph("styled two")
        p2._p.get_or_add_pPr().append(
            parse_xml(f"<w:pStyle {W} w:val='XImported'/>")
        )
        destination = docx.Document(str(fixture_path(MINIMAL)))
        destination.styles.element.append(
            parse_xml(
                f"<w:style {W} w:type='paragraph' w:styleId='X'>"
                "<w:name w:val='DestX'/><w:rPr><w:i/></w:rPr></w:style>"
            )
        )
        insert_blocks_from(
            destination, source, "styled one", count=2,
            anchor="Second body paragraph",
        )
        ids = [
            s.get(qn("w:styleId"))
            for s in destination.styles.element.findall(qn("w:style"))
        ]
        assert len(ids) == len(set(ids)), ids

    def it_remaps_refs_split_across_instr_text_runs(self, tmp_path):
        from docx.composition import insert_blocks_from

        source = docx.Document()
        target = source.add_paragraph()
        target._p.append(parse_xml(f"<w:bookmarkStart {W} w:id='1' w:name='Target'/>"))
        target._p.append(parse_xml(f"<w:r {W}><w:t>the target text</w:t></w:r>"))
        target._p.append(parse_xml(f"<w:bookmarkEnd {W} w:id='1'/>"))
        ref = source.add_paragraph()
        ref._p.append(parse_xml(f"<w:r {W}><w:fldChar w:fldCharType='begin'/></w:r>"))
        ref._p.append(
            parse_xml(f"<w:r {W}><w:instrText xml:space='preserve'> REF </w:instrText></w:r>")
        )
        ref._p.append(
            parse_xml(
                f"<w:r {W}><w:instrText xml:space='preserve'>Target \\h </w:instrText></w:r>"
            )
        )
        ref._p.append(parse_xml(f"<w:r {W}><w:fldChar w:fldCharType='separate'/></w:r>"))
        ref._p.append(parse_xml(f"<w:r {W}><w:t>(ref)</w:t></w:r>"))
        ref._p.append(parse_xml(f"<w:r {W}><w:fldChar w:fldCharType='end'/></w:r>"))
        destination = docx.Document(str(fixture_path(MINIMAL)))
        anchor_p = destination.paragraphs[1]._p
        anchor_p.append(parse_xml(f"<w:bookmarkStart {W} w:id='7' w:name='Target'/>"))
        anchor_p.append(parse_xml(f"<w:bookmarkEnd {W} w:id='7'/>"))
        report = insert_blocks_from(
            destination, source, "the target text", count=2,
            anchor="Second body paragraph",
        )
        renamed = report.bookmarks_renamed["Target"]
        instr = "".join(
            node.text or ""
            for node in destination.element.body.iter(qn("w:instrText"))
        )
        assert f"REF {renamed} " in instr
        assert " REF Target " not in instr  # copied ref no longer hits dest's

    def it_refuses_unloadable_images_before_mutating_anything(self, tmp_path):
        from docx.composition import insert_blocks_from

        source = docx.Document()
        source.add_paragraph("before image")
        source.add_picture(io.BytesIO(TINY_PNG))
        image_part = next(
            rel.target_part
            for rel in source.part.rels.values()
            if not rel.is_external and "image" in rel.reltype
        )
        image_part._blob = b"not an image at all"  # e.g. an EMF blob
        destination = docx.Document(str(fixture_path(MINIMAL)))
        styles_before = destination.styles.element.xml
        with pytest.raises(UnsupportedStructureError, match="image format"):
            insert_blocks_from(
                destination, source, "before image", count=2,
                anchor="Second body paragraph",
            )
        assert destination.styles.element.xml == styles_before


class DescribeBookmarkFieldRegressions:
    def it_refuses_deleting_a_bookmark_referenced_by_a_split_instruction(self):
        from docx.bookmarks import delete_bookmark

        document = docx.Document(
            str(fixture_path("generated/feature-isolated/bookmarks.docx"))
        )
        paragraph = document.add_paragraph()
        paragraph._p.append(parse_xml(f"<w:r {W}><w:fldChar w:fldCharType='begin'/></w:r>"))
        paragraph._p.append(
            parse_xml(f"<w:r {W}><w:instrText xml:space='preserve'> REF </w:instrText></w:r>")
        )
        paragraph._p.append(
            parse_xml(
                f"<w:r {W}><w:instrText xml:space='preserve'>DefinedTerm \\h </w:instrText></w:r>"
            )
        )
        paragraph._p.append(parse_xml(f"<w:r {W}><w:fldChar w:fldCharType='end'/></w:r>"))
        with pytest.raises(UnsupportedStructureError, match="referenced"):
            delete_bookmark(document, "DefinedTerm")

    def it_counts_noteref_and_hyperlink_anchor_references(self):
        from docx.bookmarks import delete_bookmark

        document = docx.Document(
            str(fixture_path("generated/feature-isolated/bookmarks.docx"))
        )
        paragraph = document.add_paragraph()
        paragraph._p.append(
            parse_xml(
                f"<w:hyperlink {W} w:anchor='DefinedTerm'>"
                "<w:r><w:t>jump</w:t></w:r></w:hyperlink>"
            )
        )
        with pytest.raises(UnsupportedStructureError, match="referenced"):
            delete_bookmark(document, "DefinedTerm")

    def it_refuses_bookmarking_inside_a_field_result(self):
        from docx.bookmarks import create_bookmark
        from docx.search import find_one

        document = docx.Document(
            str(fixture_path("generated/feature-isolated/fields.docx"))
        )
        span = find_one(document, "June 1, 2026")
        assert span.in_field
        with pytest.raises(UnsupportedStructureError, match="field"):
            create_bookmark(document, span, "DoomedMark")

    def it_round_trips_cross_paragraph_bookmark_text(self):
        from docx.bookmarks import list_bookmarks

        document = docx.Document(str(fixture_path(MINIMAL)))
        first = document.paragraphs[1]._p
        second = document.paragraphs[2]._p
        first.insert(  # BEFORE paragraph 1's runs, so text spans both
            list(first).index(first.find(qn("w:r"))),
            parse_xml(f"<w:bookmarkStart {W} w:id='11' w:name='Across'/>"),
        )
        second.append(parse_xml(f"<w:bookmarkEnd {W} w:id='11'/>"))
        across = next(b for b in list_bookmarks(document) if b.name == "Across")
        assert "\n" in across.text  # paragraph boundary, Span.text semantics

    def it_deletes_every_pair_sharing_a_duplicate_name(self):
        from docx.bookmarks import delete_bookmark, list_bookmarks

        document = docx.Document(str(fixture_path(MINIMAL)))
        for index, bookmark_id in enumerate((21, 22)):
            paragraph = document.paragraphs[1 + index]._p
            paragraph.append(
                parse_xml(f"<w:bookmarkStart {W} w:id='{bookmark_id}' w:name='Dup'/>")
            )
            paragraph.append(parse_xml(f"<w:bookmarkEnd {W} w:id='{bookmark_id}'/>"))
        delete_bookmark(document, "Dup")
        assert all(b.name != "Dup" for b in list_bookmarks(document))


class DescribeFormattingRegressions:
    def it_resolves_the_paragraph_style_for_runs_inside_hyperlinks(self):
        from docx.enum.style import WD_STYLE_TYPE
        from docx.formatting import format_of
        from docx.shared import Pt
        from docx.text.run import Run

        document = docx.Document(str(fixture_path(MINIMAL)))
        style = document.styles.add_style("LinkyPara", WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(17)
        paragraph = document.add_paragraph(style="LinkyPara")
        paragraph._p.append(
            parse_xml(
                f"<w:hyperlink {W} w:anchor='DefinedTerm'>"
                "<w:r><w:t>linked text</w:t></w:r></w:hyperlink>"
            )
        )
        run_element = paragraph._p.find(f"{qn('w:hyperlink')}/{qn('w:r')}")
        resolved = format_of(Run(run_element, paragraph))
        assert resolved["size_pt"].value == 17
        assert resolved["size_pt"].source == "paragraph_style:LinkyPara"

    def it_uses_hyperlink_runs_for_surrounding_format(self):
        from docx.enum.style import WD_STYLE_TYPE
        from docx.formatting import surrounding_format
        from docx.shared import Pt

        document = docx.Document(str(fixture_path(MINIMAL)))
        style = document.styles.add_style("LinkOnly", WD_STYLE_TYPE.PARAGRAPH)
        style.font.size = Pt(19)
        paragraph = document.add_paragraph(style="LinkOnly")
        paragraph._p.append(
            parse_xml(
                f"<w:hyperlink {W} w:anchor='x'>"
                "<w:r><w:rPr><w:b/></w:rPr><w:t>only linked text here</w:t></w:r>"
                "</w:hyperlink>"
            )
        )
        resolved = surrounding_format(document, "only linked text here")
        assert resolved["bold"].value is True  # the RUN resolved, not a fallback
        assert resolved["size_pt"].value == 19

    def it_reports_agreeing_values_from_different_layers_honestly(self):
        from docx.enum.style import WD_STYLE_TYPE
        from docx.formatting import format_of
        from docx.search import find_one

        document = docx.Document(str(fixture_path(MINIMAL)))
        char_style = document.styles.add_style("AlsoBold", WD_STYLE_TYPE.CHARACTER)
        char_style.font.bold = True
        paragraph = document.add_paragraph()
        direct = paragraph.add_run("direct half")
        direct.bold = True
        styled = paragraph.add_run(" styled half")
        styled.style = char_style
        span = find_one(document, "direct half styled half")
        resolved = format_of(span)
        assert resolved["bold"].value is True
        assert resolved["bold"].source == "agreeing_layers"
        assert "direct" in resolved["bold"].chain
        assert any(s.startswith("character_style:") for s in resolved["bold"].chain)
