"""Focused regressions for guarded secondary editing behavior."""

from __future__ import annotations

import pytest

import docx
from docx.bookmarks import delete_bookmark
from docx.commentops import anchored_text
from docx.errors import UnsupportedStructureError
from docx.numbering import ensure_bullet_definition, restart_numbering
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement
from docx.search import find_one


def it_refuses_numbering_restarts_that_would_drop_level_overrides():
    document = docx.Document()
    num_id = ensure_bullet_definition(document)
    numbering = document.part.numbering_part.element
    definition = next(
        node
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) == str(num_id)
    )
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "1")
    definition.append(override)
    before = numbering.xml

    with pytest.raises(UnsupportedStructureError, match="level overrides"):
        restart_numbering(document, num_id=num_id)

    assert numbering.xml == before


def it_refuses_duplicate_bookmark_marker_ids_atomically():
    document = docx.Document()
    paragraph = document.add_paragraph("bookmark targets")._p
    for bookmark_name, bookmark_id in (("target", "5"), ("unrelated", "05")):
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), bookmark_id)
        start.set(qn("w:name"), bookmark_name)
        paragraph.append(start)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), bookmark_id)
        paragraph.append(end)
    before = document.element.xml

    with pytest.raises(UnsupportedStructureError, match="duplicate"):
        delete_bookmark(document, "target")

    assert document.element.xml == before


def it_preserves_paragraph_boundaries_in_anchored_comment_text():
    document = docx.Document()
    document.add_paragraph("Alpha")
    document.add_paragraph("Beta")
    comment = find_one(document, "Alpha Beta").comment("Review", author="Reviewer")

    assert anchored_text(document, comment) == "Alpha\nBeta"
