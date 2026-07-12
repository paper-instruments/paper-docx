"""Bookmarks and field authoring.

Fields are formulas; this package authors them and never computes their
values — every insertion carries placeholder text and arms
update-fields-on-open for the renderer.
"""

from __future__ import annotations

from pathlib import Path

import pytest

import docx
from docx.bookmarks import create_bookmark, delete_bookmark, list_bookmarks
from docx.errors import (
    DocumentProtectedError,
    TargetNotFoundError,
    UnsupportedStructureError,
)
from docx.fields import (
    add_date_field,
    add_page_count_field,
    add_page_number_field,
    add_reference_field,
    insert_toc_after,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.search import find_one

from .harness.contract import save_and_reopen
from .harness.paths import fixture_path

MINIMAL = "generated/minimal-clean/minimal.docx"
BOOKMARKS = "generated/feature-isolated/bookmarks.docx"


def _doc(relpath: str = MINIMAL):
    return docx.Document(str(fixture_path(relpath)))


class DescribeBookmarks:
    def it_lists_bookmarks_with_their_wrapped_text(self):
        found = list_bookmarks(_doc(BOOKMARKS))
        named = {b.name: b for b in found}
        assert named["DefinedTerm"].text == "the Master Agreement"
        assert named["DefinedTerm"].is_point is False
        assert named["_GoBack"].is_point is True

    def it_creates_a_bookmark_on_an_exact_span(self, tmp_path: Path):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        info = create_bookmark(document, span, "KeyPhrase")
        assert info.text == "perfectly ordinary"
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        named = {b.name: b for b in list_bookmarks(reopened)}
        assert named["KeyPhrase"].text == "perfectly ordinary"
        # the surrounding text is untouched
        assert (
            "First body paragraph with perfectly ordinary text."
            in [p.text for p in reopened.paragraphs]
        )

    def it_allocates_globally_unique_ids(self):
        document = _doc(BOOKMARKS)
        info = create_bookmark(
            document, find_one(document, "for definitions"), "SecondMark"
        )
        existing_ids = [b.bookmark_id for b in list_bookmarks(document)]
        assert existing_ids.count(info.bookmark_id) == 1

    def it_refuses_duplicate_names(self):
        document = _doc(BOOKMARKS)
        with pytest.raises(UnsupportedStructureError, match="already exists"):
            create_bookmark(
                document, find_one(document, "for definitions"), "DefinedTerm"
            )

    def it_treats_bookmark_names_as_case_insensitive(self):
        document = docx.Document()
        document.add_paragraph("first second")
        create_bookmark(document, find_one(document, "first"), "DefinedTerm")

        with pytest.raises(UnsupportedStructureError, match="already exists"):
            create_bookmark(document, find_one(document, "second"), "definedterm")

    def it_refuses_an_unrelated_malformed_bookmark_before_deletion(self):
        document = docx.Document()
        document.add_paragraph("target")
        create_bookmark(document, find_one(document, "target"), "Target")
        malformed = OxmlElement("w:bookmarkStart")
        malformed.set(qn("w:id"), "bad")
        malformed.set(qn("w:name"), "Other")
        document.paragraphs[0]._p.append(malformed)

        with pytest.raises(UnsupportedStructureError, match="non-numeric"):
            delete_bookmark(document, "Target")

    def it_validates_word_legal_names(self):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        for bad in ("1starts-with-digit", "has space", "", "x" * 41):
            with pytest.raises(ValueError, match="Word-legal"):
                create_bookmark(document, span, bad)

    def it_deletes_a_bookmark_keeping_the_text(self, tmp_path: Path):
        document = _doc(BOOKMARKS)
        delete_bookmark(document, "DefinedTerm")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert "DefinedTerm" not in {b.name for b in list_bookmarks(reopened)}
        assert "See the Master Agreement for definitions." in [
            p.text for p in reopened.paragraphs
        ]

    def it_refuses_deleting_a_field_referenced_bookmark(self):
        document = _doc(BOOKMARKS)
        paragraph = document.add_paragraph("See also: ")
        add_reference_field(paragraph, bookmark="DefinedTerm")
        with pytest.raises(UnsupportedStructureError, match="referenced by"):
            delete_bookmark(document, "DefinedTerm")

    def it_refuses_deleting_an_implicitly_referenced_bookmark(self):
        document = _doc(BOOKMARKS)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), " DefinedTerm \\* MERGEFORMAT ")
        document.add_paragraph()._p.append(field)

        with pytest.raises(UnsupportedStructureError, match="referenced by"):
            delete_bookmark(document, "DefinedTerm")

    def it_never_treats_quote_literal_text_as_a_bookmark_reference(self):
        document = _doc(BOOKMARKS)
        field = OxmlElement("w:fldSimple")
        instruction = ' QUOTE "literal REF DefinedTerm text" '
        field.set(qn("w:instr"), instruction)
        document.add_paragraph()._p.append(field)

        delete_bookmark(document, "DefinedTerm")

        assert field.get(qn("w:instr")) == instruction

    def it_ignores_a_reference_to_a_different_bookmark(self):
        document = _doc(BOOKMARKS)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), " REF OtherBookmark ")
        document.add_paragraph()._p.append(field)

        delete_bookmark(document, "DefinedTerm")

    def it_does_not_mistake_an_embed_field_for_an_implicit_reference(self):
        document = docx.Document()
        document.add_paragraph("embedded")
        create_bookmark(document, find_one(document, "embedded"), "EMBED")
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), " EMBED Word.Document.12 ")
        document.add_paragraph()._p.append(field)

        delete_bookmark(document, "EMBED")

    def it_refuses_unknown_names(self):
        with pytest.raises(TargetNotFoundError):
            delete_bookmark(_doc(), "Nonexistent")

    def it_refuses_on_protected_documents(self):
        document = _doc("generated/feature-isolated/protected-readonly.docx")
        span = find_one(document, "locked")
        with pytest.raises(DocumentProtectedError):
            create_bookmark(document, span, "Nope")


class DescribeFieldAuthoring:
    def _settings_xml(self, document) -> str:
        from lxml import etree

        return etree.tostring(document.settings.element).decode()

    def it_authors_page_and_count_fields_in_a_footer(self, tmp_path: Path):
        document = _doc()
        footer = document.sections[0].footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0]
        paragraph.text = "Page "
        add_page_number_field(paragraph)
        paragraph.add_run(" of ")
        add_page_count_field(paragraph)
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        footer_xml = reopened.sections[0].footer.paragraphs[0]._p.xml
        assert ' PAGE ' in footer_xml and ' NUMPAGES ' in footer_xml
        assert 'w:updateFields' in self._settings_xml(reopened)

    def it_authors_a_date_field_without_computing_a_date(self, tmp_path: Path):
        document = _doc()
        paragraph = document.add_paragraph("Dated: ")
        add_date_field(paragraph, date_format="MMMM d, yyyy")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        xml = reopened.element.xml
        # quotes inside the w:instr attribute serialize as &quot;
        assert "DATE \\@ &quot;MMMM d, yyyy&quot;" in xml
        assert "(date)" in xml  # a placeholder, never a computed value

    def it_authors_cross_references_in_three_kinds(self):
        document = _doc(BOOKMARKS)
        paragraph = document.add_paragraph("Reference: ")
        add_reference_field(paragraph, bookmark="DefinedTerm", kind="text")
        add_reference_field(paragraph, bookmark="DefinedTerm", kind="page")
        add_reference_field(paragraph, bookmark="DefinedTerm", kind="number")
        xml = document.element.xml
        assert " REF DefinedTerm \\h " in xml
        assert " PAGEREF DefinedTerm \\h " in xml
        assert " REF DefinedTerm \\r \\h " in xml

    def it_refuses_references_to_unknown_bookmarks(self):
        document = _doc()
        paragraph = document.add_paragraph("Reference: ")
        with pytest.raises(TargetNotFoundError, match="no bookmark"):
            add_reference_field(paragraph, bookmark="Ghost")

    def it_validates_the_reference_kind(self):
        document = _doc(BOOKMARKS)
        with pytest.raises(ValueError, match="kind"):
            add_reference_field(
                document.add_paragraph(), bookmark="DefinedTerm", kind="chapter"
            )

    def it_inserts_a_toc_complex_field(self, tmp_path: Path):
        document = _doc()
        insert_toc_after(document, "Minimal Clean Document", levels=(1, 3))
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        xml = reopened.element.xml
        assert ' TOC \\o "1-3" \\h \\z \\u ' in xml
        assert 'w:fldCharType="begin"' in xml and 'w:fldCharType="end"' in xml
        assert 'w:dirty="true"' in xml
        assert "w:updateFields" in self._settings_xml(reopened)

    def it_validates_toc_levels(self):
        with pytest.raises(ValueError, match="levels"):
            insert_toc_after(_doc(), "Minimal Clean Document", levels=(3, 1))

    def it_refuses_field_authoring_on_protected_documents(self):
        document = _doc("generated/feature-isolated/protected-readonly.docx")
        with pytest.raises(DocumentProtectedError):
            add_page_number_field(document.paragraphs[0])


class DescribeFieldSelfConsistency:
    """The in_field guard must recognize the fields this package authors."""

    def it_refuses_edits_inside_our_own_toc(self):
        document = _doc()
        insert_toc_after(document, "Minimal Clean Document")
        span = find_one(document, "Table of contents placeholder")
        assert span.in_field
        with pytest.raises(UnsupportedStructureError):
            span.replace("hand-edited toc")

    def it_refuses_edits_inside_our_own_simple_fields(self):
        document = _doc()
        paragraph = document.add_paragraph("Dated: ")
        add_date_field(paragraph)
        span = find_one(document, "(date)")
        assert span.in_field
        with pytest.raises(UnsupportedStructureError):
            span.replace("July 8, 2026")
