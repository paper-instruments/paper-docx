"""Tests for docx.blocks — anchor-relative block operations."""

from __future__ import annotations

import datetime as dt
import shutil
from pathlib import Path

import pytest

import docx
from docx.blocks import (
    insert_section_after,
    tracked_delete_paragraphs,
    tracked_replace_paragraphs,
)
from docx.errors import (
    BoundaryViolationError,
    TargetNotFoundError,
    UnsupportedStructureError,
)
from docx.oxml.ns import qn
from docx.story import iter_blocks, outline

from .harness.contract import assert_changed_parts, assert_refusal_atomic, save_and_reopen
from .harness.paths import fixture_path

MINIMAL = "generated/minimal-clean/minimal.docx"
TABLES = "generated/feature-isolated/table-merged-nested.docx"
CONTROLS = "generated/feature-isolated/content-control.docx"
FRAGMENTED = "generated/feature-isolated/fragmented-runs.docx"
GAUNTLET = "generated/gauntlet/gauntlet.docx"

FROZEN = dt.datetime(2026, 7, 7, 12, 0, 0, tzinfo=dt.timezone.utc)


def _doc(relpath: str):
    return docx.Document(str(fixture_path(relpath)))


def _texts(document, view: str = "current"):
    return [b.text for b in iter_blocks(document, view=view)]


class DescribeInsertSectionAfter:
    def it_inserts_a_heading_and_body_after_a_string_anchor(self, tmp_path: Path):
        document = _doc(MINIMAL)
        result = insert_section_after(
            document,
            "First body paragraph",
            heading="New Section",
            paragraphs=["Alpha body.", "Beta body."],
        )
        assert result.inserted_blocks == 3
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        blocks = list(iter_blocks(reopened))
        assert [b.text for b in blocks[1:5]] == [
            "First body paragraph with perfectly ordinary text.",
            "New Section",
            "Alpha body.",
            "Beta body.",
        ]
        assert blocks[2].style_id == "Heading2"

    def it_accepts_a_block_anchor_object(self):
        document = _doc(MINIMAL)
        anchor = outline(document).blocks[1].anchor
        insert_section_after(
            document, anchor, heading="Anchored Section", paragraphs=["Body."]
        )
        assert "Anchored Section" in _texts(document)

    def it_refuses_a_stale_block_anchor(self):
        document = _doc(MINIMAL)
        anchor = outline(document).blocks[1].anchor
        from docx.search import find_one

        find_one(document, "perfectly ordinary").replace("entirely different")
        with pytest.raises(TargetNotFoundError, match="stale"):
            insert_section_after(document, anchor, heading="X", paragraphs=[])

    def it_validates_style_ids_before_mutating(self):
        document = _doc(MINIMAL)
        assert_refusal_atomic(
            document,
            lambda doc: insert_section_after(
                doc, "First body paragraph", heading="X", paragraphs=["Y"],
                heading_style="NoSuchStyle99",
            ),
            TargetNotFoundError,
        )

    def it_requires_an_author_when_tracked(self):
        document = _doc(MINIMAL)
        with pytest.raises(ValueError, match="author"):
            insert_section_after(
                document, "First body paragraph", heading="X", paragraphs=[],
                tracked=True,
            )

    def it_makes_tracked_insertions_reject_away_completely(self):
        """The paragraph-mark stamp means rejection leaves no husk behind."""
        document = _doc(MINIMAL)
        before = _texts(document)
        insert_section_after(
            document,
            "First body paragraph",
            heading="Tracked Section",
            paragraphs=["Tracked body."],
            tracked=True,
            author="Carol QA",
            date=FROZEN,
        )
        assert "Tracked Section" in _texts(document)
        document.revisions.reject_all()
        assert _texts(document) == before

    def it_keeps_the_changed_part_budget(self, tmp_path: Path):
        source = fixture_path(MINIMAL)
        working = tmp_path / "work.docx"
        shutil.copyfile(source, working)
        document = docx.Document(str(working))
        insert_section_after(
            document, "First body paragraph", heading="Budget", paragraphs=["B."]
        )
        out = tmp_path / "out.docx"
        docx.package.patch_save(working, document, out)
        assert_changed_parts(working, out, {"word/document.xml"})


class DescribeTrackedDeleteParagraphs:
    def it_preserves_each_runs_formatting_inside_the_deletion(self):
        """Runs move into w:del with their own rPr (reject restores exactly)."""
        document = _doc(FRAGMENTED)
        result = tracked_delete_paragraphs(
            document, "Consulting rate", count=1, author="Carol QA", date=FROZEN
        )
        assert result.deleted_blocks == 1
        assert result.deleted_text[0].startswith("Consulting rate: $75–100/hr")
        (deletion,) = document.element.body.xpath("//w:p/w:del")
        runs = deletion.findall(qn("w:r"))
        assert len(runs) == 8, "each original run must survive inside w:del"
        assert deletion.xpath(".//w:delText") and not deletion.xpath(".//w:t")
        bold_runs = [r for r in runs if r.find(qn("w:rPr")) is not None
                     and r.find(qn("w:rPr")).find(qn("w:b")) is not None]
        assert len(bold_runs) == 3, "bold formatting must survive deletion markup"

    def it_stamps_the_paragraph_mark_so_accept_removes_the_paragraph(self):
        document = _doc(MINIMAL)
        tracked_delete_paragraphs(
            document, "First body paragraph", count=1, author="Carol QA", date=FROZEN
        )
        document.revisions.accept_all()
        texts = _texts(document)
        assert "First body paragraph with perfectly ordinary text." not in texts
        assert "" not in texts, "accepting a deletion must not leave an empty husk"

    def it_selects_a_range_by_end_anchor(self):
        document = _doc(MINIMAL)
        result = tracked_delete_paragraphs(
            document,
            "First body paragraph",
            end_anchor="Second body paragraph",
            author="Carol QA",
            date=FROZEN,
        )
        assert result.deleted_blocks == 2

    def it_refuses_ranges_that_do_not_share_one_parent(self):
        """Body paragraph -> table-cell paragraph: never silently corrupted."""
        document = _doc(TABLES)
        assert_refusal_atomic(
            document,
            lambda doc: tracked_delete_paragraphs(
                doc, "Paragraph before the merged",
                end_anchor="N11",  # nested-table cell paragraph
                author="Carol QA", date=FROZEN,
            ),
            BoundaryViolationError,
        )

    def it_counts_ranges_among_siblings_so_tables_are_bracketed_not_selected(self):
        """count=2 from the paragraph before a table selects the paragraphs
        AROUND it (siblings); the table itself is untouched."""
        document = _doc(TABLES)
        result = tracked_delete_paragraphs(
            document, "Paragraph before the merged", count=2,
            author="Carol QA", date=FROZEN,
        )
        assert result.deleted_blocks == 2
        assert result.deleted_text[1] == "Paragraph after the tables."
        assert len(document.tables) == 1  # table survives

    def it_refuses_paragraphs_with_inline_controls(self):
        document = _doc(CONTROLS)
        assert_refusal_atomic(
            document,
            lambda doc: tracked_delete_paragraphs(
                doc, "Inline control follows", count=1,
                author="Carol QA", date=FROZEN,
            ),
            UnsupportedStructureError,
        )

    def it_refuses_cross_story_ranges(self):
        document = _doc(GAUNTLET)
        with pytest.raises(BoundaryViolationError, match="different story"):
            tracked_delete_paragraphs(
                document,
                "Gauntlet section two body",
                end_anchor="Gauntlet header, section one",
                author="Carol QA",
                date=FROZEN,
            )

    def it_refuses_counts_past_the_last_paragraph(self):
        document = _doc(MINIMAL)
        with pytest.raises(TargetNotFoundError, match="past the last"):
            tracked_delete_paragraphs(
                document, "Second body paragraph", count=10,
                author="Carol QA", date=FROZEN,
            )


class DescribeTrackedReplaceParagraphs:
    def it_deletes_the_range_and_inserts_replacements(self, tmp_path: Path):
        document = _doc(MINIMAL)
        result = tracked_replace_paragraphs(
            document,
            "Second body paragraph",
            ["Replacement one.", "Replacement two."],
            author="Carol QA",
            date=FROZEN,
        )
        assert result.deleted_blocks == 1 and result.inserted_blocks == 2
        assert len(result.revision_ids) == 3
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        texts = _texts(reopened)
        assert "Replacement one." in texts and "Replacement two." in texts

    def it_keeps_the_changed_part_budget(self, tmp_path: Path):
        source = fixture_path(MINIMAL)
        working = tmp_path / "work.docx"
        shutil.copyfile(source, working)
        document = docx.Document(str(working))
        tracked_replace_paragraphs(
            document, "Second body paragraph", ["R."], author="Carol QA", date=FROZEN
        )
        out = tmp_path / "out.docx"
        docx.package.patch_save(working, document, out)
        assert_changed_parts(working, out, {"word/document.xml"})

    @pytest.mark.lo_smoke
    def it_produces_output_libreoffice_can_open(self, tmp_path: Path):
        from .harness.lo import assert_libreoffice_opens

        document = _doc(MINIMAL)
        tracked_replace_paragraphs(
            document, "Second body paragraph", ["LO check."],
            author="Carol QA", date=FROZEN,
        )
        out = tmp_path / "out.docx"
        document.save(str(out))
        assert_libreoffice_opens(out)
