"""Directory-prefix collisions, pinned to measured Microsoft Word verdicts.

Word was shown three spellings of the same collision: a zero-length member named
``word`` sitting beside members under ``word/``. It refused all three -- plain,
with the MS-DOS directory bit set, and with a unix directory mode -- so the
colliding *name* is the defect and the attributes are incidental. Word opens a
package carrying ordinary zero-length ``word/`` folder records, so the rule must
apply to members only.

Ledger: ``verifying-against-word/WORD-VERDICTS.md``, the ``prefix_collision_*``
rows.
"""

from __future__ import annotations

import stat
import zipfile
from pathlib import Path

import pytest

import docx
from docx import _paperpkg
from docx.errors import MalformedPackageError

_COLLIDING_MEMBERS = ("word/document.xml", "word/_rels/document.xml.rels")


def _write_collision(path: Path, external_attr: int) -> None:
    """Write an archive whose member ``word`` collides with the ``word/`` tree.

    ``word`` is written first, as it is in every recorded fixture: that ordering
    is what makes a per-member attribute check report the wrong defect.
    """
    with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_STORED) as archive:
        info = zipfile.ZipInfo("word")
        info.external_attr = external_attr
        archive.writestr(info, b"")
        for name in _COLLIDING_MEMBERS:
            archive.writestr(name, b"payload")


def _repack_with_folder_records(source: Path, target: Path) -> None:
    """Copy a package, inserting zero-length folder records ahead of it.

    ``ZIP_STORED`` is load-bearing: a deflated empty member carries a two-byte
    payload and is refused as a directory entry carrying data, which is a
    different defect and would make this control misleading.
    """
    with zipfile.ZipFile(source) as original:
        members = [(info.filename, original.read(info.filename)) for info in original.infolist()]

    with zipfile.ZipFile(target, "w", compression=zipfile.ZIP_STORED) as archive:
        for folder in ("word/", "word/_rels/", "docProps/", "word/media/"):
            info = zipfile.ZipInfo(folder)
            info.external_attr = (stat.S_IFDIR | 0o755) << 16 | 0x10
            archive.writestr(info, b"")
        for name, data in members:
            archive.writestr(name, data)


class DescribeWordVerdictPrefixCollisions:
    """All three shapes refuse, and all three name the collision."""

    @pytest.mark.parametrize(
        ("shape", "external_attr"),
        [
            ("plain", 0),
            ("dosbit", 0x10),
            ("dirmode", (stat.S_IFDIR | 0o755) << 16),
        ],
    )
    def it_refuses_every_measured_collision_shape(
        self, tmp_path: Path, shape: str, external_attr: int
    ):
        # Word: REFUSES, for all three.
        path = tmp_path / f"collision-{shape}.docx"
        _write_collision(path, external_attr)

        with pytest.raises(MalformedPackageError) as exc:
            _paperpkg._read_zip(path)

        message = str(exc.value)
        assert "'word'" in message
        assert "directory prefix of member 'word/" in message
        assert "one name cannot denote both a file and a directory" in message
        assert "Remove the member named 'word'." in message

    @pytest.mark.parametrize(
        ("shape", "external_attr"),
        [
            ("plain", 0),
            ("dosbit", 0x10),
            ("dirmode", (stat.S_IFDIR | 0o755) << 16),
        ],
    )
    def it_names_the_collision_rather_than_the_attributes(
        self, tmp_path: Path, shape: str, external_attr: int
    ):
        # The pre-pass placement is what makes this true: inside the per-member
        # loop the shallow member's attribute check fires first for two of the
        # three shapes and reports the wrong defect.
        path = tmp_path / f"collision-{shape}.docx"
        _write_collision(path, external_attr)

        with pytest.raises(MalformedPackageError) as exc:
            _paperpkg._read_zip(path)

        message = str(exc.value)
        assert "is a directory entry" not in message
        assert "unsupported filesystem entry type" not in message

    def it_refuses_a_collision_nested_deeper_than_the_top_level(self, tmp_path: Path):
        path = tmp_path / "deep.docx"
        with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_STORED) as archive:
            archive.writestr("word/media", b"")
            archive.writestr("word/media/image1.png", b"payload")

        with pytest.raises(MalformedPackageError) as exc:
            _paperpkg._read_zip(path)

        assert "'word/media'" in str(exc.value)

    def it_reads_a_package_whose_names_merely_share_a_prefix_string(self, tmp_path: Path):
        # "word" is a string prefix of "wordy.xml" but not a directory prefix.
        path = tmp_path / "sharedstring.docx"
        with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_STORED) as archive:
            archive.writestr("word", b"")
            archive.writestr("wordy.xml", b"payload")

        parts, _ = _paperpkg._read_zip(path)

        assert set(parts) == {"word", "wordy.xml"}


class DescribeFolderRecordsStillOpen:
    """Zero-length ``word/`` folder records are confirmed OPENS in Word."""

    def it_opens_a_package_carrying_zero_length_folder_records(self, tmp_path: Path):
        source = tmp_path / "plain.docx"
        document = docx.Document()
        document.add_paragraph("folder records")
        document.save(source)
        repacked = tmp_path / "foldered.docx"
        _repack_with_folder_records(source, repacked)

        reopened = docx.Document(repacked)

        assert [paragraph.text for paragraph in reopened.paragraphs] == ["folder records"]

    def it_opens_an_ordinary_package_unchanged(self, tmp_path: Path):
        path = tmp_path / "ordinary.docx"
        document = docx.Document()
        document.add_paragraph("ordinary")
        document.save(path)

        reopened = docx.Document(path)

        assert [paragraph.text for paragraph in reopened.paragraphs] == ["ordinary"]


class DescribeNonCollidingAttributedMembers:
    """The two attribute checks stay; this phase moves exactly one fixture.

    A directory-attributed member that collides with nothing has no Word verdict.
    Both checks are kept, so both refusals are pinned here as today's behaviour
    rather than left incidental.
    """

    def it_still_refuses_a_ms_dos_directory_bit_on_a_non_colliding_member(self, tmp_path: Path):
        path = tmp_path / "dosbit.docx"
        with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_STORED) as archive:
            info = zipfile.ZipInfo("notes")
            info.external_attr = 0x10
            archive.writestr(info, b"")

        with pytest.raises(MalformedPackageError) as exc:
            _paperpkg._read_zip(path)

        assert "is a directory entry" in str(exc.value)

    def it_still_refuses_a_unix_directory_mode_on_a_non_colliding_member(self, tmp_path: Path):
        path = tmp_path / "dirmode.docx"
        with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_STORED) as archive:
            info = zipfile.ZipInfo("notes")
            info.external_attr = (stat.S_IFDIR | 0o755) << 16
            archive.writestr(info, b"")

        with pytest.raises(MalformedPackageError) as exc:
            _paperpkg._read_zip(path)

        assert "unsupported filesystem entry type" in str(exc.value)

    def it_still_reads_a_plain_non_colliding_member(self, tmp_path: Path):
        path = tmp_path / "plainmember.docx"
        with zipfile.ZipFile(path, "w", compression=zipfile.ZIP_STORED) as archive:
            archive.writestr("notes", b"payload")

        parts, _ = _paperpkg._read_zip(path)

        assert parts == {"notes": b"payload"}
