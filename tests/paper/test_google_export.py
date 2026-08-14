"""Smoke a Google Docs → Word export through open, inspect, and replace."""

from __future__ import annotations

from pathlib import Path

import docx
from docx.formatting import format_of
from docx.numbering import list_numbering
from docx.package import patch_save
from docx.search import find_one

from .harness.paths import fixture_path

GOOGLE_EXPORT = "google/minimal-clean/google-docs-export.docx"


def _doc():
    return docx.Document(str(fixture_path(GOOGLE_EXPORT)))


class DescribeGoogleDocsExport:
    def it_opens_and_reads_the_visible_paragraphs(self):
        texts = [p.text for p in _doc().paragraphs if p.text]
        assert texts == [
            "Title",
            "Heading 1",
            "Heading 2",
            "Heading 3",
            "This is a document created in Google Docs, and exported to Microsoft Word.",
            "Bold",
            "Italic",
            "Underlined",
            "Bullet point 1",
            "Bullet point 2",
            "Bullet point 3",
            "Blue text",
            "Numbered list 1",
            "Numbered list 2",
            "Numbered list 3",
        ]
        document = _doc()
        assert document.paragraphs[0].style.name == "Title"
        assert document.paragraphs[1].style.name == "Heading 1"

    def it_sees_google_direct_formatting(self):
        document = _doc()
        assert format_of(find_one(document, "Bold"))["bold"].value is True
        assert format_of(find_one(document, "Italic"))["italic"].value is True
        assert format_of(find_one(document, "Underlined"))["underline"].value == "single"
        assert format_of(find_one(document, "Blue text"))["color_rgb"].value == "0000ff"

    def it_sees_google_bullet_and_numbered_lists(self):
        report = list_numbering(_doc())
        bullets = [p.text for p in report.numbered_paragraphs if p.num_id == 2]
        numbers = [p.text for p in report.numbered_paragraphs if p.num_id == 1]
        assert bullets == ["Bullet point 1", "Bullet point 2", "Bullet point 3"]
        assert numbers == ["Numbered list 1", "Numbered list 2", "Numbered list 3"]
        by_id = {item.num_id: item for item in report.definitions}
        assert by_id[2].levels[0].num_fmt == "bullet"
        assert by_id[1].levels[0].num_fmt == "decimal"

    def it_replaces_a_quoted_phrase_and_saves(self, tmp_path: Path):
        source = fixture_path(GOOGLE_EXPORT)
        document = docx.Document(str(source))
        find_one(
            document,
            "exported to Microsoft Word",
        ).replace("exported as a .docx")
        out = tmp_path / "edited.docx"
        result = patch_save(source, document, out)
        assert "word/document.xml" in result.changed_parts
        reopened = docx.Document(str(out))
        assert any(
            "exported as a .docx" in p.text for p in reopened.paragraphs
        )
        assert any("Google Docs" in p.text for p in reopened.paragraphs)
