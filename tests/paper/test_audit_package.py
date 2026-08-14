"""Adversarial package-I/O tests for bounded, atomic DOCX reads and writes."""

from __future__ import annotations

import io
import os
import shutil
import stat
import struct
import warnings
import zipfile
import zlib
from pathlib import Path

import pytest

import docx
from docx._zipguard import GuardedZipReader
from docx.errors import PackageLimitError, PaperRefusal
from docx.package import diagnose, diff_package, patch_save

from .harness.paths import fixture_path

MINIMAL = "generated/minimal-clean/minimal.docx"


def _write_archive(
    path: Path,
    members: tuple[tuple[str, bytes], ...],
    *,
    compression: int = zipfile.ZIP_STORED,
) -> None:
    with zipfile.ZipFile(path, "w", compression=compression) as archive:
        for name, data in members:
            archive.writestr(name, data)


def _guarded_read(path: Path) -> dict[str, bytes]:
    with zipfile.ZipFile(path) as archive:
        parts, _ = GuardedZipReader(archive).read_all()
    return parts


def _set_encrypted_flags(path: Path) -> None:
    data = bytearray(path.read_bytes())
    local_offset = data.index(b"PK\x03\x04")
    central_offset = data.index(b"PK\x01\x02")
    local_flags = struct.unpack_from("<H", data, local_offset + 6)[0]
    central_flags = struct.unpack_from("<H", data, central_offset + 8)[0]
    struct.pack_into("<H", data, local_offset + 6, local_flags | 0x0001)
    struct.pack_into("<H", data, central_offset + 8, central_flags | 0x0001)
    path.write_bytes(data)


def _understate_expanded_size(path: Path, declared_size: int) -> None:
    data = bytearray(path.read_bytes())
    local_offset = data.index(b"PK\x03\x04")
    central_offset = data.index(b"PK\x01\x02")
    struct.pack_into("<I", data, local_offset + 22, declared_size)
    struct.pack_into("<I", data, central_offset + 24, declared_size)
    path.write_bytes(data)


def _understate_stored_member(path: Path, declared_data: bytes) -> None:
    data = bytearray(path.read_bytes())
    local_offset = data.index(b"PK\x03\x04")
    central_offset = data.index(b"PK\x01\x02")
    declared_size = len(declared_data)
    crc = zlib.crc32(declared_data) & 0xFFFFFFFF
    struct.pack_into("<III", data, local_offset + 14, crc, declared_size, declared_size)
    struct.pack_into("<III", data, central_offset + 16, crc, declared_size, declared_size)
    path.write_bytes(data)


class DescribeZipStructureGuard:
    @pytest.mark.parametrize(
        "kind", ["duplicate", "noncanonical", "encrypted", "compression", "symlink"]
    )
    def it_refuses_ambiguous_encrypted_and_unsupported_entries(self, kind: str, tmp_path: Path):
        path = tmp_path / f"{kind}.docx"
        if kind == "duplicate":
            with warnings.catch_warnings():
                warnings.simplefilter("ignore", UserWarning)
                _write_archive(
                    path,
                    (("word/document.xml", b"<one/>"), ("word/document.xml", b"<two/>")),
                )
        elif kind == "noncanonical":
            _write_archive(path, (("../word/document.xml", b"<document/>"),))
        elif kind == "encrypted":
            _write_archive(path, (("word/document.xml", b"<document/>"),))
            _set_encrypted_flags(path)
        elif kind == "compression":
            _write_archive(
                path,
                (("word/document.xml", b"<document/>"),),
                compression=zipfile.ZIP_BZIP2,
            )
        else:
            info = zipfile.ZipInfo("word/document.xml")
            info.external_attr = (stat.S_IFLNK | 0o777) << 16
            with zipfile.ZipFile(path, "w") as archive:
                archive.writestr(info, b"elsewhere.xml")

        with pytest.raises(PackageLimitError):
            _guarded_read(path)

    def it_rejects_case_ambiguous_names(self, tmp_path: Path):
        path = tmp_path / "case-collision.docx"
        _write_archive(
            path,
            (("word/document.xml", b"<one/>"), ("WORD/document.xml", b"<two/>")),
        )
        with pytest.raises(PackageLimitError, match="case-ambiguous"):
            _guarded_read(path)


class DescribeZipInflationHonesty:
    def it_opens_a_highly_compressible_member(self, tmp_path: Path):
        path = tmp_path / "repetitive.docx"
        repetitive = b"<w:p><w:r><w:t>same line</w:t></w:r></w:p>" * 8_192
        _write_archive(
            path,
            (("word/document.xml", repetitive),),
            compression=zipfile.ZIP_DEFLATED,
        )
        with zipfile.ZipFile(path) as archive:
            info = archive.getinfo("word/document.xml")
            assert info.file_size > info.compress_size * 100
        assert _guarded_read(path)["word/document.xml"] == repetitive

    def it_enforces_actual_inflated_bytes_when_metadata_understates_them(
        self, tmp_path: Path
    ):
        path = tmp_path / "forged-size.docx"
        _write_archive(
            path,
            (("word/document.xml", b"0123456789abcdef" * 64),),
            compression=zipfile.ZIP_DEFLATED,
        )
        _understate_expanded_size(path, declared_size=32)

        with pytest.raises(PackageLimitError, match="declared size"):
            _guarded_read(path)

    def it_rejects_undeclared_stored_bytes_even_with_matching_forged_crc(self, tmp_path: Path):
        path = tmp_path / "forged-stored-size.docx"
        _write_archive(path, (("word/media/data.bin", b"abcdefgh"),))
        _understate_stored_member(path, declared_data=b"abcd")

        with pytest.raises(PackageLimitError, match="undeclared trailing data"):
            _guarded_read(path)


class DescribeGuardedPackageApis:
    def it_refuses_open_diff_and_patch_without_touching_the_destination(self, tmp_path: Path):
        unsafe = tmp_path / "unsafe.docx"
        _write_archive(unsafe, (("word/document.xml", b"<document/>"),))
        _set_encrypted_flags(unsafe)
        source = fixture_path(MINIMAL)
        document = docx.Document(str(source))
        out = tmp_path / "out.docx"
        sentinel = b"existing destination"
        out.write_bytes(sentinel)
        out.chmod(0o640)

        assert issubclass(PackageLimitError, PaperRefusal)
        with pytest.raises(PackageLimitError):
            docx.Document(str(unsafe))
        with pytest.raises(PackageLimitError):
            diff_package(source, unsafe)
        with pytest.raises(PackageLimitError):
            patch_save(unsafe, document, out)

        assert out.read_bytes() == sentinel
        assert stat.S_IMODE(out.stat().st_mode) == 0o640
        assert not list(tmp_path.glob("out.docx.*.partial"))

    def it_diagnoses_a_limit_refusal_as_an_unsafe_archive(self, tmp_path: Path):
        unsafe = tmp_path / "unsafe.docx"
        _write_archive(unsafe, (("word/document.xml", b"<document/>"),))
        _set_encrypted_flags(unsafe)

        report = diagnose(unsafe)

        assert not report.readable
        assert report.kind == "unsafe-archive"
        assert "encrypted" in report.problems[0]


class DescribePatchSavePermissions:
    def it_preserves_an_existing_destination_mode(self, tmp_path: Path):
        source = fixture_path(MINIMAL)
        document = docx.Document(str(source))
        document.add_paragraph("A bounded edit.")
        out = tmp_path / "existing.docx"
        out.write_bytes(b"old destination")
        out.chmod(0o640)

        patch_save(source, document, out)

        assert stat.S_IMODE(out.stat().st_mode) == 0o640
        assert docx.Document(str(out)).paragraphs[-1].text == "A bounded edit."

    def it_uses_the_original_mode_for_a_new_destination(self, tmp_path: Path):
        source = tmp_path / "source.docx"
        shutil.copyfile(fixture_path(MINIMAL), source)
        source.chmod(0o604)
        out = tmp_path / "new.docx"

        result = patch_save(source, docx.Document(str(source)), out)

        assert result.verbatim_copy
        assert out.read_bytes() == source.read_bytes()
        assert stat.S_IMODE(out.stat().st_mode) == 0o604

    def it_refuses_if_the_destination_symlink_changes_during_save(
        self, tmp_path: Path, monkeypatch
    ):
        source = fixture_path(MINIMAL)
        document = docx.Document(str(source))
        document.add_paragraph("must not land")
        old_target = tmp_path / "old.docx"
        new_target = tmp_path / "new.docx"
        old_target.write_bytes(b"old target")
        new_target.write_bytes(b"new target")
        link = tmp_path / "link.docx"
        link.symlink_to(old_target.name)
        original_lstat = os.lstat
        original_fsync = os.fsync
        staged = {"done": False}

        def fsync(fd):
            staged["done"] = True
            return original_fsync(fd)

        def lstat(path):
            path_s = os.path.abspath(os.fspath(path))
            if (
                staged["done"]
                and path_s == os.path.abspath(link)
                and os.path.basename(os.readlink(link)) == old_target.name
            ):
                link.unlink()
                link.symlink_to(new_target.name)
            return original_lstat(path)

        monkeypatch.setattr(os, "fsync", fsync)
        monkeypatch.setattr(os, "lstat", lstat)

        with pytest.raises(OSError, match="symlink changed"):
            patch_save(source, document, link)

        assert old_target.read_bytes() == b"old target"
        assert new_target.read_bytes() == b"new target"
        assert not list(tmp_path.glob("*.partial"))

    def it_follows_a_destination_symlink_and_keeps_the_link(self, tmp_path: Path):
        source = fixture_path(MINIMAL)
        document = docx.Document(str(source))
        document.add_paragraph("saved through link")
        target = tmp_path / "target.docx"
        target.write_bytes(b"old")
        target.chmod(0o640)
        link = tmp_path / "link.docx"
        link.symlink_to(target.name)

        patch_save(source, document, link)

        assert link.is_symlink()
        assert stat.S_IMODE(target.stat().st_mode) == 0o640
        assert docx.Document(str(target)).paragraphs[-1].text == "saved through link"


class DescribeGuardedStreamReads:
    def it_accepts_a_healthy_seekable_stream(self):
        data = fixture_path(MINIMAL).read_bytes()
        document = docx.Document(io.BytesIO(data))
        assert document.paragraphs

    def it_accepts_valid_data_descriptors_from_a_nonseekable_writer(self):
        class NonseekableSink(io.BytesIO):
            def seek(self, _offset: int, _whence: int = 0) -> int:
                raise OSError("not seekable")

        sink = NonseekableSink()
        with zipfile.ZipFile(sink, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("word/document.xml", b"<document/>" * 100)

        with zipfile.ZipFile(io.BytesIO(sink.getvalue())) as archive:
            parts, _ = GuardedZipReader(archive).read_all()

        assert parts["word/document.xml"] == b"<document/>" * 100


class DescribeDirectoryEntries:
    """Folder records (zip -r style) are inert in OPC: ignored, never parts."""

    def it_reads_a_package_carrying_inert_directory_entries(self, tmp_path: Path):
        source = fixture_path(MINIMAL)
        path = tmp_path / "rezipped.docx"
        with zipfile.ZipFile(source) as zin, zipfile.ZipFile(
            path, "w", zipfile.ZIP_DEFLATED
        ) as zout:
            added: set[str] = set()
            for name in zin.namelist():
                prefix_parts = name.split("/")[:-1]
                for depth in range(1, len(prefix_parts) + 1):
                    folder = "/".join(prefix_parts[:depth]) + "/"
                    if folder not in added:
                        added.add(folder)
                        zout.writestr(
                            zipfile.ZipInfo(folder, date_time=(1980, 1, 1, 0, 0, 0)),
                            b"",
                        )
                zout.writestr(name, zin.read(name))
        assert added  # the rezip really did interleave folder records

        document = docx.Document(str(path))
        assert document.paragraphs is not None
        parts = _guarded_read(path)
        assert not any(name.endswith("/") for name in parts)

    def and_it_refuses_a_directory_entry_carrying_data(self, tmp_path: Path):
        source = fixture_path(MINIMAL)
        path = tmp_path / "smuggled.docx"
        with zipfile.ZipFile(source) as zin, zipfile.ZipFile(
            path, "w", zipfile.ZIP_DEFLATED
        ) as zout:
            zout.writestr(
                zipfile.ZipInfo("word/", date_time=(1980, 1, 1, 0, 0, 0)),
                b"hidden payload",
            )
            for name in zin.namelist():
                zout.writestr(name, zin.read(name))
        with pytest.raises(PackageLimitError, match="carries data"):
            docx.Document(str(path))

    def and_it_validates_local_records_in_a_directory_only_archive(
        self, tmp_path: Path
    ):
        path = tmp_path / "directory-only.zip"
        with zipfile.ZipFile(path, "w") as archive:
            archive.writestr(zipfile.ZipInfo("word/"), b"")
        data = bytearray(path.read_bytes())
        local_offset = data.index(b"PK\x03\x04")
        data[local_offset : local_offset + 4] = b"BAD!"
        path.write_bytes(data)

        with pytest.raises(PackageLimitError, match="invalid header"):
            _guarded_read(path)
