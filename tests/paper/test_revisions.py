"""Tests for docx.revision / Document.revisions (Phase 8), including the
tracked-edit algebra invariants that cross-check Phases 5, 6 and 8
(CONVENTIONS §4 — the highest-value tests in the repo)."""

from __future__ import annotations

import datetime as dt
import json
from pathlib import Path

import pytest

import docx
from docx.blocks import tracked_delete_paragraphs, tracked_replace_paragraphs
from docx.errors import UnsupportedStructureError
from docx.search import find_one
from docx.story import iter_blocks

from .harness.contract import save_and_reopen
from .harness.paths import fixture_path, sidecar_path

TRACKED = "generated/feature-isolated/tracked-ins-del.docx"
TRACKED_LO = "libreoffice/feature-isolated/tracked-ins-del.docx"
FRAGMENTED = "generated/feature-isolated/fragmented-runs.docx"
MINIMAL = "generated/minimal-clean/minimal.docx"

FROZEN = dt.datetime(2026, 7, 7, 12, 0, 0, tzinfo=dt.timezone.utc)


def _doc(relpath: str):
    return docx.Document(str(fixture_path(relpath)))


def _texts(document, view: str = "current"):
    return [b.text for b in iter_blocks(document, view=view)]


def _ground_truth(relpath: str) -> dict:
    return json.loads(sidecar_path(fixture_path(relpath)).read_text(encoding="utf-8"))[
        "ground_truth"
    ]


class DescribeEnumeration:
    @pytest.mark.parametrize("relpath", [TRACKED, TRACKED_LO])
    def it_lists_every_revision_with_metadata(self, relpath: str):
        truth = _ground_truth(relpath)
        revisions = _doc(relpath).revisions
        insertions = [r for r in revisions if r.revision_type == "insertion"]
        deletions = [r for r in revisions if r.revision_type == "deletion"]
        assert [r.author for r in insertions] == truth["tracked_insertions"]["authors"]
        assert [r.text for r in insertions] == truth["tracked_insertions"]["inserted_texts"]
        assert [r.text for r in deletions] == truth["tracked_deletions"]["deleted_texts"]
        assert [r.author for r in deletions] == truth["tracked_deletions"]["authors"]

    def it_parses_dates(self):
        revisions = _doc(TRACKED).revisions
        dates = sorted(r.date for r in revisions if r.date)
        assert dates[0] == dt.datetime(2026, 6, 1, 9, 30, tzinfo=dt.timezone.utc)

    def it_carries_block_anchors(self):
        revision = _doc(TRACKED).revisions[0]
        assert revision.anchor.story == "word/document.xml"
        assert revision.story == "word/document.xml"

    def it_serializes_deterministically(self):
        payload_1 = json.dumps(_doc(TRACKED).revisions.to_dict())
        payload_2 = json.dumps(_doc(TRACKED).revisions.to_dict())
        assert payload_1 == payload_2
        parsed = json.loads(payload_1)
        assert parsed["schema"] == "paper_revisions" and parsed["version"] == 2
        assert parsed["remaining_unsupported"] == {}


class DescribeAcceptReject:
    def it_accepts_all_revisions(self, tmp_path: Path):
        truth = _ground_truth(TRACKED)
        document = _doc(TRACKED)
        count = document.revisions.accept_all()
        assert count == 3
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        texts = _texts(reopened)
        assert truth["revised_paragraph"]["visible_text_current"] in texts
        assert "This whole sentence was inserted with tracking on." in texts
        assert not reopened.revisions, "accepting all must leave zero revisions"

    def it_rejects_all_revisions(self, tmp_path: Path):
        truth = _ground_truth(TRACKED)
        document = _doc(TRACKED)
        count = document.revisions.reject_all()
        assert count == 3
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        texts = _texts(reopened)
        assert truth["revised_paragraph"]["visible_text_original"] in texts
        assert "This whole sentence was inserted with tracking on." not in texts

    def it_filters_by_author(self):
        document = _doc(TRACKED)
        count = document.revisions.accept_all(author="Alice Editor")
        assert count == 1  # only Alice's insertion
        remaining = document.revisions
        assert {r.author for r in remaining} == {"Bob Reviewer"}

    def it_resolves_individual_revisions(self):
        document = _doc(TRACKED)
        deletion = next(r for r in document.revisions if r.revision_type == "deletion")
        deletion.reject()
        assert "forty-two" in _texts(document)[1]


class DescribeTrackedEditAlgebra:
    """accept(tracked X) ≡ plain X; reject(tracked X) ≡ original."""

    def it_makes_accepted_span_replace_equal_plain_replace(self, tmp_path: Path):
        tracked_doc = _doc(FRAGMENTED)
        find_one(tracked_doc, "$75-100/hr").replace(
            "$85–110/hr", tracked=True, author="Carol QA", date=FROZEN
        )
        tracked_doc.revisions.accept_all()

        plain_doc = _doc(FRAGMENTED)
        find_one(plain_doc, "$75-100/hr").replace("$85–110/hr")

        tracked_out = save_and_reopen(tracked_doc, tmp_path / "tracked.docx")
        plain_out = save_and_reopen(plain_doc, tmp_path / "plain.docx")
        assert _texts(tracked_out) == _texts(plain_out)

    def it_makes_rejected_span_replace_equal_the_original(self, tmp_path: Path):
        document = _doc(FRAGMENTED)
        pristine_texts = _texts(document)
        find_one(document, "$75-100/hr").replace(
            "$85–110/hr", tracked=True, author="Carol QA", date=FROZEN
        )
        document.revisions.reject_all()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _texts(reopened) == pristine_texts
        bold_text = "".join(
            r.text for r in reopened.paragraphs[0].runs if r.bold
        )
        assert bold_text == "$75–100/hr", "formatting must survive the reject"

    def it_makes_accepted_paragraph_delete_equal_removal(self, tmp_path: Path):
        document = _doc(MINIMAL)
        tracked_delete_paragraphs(
            document, "First body paragraph", count=1, author="Carol QA", date=FROZEN
        )
        document.revisions.accept_all()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _texts(reopened) == [
            "Minimal Clean Document",
            "Second body paragraph, equally unremarkable.",
        ]

    def it_makes_rejected_paragraph_delete_equal_the_original(self, tmp_path: Path):
        document = _doc(MINIMAL)
        pristine_texts = _texts(document)
        tracked_delete_paragraphs(
            document, "First body paragraph", count=2, author="Carol QA", date=FROZEN
        )
        document.revisions.reject_all()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _texts(reopened) == pristine_texts

    def it_makes_rejected_paragraph_replace_equal_the_original(self, tmp_path: Path):
        document = _doc(MINIMAL)
        pristine_texts = _texts(document)
        tracked_replace_paragraphs(
            document, "Second body paragraph", ["Replacement."],
            author="Carol QA", date=FROZEN,
        )
        document.revisions.reject_all()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _texts(reopened) == pristine_texts

    def it_makes_accepted_paragraph_replace_equal_the_replacement(self, tmp_path: Path):
        document = _doc(MINIMAL)
        tracked_replace_paragraphs(
            document, "Second body paragraph", ["Replacement one.", "Replacement two."],
            author="Carol QA", date=FROZEN,
        )
        document.revisions.accept_all()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _texts(reopened) == [
            "Minimal Clean Document",
            "First body paragraph with perfectly ordinary text.",
            "Replacement one.",
            "Replacement two.",
        ]


class DescribeUnresolvableRevisions:
    """v0.1 honesty recall: moves and format changes are seen, counted, and
    refused — never half-resolved, never silently omitted."""

    MOVES = "generated/feature-isolated/tracked-moves.docx"
    FORMAT_CHANGES = "generated/feature-isolated/format-changes.docx"

    def it_reports_a_census_of_unsupported_revisions(self):
        revisions = _doc(self.MOVES).revisions
        assert revisions.remaining_unsupported() == {"move_from": 1, "move_to": 1}
        assert _doc(self.FORMAT_CHANGES).revisions.remaining_unsupported() == {
            "format_change": 2
        }

    def it_refuses_individual_resolution_of_a_move(self):
        revision = next(
            r for r in _doc(self.MOVES).revisions if r.revision_type == "move_from"
        )
        with pytest.raises(UnsupportedStructureError, match="not yet"):
            revision.accept()

    def it_resolves_an_author_filtered_clean_subset_alongside_moves(self, tmp_path):
        """Selected-set semantics: Carol's plain insertion resolves even though
        Alice's move stays pending — the census carries the rest."""
        import shutil

        path = tmp_path / "mixed.docx"
        shutil.copyfile(fixture_path(self.MOVES), path)
        document = docx.Document(str(path))
        find_one(document, "Paragraph before the tracked move.").replace(
            "Paragraph ahead of the tracked move.",
            tracked=True, author="Carol QA",
        )
        resolved = document.revisions.accept_all(author="Carol QA")
        assert resolved > 0
        assert document.revisions.remaining_unsupported() == {
            "move_from": 1, "move_to": 1
        }
        with pytest.raises(UnsupportedStructureError):
            document.revisions.accept_all()
