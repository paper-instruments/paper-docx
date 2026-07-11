"""Regressions for compound revision resolution and refusal atomicity."""

from __future__ import annotations

import copy

import pytest

import docx
from docx.errors import UnsupportedStructureError
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

from .harness.paths import fixture_path

NOTES = "generated/feature-isolated/footnotes-endnotes.docx"
ROW_REVISIONS = "generated/feature-isolated/row-revisions.docx"
W = nsdecls("w")


def _revision(tag: str, inner_xml: str, revision_id: int, author: str = "Audit"):
    return parse_xml(
        f'<w:{tag} {W} w:id="{revision_id}" w:author="{author}">'
        f"{inner_xml}</w:{tag}>"
    )


def _wrap_in_insertion(element, revision_id: int) -> None:
    wrapper = _revision("ins", "", revision_id)
    element.addprevious(wrapper)
    wrapper.append(element)


def _relationship_snapshot(document) -> tuple:
    return tuple(
        sorted(
            (r_id, rel.reltype, rel.target_ref)
            for r_id, rel in document.part.rels.items()
        )
    )


def _opaque_part(document, partname: str, content_type: str, blob: bytes) -> Part:
    package = document.part.package
    assert package is not None
    return Part(PackURI(partname), content_type, blob, package)


def _footnote_ids(document) -> set[int]:
    root = document.part.part_related_by(RT.FOOTNOTES)._element  # noqa: SLF001
    return {
        int(note.get(qn("w:id")))
        for note in root
        if note.tag == qn("w:footnote")
        and note.get(qn("w:type")) is None
        and note.get(qn("w:id")) is not None
    }


class DescribeCleanupPreflight:
    def it_refuses_duplicate_note_relationships_before_mutation(self):
        document = docx.Document(str(fixture_path(NOTES)))
        reference = next(document.element.iter(qn("w:footnoteReference")))
        _wrap_in_insertion(reference.getparent(), 10)
        duplicate = _opaque_part(
            document,
            "/word/footnotes-audit.xml",
            CT.WML_FOOTNOTES,
            b"<audit/>",
        )
        document.part.relate_to(duplicate, RT.FOOTNOTES)
        before_xml = document.element.xml
        before_rels = _relationship_snapshot(document)

        with pytest.raises(UnsupportedStructureError, match="multiple.*footnotes"):
            document.revisions.reject_all()

        assert document.element.xml == before_xml
        assert _relationship_snapshot(document) == before_rels

    def it_refuses_an_opaque_note_part_before_mutation(self):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        paragraph.append(
            _revision(
                "ins",
                '<w:r><w:footnoteReference w:id="1"/></w:r>',
                11,
            )
        )
        opaque = _opaque_part(
            document,
            "/word/footnotes.xml",
            CT.WML_FOOTNOTES,
            b"not live XML",
        )
        document.part.relate_to(opaque, RT.FOOTNOTES)
        before_xml = document.element.xml
        before_rels = _relationship_snapshot(document)

        with pytest.raises(UnsupportedStructureError, match="footnotes.*live XML"):
            document.revisions.reject_all()

        assert document.element.xml == before_xml
        assert _relationship_snapshot(document) == before_rels

    def it_refuses_duplicate_comment_relationships_before_mutation(self):
        document = docx.Document()
        paragraph = document.add_paragraph("commented")
        document.add_comment(paragraph.runs, text="comment", author="Audit")
        reference = next(document.element.iter(qn("w:commentReference")))
        _wrap_in_insertion(reference.getparent(), 12)
        duplicate = _opaque_part(
            document,
            "/word/comments-audit.xml",
            CT.WML_COMMENTS,
            b"<audit/>",
        )
        document.part.relate_to(duplicate, RT.COMMENTS)
        before_xml = document.element.xml
        before_rels = _relationship_snapshot(document)

        with pytest.raises(UnsupportedStructureError, match="multiple.*comments"):
            document.revisions.reject_all()

        assert document.element.xml == before_xml
        assert _relationship_snapshot(document) == before_rels

    def it_refuses_an_opaque_comment_part_before_mutation(self):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        paragraph.append(
            _revision(
                "ins",
                '<w:r><w:commentReference w:id="1"/></w:r>',
                13,
            )
        )
        opaque = _opaque_part(
            document,
            "/word/comments.xml",
            CT.WML_COMMENTS,
            b"not live XML",
        )
        document.part.relate_to(opaque, RT.COMMENTS)
        before_xml = document.element.xml
        before_rels = _relationship_snapshot(document)

        with pytest.raises(UnsupportedStructureError, match="comments.*live XML"):
            document.revisions.reject_all()

        assert document.element.xml == before_xml
        assert _relationship_snapshot(document) == before_rels

    @pytest.mark.parametrize("bad_id", ["not-a-number", ""])
    def it_refuses_malformed_note_reference_ids_before_any_mutation(self, bad_id: str):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        paragraph.append(
            _revision("ins", "<w:r><w:t>first</w:t></w:r>", 14)
        )
        paragraph.append(
            _revision(
                "ins",
                f'<w:r><w:footnoteReference w:id="{bad_id}"/></w:r>',
                15,
            )
        )
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match=r"malformed.*w:id"):
            document.revisions.reject_all()

        assert document.element.xml == before

    def it_compares_equivalent_note_ids_numerically(self):
        document = docx.Document(str(fixture_path(NOTES)))
        reference = next(document.element.iter(qn("w:footnoteReference")))
        note_id = int(reference.get(qn("w:id")))
        _wrap_in_insertion(reference.getparent(), 16)
        surviving_run = document.add_paragraph().add_run()._r
        surviving_run.append(
            parse_xml(
                f'<w:footnoteReference {W} w:id="{note_id:02d}"/>'
            )
        )

        assert document.revisions.reject_all() == 1
        assert note_id in _footnote_ids(document)

    def it_refuses_malformed_note_body_ids_before_mutation(self):
        document = docx.Document(str(fixture_path(NOTES)))
        reference = next(document.element.iter(qn("w:footnoteReference")))
        _wrap_in_insertion(reference.getparent(), 17)
        notes_root = document.part.part_related_by(RT.FOOTNOTES)._element  # noqa: SLF001
        note = next(
            item
            for item in notes_root
            if item.tag == qn("w:footnote")
            and item.get(qn("w:type")) is None
        )
        note.set(qn("w:id"), "broken")
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match=r"malformed.*w:id"):
            document.revisions.reject_all()

        assert document.element.xml == before

    def it_refuses_removing_one_of_multiple_comment_references(self):
        document = docx.Document()
        paragraph = document.add_paragraph("commented")
        comment = document.add_comment(
            paragraph.runs, text="comment", author="Audit"
        )
        reference = next(document.element.iter(qn("w:commentReference")))
        _wrap_in_insertion(reference.getparent(), 18)
        duplicate_run = document.add_paragraph().add_run()._r
        duplicate_run.append(copy.deepcopy(reference))
        comments_root = document.part.part_related_by(RT.COMMENTS)._element  # noqa: SLF001
        before_document = document.element.xml
        before_comments = comments_root.xml

        with pytest.raises(
            UnsupportedStructureError, match="reference marks|exactly one"
        ):
            document.revisions.reject_all()

        assert document.element.xml == before_document
        assert comments_root.xml == before_comments
        assert document.comments.get(comment.comment_id) is not None

    def it_refuses_removing_only_one_comment_range_boundary(self):
        document = docx.Document()
        paragraph = document.add_paragraph("commented")
        document.add_comment(paragraph.runs, text="comment", author="Audit")
        range_start = next(document.element.iter(qn("w:commentRangeStart")))
        _wrap_in_insertion(range_start, 19)
        before = document.element.xml

        with pytest.raises(
            UnsupportedStructureError, match="range markers.*reference"
        ):
            document.revisions.reject_all()

        assert document.element.xml == before


class DescribeIndividualDestructiveClosure:
    def it_refuses_a_row_change_with_independently_authored_nested_markup(self):
        document = docx.Document(str(fixture_path(ROW_REVISIONS)))
        row_marker = next(
            revision
            for revision in document.revisions
            if revision.revision_type == "row_insertion"
        )._element  # noqa: SLF001
        tr_pr = row_marker.getparent()
        assert tr_pr is not None
        row = tr_pr.getparent()
        assert row is not None
        nested = next(node for node in row.iter(qn("w:ins")) if node is not row_marker)
        nested.set(qn("w:author"), "Independent Reviewer")
        row_insertion = next(
            revision
            for revision in document.revisions
            if revision.revision_type == "row_insertion"
        )
        before = document.element.xml

        with pytest.raises(
            UnsupportedStructureError,
            match=r"cannot reject.*row_insertion.*Independent Reviewer",
        ):
            row_insertion.reject()

        assert document.element.xml == before

    def it_refuses_accepting_an_outer_deletion_with_a_nested_revision(self):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        paragraph.append(
            _revision(
                "del",
                '<w:r><w:delText>outer</w:delText></w:r>'
                '<w:ins w:id="21" w:author="Bob">'
                "<w:r><w:t>nested</w:t></w:r></w:ins>",
                20,
                author="Alice",
            )
        )
        revisions = document.revisions
        outer = next(r for r in revisions if r.revision_type == "deletion")
        before = document.element.xml

        with pytest.raises(
            UnsupportedStructureError,
            match=r"cannot accept.*deletion.*unselected.*insertion",
        ):
            outer.accept()

        assert document.element.xml == before
        assert [
            (revision.revision_type, revision.author)
            for revision in document.revisions
        ] == [("deletion", "Alice"), ("insertion", "Bob")]
        assert document.revisions.accept_all() == 2
        assert not document.revisions

    def it_refuses_rejecting_an_outer_insertion_with_a_nested_revision(self):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        paragraph.append(
            _revision(
                "ins",
                '<w:r><w:t>outer</w:t></w:r>'
                '<w:del w:id="23" w:author="Bob">'
                "<w:r><w:delText>nested</w:delText></w:r></w:del>",
                22,
                author="Alice",
            )
        )
        revisions = document.revisions
        outer = next(r for r in revisions if r.revision_type == "insertion")
        before = document.element.xml

        with pytest.raises(
            UnsupportedStructureError,
            match=r"cannot reject.*insertion.*unselected.*deletion",
        ):
            outer.reject()

        assert document.element.xml == before
        assert [
            (revision.revision_type, revision.author)
            for revision in document.revisions
        ] == [("insertion", "Alice"), ("deletion", "Bob")]
        assert document.revisions.reject_all() == 2
        assert not document.revisions


class DescribeFilteredDestructiveClosure:
    def it_refuses_an_outer_revision_that_would_consume_another_author(self):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        paragraph.append(
            _revision(
                "del",
                '<w:r><w:delText>outer</w:delText></w:r>'
                '<w:ins w:id="21" w:author="Bob">'
                "<w:r><w:t>nested</w:t></w:r></w:ins>",
                20,
                author="Alice",
            )
        )
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match="Bob|another author"):
            document.revisions.accept_all(author="Alice")

        assert document.element.xml == before
        assert {revision.author for revision in document.revisions} == {
            "Alice",
            "Bob",
        }

    def it_refuses_a_row_removal_that_would_consume_another_author(self):
        document = docx.Document()
        table = parse_xml(
            f"<w:tbl {W}><w:tblPr/><w:tblGrid><w:gridCol/></w:tblGrid>"
            "<w:tr><w:tc><w:tcPr/><w:p/></w:tc></w:tr>"
            "<w:tr><w:trPr><w:del w:id='30' w:author='Alice'/></w:trPr>"
            "<w:tc><w:tcPr/><w:p>"
            "<w:ins w:id='31' w:author='Bob'><w:r><w:t>nested</w:t></w:r></w:ins>"
            "</w:p></w:tc></w:tr></w:tbl>"
        )
        document.element.body.insert(0, table)
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match="Bob|another author"):
            document.revisions.accept_all(author="Alice")

        assert document.element.xml == before
        assert len(table.findall(qn("w:tr"))) == 2

    def it_refuses_a_paragraph_join_that_would_consume_another_author(self):
        document = docx.Document()
        first = parse_xml(
            f"<w:p {W}><w:pPr><w:rPr>"
            "<w:del w:id='32' w:author='Alice'/>"
            "<w:ins w:id='33' w:author='Bob'/>"
            "</w:rPr></w:pPr><w:r><w:t>first</w:t></w:r></w:p>"
        )
        following = parse_xml(
            f"<w:p {W}><w:r><w:t>following</w:t></w:r></w:p>"
        )
        document.element.body.insert(0, following)
        document.element.body.insert(0, first)
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match="Bob|another author"):
            document.revisions.accept_all(author="Alice")

        assert document.element.xml == before

    def it_refuses_row_removal_that_would_consume_marker_only_move_markup(self):
        document = docx.Document()
        table = parse_xml(
            f"<w:tbl {W}><w:tblPr/><w:tblGrid><w:gridCol/></w:tblGrid>"
            "<w:tr><w:tc><w:tcPr/><w:p/></w:tc></w:tr>"
            "<w:tr><w:trPr><w:del w:id='34' w:author='Alice'/></w:trPr>"
            "<w:tc><w:tcPr/><w:p>"
            "<w:moveFromRangeStart w:id='35' w:name='nested' w:author='Bob'/>"
            "<w:moveFromRangeEnd w:id='35'/>"
            "<w:moveToRangeStart w:id='36' w:name='nested' w:author='Bob'/>"
            "<w:moveToRangeEnd w:id='36'/>"
            "</w:p></w:tc></w:tr></w:tbl>"
        )
        document.element.body.insert(0, table)
        before = document.element.xml

        with pytest.raises(
            UnsupportedStructureError, match="unselected revision markup"
        ):
            document.revisions.accept_all(author="Alice")

        assert document.element.xml == before
        assert len(table.findall(qn("w:tr"))) == 2


class DescribeMoveAndParagraphCompounds:
    @pytest.mark.parametrize("method", ["accept_all", "reject_all"])
    def it_refuses_a_complete_marker_only_move_unit(self, method: str):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        for tag, revision_id in (
            ("moveFromRangeStart", 40),
            ("moveFromRangeEnd", 40),
            ("moveToRangeStart", 41),
            ("moveToRangeEnd", 41),
        ):
            name = ' w:name="marker-only"' if tag.endswith("Start") else ""
            paragraph.append(
                parse_xml(f'<w:{tag} {W} w:id="{revision_id}"{name}/>')
            )
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match="marker-only|resolvable"):
            getattr(document.revisions, method)()

        assert document.element.xml == before

    @pytest.mark.parametrize(
        "barrier",
        [
            f"<w:customXml {W}><w:p><w:r><w:t>middle</w:t></w:r></w:p></w:customXml>",
            f'<w:altChunk {nsdecls("w", "r")} r:id="rIdMissing"/>',
        ],
        ids=["custom-xml", "alt-chunk"],
    )
    def it_refuses_a_paragraph_join_across_unrecognized_block_content(
        self, barrier: str
    ):
        document = docx.Document()
        first = parse_xml(
            f"<w:p {W}><w:pPr><w:rPr>"
            "<w:del w:id='50' w:author='Audit'/>"
            "</w:rPr></w:pPr><w:r><w:t>first</w:t></w:r></w:p>"
        )
        following = parse_xml(
            f"<w:p {W}><w:r><w:t>following</w:t></w:r></w:p>"
        )
        document.element.body.insert(0, following)
        document.element.body.insert(0, parse_xml(barrier))
        document.element.body.insert(0, first)
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match="paragraph|block"):
            document.revisions.accept_all()

        assert document.element.xml == before
