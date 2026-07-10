"""The effective-format resolver.

Read-only, provenance-bearing: every value names the layer it came from,
toggles XOR through style layers (the nested-bold gotcha), and what cannot be
resolved is declared, never guessed.
"""

from __future__ import annotations

import pytest

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.formatting import format_of, surrounding_format
from docx.search import find_one
from docx.shared import Pt

from .harness.paths import fixture_path

MINIMAL = "generated/minimal-clean/minimal.docx"


def _doc():
    return docx.Document(str(fixture_path(MINIMAL)))


class DescribeDirectFormatting:
    def it_resolves_direct_run_formatting_with_provenance(self):
        document = _doc()
        run = document.add_paragraph().add_run("bolded")
        run.bold = True
        run.font.size = Pt(14)
        resolved = format_of(run)
        assert resolved["bold"].value is True
        assert resolved["bold"].source == "direct"
        assert resolved["size_pt"].value == 14
        assert resolved["size_pt"].source == "direct"

    def it_reports_unspecified_toggles_as_false_from_none(self):
        document = _doc()
        run = document.add_paragraph().add_run("plain")
        resolved = format_of(run)
        assert resolved["italic"].value is False
        assert resolved["italic"].source == "none"

    def it_reads_doc_defaults(self):
        document = _doc()
        run = document.add_paragraph().add_run("default-sized")
        resolved = format_of(run)
        assert resolved["size_pt"].value == 11
        assert resolved["size_pt"].source == "doc_defaults"
        # the template's docDefaults use a THEME font reference; the honest
        # answer is the token, never a guessed literal
        assert resolved["font_name"].value == "theme:minorHAnsi"
        assert "theme_font_resolution" in resolved.unresolved

    def it_declares_what_it_cannot_resolve(self):
        document = _doc()
        resolved = format_of(document.paragraphs[1].runs[0])
        assert "table_style_conditional_formatting" in resolved.unresolved
        payload = resolved.to_dict()
        assert payload["schema"] == "paper_effective_format"


class DescribeStyleChainResolution:
    def it_resolves_through_the_paragraph_style_chain(self):
        document = _doc()
        base = document.styles.add_style("ChainBase", WD_STYLE_TYPE.PARAGRAPH)
        base.font.size = Pt(16)
        leaf = document.styles.add_style("ChainLeaf", WD_STYLE_TYPE.PARAGRAPH)
        leaf.base_style = base
        paragraph = document.add_paragraph("chained text", style="ChainLeaf")
        resolved = format_of(paragraph.runs[0])
        assert resolved["size_pt"].value == 16
        assert resolved["size_pt"].source == "paragraph_style:ChainBase"

    def it_lets_the_nearer_layer_win_for_ordinary_properties(self):
        document = _doc()
        base = document.styles.add_style("SizedBase", WD_STYLE_TYPE.PARAGRAPH)
        base.font.size = Pt(16)
        leaf = document.styles.add_style("SizedLeaf", WD_STYLE_TYPE.PARAGRAPH)
        leaf.base_style = base
        leaf.font.size = Pt(9)
        paragraph = document.add_paragraph("near wins", style="SizedLeaf")
        resolved = format_of(paragraph.runs[0])
        assert resolved["size_pt"].value == 9
        assert resolved["size_pt"].source == "paragraph_style:SizedLeaf"

    def it_resolves_character_styles_above_paragraph_styles(self):
        document = _doc()
        para_style = document.styles.add_style("ColoredPara", WD_STYLE_TYPE.PARAGRAPH)
        para_style.font.size = Pt(15)
        char_style = document.styles.add_style("Emph", WD_STYLE_TYPE.CHARACTER)
        char_style.font.size = Pt(8)
        paragraph = document.add_paragraph(style="ColoredPara")
        run = paragraph.add_run("char styled")
        run.style = char_style
        resolved = format_of(run)
        assert resolved["size_pt"].value == 8
        assert resolved["size_pt"].source == "character_style:Emph"

    def it_survives_a_based_on_cycle(self):
        from docx.oxml.parser import parse_xml
        from docx.oxml.ns import nsdecls

        document = _doc()
        styles = document.styles.element
        w = nsdecls("w")
        styles.append(parse_xml(
            f'<w:style {w} w:type="paragraph" w:styleId="CycleA">'
            '<w:name w:val="Cycle A"/><w:basedOn w:val="CycleB"/></w:style>'
        ))
        styles.append(parse_xml(
            f'<w:style {w} w:type="paragraph" w:styleId="CycleB">'
            '<w:name w:val="Cycle B"/><w:basedOn w:val="CycleA"/>'
            "<w:rPr><w:sz w:val=\"40\"/></w:rPr></w:style>"
        ))
        paragraph = document.add_paragraph("cyclic")
        paragraph.style = document.styles["Cycle A"]
        resolved = format_of(paragraph.runs[0])  # must not recurse forever
        assert resolved["size_pt"].value == 20


class DescribeToggleSemantics:
    """The famous gotcha: nested toggles XOR through style layers."""

    def it_xors_bold_when_paragraph_and_character_styles_both_set_it(self):
        document = _doc()
        para_style = document.styles.add_style("BoldPara", WD_STYLE_TYPE.PARAGRAPH)
        para_style.font.bold = True
        char_style = document.styles.add_style("BoldChar", WD_STYLE_TYPE.CHARACTER)
        char_style.font.bold = True
        paragraph = document.add_paragraph(style="BoldPara")
        run = paragraph.add_run("double bold cancels")
        run.style = char_style
        resolved = format_of(run)
        assert resolved["bold"].value is False  # True XOR True
        assert resolved["bold"].source == "toggle_xor"
        assert resolved["bold"].chain == (
            "paragraph_style:BoldPara",
            "character_style:BoldChar",
        )

    def it_keeps_bold_from_a_single_style_layer(self):
        document = _doc()
        para_style = document.styles.add_style("JustBold", WD_STYLE_TYPE.PARAGRAPH)
        para_style.font.bold = True
        paragraph = document.add_paragraph("single layer", style="JustBold")
        resolved = format_of(paragraph.runs[0])
        assert resolved["bold"].value is True
        assert resolved["bold"].source == "paragraph_style:JustBold"

    def it_lets_direct_formatting_override_toggles_absolutely(self):
        document = _doc()
        para_style = document.styles.add_style("BoldBase", WD_STYLE_TYPE.PARAGRAPH)
        para_style.font.bold = True
        paragraph = document.add_paragraph(style="BoldBase")
        run = paragraph.add_run("explicitly unbold")
        run.bold = False
        resolved = format_of(run)
        assert resolved["bold"].value is False
        assert resolved["bold"].source == "direct"


class DescribeParagraphAndSpanTargets:
    def it_resolves_paragraph_alignment_with_provenance(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        document = _doc()
        paragraph = document.add_paragraph("centered")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        resolved = format_of(paragraph)
        assert resolved["alignment"].value == "center"
        assert resolved["alignment"].source == "direct"
        assert resolved["style_name"].value == "Normal"

    def it_resolves_a_span_and_reports_disagreement_as_mixed(self):
        document = _doc()
        paragraph = document.add_paragraph()
        paragraph.add_run("bold half").bold = True
        paragraph.add_run(" plain half")
        span = find_one(document, "bold half plain half")
        resolved = format_of(span)
        assert resolved["bold"].value is None
        assert resolved["bold"].source == "mixed"
        assert resolved["size_pt"].value == 11  # layers agree

    def it_rejects_unsupported_targets(self):
        with pytest.raises(TypeError, match="format_of"):
            format_of("just a string")


class DescribeSurroundingFormat:
    def it_reports_the_anchor_neighborhood_for_insertions(self):
        document = _doc()
        resolved = surrounding_format(document, "Minimal Clean Document")
        assert resolved["style_name"].value == "Heading 1"
        # Heading 1 in the default template resolves sz through its chain
        assert resolved["size_pt"].value is not None
