"""Audit regressions for XML text and comment refusal atomicity."""

from __future__ import annotations

import io

import pytest

import docx
from docx.commentops import reply, resolve
from docx.errors import PaperRefusal
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml.ns import qn
from docx.search import find_one


def _package_bytes(document) -> bytes:
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


class DescribeXmlCharacterValidation:
    @pytest.mark.parametrize("tracked", [False, True])
    def it_refuses_xml_illegal_replacement_before_mutation(self, tracked: bool):
        document = docx.Document()
        document.add_paragraph("replace target")
        span = find_one(document, "replace target")
        before = _package_bytes(document)

        with pytest.raises(ValueError, match="XML 1.0"):
            span.replace(
                "invalid\x00text",
                tracked=tracked,
                author="Reviewer" if tracked else None,
            )

        assert _package_bytes(document) == before

    def it_refuses_invalid_comment_text_before_splitting_runs(self):
        document = docx.Document()
        document.add_paragraph("prefix target suffix")
        span = find_one(document, "target")
        before = _package_bytes(document)

        with pytest.raises(ValueError, match="XML 1.0"):
            span.comment("invalid\x00comment", author="Reviewer")

        assert _package_bytes(document) == before

    def it_refuses_invalid_reply_text_before_thread_metadata_changes(self):
        document = docx.Document()
        document.add_paragraph("comment target")
        comment = find_one(document, "comment target").comment(
            "Parent comment", author="Reviewer"
        )
        before = _package_bytes(document)

        with pytest.raises(ValueError, match="XML 1.0"):
            reply(document, comment, "invalid\x00reply", author="Second reviewer")

        assert _package_bytes(document) == before

    def it_refuses_malformed_existing_comment_ids_before_splitting_runs(self):
        document = docx.Document()
        document.add_paragraph("existing comment")
        existing = find_one(document, "existing comment").comment(
            "Existing", author="Reviewer"
        )
        existing._comment_elm.set(qn("w:id"), "broken")  # noqa: SLF001
        document.add_paragraph("prefix target suffix")
        span = find_one(document, "target")
        before = _package_bytes(document)

        with pytest.raises(PaperRefusal, match="malformed comment"):
            span.comment("New", author="Reviewer")

        assert _package_bytes(document) == before

    def it_refuses_malformed_unrelated_comment_ids_before_reply_metadata(self):
        document = docx.Document()
        document.add_paragraph("parent")
        parent = find_one(document, "parent").comment("Parent", author="Reviewer")
        document.add_paragraph("other")
        other = find_one(document, "other").comment("Other", author="Reviewer")
        other._comment_elm.set(qn("w:id"), "broken")  # noqa: SLF001
        before = _package_bytes(document)

        with pytest.raises(PaperRefusal, match="malformed comment"):
            reply(document, parent, "Reply", author="Second reviewer")

        assert _package_bytes(document) == before

    def it_refuses_a_colliding_comments_extended_part_before_resolution(self):
        document = docx.Document()
        document.add_paragraph("comment target")
        comment = find_one(document, "comment target").comment(
            "Comment", author="Reviewer"
        )
        package = document.part.package
        assert package is not None
        collision = Part(
            PackURI("/word/commentsExtended.xml"),
            "application/octet-stream",
            b"collision",
            package,
        )
        document.part.relate_to(collision, "urn:paper-docx:audit-collision")
        before = _package_bytes(document)

        with pytest.raises(PaperRefusal, match="already occupies"):
            resolve(document, comment)

        assert _package_bytes(document) == before


class DescribeTrackedIdentityPreflight:
    """w:author/w:date are stamped after mutation begins, so malformed
    identity must refuse while nothing has changed."""

    def it_refuses_a_malformed_tracked_author_before_mutating(self):
        document = docx.Document()
        document.add_paragraph("replace target here")
        span = find_one(document, "target")
        before = _package_bytes(document)
        with pytest.raises(ValueError, match="author"):
            span.replace("changed", tracked=True, author="bad\x00author")
        assert _package_bytes(document) == before

    def and_it_refuses_a_string_date_before_mutating(self):
        document = docx.Document()
        document.add_paragraph("replace target here")
        span = find_one(document, "target")
        before = _package_bytes(document)
        with pytest.raises(TypeError, match="datetime"):
            span.replace("changed", tracked=True, author="A", date="2026-01-01")
        assert _package_bytes(document) == before
