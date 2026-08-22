"""Regression tests for previously fixed defects — one test (at least) per
defect, named for the failure mode."""

from __future__ import annotations

import datetime as dt
from pathlib import Path

import pytest
from lxml import etree

import docx
from docx.errors import TargetNotFoundError, UnsupportedStructureError
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.package import diff_package, xml_equivalent
from docx.search import find_one, find_text
from docx.story import iter_blocks

from .harness.paths import fixture_path

FROZEN = dt.datetime(2026, 7, 7, 12, 0, 0, tzinfo=dt.timezone.utc)
MINIMAL = "generated/minimal-clean/minimal.docx"
W = nsdecls("w")


def _doc(relpath: str = MINIMAL):
    return docx.Document(str(fixture_path(relpath)))


class DescribeKernelComparison:
    def it_distinguishes_processing_instruction_targets(self):
        a = b'<a><?foo x="1"?></a>'
        b = b'<a><?bar x="1"?></a>'
        assert not xml_equivalent(a, b)
        assert xml_equivalent(a, a)

    def it_compares_prolog_comments_and_pis(self):
        with_pi = b'<?mso-application progid="Word.Document"?><a/>'
        assert not xml_equivalent(with_pi, b"<a/>")
        assert not xml_equivalent(b"<!--x--><a/>", b"<a/>")
        assert xml_equivalent(with_pi, with_pi)

    def it_refuses_dtd_bearing_parts_loudly(self):
        from docx._paperpkg import UnsupportedXmlError

        safe = b'<!DOCTYPE a [<!ENTITY e "SAFE">]><a>&e;</a>'
        evil = b'<!DOCTYPE a [<!ENTITY e "EVIL">]><a>&e;</a>'
        with pytest.raises(UnsupportedXmlError, match="DTD"):
            xml_equivalent(safe, evil)

    def it_treats_changed_dtd_parts_as_semantic_change_in_diff(self, tmp_path: Path):
        import zipfile

        source = fixture_path(MINIMAL)
        for label, payload in (
            ("a", b'<!DOCTYPE x [<!ENTITY e "SAFE">]><x>&e;</x>'),
            ("b", b'<!DOCTYPE x [<!ENTITY e "EVIL">]><x>&e;</x>'),
        ):
            with zipfile.ZipFile(source) as zin:
                with zipfile.ZipFile(tmp_path / f"{label}.docx", "w") as zout:
                    for name in zin.namelist():
                        zout.writestr(name, zin.read(name))
                    zout.writestr("word/custom.xml", payload)
        diff = diff_package(tmp_path / "a.docx", tmp_path / "b.docx")
        assert "word/custom.xml" in diff.semantic_changed_parts()


class DescribeCasefoldExpansion:
    def it_matches_text_containing_expanding_casefolds(self):
        document = _doc()
        document.add_paragraph("Straße und Haus")
        assert find_one(document, "Straße").text == "Straße"
        # matches AFTER the expansion point must stay aligned
        assert find_one(document, "Haus").text == "Haus"
        assert find_one(document, "strasse").text == "Straße"  # casefold equal


class DescribeTabAndBreakMatching:
    def it_matches_across_a_tab_but_refuses_to_replace_over_it(self):
        document = _doc()
        paragraph = document.add_paragraph()
        paragraph.add_run("alpha")
        paragraph.add_run().add_tab()
        paragraph.add_run("beta")
        span = find_one(document, "alpha beta")  # tab normalizes to a space
        with pytest.raises(UnsupportedStructureError, match="tab or line break"):
            span.replace("gamma delta")

    def it_still_replaces_segments_beside_the_tab(self, tmp_path: Path):
        document = _doc()
        paragraph = document.add_paragraph()
        paragraph.add_run("alpha")
        paragraph.add_run().add_tab()
        paragraph.add_run("beta")
        find_one(document, "beta").replace("gamma")
        assert paragraph.text == "alpha\tgamma"


class DescribeOriginalViewNestedDeletions:
    def it_excludes_deletions_nested_inside_pending_insertions(self):
        """w:delText inside w:ins never existed in the original document."""
        document = _doc()
        document.paragraphs[0].add_run("tail")
        body_p = document.paragraphs[-1]._p
        body_p.append(
            parse_xml(
                f'<w:ins {W} w:id="900" w:author="A" w:date="2026-06-01T09:30:00Z">'
                '<w:del w:id="901" w:author="B" w:date="2026-06-02T09:30:00Z">'
                "<w:r><w:delText>never-existed</w:delText></w:r></w:del>"
                "<w:r><w:t>added-later</w:t></w:r></w:ins>"
            )
        )
        original = "\n".join(b.text for b in iter_blocks(document, view="original"))
        assert "never-existed" not in original
        assert "added-later" not in original
        everything = "\n".join(b.text for b in iter_blocks(document, view="all"))
        assert "never-existed" in everything
        assert find_text(document, "never-existed", view="original") == []


class DescribeConsumedAndDetachedSpans:
    def it_refuses_reuse_of_a_span_consumed_by_tracked_replace(self):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        span.replace("quite mundane", tracked=True, author="Carol QA", date=FROZEN)
        with pytest.raises(TargetNotFoundError, match="consumed"):
            span.replace("again", tracked=True, author="Carol QA", date=FROZEN)

    def it_refuses_spans_whose_structure_was_removed(self):
        """A detached span must refuse, not report success into an orphan."""
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        paragraph = span._atoms[0].paragraph  # noqa: SLF001
        paragraph.getparent().remove(paragraph)
        with pytest.raises(TargetNotFoundError, match="removed"):
            span.replace("lost edit")

    def it_consumes_a_span_after_exact_structure_replacement(self):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        span.replace("thoroughly mundane", preserve_structure=True)
        with pytest.raises(TargetNotFoundError, match="structure-preserving"):
            span.replace("again")
        assert find_one(document, "thoroughly mundane").text == "thoroughly mundane"


class DescribePreservedRevisionAncestry:
    def it_refuses_a_current_match_that_bridges_hidden_nested_history(self):
        document = _doc()
        paragraph = document.add_paragraph()._p
        paragraph.append(
            parse_xml(
                f'<w:ins {W} w:id="700" w:author="Alice">'
                "<w:r><w:t>alpha</w:t></w:r>"
                '<w:del w:id="701" w:author="Bob">'
                "<w:r><w:delText>hidden</w:delText></w:r></w:del>"
                "<w:r><w:t>beta</w:t></w:r></w:ins>"
            )
        )
        span = find_one(document, "alphabeta")
        before = document.element.xml
        with pytest.raises(UnsupportedStructureError, match="nested"):
            span.replace("updated", preserve_revision=True)
        assert document.element.xml == before

    def it_refuses_a_span_crossing_two_insertions(self):
        document = _doc()
        paragraph = document.add_paragraph()._p
        for revision_id, text in ((710, "alpha"), (711, "beta")):
            paragraph.append(
                parse_xml(
                    f'<w:ins {W} w:id="{revision_id}" w:author="Alice">'
                    f"<w:r><w:t>{text}</w:t></w:r></w:ins>"
                )
            )
        with pytest.raises(UnsupportedStructureError, match="multiple"):
            find_one(document, "alphabeta").replace(
                "updated", preserve_revision=True
            )

    def it_refuses_move_destinations_and_noncurrent_views(self):
        moved = _doc()
        moved.add_paragraph()._p.append(
            parse_xml(
                f'<w:moveTo {W} w:id="720" w:author="Alice">'
                "<w:r><w:t>moved text</w:t></w:r></w:moveTo>"
            )
        )
        with pytest.raises(UnsupportedStructureError, match="moves"):
            find_one(moved, "moved text").replace(
                "updated", preserve_revision=True
            )

        inserted = _doc()
        inserted.add_paragraph()._p.append(
            parse_xml(
                f'<w:ins {W} w:id="721" w:author="Alice">'
                "<w:r><w:t>inserted text</w:t></w:r></w:ins>"
            )
        )
        with pytest.raises(UnsupportedStructureError, match="view='current'"):
            find_one(inserted, "inserted text", view="all").replace(
                "updated", preserve_revision=True
            )

    def it_keeps_structure_preservation_orthogonal_to_revision_authorization(self):
        document = _doc()
        document.add_paragraph()._p.append(
            parse_xml(
                f'<w:ins {W} w:id="730" w:author="Alice">'
                "<w:r><w:t>inserted text</w:t></w:r></w:ins>"
            )
        )
        with pytest.raises(UnsupportedStructureError, match="pending tracked insertion"):
            find_one(document, "inserted text").replace(
                "updated", preserve_structure=True
            )


class DescribeLayeredTrackedEdits:
    def it_refuses_spans_straddling_a_pending_insertion(self):
        """A w:del claiming inserted text was base content fabricates
        history and corrupts reject/original views."""
        document = _doc()
        find_one(document, "perfectly ordinary").replace(
            "thoroughly modern", tracked=True, author="Alice Editor", date=FROZEN
        )
        straddling = find_one(document, "with thoroughly modern text")
        with pytest.raises(UnsupportedStructureError, match="pending tracked insertion"):
            straddling.replace("x", tracked=True, author="Bob Reviewer", date=FROZEN)

    def it_allows_edits_fully_inside_one_insertion(self):
        document = _doc()
        pristine = [b.text for b in iter_blocks(document)]
        find_one(document, "perfectly ordinary").replace(
            "thoroughly modern", tracked=True, author="Alice Editor", date=FROZEN
        )
        inner = find_one(document, "thoroughly")
        inner.replace("startlingly", tracked=True, author="Bob Reviewer", date=FROZEN)
        document.revisions.reject_all()
        assert [b.text for b in iter_blocks(document)] == pristine


class DescribeTrackedReplaceRunIntegrity:
    def it_keeps_run_content_after_the_match_in_visible_order(self):
        """Content following the matched w:t inside the same run (a tab, a
        second w:t) must not end up before the emitted revision."""
        document = _doc()
        paragraph = document.add_paragraph()
        run = paragraph.add_run("Hello")
        run.add_tab()
        run.add_text("World")
        find_one(document, "ell").replace(
            "xyz", tracked=True, author="Carol QA", date=FROZEN
        )
        # upstream paragraph.text is blind to w:ins; the story view sees all
        current = next(
            b.text for b in iter_blocks(document) if "World" in b.text
        )
        assert current == "Hxyzo\tWorld"
        original = next(
            b.text for b in iter_blocks(document, view="original")
            if "World" in b.text
        )
        assert original == "Hello\tWorld"

    def it_never_moves_kept_characters_into_another_runs_formatting(self):
        """Common-affix trim is clamped to run boundaries: characters the
        edit does not change must keep their formatting."""
        document = _doc()
        paragraph = document.add_paragraph()
        paragraph.add_run("AB")
        paragraph.add_run("CDEF").bold = True
        span = find_one(document, "BCDE")
        span.replace("BCDXE", tracked=True, author="Carol QA", date=FROZEN)
        document.revisions.accept_all()
        bold_text = "".join(r.text for r in paragraph.runs if r.bold)
        assert "CD" in bold_text, "unchanged bold 'CD' migrated to a plain run"
        plain_text = "".join(r.text for r in paragraph.runs if not r.bold)
        assert "CD" not in plain_text

    def it_works_on_documents_with_footnotes(self):
        """Revision-id allocation must not crash on story roots without
        registered oxml classes (w:footnotes has no .xpath prefix support)."""
        document = _doc("generated/feature-isolated/footnotes-endnotes.docx")
        result = find_one(document, "carries a footnote").replace(
            "bears a footnote", tracked=True, author="Carol QA", date=FROZEN
        )
        assert result.tracked and result.revision_ids


class DescribeParagraphMarkResolution:
    def _doc_with_mark_deleted_pilcrow(self):
        """Word's shape for a deleted paragraph BREAK: mark deleted, content kept."""
        document = _doc()
        first_body_p = document.paragraphs[1]._p
        p_pr = first_body_p.get_or_add_pPr()
        r_pr = parse_xml(
            f'<w:rPr {W}>'
            '<w:del w:id="800" w:author="Carol QA" w:date="2026-07-07T12:00:00Z"/>'
            "</w:rPr>"
        )
        p_pr.insert(0, r_pr)
        return document

    def it_merges_content_into_the_next_paragraph_on_accept(self):
        document = self._doc_with_mark_deleted_pilcrow()
        count = document.revisions.accept_all()
        assert count == 1
        texts = [b.text for b in iter_blocks(document)]
        assert (
            "First body paragraph with perfectly ordinary text."
            "Second body paragraph, equally unremarkable."
        ) in texts

    def it_restores_the_break_on_reject(self):
        document = self._doc_with_mark_deleted_pilcrow()
        pristine = [
            "Minimal Clean Document",
            "First body paragraph with perfectly ordinary text.",
            "Second body paragraph, equally unremarkable.",
        ]
        document.revisions.reject_all()
        assert [b.text for b in iter_blocks(document)] == pristine

    def it_never_empties_a_container_that_needs_a_block(self):
        """Accepting the deletion of a cell's only paragraph must leave a
        (possibly empty) paragraph — a block-less w:tc is schema-invalid."""
        from docx.blocks import tracked_delete_paragraphs
        from docx.search import find_one as find

        document = _doc()
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "only cell content"
        tracked_delete_paragraphs(
            document, "only cell content", count=1, author="Carol QA", date=FROZEN
        )
        document.revisions.accept_all()
        cell_tc = table.cell(0, 0)._tc
        assert cell_tc.find(qn("w:p")) is not None, "cell lost its last block"


class DescribeRevisionAnchors:
    def it_carries_block_anchors_usable_for_block_operations(self):
        from docx.blocks import insert_section_after

        document = _doc("generated/feature-isolated/tracked-ins-del.docx")
        revision = next(
            r for r in document.revisions if not r.is_paragraph_mark
        )
        insert_section_after(
            document, revision.anchor, heading="After Revision", paragraphs=[]
        )
        assert "After Revision" in [b.text for b in iter_blocks(document)]


class DescribeTableGuardsAndFormatting:
    def it_refuses_cells_whose_only_complexity_is_a_nested_table(self):
        """Regression for the reference's `.//w:tbl//w:tbl` bug: singly-nested
        tables must be detected. The guard is narrowed to the target cell, so
        the refusal fires on the nesting cell itself while its plain neighbors
        stay editable."""
        from docx.tableops import update_cell

        document = _doc()
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "host"
        table.cell(0, 1).add_table(rows=1, cols=1)  # nested, no merges
        with pytest.raises(UnsupportedStructureError, match="nested"):
            update_cell(table, 0, 1, "x")
        update_cell(table, 0, 0, "plain neighbor updates fine")
        assert table.cell(0, 0).text == "plain neighbor updates fine"

    def it_copies_run_formatting_into_inserted_rows(self):
        from docx.tableops import insert_row_after

        document = _doc()
        table = document.add_table(rows=1, cols=1)
        template_run = table.cell(0, 0).paragraphs[0].add_run("bold cell")
        template_run.bold = True
        insert_row_after(table, 0, ["copied"])
        new_runs = table.rows[1].cells[0].paragraphs[0].runs
        assert [r.text for r in new_runs] == ["copied"]
        assert new_runs[0].bold, "template run formatting was dropped"


class DescribeParagraphMarkStampOrder:
    def it_puts_the_mark_first_in_the_paragraph_rpr(self):
        """CT_ParaRPr requires w:ins/w:del before run properties."""
        from docx.blocks import insert_section_after

        document = _doc()
        insert_section_after(
            document, "First body paragraph", heading="H", paragraphs=[],
            tracked=True, author="Carol QA", date=FROZEN,
        )
        (mark,) = document.element.body.xpath("//w:pPr/w:rPr/w:ins")
        r_pr = mark.getparent()
        assert list(r_pr)[0] is mark
