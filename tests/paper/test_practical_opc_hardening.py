"""Focused regressions for practical OPC read/write hardening."""

from __future__ import annotations

import copy
import io
import stat
import struct
import zipfile
from pathlib import Path

import pytest

import docx
import docx._paperpkg as paperpkg_module
import docx.opc.package as package_module
from docx._transaction import rollback_on_error
from docx.errors import MalformedPackageError
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.opc.pkgreader import _SerializedRelationships
from docx.package import patch_save
from docx.search import find_one
from docx.story import story_parts


def _document_bytes() -> bytes:
    stream = io.BytesIO()
    docx.Document().save(stream)
    return stream.getvalue()


def _rewrite_package(data: bytes, transform) -> bytes:
    destination = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(data)) as source, zipfile.ZipFile(
        destination, "w", zipfile.ZIP_DEFLATED
    ) as output:
        for name in source.namelist():
            replacement = transform(name, source.read(name))
            if replacement is not None:
                output.writestr(*replacement)
    return destination.getvalue()


def it_saves_paths_atomically_preserving_mode_and_following_destination_symlink(
    tmp_path: Path,
):
    target = tmp_path / "target.docx"
    target.write_bytes(b"old")
    target.chmod(0o640)
    link = tmp_path / "link.docx"
    link.symlink_to(target.name)

    document = docx.Document()
    document.add_paragraph("saved through link")
    document.save(link)

    assert link.is_symlink()
    assert stat.S_IMODE(target.stat().st_mode) == 0o640
    assert docx.Document(target).paragraphs[-1].text == "saved through link"
    assert not list(tmp_path.glob("*.partial"))


def it_preserves_an_existing_files_mode_without_fchmod(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    target = tmp_path / "target.docx"
    docx.Document().save(target)
    target.chmod(0o640)
    monkeypatch.delattr(package_module.os, "fchmod", raising=False)

    document = docx.Document()
    document.add_paragraph("portable save")
    document.save(target)

    assert stat.S_IMODE(target.stat().st_mode) == 0o640
    assert docx.Document(target).paragraphs[-1].text == "portable save"


def it_refuses_a_dtd_before_parsing_relationship_records():
    relationships = b"""<!DOCTYPE Relationships [<!ENTITY role "unsafe">]>
    <Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
      <Relationship Id="rId1" Type="&role;" Target="word/document.xml"/>
    </Relationships>"""

    with pytest.raises(MalformedPackageError, match="prohibited DTD"):
        _SerializedRelationships.load_from_xml("/", relationships)


def it_validates_patch_save_after_final_recompression(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
):
    source = tmp_path / "source.docx"
    original = docx.Document()
    original.add_paragraph("before")
    original.save(source)
    document = docx.Document(source)
    document.paragraphs[0].text = "after"
    destination = tmp_path / "out.docx"
    destination.write_bytes(b"existing destination")

    unsafe = io.BytesIO()
    with zipfile.ZipFile(unsafe, "w", zipfile.ZIP_STORED) as archive:
        archive.writestr("word/document.xml", b"<document/>")
    payload = bytearray(unsafe.getvalue())
    local_offset = payload.index(b"PK\x03\x04")
    central_offset = payload.index(b"PK\x01\x02")
    local_flags = struct.unpack_from("<H", payload, local_offset + 6)[0]
    central_flags = struct.unpack_from("<H", payload, central_offset + 8)[0]
    struct.pack_into("<H", payload, local_offset + 6, local_flags | 0x0001)
    struct.pack_into("<H", payload, central_offset + 8, central_flags | 0x0001)
    monkeypatch.setattr(
        paperpkg_module,
        "_deterministic_zip_bytes",
        lambda _parts, _order: bytes(payload),
    )

    with pytest.raises(MalformedPackageError, match="encrypted"):
        patch_save(source, document, destination)

    assert destination.read_bytes() == b"existing destination"


def it_refuses_if_the_destination_symlink_changes_during_save(
    tmp_path: Path, monkeypatch
):
    old_target = tmp_path / "old.docx"
    new_target = tmp_path / "new.docx"
    old_target.write_bytes(b"old target")
    new_target.write_bytes(b"new target")
    link = tmp_path / "link.docx"
    link.symlink_to(old_target.name)
    validate = package_module._validate_serialized_output

    def validate_and_repoint(source):
        validate(source)
        link.unlink()
        link.symlink_to(new_target.name)

    monkeypatch.setattr(
        package_module, "_validate_serialized_output", validate_and_repoint
    )

    with pytest.raises(OSError, match="symlink changed"):
        docx.Document().save(link)

    assert old_target.read_bytes() == b"old target"
    assert new_target.read_bytes() == b"new target"


def it_restores_a_seekable_stream_after_commit_error():
    class FailOnce(io.BytesIO):
        failed = False

        def write(self, data):
            if not self.failed:
                self.failed = True
                super().write(data[:7])
                raise OSError("commit failed")
            return super().write(data)

    original = b"prefix and original suffix"
    stream = FailOnce(original)
    stream.seek(0)

    with pytest.raises(OSError, match="commit failed"):
        docx.Document().save(stream)

    # -- the subject is commit-failure rollback, not the cursor: a nonzero position is now
    # -- refused outright, so this exercises snapshot/restore at offset 0
    assert stream.getvalue() == original
    assert stream.tell() == 0


def it_refuses_an_append_mode_stream_without_changing_it(tmp_path: Path):
    destination = tmp_path / "append.docx"
    original = b"existing destination"
    destination.write_bytes(original)

    with destination.open("a+b") as stream:
        stream.seek(0)
        with pytest.raises(OSError, match="append-mode"):
            docx.Document().save(stream)

    assert destination.read_bytes() == original


def it_refuses_a_nonzero_position_stream_without_changing_it():
    """Word refuses a package with bytes in front of it.

    This test previously asserted the opposite -- that the save succeeded, preserved the
    caller's prefix, and that the whole stream reopened as a Document. Word for Mac
    (2026-08-18) refuses exactly that file, and the extractable slice is offset-skewed by the
    prefix length, so the operation has no usable output by either route.
    """
    prefix = b"real destination prefix"
    original = prefix + b"existing suffix"
    stream = io.BytesIO(original)
    stream.seek(len(prefix))

    with pytest.raises(OSError, match="positioned at offset"):
        docx.Document().save(stream)

    assert stream.getvalue() == original
    assert stream.tell() == len(prefix)


def it_truncates_an_offset_zero_stream_holding_a_longer_document(tmp_path: Path):
    """The behaviour the truncation gate must preserve.

    A short document written over a longer one at offset 0 must leave no tail of the previous
    document: undeclared trailing bytes are a shape Word refuses.
    """
    reference = tmp_path / "fresh.docx"
    docx.Document().save(str(reference))
    fresh_bytes = reference.read_bytes()

    stream = io.BytesIO(b"Z" * (len(fresh_bytes) * 3))
    stream.seek(0)
    docx.Document().save(stream)

    assert stream.getvalue() == fresh_bytes


def it_refuses_before_overwriting_an_unreadable_seekable_stream():
    class Unreadable(io.BytesIO):
        def read(self, size=-1):
            raise OSError("existing bytes are unavailable")

    original = b"existing destination"
    stream = Unreadable(original)

    with pytest.raises(OSError, match="could not be snapshotted"):
        docx.Document().save(stream)

    assert stream.getvalue() == original
    assert stream.tell() == 0


def it_restores_mutables_base_uri_and_detaches_corrupt_graph_residuals():
    class Participant:
        def __init__(self):
            self.values = ["original"]

    document = docx.Document()
    participant = Participant()
    values = participant.values
    relationships = document.part.rels
    base_uri = relationships._baseURI
    rogue = Part(
        PackURI("/word/rogue.bin"), "application/octet-stream", b"rogue", document.part.package
    )

    def fail_late():
        with rollback_on_error(document, participant):
            values.append("changed")
            relationships._baseURI = "/corrupt"
            relationships._target_parts_by_rId["rId-corrupt"] = rogue
            raise ValueError("late failure")

    with pytest.raises(ValueError, match="late failure"):
        fail_late()

    assert participant.values is values
    assert values == ["original"]
    assert relationships._baseURI == base_uri
    assert "rId-corrupt" not in relationships._target_parts_by_rId
    assert rogue.package is None


def it_resolves_unambiguous_physical_members_without_case_sensitivity():
    source = docx.Document()
    source.add_paragraph("case-insensitive story")
    stream = io.BytesIO()
    source.save(stream)
    data = _rewrite_package(
        stream.getvalue(),
        lambda name, blob: (
            ("WORD/DOCUMENT.XML", blob) if name == "word/document.xml" else (name, blob)
        ),
    )

    document = docx.Document(io.BytesIO(data))

    assert story_parts(document)[0] == "word/document.xml"
    assert find_one(document, "case-insensitive story").story == "word/document.xml"


def it_refuses_a_symlinked_member_in_an_expanded_package(tmp_path: Path):
    expanded = tmp_path / "expanded"
    with zipfile.ZipFile(io.BytesIO(_document_bytes())) as package:
        package.extractall(expanded)
    member = expanded / "word" / "document.xml"
    outside = tmp_path / "outside-document.xml"
    member.replace(outside)
    member.symlink_to(outside)

    with pytest.raises(MalformedPackageError, match="symbolic link"):
        docx.Document(expanded)


def it_refuses_duplicate_relationship_ids():
    relationships = b"""<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
      <Relationship Id="rId1" Type="urn:one" Target="one.xml"/>
      <Relationship Id="rId1" Type="urn:two" Target="two.xml"/>
    </Relationships>"""

    with pytest.raises(MalformedPackageError, match="two different targets"):
        _SerializedRelationships.load_from_xml("/word", relationships)


def it_refuses_multiple_main_document_relationships_before_replacing_a_path(
    tmp_path: Path,
):
    destination = tmp_path / "destination.docx"
    destination.write_bytes(b"original")
    document = docx.Document()
    document.part.package.load_rel(
        RT.OFFICE_DOCUMENT, document.part, "rIdDuplicateMain"
    )

    with pytest.raises(MalformedPackageError, match="multiple officeDocument"):
        document.save(destination)

    assert destination.read_bytes() == b"original"


def it_accepts_xml_comments_in_relationship_parts():
    def transform(name: str, blob: bytes):
        if name == "_rels/.rels":
            blob = blob.replace(b"<Relationship ", b"<!-- retained --><Relationship ", 1)
        return name, blob

    document = docx.Document(
        io.BytesIO(_rewrite_package(_document_bytes(), transform))
    )

    assert document.paragraphs is not None


def it_refuses_unexpected_relationship_elements_instead_of_dropping_them():
    def transform(name: str, blob: bytes):
        if name == "word/_rels/document.xml.rels":
            from lxml import etree

            root = etree.fromstring(blob)
            relationship = next(
                child for child in root if (child.get("Type") or "").endswith("/styles")
            )
            relationship.tag = relationship.tag.replace("Relationship", "Unexpected")
            blob = etree.tostring(root, encoding="UTF-8", standalone=True)
        return name, blob

    with pytest.raises(MalformedPackageError, match="unexpected element"):
        docx.Document(io.BytesIO(_rewrite_package(_document_bytes(), transform)))


def it_validates_case_insensitive_content_types_without_losing_duplicates():
    def transform(name: str, blob: bytes):
        if name != "[Content_Types].xml":
            return name, blob
        from lxml import etree

        root = etree.fromstring(blob)
        override = next(child for child in root if child.get("PartName"))
        duplicate = copy.deepcopy(override)
        duplicate.set("PartName", override.get("PartName").upper())
        root.append(duplicate)
        return (
            "[CONTENT_TYPES].XML",
            etree.tostring(root, encoding="UTF-8", standalone=True),
        )

    with pytest.raises(MalformedPackageError, match="ambiguous Override"):
        docx.Document(io.BytesIO(_rewrite_package(_document_bytes(), transform)))


# The "known relationship role declared with the wrong content type" mechanism is
# covered by test_word_media_content_types.py, which mutates the styles role -- the
# case Word actually verified -- and asserts the declaration changed before opening.
# This file keeps only the main-document exception, below.


def it_leaves_the_main_document_content_type_check_to_upstream():
    """The officeDocument role is not validated during load, on purpose.

    `api.Document` already checks the main part's content type and raises `ValueError`
    naming the file and the type it found. That check is upstream's, byte-identical,
    and a caller migrating from python-docx may rely on it. Validating the same
    condition during package load only pre-empted it with a different exception type.
    """

    def transform(name: str, blob: bytes):
        if name == "[Content_Types].xml":
            blob = blob.replace(
                b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                b"application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml",
            )
        return name, blob

    with pytest.raises(ValueError, match="is not a Word file, content type is"):
        docx.Document(io.BytesIO(_rewrite_package(_document_bytes(), transform)))


def it_refuses_a_relationship_to_a_missing_package_part():
    data = _rewrite_package(
        _document_bytes(),
        lambda name, blob: None if name == "word/styles.xml" else (name, blob),
    )

    with pytest.raises(MalformedPackageError, match="targets missing package part"):
        docx.Document(io.BytesIO(data))


def it_refuses_a_relationship_target_without_a_declared_content_type():
    def transform(name: str, blob: bytes):
        if name == "[Content_Types].xml":
            from lxml import etree

            root = etree.fromstring(blob)
            for declaration in tuple(root):
                if declaration.get("PartName") == "/word/styles.xml" or declaration.get(
                    "Extension"
                ) == "xml":
                    root.remove(declaration)
            blob = etree.tostring(root, encoding="UTF-8", standalone=True)
        return name, blob

    with pytest.raises(MalformedPackageError, match="no declared content type"):
        docx.Document(io.BytesIO(_rewrite_package(_document_bytes(), transform)))


def it_retains_default_creation_for_missing_optional_core_properties():
    def transform(name: str, blob: bytes):
        if name == "docProps/core.xml":
            return None
        if name == "_rels/.rels":
            from lxml import etree

            root = etree.fromstring(blob)
            for relationship in tuple(root):
                if relationship.get("Target") == "docProps/core.xml":
                    root.remove(relationship)
            blob = etree.tostring(root, encoding="UTF-8", standalone=True)
        return name, blob

    document = docx.Document(io.BytesIO(_rewrite_package(_document_bytes(), transform)))

    assert document.core_properties.title == "Word Document"
