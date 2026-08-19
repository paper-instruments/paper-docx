"""The protection gate, cell by cell against Word's measured behaviour.

Word's Restrict Editing rules are per-mode AND per-operation-class. The verdict
set `2-refused-operations` (Word for Mac, 2026-08-18) measured them:

| `w:edit`         | comment | form-field value | body content |
|------------------|---------|------------------|--------------|
| `readOnly`       | refuse  | refuse           | refuse       |
| `comments`       | PERMIT  | refuse           | refuse       |
| `trackedChanges` | PERMIT  | refuse           | PERMIT (u)   |
| `forms`          | refuse  | PERMIT           | refuse       |
| formatting-only  | PERMIT  | PERMIT           | PERMIT (u)   |

"(u)" marks the two cells Word has NOT been asked about. They ship permissive:
a refusal with no Word verdict behind it is an unjustified regression, and this
gate is a policy mirror of Word's UI rather than a corruption guard, so
permitting cannot produce a file Word refuses. See
`it_permits_the_two_unmeasured_body_content_cells` — measuring one later should
change that test deliberately.
"""

from __future__ import annotations

import ast
import functools
import io
import pathlib
from typing import Any, Callable, Dict, Optional, Tuple

import pytest

import docx
from docx.controls import iter_controls
from docx.drawing import Drawing
from docx.errors import DocumentProtectedError
from docx.links import add_hyperlink
from docx.notes import add_footnote
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import OxmlElement, parse_xml
from docx.protection import (
    _DOCUMENT_PROTECTION_SUCCESSORS,
    OP_BODY,
    OP_COMMENT,
    OP_FORM_FIELD,
    OPERATION_CLASSES,
    acknowledge_protection,
)
from docx.search import find_one

from .test_composition import TINY_PNG

# -- fixture helpers ---------------------------------------------------------

BODY_PHRASE = "an ordinary sentence of body text"


def _protect(
    document: "docx.document.Document",
    *,
    edit: Optional[str] = None,
    formatting: bool = False,
    enforcement: bool = True,
) -> "docx.document.Document":
    """Write `w:documentProtection` at its schema position.

    `_DOCUMENT_PROTECTION_SUCCESSORS` is the CT_Settings child sequence that
    follows `w:documentProtection`, so inserting before the first of them puts
    the element ahead of `w:defaultTabStop` where the schema requires it.
    """
    element = OxmlElement("w:documentProtection")
    if edit is not None:
        element.set(qn("w:edit"), edit)
    if formatting:
        element.set(qn("w:formatting"), "1")
    if enforcement:
        element.set(qn("w:enforcement"), "1")
    settings = document.settings.element
    settings.insert_element_before(element, *_DOCUMENT_PROTECTION_SUCCESSORS)
    return document


def _document() -> "docx.document.Document":
    document = docx.Document()
    document.add_paragraph(BODY_PHRASE)
    return document


def _control_document() -> Tuple["docx.document.Document", None]:
    """A document carrying one plain inline content control (a form field)."""
    document = _document()
    paragraph = parse_xml(
        f"<w:p {nsdecls('w')}>"
        "<w:sdt>"
        "<w:sdtPr>"
        '<w:alias w:val="Field"/><w:tag w:val="field-1"/><w:id w:val="4242"/>'
        "</w:sdtPr>"
        "<w:sdtContent><w:r><w:t>controlled text</w:t></w:r></w:sdtContent>"
        "</w:sdt>"
        "</w:p>"
    )
    body = document.element.body
    sect_pr = body.find(qn("w:sectPr"))
    if sect_pr is not None:
        sect_pr.addprevious(paragraph)
    else:
        body.append(paragraph)
    return document, None


# -- the operations, as real API calls ----------------------------------------
#
# Each entry is (setup, action). `setup` returns the fully populated document
# and whatever handle `action` needs; protection is applied BETWEEN the two, so
# the setup itself never trips the gate.


def _plain_document() -> Tuple["docx.document.Document", None]:
    return _document(), None


def _document_with_picture() -> Tuple["docx.document.Document", Drawing]:
    document = _document()
    run = document.add_paragraph().add_run()
    run.add_picture(io.BytesIO(TINY_PNG))
    drawing = next(item for item in run.iter_inner_content() if isinstance(item, Drawing))
    return document, drawing


def _document_with_hyperlink() -> Tuple["docx.document.Document", Any]:
    document = _document()
    link = add_hyperlink(document, find_one(document, "ordinary"), "https://example.com/a")
    return document, link


def _retarget(document: "docx.document.Document", link: Any) -> None:
    link.address = "https://example.com/b"


OPERATIONS: Dict[str, Tuple[Callable[[], Tuple[Any, Any]], Callable[[Any, Any], Any]]] = {
    # `comments.py` -> `_refuse_document_protection` -> the gate
    "comment": (
        _plain_document,
        lambda document, _: document.comments.add_comment("note", author="A"),
    ),
    # `search.py:Span.comment` and the `commentops` range funnel
    "anchored-comment": (
        _plain_document,
        lambda document, _: find_one(document, BODY_PHRASE).comment("note", author="A"),
    ),
    # `controls.py:Control.set_value` — the single form-field call site
    "form-field": (
        _control_document,
        lambda document, _: next(iter_controls(document)).set_value("typed in"),
    ),
    # `search.py:Span.replace` — plain body-content editing
    "body": (
        _plain_document,
        lambda document, _: find_one(document, BODY_PHRASE).replace("something else"),
    ),
    # `notes.py` builds its operation string at runtime — body content
    "note": (
        _plain_document,
        lambda document, _: add_footnote(document, find_one(document, BODY_PHRASE), "n"),
    ),
    # `drawing/__init__.py` — a subpackage site earlier counts missed
    "picture": (
        _document_with_picture,
        lambda document, drawing: drawing.replace_picture(io.BytesIO(TINY_PNG)),
    ),
    # `text/hyperlink.py` — the other subpackage site earlier counts missed
    "hyperlink": (_document_with_hyperlink, _retarget),
}


def _build(operation: str, **protection) -> Tuple["docx.document.Document", Callable[[], object]]:
    """Populate a document, protect it, and hand back the operation as a thunk."""
    setup, action = OPERATIONS[operation]
    document, handle = setup()
    _protect(document, **protection)
    return document, lambda: action(document, handle)


# -- the matrix --------------------------------------------------------------

MODES: Dict[str, dict] = {
    "readOnly": {"edit": "readOnly"},
    "comments": {"edit": "comments"},
    "trackedChanges": {"edit": "trackedChanges"},
    "forms": {"edit": "forms"},
    "formatting-only": {"formatting": True},
}

#: True = Word permits it, so the gate must permit it. Every cell not listed in
#: `UNMEASURED` was measured in Word for Mac.
MATRIX: Dict[Tuple[str, str], bool] = {
    ("readOnly", OP_COMMENT): False,
    ("readOnly", OP_FORM_FIELD): False,
    ("readOnly", OP_BODY): False,
    ("comments", OP_COMMENT): True,
    ("comments", OP_FORM_FIELD): False,
    ("comments", OP_BODY): False,
    ("trackedChanges", OP_COMMENT): True,
    ("trackedChanges", OP_FORM_FIELD): False,
    ("trackedChanges", OP_BODY): True,
    ("forms", OP_COMMENT): False,
    ("forms", OP_FORM_FIELD): True,
    ("forms", OP_BODY): False,
    ("formatting-only", OP_COMMENT): True,
    ("formatting-only", OP_FORM_FIELD): True,
    ("formatting-only", OP_BODY): True,
}

#: The two cells Word has NOT been asked about; see the module docstring.
UNMEASURED: Tuple[Tuple[str, str], ...] = (
    ("trackedChanges", OP_BODY),
    ("formatting-only", OP_BODY),
)

#: the operation exercising each class in the matrix
BUILDERS: Dict[str, str] = {
    OP_COMMENT: "comment",
    OP_FORM_FIELD: "form-field",
    OP_BODY: "body",
}


def _assert_matches(mode: str, operation_class: str, operate: Callable[[], object]) -> None:
    """Run `operate`, expecting exactly the verdict the matrix records."""
    if MATRIX[(mode, operation_class)]:
        operate()  # Word permits it, so paper-docx must not refuse
        return
    with pytest.raises(DocumentProtectedError):
        operate()


class DescribeTheProtectionMatrix:
    """One test per populated cell — Fix 4 of the word-oracle-fixes spec."""

    @pytest.mark.parametrize(("mode", "operation_class"), sorted(MATRIX))
    def it_matches_words_measured_verdict_for_every_cell(
        self, mode: str, operation_class: str
    ):
        _, operate = _build(BUILDERS[operation_class], **MODES[mode])
        _assert_matches(mode, operation_class, operate)

    @pytest.mark.parametrize("mode", [mode for mode, _ in UNMEASURED])
    def it_permits_the_two_unmeasured_body_content_cells(self, mode: str):
        """Body content under `trackedChanges` and under formatting-only.

        NEITHER CELL IS MEASURED IN WORD. They ship permissive because a
        refusal with no verdict behind it is an unjustified regression, and the
        gate is a policy mirror rather than a corruption guard. If someone
        measures either cell in Word and it turns out blocked, this test is the
        one that must change, deliberately.
        """
        assert MATRIX[(mode, OP_BODY)] is True
        _, operate = _build("body", **MODES[mode])
        operate()

    @pytest.mark.parametrize("operation_class", sorted(OPERATION_CLASSES))
    @pytest.mark.parametrize("token", ["lockedTighterThanThis", "formatting-only"])
    def it_refuses_every_class_for_an_unrecognised_edit_token(
        self, token: str, operation_class: str
    ):
        """An unknown restriction is not a licence.

        `formatting-only` is the name this module gives the row for a document
        that declares `w:formatting` and no `w:edit`. A document declaring
        `w:edit="formatting-only"` is a different thing — an unrecognised token —
        and must not reach that permissive row by spelling its label.
        """
        _, operate = _build(BUILDERS[operation_class], edit=token)
        with pytest.raises(DocumentProtectedError):
            operate()

    def it_ignores_protection_that_is_not_enforced(self):
        _, operate = _build("body", edit="readOnly", enforcement=False)
        operate()

    @pytest.mark.parametrize(("mode", "operation_class"), sorted(MATRIX))
    def it_lets_the_acknowledgement_flag_override_every_refused_cell(
        self, mode: str, operation_class: str
    ):
        if MATRIX[(mode, operation_class)]:
            pytest.skip("cell already permits; nothing to override")
        document, operate = _build(BUILDERS[operation_class], **MODES[mode])
        acknowledge_protection(document)
        operate()


class DescribeCommentClassCallSites:
    """The comment-class call lines, reached through both funnels."""

    @pytest.mark.parametrize("mode", sorted(MODES))
    def it_matches_the_comment_verdict_when_anchoring_a_comment(self, mode: str):
        _, operate = _build("anchored-comment", **MODES[mode])
        _assert_matches(mode, OP_COMMENT, operate)

    @pytest.mark.parametrize(
        "mode", [mode for mode in sorted(MODES) if MATRIX[(mode, OP_COMMENT)]]
    )
    def it_permits_the_inherited_comment_mutators(self, mode: str):
        document = _document()
        comment = document.comments.add_comment("first", author="A")
        _protect(document, **MODES[mode])
        comment.add_paragraph("second")


class DescribeThePreviouslyMissedSubpackageSites:
    """`drawing/__init__.py` and `text/hyperlink.py`.

    Both live in subpackages and an earlier `src/docx/*.py` glob skipped them.
    Left unclassified they would keep the old blanket gate while their siblings
    became mode-aware — silently inconsistent rather than broken.
    """

    @pytest.mark.parametrize("operation", ["picture", "hyperlink"])
    @pytest.mark.parametrize("mode", sorted(MODES))
    def it_matches_the_body_content_verdict(self, operation: str, mode: str):
        _, operate = _build(operation, **MODES[mode])
        _assert_matches(mode, OP_BODY, operate)


class DescribeTheRuntimeOperationStrings:
    """Four call expressions pass `operation` as a variable, not a literal."""

    @pytest.mark.parametrize("mode", sorted(MODES))
    def it_matches_the_body_content_verdict_for_the_runtime_note_operation(self, mode: str):
        _, operate = _build("note", **MODES[mode])
        _assert_matches(mode, OP_BODY, operate)

    def it_names_the_runtime_operation_in_the_message(self):
        _, operate = _build("note", edit="readOnly")
        with pytest.raises(DocumentProtectedError) as exc:
            operate()
        assert "add a footnote" in str(exc.value)


class DescribeTheRefusalMessage:
    """Fix 4's wording half: a policy restriction, not a broken document."""

    def it_says_word_would_not_permit_the_edit(self):
        _, operate = _build("body", edit="readOnly")
        with pytest.raises(DocumentProtectedError) as exc:
            operate()
        message = str(exc.value)
        assert "Word" in message
        assert "would not permit" in message
        assert "acknowledge_protection" in message
        assert "'readOnly' editing" in message

    def it_no_longer_implies_the_document_would_break(self):
        """Every protected document carrying an upstream-added comment opened
        cleanly in Word with the comment visible: the artifact is never
        damaged, so the message must not suggest it is."""
        _, operate = _build("body", edit="readOnly")
        with pytest.raises(DocumentProtectedError) as exc:
            operate()
        message = str(exc.value).lower()
        for word in ("corrupt", "damage", "broken", "unreadable", "invalid"):
            assert word not in message

    def it_names_the_edit_token_when_formatting_is_also_restricted(self):
        _, operate = _build("form-field", edit="lockedTighterThanThis", formatting=True)
        with pytest.raises(DocumentProtectedError) as exc:
            operate()
        assert "'lockedTighterThanThis' editing" in str(exc.value)


# -- coverage assertion ------------------------------------------------------
#
# Enumerated with an AST walk, NOT a regex, for two reasons: the operation class
# of a call site is a keyword argument that has to be resolved back to its
# constant, and `comments.py` reaches the gate through a one-line forwarder, so
# counting spellings of the gate's own name would miss its two call sites.

_SRC = pathlib.Path(docx.__file__).parent
_GATE = "_refuse_if_protected"

#: The one function whose whole body is a gate call: calling it is calling the
#: gate. A validation routine that merely happens to call the gate (like
#: `commentops._preflight_comment_range`) is NOT one — it does other work, and
#: its callers are not gate call expressions.
_FORWARDER = "_refuse_document_protection"


def _callee(node: ast.AST) -> Optional[str]:
    if isinstance(node, ast.Name):
        return node.id
    if isinstance(node, ast.Attribute):
        return node.attr
    return None


def _class_of(call: ast.Call, constants: Dict[str, str]) -> str:
    """The operation class a gate call resolves to; the default is body content."""
    for keyword in call.keywords:
        if keyword.arg != "operation_class":
            continue
        value = keyword.value
        if isinstance(value, ast.Constant):
            return str(value.value)
        name = _callee(value)
        assert name in constants, f"operation_class is not a known constant: {name!r}"
        return constants[name]
    return OP_BODY


def _forwarder_class(trees: Dict[str, ast.Module], constants: Dict[str, str]) -> str:
    """The class `_FORWARDER` passes on, read off its single gate call."""
    definitions = [
        node
        for tree in trees.values()
        for node in ast.walk(tree)
        if isinstance(node, ast.FunctionDef) and node.name == _FORWARDER
    ]
    assert len(definitions) == 1, f"expected exactly one {_FORWARDER} definition"
    call = next(
        node
        for node in ast.walk(definitions[0])
        if isinstance(node, ast.Call) and _callee(node.func) == _GATE
    )
    return _class_of(call, constants)


@functools.lru_cache(maxsize=None)
def _census() -> Dict[str, Dict[str, int]]:
    """module -> operation class -> number of call expressions reaching the gate."""
    import docx.protection as protection_module

    constants = {
        name: getattr(protection_module, name)
        for name in ("OP_COMMENT", "OP_FORM_FIELD", "OP_BODY")
    }
    trees = {
        str(path.relative_to(_SRC)): ast.parse(path.read_text(encoding="utf-8"))
        for path in sorted(_SRC.rglob("*.py"))
    }
    forwarded = _forwarder_class(trees, constants)

    census: Dict[str, Dict[str, int]] = {}
    for module, tree in trees.items():
        for node in ast.walk(tree):
            if not isinstance(node, ast.Call):
                continue
            name = _callee(node.func)
            if name == _GATE:
                operation_class = _class_of(node, constants)
            elif name == _FORWARDER:
                operation_class = forwarded
            else:
                continue
            counts = census.setdefault(module, {})
            counts[operation_class] = counts.get(operation_class, 0) + 1
    return census


def _sites(operation_class: str) -> Dict[str, int]:
    """The call-site count per module for one operation class."""
    return {
        module: counts[operation_class]
        for module, counts in _census().items()
        if operation_class in counts
    }


#: The census phase 4 settled by AST parse: 36 call expressions across 15
#: modules. The totals are pinned rather than the whole per-module table
#: because body content is the default, so pinning where every body site lands
#: would make moving a gated function between modules fail a protection test
#: for a reason that has nothing to do with protection. The totals still do the
#: job the pin is for: a gate call added anywhere fails here, so it has to be
#: classified deliberately instead of silently inheriting the default.
TOTAL_SITES = 36
TOTAL_MODULES = 15


class DescribeGateCallSiteCoverage:
    """Every call expression reaching the gate resolves to a declared class."""

    def it_classifies_every_call_expression_in_every_module(self):
        census = _census()
        assert census, "no gate call sites found at all"
        seen = {
            operation_class for counts in census.values() for operation_class in counts
        }
        assert seen <= set(OPERATION_CLASSES), seen - set(OPERATION_CLASSES)
        assert len(census) == TOTAL_MODULES, sorted(census)
        total = sum(count for counts in census.values() for count in counts.values())
        assert total == TOTAL_SITES, census

    def it_declares_every_comment_class_call_site(self):
        assert _sites(OP_COMMENT) == {
            "commentops.py": 4,
            "comments.py": 3,
            "search.py": 1,
        }

    def it_declares_the_single_form_field_call_site(self):
        assert _sites(OP_FORM_FIELD) == {"controls.py": 1}

    def it_covers_the_two_subpackage_modules_earlier_counts_missed(self):
        census = _census()
        assert census["drawing/__init__.py"] == {OP_BODY: 1}
        assert census["text/hyperlink.py"] == {OP_BODY: 1}

    def it_defaults_an_unclassified_call_to_body_content(self):
        """The strictest class, so a new site cannot become permissive by
        omission."""
        tree = ast.parse("_refuse_if_protected(document, 'do a thing')\n")
        call = next(node for node in ast.walk(tree) if isinstance(node, ast.Call))
        assert _class_of(call, {}) == OP_BODY
