"""Missing-table APIs: comments, lock, pictures, links, captions, notes."""

from __future__ import annotations

import io
import struct
from pathlib import Path

import pytest
from lxml import etree

import docx
from docx.bookmarks import create_bookmark
from docx.commentops import COMMENTS_IDS_RELATIONSHIP_TYPE, delete_comment
from docx.composition import CompositionReport, _copy_letterhead, append_document
from docx.controls import _part_root, get_control, set_control_value
from docx.drawing import Drawing
from docx.errors import (
    BoundaryViolationError,
    DocumentProtectedError,
    TargetNotFoundError,
    UnsupportedStructureError,
)
from docx.fields import add_caption
from docx.enum.style import WD_STYLE_TYPE
from docx.links import add_hyperlink
from docx.notes import add_endnote, add_footnote
from docx.numbering import apply_numbering, ensure_decimal_definition, list_numbering
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import OxmlElement, parse_xml
from docx.protection import acknowledge_protection, set_protection
from docx.search import find_one

from .harness.contract import save_and_reopen
from .harness.paths import fixture_path

MINIMAL = "generated/minimal-clean/minimal.docx"
STORE_ID = "{11111111-1111-1111-1111-111111111111}"


def _doc(relpath: str = MINIMAL):
    return docx.Document(str(fixture_path(relpath)))


def _bmp(red: int, green: int, blue: int) -> bytes:
    return (
        struct.pack("<2sIHHI", b"BM", 58, 0, 0, 54)
        + struct.pack("<IiiHHIIiiII", 40, 1, 1, 1, 24, 0, 4, 2835, 2835, 0, 0)
        + bytes((blue, green, red, 0))
    )


def _blip_embed(drawing: Drawing) -> str:
    return drawing._drawing.xpath(".//pic:blipFill/a:blip/@r:embed")[0]


def _attach_bound_control(
    document,
    tag: str,
    *,
    locked: bool = False,
    extra_pr: str = "",
    xpath: str = "/ns0:root/ns0:name",
    prefix_mappings: str | None = "xmlns:ns0='http://example.com/form'",
    item_xml: str = (
        '<ns0:root xmlns:ns0="http://example.com/form">'
        "<ns0:name>old</ns0:name></ns0:root>"
    ),
) -> None:
    item = parse_xml(item_xml)
    item_part = XmlPart(
        PackURI("/customXml/item2.xml"),
        "application/xml",
        item,
        document.part.package,
    )
    props = parse_xml(
        '<ds:datastoreItem xmlns:ds="http://schemas.openxmlformats.org/'
        'officeDocument/2006/customXml"'
        f' ds:itemID="{STORE_ID}"/>'
    )
    props_part = XmlPart(
        PackURI("/customXml/itemProps2.xml"),
        CT.OFC_CUSTOM_XML_PROPERTIES,
        props,
        document.part.package,
    )
    document.part.relate_to(item_part, RT.CUSTOM_XML)
    item_part.relate_to(props_part, RT.CUSTOM_XML_PROPS)
    lock = '<w:lock w:val="contentLocked"/>' if locked else ""
    mapping_attr = (
        f" w:prefixMappings=\"{prefix_mappings}\"" if prefix_mappings is not None else ""
    )
    document.element.body.insert(
        len(document.element.body) - 1,
        parse_xml(
            f'<w:p {nsdecls("w", "w14")}>'
            f'<w:sdt><w:sdtPr><w:tag w:val="{tag}"/>{lock}{extra_pr}'
            f"<w:dataBinding{mapping_attr}"
            f' w:xpath="{xpath}"'
            f' w:storeItemID="{STORE_ID}"/></w:sdtPr>'
            "<w:sdtContent><w:r><w:t>old</w:t></w:r></w:sdtContent>"
            "</w:sdt></w:p>"
        ),
    )


class DescribeCommentDeleteAndIdentity:
    def it_writes_modern_comment_identity_parts(self):
        document = _doc()
        find_one(document, "perfectly ordinary").comment(text="note", author="Ada")
        names = {str(part.partname) for part in document.part.package.iter_parts()}
        assert "/word/commentsIds.xml" in names
        assert "/word/commentsExtensible.xml" in names

    def it_deletes_one_comment_and_leaves_the_rest(self, tmp_path: Path):
        document = _doc()
        first = find_one(document, "perfectly ordinary").comment(text="keep", author="Ada")
        second = find_one(document, "First body paragraph").comment(
            text="drop", author="Ada"
        )
        delete_comment(document, second)
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        texts = [comment.text for comment in reopened.comments]
        assert "keep" in texts
        assert "drop" not in texts
        assert first.comment_id in {comment.comment_id for comment in reopened.comments}

    def it_retargets_identity_when_the_last_paragraph_moves(self):
        document = _doc()
        comment = find_one(document, "perfectly ordinary").comment(
            text="note", author="Ada"
        )
        old_para = comment.paragraphs[-1]._p.get(qn("w14:paraId"))
        comment.add_paragraph("more")
        new_para = comment.paragraphs[-1]._p.get(qn("w14:paraId"))
        assert new_para and new_para != old_para
        ids_root = document.part.part_related_by(COMMENTS_IDS_RELATIONSHIP_TYPE)._element
        para_ids = [
            entry.get(
                "{http://schemas.microsoft.com/office/word/2016/wordml/cid}paraId"
            )
            for entry in ids_root
        ]
        assert new_para in para_ids
        assert old_para not in para_ids
        delete_comment(document, comment)
        assert list(ids_root) == []


class DescribeProtectionSetter:
    def it_locks_the_file_for_the_recipient(self):
        document = _doc()
        status = set_protection(document, edit="readOnly")
        assert status.edit == "readOnly"
        assert status.enforced
        with pytest.raises(DocumentProtectedError):
            find_one(document, "perfectly ordinary").replace("nope")

    def it_clears_an_in_memory_ack_when_locking_again(self):
        document = _doc()
        set_protection(document, edit="readOnly")
        acknowledge_protection(document)
        find_one(document, "perfectly ordinary").replace("ok")
        set_protection(document, edit="comments")
        with pytest.raises(DocumentProtectedError):
            find_one(document, "ok").replace("nope")

    def it_inserts_protection_before_later_settings_children(self):
        document = _doc()
        settings = document.settings.element
        for child in list(settings):
            if child.tag == qn("w:defaultTabStop"):
                settings.remove(child)
        footnote_pr = OxmlElement("w:footnotePr")
        compat = settings.find(qn("w:compat"))
        if compat is not None:
            compat.addprevious(footnote_pr)
        else:
            settings.append(footnote_pr)
        set_protection(document, edit="readOnly")
        tags = [child.tag for child in settings]
        assert tags.index(qn("w:documentProtection")) < tags.index(qn("w:footnotePr"))


class DescribePictureReplace:
    def it_swaps_bytes_without_sharing_the_part(self):
        document = _doc()
        run = document.add_paragraph().add_run()
        run.add_picture(io.BytesIO(_bmp(255, 0, 0)))
        other = document.add_paragraph().add_run()
        other.add_picture(io.BytesIO(_bmp(255, 0, 0)))
        drawing = next(
            item for item in run.iter_inner_content() if isinstance(item, Drawing)
        )
        other_drawing = next(
            item for item in other.iter_inner_content() if isinstance(item, Drawing)
        )
        shared = _blip_embed(drawing)
        assert _blip_embed(other_drawing) == shared
        extent = drawing._drawing.xpath(".//wp:extent")[0]
        cx, cy = extent.get("cx"), extent.get("cy")
        drawing.replace_picture(io.BytesIO(_bmp(0, 0, 255)))
        assert _blip_embed(drawing) != shared
        assert _blip_embed(other_drawing) == shared
        assert extent.get("cx") == cx
        assert extent.get("cy") == cy

    def it_refuses_a_linked_picture(self):
        document = _doc()
        run = document.add_paragraph().add_run()
        run.add_picture(io.BytesIO(_bmp(255, 0, 0)))
        drawing = next(
            item for item in run.iter_inner_content() if isinstance(item, Drawing)
        )
        blip = drawing._drawing.xpath(".//pic:blipFill/a:blip")[0]
        blip.set(qn("r:link"), "rId99")
        with pytest.raises(UnsupportedStructureError, match="linked picture"):
            drawing.replace_picture(io.BytesIO(_bmp(0, 0, 255)))


class DescribeHyperlinks:
    def it_wraps_a_phrase_and_can_retarget(self, tmp_path: Path):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        link = add_hyperlink(document, span, "https://example.com/a")
        assert link.address == "https://example.com/a"
        link.address = "https://example.com/b"
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        addresses = [
            hyperlink.address
            for paragraph in reopened.paragraphs
            for hyperlink in paragraph.hyperlinks
        ]
        assert "https://example.com/b" in addresses

    def it_puts_the_relationship_on_the_header_part(self, tmp_path: Path):
        document = _doc()
        header = document.sections[0].header
        header.paragraphs[0].text = "HeaderLinkUnique"
        span = find_one(document, "HeaderLinkUnique", story="word/header1.xml")
        link = add_hyperlink(document, span, "https://example.com/header")
        header_part = document.sections[0].header.part
        r_id = link._hyperlink.get(qn("r:id"))
        assert r_id in header_part.rels
        assert header_part.rels[r_id].reltype == RT.HYPERLINK
        assert header_part.rels[r_id].target_ref == "https://example.com/header"
        assert link.address == "https://example.com/header"
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        addresses = [
            hyperlink.address
            for paragraph in reopened.sections[0].header.paragraphs
            for hyperlink in paragraph.hyperlinks
        ]
        assert "https://example.com/header" in addresses

    def it_refuses_to_nest_a_hyperlink(self):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        add_hyperlink(document, span, "https://example.com/a")
        again = find_one(document, "perfectly ordinary")
        with pytest.raises(UnsupportedStructureError, match="already inside a hyperlink"):
            add_hyperlink(document, again, "https://example.com/b")

    def it_refuses_intervening_markup_between_runs(self):
        document = _doc()
        paragraph = document.add_paragraph()
        first = paragraph.add_run("hello ")
        second = paragraph.add_run("world")
        first._r.addnext(
            parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="1" w:name="Term"/>')
        )
        second._r.addnext(parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="1"/>'))
        before = paragraph._p.xml
        with pytest.raises(UnsupportedStructureError, match="intervening markup"):
            add_hyperlink(document, find_one(document, "hello world"), "https://example.com/a")
        assert paragraph._p.xml == before

    def it_clears_a_stale_internal_anchor_when_retargeting(self):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        link = add_hyperlink(document, span, "https://example.com/a")
        link._hyperlink.anchor = "_TocOld"
        link.address = "https://example.com/b"
        assert link._hyperlink.anchor is None
        assert link.address == "https://example.com/b"

    def it_refuses_a_span_from_another_document(self):
        document = _doc()
        foreign = find_one(_doc(), "perfectly ordinary")
        with pytest.raises(BoundaryViolationError, match="different document"):
            add_hyperlink(document, foreign, "https://example.com/a")

    def it_refuses_a_stale_span(self):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        find_one(document, "perfectly ordinary").replace("changed")
        with pytest.raises(TargetNotFoundError, match="stale"):
            add_hyperlink(document, span, "https://example.com/a")

    def it_keeps_a_shared_relationship_when_retargeting_one(self):
        document = _doc()
        first = add_hyperlink(
            document, find_one(document, "perfectly ordinary"), "https://example.com/a"
        )
        second = add_hyperlink(
            document, find_one(document, "First body paragraph"), "https://example.com/a"
        )
        assert first._hyperlink.get(qn("r:id")) == second._hyperlink.get(qn("r:id"))
        first.address = "https://example.com/b"
        assert first.address == "https://example.com/b"
        assert second.address == "https://example.com/a"

    def it_refuses_a_field_result_span(self):
        document = _doc("generated/feature-isolated/fields.docx")
        span = find_one(document, "June 1, 2026")
        assert span.in_field
        with pytest.raises(UnsupportedStructureError, match="field result"):
            add_hyperlink(document, span, "https://example.com/a")

    def it_defines_the_hyperlink_character_style(self):
        document = _doc()
        assert "Hyperlink" not in document.styles
        add_hyperlink(document, find_one(document, "perfectly ordinary"), "https://example.com/a")
        assert "Hyperlink" in document.styles

    def it_refuses_a_data_bound_span(self):
        document = _doc()
        document.element.body.insert(
            len(document.element.body) - 1,
            parse_xml(
                f'<w:p {nsdecls("w")}>'
                '<w:sdt><w:sdtPr><w:tag w:val="bound"/>'
                '<w:dataBinding w:xpath="/x" w:storeItemID="{11111111-1111-1111-1111-111111111111}"/>'
                "</w:sdtPr>"
                "<w:sdtContent><w:r><w:t>BoundLink</w:t></w:r></w:sdtContent>"
                "</w:sdt></w:p>"
            ),
        )
        with pytest.raises(UnsupportedStructureError, match="data-bound"):
            add_hyperlink(document, find_one(document, "BoundLink"), "https://example.com/a")

    def it_refuses_a_locked_control_span(self):
        document = _doc()
        document.element.body.insert(
            len(document.element.body) - 1,
            parse_xml(
                f'<w:p {nsdecls("w")}>'
                '<w:sdt><w:sdtPr><w:tag w:val="locked"/>'
                '<w:lock w:val="contentLocked"/></w:sdtPr>'
                "<w:sdtContent><w:r><w:t>LockedLink</w:t></w:r></w:sdtContent>"
                "</w:sdt></w:p>"
            ),
        )
        with pytest.raises(UnsupportedStructureError, match="locked"):
            add_hyperlink(document, find_one(document, "LockedLink"), "https://example.com/a")

    def it_wraps_text_inside_an_unlocked_control(self):
        document = _doc()
        document.element.body.insert(
            len(document.element.body) - 1,
            parse_xml(
                f'<w:p {nsdecls("w")}>'
                '<w:sdt><w:sdtPr><w:tag w:val="plain"/></w:sdtPr>'
                "<w:sdtContent><w:r><w:t>PlainLink</w:t></w:r></w:sdtContent>"
                "</w:sdt></w:p>"
            ),
        )
        link = add_hyperlink(
            document, find_one(document, "PlainLink"), "https://example.com/a"
        )
        assert link.address == "https://example.com/a"

    def it_refuses_a_plain_text_control_span(self):
        document = _doc()
        document.element.body.insert(
            len(document.element.body) - 1,
            parse_xml(
                f'<w:p {nsdecls("w")}>'
                '<w:sdt><w:sdtPr><w:tag w:val="plain-text"/><w:text/></w:sdtPr>'
                "<w:sdtContent><w:r><w:t>PlainTextLink</w:t></w:r></w:sdtContent>"
                "</w:sdt></w:p>"
            ),
        )
        with pytest.raises(UnsupportedStructureError, match="plain-text"):
            add_hyperlink(
                document, find_one(document, "PlainTextLink"), "https://example.com/a"
            )


class DescribeCaptions:
    def it_inserts_a_seq_field(self):
        document = _doc()
        paragraph = document.add_paragraph()
        add_caption(paragraph, label="Figure", description="Diagram")
        xml = paragraph._p.xml
        assert "SEQ Figure" in xml
        assert paragraph._p.style == "Caption"


class DescribeNotes:
    def it_adds_a_footnote_and_endnote(self, tmp_path: Path):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        footnote_id = add_footnote(document, span, "A footnote.")
        endnote_id = add_endnote(document, span, "An endnote.")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        footnotes = reopened.part.part_related_by(RT.FOOTNOTES)._element
        endnotes = reopened.part.part_related_by(RT.ENDNOTES)._element
        assert str(footnote_id) in {
            note.get(qn("w:id")) for note in footnotes.findall(qn("w:footnote"))
        }
        assert str(endnote_id) in {
            note.get(qn("w:id")) for note in endnotes.findall(qn("w:endnote"))
        }

    def it_keeps_ampersands_and_angles_as_text(self, tmp_path: Path):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        add_footnote(document, span, "A & B <C>")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        footnotes = reopened.part.part_related_by(RT.FOOTNOTES)._element
        texts = [node.text or "" for node in footnotes.findall(".//" + qn("w:t"))]
        assert any("A & B <C>" in text for text in texts)

    def it_refuses_a_note_outside_the_body(self):
        document = _doc()
        document.sections[0].header.paragraphs[0].text = "HeaderNoteUnique"
        span = find_one(document, "HeaderNoteUnique", story="word/header1.xml")
        with pytest.raises(UnsupportedStructureError, match="main document body"):
            add_footnote(document, span, "nope")

    def it_refuses_a_span_from_another_document(self):
        document = _doc()
        foreign = find_one(_doc(), "perfectly ordinary")
        with pytest.raises(BoundaryViolationError, match="different document"):
            add_footnote(document, foreign, "nope")

    def it_refuses_a_stale_span(self):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        find_one(document, "perfectly ordinary").replace("changed")
        with pytest.raises(TargetNotFoundError, match="stale"):
            add_footnote(document, span, "nope")

    def it_refuses_a_note_inside_a_text_box(self):
        document = _doc("generated/feature-isolated/textbox.docx")
        span = find_one(document, "living inside the text box")
        assert span.in_text_box
        with pytest.raises(UnsupportedStructureError, match="text box"):
            add_footnote(document, span, "nope")

    def it_refuses_a_note_inside_a_field_result(self):
        document = _doc("generated/feature-isolated/fields.docx")
        span = find_one(document, "June 1, 2026")
        assert span.in_field
        with pytest.raises(UnsupportedStructureError, match="field result"):
            add_footnote(document, span, "nope")

    def it_keeps_stacked_note_marks_in_insertion_order(self):
        document = _doc()
        span = find_one(document, "perfectly ordinary")
        first = add_footnote(document, span, "First")
        second = add_footnote(document, span, "Second")
        refs = [
            node.get(qn("w:id"))
            for node in document.element.iter(qn("w:footnoteReference"))
        ]
        assert refs == [str(first), str(second)]

    def it_refuses_a_data_bound_span(self):
        document = _doc()
        document.element.body.insert(
            len(document.element.body) - 1,
            parse_xml(
                f'<w:p {nsdecls("w")}>'
                '<w:sdt><w:sdtPr><w:tag w:val="bound"/>'
                '<w:dataBinding w:xpath="/x" w:storeItemID="{11111111-1111-1111-1111-111111111111}"/>'
                "</w:sdtPr>"
                "<w:sdtContent><w:r><w:t>BoundNote</w:t></w:r></w:sdtContent>"
                "</w:sdt></w:p>"
            ),
        )
        with pytest.raises(UnsupportedStructureError, match="data-bound"):
            add_footnote(document, find_one(document, "BoundNote"), "nope")

    def it_refuses_a_locked_control_span(self):
        document = _doc()
        document.element.body.insert(
            len(document.element.body) - 1,
            parse_xml(
                f'<w:p {nsdecls("w")}>'
                '<w:sdt><w:sdtPr><w:tag w:val="locked"/>'
                '<w:lock w:val="contentLocked"/></w:sdtPr>'
                "<w:sdtContent><w:r><w:t>LockedNote</w:t></w:r></w:sdtContent>"
                "</w:sdt></w:p>"
            ),
        )
        with pytest.raises(UnsupportedStructureError, match="locked"):
            add_footnote(document, find_one(document, "LockedNote"), "nope")

    def it_adds_a_note_inside_an_unlocked_control(self):
        document = _doc()
        document.element.body.insert(
            len(document.element.body) - 1,
            parse_xml(
                f'<w:p {nsdecls("w")}>'
                '<w:sdt><w:sdtPr><w:tag w:val="rich"/></w:sdtPr>'
                "<w:sdtContent><w:r><w:t>RichNote</w:t></w:r></w:sdtContent>"
                "</w:sdt></w:p>"
            ),
        )
        note_id = add_footnote(document, find_one(document, "RichNote"), "ok")
        refs = [
            node.get(qn("w:id"))
            for node in document.element.iter(qn("w:footnoteReference"))
        ]
        assert refs == [str(note_id)]

    def it_refuses_a_plain_text_control_span(self):
        document = _doc()
        document.element.body.insert(
            len(document.element.body) - 1,
            parse_xml(
                f'<w:p {nsdecls("w")}>'
                '<w:sdt><w:sdtPr><w:tag w:val="plain-text"/><w:text/></w:sdtPr>'
                "<w:sdtContent><w:r><w:t>PlainTextNote</w:t></w:r></w:sdtContent>"
                "</w:sdt></w:p>"
            ),
        )
        with pytest.raises(UnsupportedStructureError, match="plain-text"):
            add_footnote(document, find_one(document, "PlainTextNote"), "nope")


class DescribeDataBoundControls:
    # -- A binding's xpath is document content, not a caller argument, so every way it can be
    # -- unusable has to speak as a typed refusal. Word treats a binding it cannot evaluate as
    # -- an unbound control -- it opens the document and shows the body placeholder -- so these
    # -- refusals must not claim the document is damaged.
    @pytest.mark.parametrize(
        "xpath",
        [
            "/ns0:root/[[[",                    # malformed syntax
            "bogus-fn(/ns0:root)",              # unregistered function
            "count(/ns0:root/ns0:name)",        # evaluates to a number
            "string(/ns0:root/ns0:name)",       # evaluates to a string
            "boolean(/ns0:root)",               # evaluates to a boolean
        ],
    )
    def it_refuses_an_xpath_that_cannot_name_a_node(self, xpath: str):
        document = _doc()
        _attach_bound_control(document, "bound", xpath=xpath)
        before = etree.tostring(document.element)

        with pytest.raises(UnsupportedStructureError, match="nothing was changed"):
            set_control_value(document, "new", tag="bound")

        assert etree.tostring(document.element) == before
        assert get_control(document, tag="bound").value == "old"

    def it_says_the_binding_is_unusable_without_calling_the_document_damaged(self):
        document = _doc()
        _attach_bound_control(document, "bound", xpath="/ns0:root/[[[")

        with pytest.raises(UnsupportedStructureError) as exc:
            set_control_value(document, "new", tag="bound")

        message = str(exc.value)
        assert "not a valid XPath expression" in message
        assert "/ns0:root/[[[" in message
        # -- Word opens these documents without complaint, so the message must not imply
        # -- corruption or send the caller to repair the file
        assert "corrupt" not in message.lower()
        assert "w:dataBinding expression" in message

    def it_writes_the_custom_xml_store(self, tmp_path: Path):
        document = _doc()
        _attach_bound_control(document, "bound")
        set_control_value(document, "new", tag="bound")
        assert get_control(document, tag="bound").value == "new"
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        store = None
        for part in reopened.part.package.iter_parts():
            if str(part.partname) == "/customXml/item2.xml":
                element = getattr(part, "_element", None)
                store = element if element is not None else parse_xml(part.blob)
        assert store is not None
        ns = {"ns0": "http://example.com/form"}
        assert store.xpath("/ns0:root/ns0:name", namespaces=ns)[0].text == "new"

    def it_writes_the_store_when_the_package_has_binary_parts(self, tmp_path: Path):
        document = _doc()
        document.add_paragraph().add_run().add_picture(io.BytesIO(_bmp(255, 0, 0)))
        _attach_bound_control(document, "bound")
        set_control_value(document, "new", tag="bound")
        assert get_control(document, tag="bound").value == "new"
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        store = None
        for part in reopened.part.package.iter_parts():
            if str(part.partname) == "/customXml/item2.xml":
                element = getattr(part, "_element", None)
                store = element if element is not None else parse_xml(part.blob)
        assert store is not None
        ns = {"ns0": "http://example.com/form"}
        assert store.xpath("/ns0:root/ns0:name", namespaces=ns)[0].text == "new"

    def it_refuses_a_locked_data_bound_control(self):
        document = _doc()
        _attach_bound_control(document, "bound-locked", locked=True)
        with pytest.raises(UnsupportedStructureError, match="locked"):
            set_control_value(document, "new", tag="bound-locked")
        assert get_control(document, tag="bound-locked").value == "old"

    def it_clears_xsi_nil_when_writing_the_store(self):
        document = _doc()
        _attach_bound_control(document, "bound")
        store = None
        for part in document.part.package.iter_parts():
            if str(part.partname) == "/customXml/item2.xml":
                store = part._element
        ns = {"ns0": "http://example.com/form"}
        node = store.xpath("/ns0:root/ns0:name", namespaces=ns)[0]
        node.set("{http://www.w3.org/2001/XMLSchema-instance}nil", "true")
        set_control_value(document, "new", tag="bound")
        assert node.get("{http://www.w3.org/2001/XMLSchema-instance}nil") is None
        assert node.text == "new"

    def it_refuses_a_data_bound_checkbox(self):
        document = _doc()
        _attach_bound_control(
            document,
            "bound-box",
            extra_pr='<w14:checkbox><w14:checked w14:val="0"/></w14:checkbox>',
        )
        with pytest.raises(UnsupportedStructureError, match="checkbox"):
            set_control_value(document, "true", tag="bound-box")
        assert get_control(document, tag="bound-box").value is False

    def it_refuses_a_data_bound_picture(self):
        document = _doc()
        _attach_bound_control(document, "bound-pic", extra_pr="<w:picture/>")
        with pytest.raises(UnsupportedStructureError, match="picture"):
            set_control_value(document, "new", tag="bound-pic")

    def it_writes_an_attribute_binding(self):
        document = _doc()
        _attach_bound_control(
            document,
            "bound-attr",
            xpath="/ns0:root/ns0:name/@status",
            item_xml=(
                '<ns0:root xmlns:ns0="http://example.com/form">'
                '<ns0:name status="old">keep</ns0:name></ns0:root>'
            ),
        )
        set_control_value(document, "new", tag="bound-attr")
        store = None
        for part in document.part.package.iter_parts():
            if str(part.partname) == "/customXml/item2.xml":
                store = part._element
        ns = {"ns0": "http://example.com/form"}
        node = store.xpath("/ns0:root/ns0:name", namespaces=ns)[0]
        assert node.get("status") == "new"
        assert node.text == "keep"

    def it_writes_the_store_after_reopen(self, tmp_path: Path):
        document = _doc()
        _attach_bound_control(document, "bound")
        reopened = save_and_reopen(document, tmp_path / "in.docx")
        set_control_value(reopened, "new", tag="bound")
        saved = save_and_reopen(reopened, tmp_path / "out.docx")
        store = None
        for part in saved.part.package.iter_parts():
            if str(part.partname) == "/customXml/item2.xml":
                element = getattr(part, "_element", None)
                store = element if element is not None else parse_xml(part.blob)
        assert store is not None
        ns = {"ns0": "http://example.com/form"}
        assert store.xpath("/ns0:root/ns0:name", namespaces=ns)[0].text == "new"

    def it_parses_blob_stores_without_resolving_entities(self):
        class _BlobPart:
            content_type = "application/xml"
            blob = b'<!DOCTYPE root [<!ENTITY e "EXPANDED">]><root>&e;</root>'

        root = _part_root(_BlobPart())
        assert (root.text or "") != "EXPANDED"

    def it_writes_the_store_using_the_item_namespaces(self):
        document = _doc()
        _attach_bound_control(document, "bound", prefix_mappings=None)
        set_control_value(document, "new", tag="bound")
        store = None
        for part in document.part.package.iter_parts():
            if str(part.partname) == "/customXml/item2.xml":
                store = part._element
        ns = {"ns0": "http://example.com/form"}
        assert store.xpath("/ns0:root/ns0:name", namespaces=ns)[0].text == "new"

    def it_writes_the_store_using_the_default_namespace(self):
        document = _doc()
        _attach_bound_control(
            document,
            "bound-default",
            prefix_mappings=None,
            xpath="/root/name",
            item_xml=(
                '<root xmlns="http://example.com/form"><name>old</name></root>'
            ),
        )
        set_control_value(document, "new", tag="bound-default")
        store = None
        for part in document.part.package.iter_parts():
            if str(part.partname) == "/customXml/item2.xml":
                store = part._element
        ns = {"ns0": "http://example.com/form"}
        assert store.xpath("/ns0:root/ns0:name", namespaces=ns)[0].text == "new"

    def it_writes_the_store_when_prefix_mappings_declare_a_default_ns(self):
        document = _doc()
        _attach_bound_control(
            document,
            "bound-default-map",
            prefix_mappings="xmlns='http://example.com/form'",
            xpath="/root/name",
            item_xml=(
                '<root xmlns="http://example.com/form"><name>old</name></root>'
            ),
        )
        set_control_value(document, "new", tag="bound-default-map")
        store = None
        for part in document.part.package.iter_parts():
            if str(part.partname) == "/customXml/item2.xml":
                store = part._element
        ns = {"ns0": "http://example.com/form"}
        assert store.xpath("/ns0:root/ns0:name", namespaces=ns)[0].text == "new"


class DescribeSourceLetterhead:
    def it_copies_source_headers_when_requested(self, tmp_path: Path):
        destination = _doc()
        source = _doc()
        source.sections[0].header.paragraphs[0].text = "Source letterhead"
        append_document(destination, source, headers="source")
        reopened = save_and_reopen(destination, tmp_path / "out.docx")
        assert "Source letterhead" in reopened.sections[-1].header.paragraphs[0].text

    def it_copies_header_images_onto_the_header_part(self, tmp_path: Path):
        destination = _doc()
        source = _doc()
        source.sections[0].header.paragraphs[0].add_run().add_picture(io.BytesIO(_bmp(0, 128, 0)))
        append_document(destination, source, headers="source")
        reopened = save_and_reopen(destination, tmp_path / "out.docx")
        header = reopened.sections[-1].header
        blips = header._element.findall(".//" + qn("a:blip"))
        assert blips
        r_id = blips[0].get(qn("r:embed"))
        assert r_id in header.part.rels
        assert header.part.rels[r_id].reltype == RT.IMAGE

    def it_copies_header_hyperlinks_onto_the_header_part(self, tmp_path: Path):
        destination = _doc()
        source = _doc()
        source.sections[0].header.paragraphs[0].text = "SourceHeaderLink"
        span = find_one(source, "SourceHeaderLink", story="word/header1.xml")
        add_hyperlink(source, span, "https://example.com/letterhead")
        append_document(destination, source, headers="source")
        reopened = save_and_reopen(destination, tmp_path / "out.docx")
        header = reopened.sections[-1].header
        addresses = [
            hyperlink.address
            for paragraph in header.paragraphs
            for hyperlink in paragraph.hyperlinks
        ]
        assert "https://example.com/letterhead" in addresses
        r_ids = [
            node.get(qn("r:id"))
            for node in header._element.findall(".//" + qn("w:hyperlink"))
        ]
        assert r_ids and r_ids[0]
        assert r_ids[0] in header.part.rels
        assert header.part.rels[r_ids[0]].reltype == RT.HYPERLINK
        assert header.part.rels[r_ids[0]].target_ref == "https://example.com/letterhead"

    def it_turns_on_even_page_headers_when_the_source_uses_them(self, tmp_path: Path):
        destination = _doc()
        source = _doc()
        source.settings.odd_and_even_pages_header_footer = True
        source.sections[0].even_page_header.paragraphs[0].text = "Even letterhead"
        append_document(destination, source, headers="source")
        reopened = save_and_reopen(destination, tmp_path / "out.docx")
        assert reopened.settings.odd_and_even_pages_header_footer
        assert "Even letterhead" in reopened.sections[-1].even_page_header.paragraphs[0].text

    def it_copies_an_even_header_the_last_source_section_inherits(self, tmp_path: Path):
        destination = _doc()
        source = _doc()
        source.settings.odd_and_even_pages_header_footer = True
        source.sections[0].even_page_header.paragraphs[0].text = "Inherited even"
        source.add_section()
        assert source.sections[-1].even_page_header.is_linked_to_previous
        _copy_letterhead(
            destination, source, CompositionReport(), len(destination.sections)
        )
        reopened = save_and_reopen(destination, tmp_path / "out.docx")
        assert "Inherited even" in reopened.sections[-1].even_page_header.paragraphs[0].text

    def it_refuses_even_odd_when_the_destination_has_more_sections(self):
        destination = _doc()
        destination.add_section()
        source = _doc()
        source.settings.odd_and_even_pages_header_footer = True
        source.sections[0].even_page_header.paragraphs[0].text = "Even letterhead"
        with pytest.raises(UnsupportedStructureError, match="document-wide"):
            append_document(destination, source, headers="source")

    def it_allows_even_odd_when_only_the_source_gained_sections(self, tmp_path: Path):
        destination = _doc()
        source = _doc()
        source.settings.odd_and_even_pages_header_footer = True
        source.sections[0].even_page_header.paragraphs[0].text = "Even letterhead"
        source.add_section()
        for child in list(source.sections[0]._sectPr):
            if child.tag in (qn("w:headerReference"), qn("w:footerReference")):
                source.sections[0]._sectPr.remove(child)
        append_document(destination, source, headers="source")
        reopened = save_and_reopen(destination, tmp_path / "out.docx")
        assert reopened.settings.odd_and_even_pages_header_footer

    def it_imports_a_custom_header_style(self, tmp_path: Path):
        destination = _doc()
        source = _doc()
        style = source.styles.add_style("LetterheadBrand", WD_STYLE_TYPE.PARAGRAPH)
        style.font.italic = True
        header = source.sections[0].header.paragraphs[0]
        header.style = style
        header.text = "Source letterhead"
        report = append_document(destination, source, headers="source")
        assert "LetterheadBrand" in report.imported_styles
        reopened = save_and_reopen(destination, tmp_path / "out.docx")
        assert reopened.sections[-1].header.paragraphs[0].style.name == "LetterheadBrand"

    def it_remaps_header_numbering(self, tmp_path: Path):
        destination = _doc()
        ensure_decimal_definition(destination)
        source = _doc()
        source_num_id = ensure_decimal_definition(source)
        header = source.sections[0].header.paragraphs[0]
        header.text = "Header item"
        apply_numbering(header, num_id=source_num_id)
        report = append_document(destination, source, headers="source")
        assert source_num_id in report.numbering_map
        reopened = save_and_reopen(destination, tmp_path / "out.docx")
        copied = [
            item
            for item in list_numbering(reopened).numbered_paragraphs
            if item.text == "Header item"
        ]
        assert copied
        assert copied[0].num_id == report.numbering_map[source_num_id]

    def it_reuses_a_body_numbering_remap_in_the_header(self, tmp_path: Path):
        destination = _doc()
        ensure_decimal_definition(destination)
        source = _doc()
        source_num_id = ensure_decimal_definition(source)
        apply_numbering(source.add_paragraph("Body item"), num_id=source_num_id)
        header = source.sections[0].header.paragraphs[0]
        header.text = "Header item"
        apply_numbering(header, num_id=source_num_id)
        report = append_document(destination, source, headers="source")
        new_id = report.numbering_map[source_num_id]
        reopened = save_and_reopen(destination, tmp_path / "out.docx")
        numbered = list_numbering(reopened).numbered_paragraphs
        body_item = next(item for item in numbered if "Body item" in item.text)
        header_item = next(item for item in numbered if item.text == "Header item")
        assert body_item.num_id == header_item.num_id == new_id

    def it_refuses_a_note_mark_in_the_source_letterhead(self):
        destination = _doc()
        source = _doc()
        source.sections[0].header.paragraphs[0].text = "Source letterhead"
        source.sections[0].header.paragraphs[0].add_run()._r.append(
            parse_xml(f'<w:footnoteReference {nsdecls("w")} w:id="1"/>')
        )
        with pytest.raises(UnsupportedStructureError, match="footnote"):
            append_document(destination, source, headers="source")

    def it_renames_a_header_bookmark_that_collides_with_the_destination(self):
        destination = _doc()
        create_bookmark(destination, find_one(destination, "perfectly ordinary"), "SharedMark")
        source = _doc()
        source.sections[0].header.paragraphs[0].text = "HeaderMark"
        create_bookmark(
            source,
            find_one(source, "HeaderMark", story="word/header1.xml"),
            "SharedMark",
        )
        report = append_document(destination, source, headers="source")
        assert report.bookmarks_renamed["SharedMark"] != "SharedMark"
