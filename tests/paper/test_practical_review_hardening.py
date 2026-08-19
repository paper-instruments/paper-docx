"""Focused regressions for the practical review-hardening proposal."""

import copy

import pytest

import docx
import docx.search as search_module
from docx.commentops import (
    _comments_extended_root,
    is_resolved,
    reply,
    resolve,
)
from docx.controls import iter_controls
from docx.errors import DocumentProtectedError, TargetNotFoundError, UnsupportedStructureError
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.search import Span, find_one, replace_all


class DescribePracticalReviewHardening:
    def it_applies_comment_range_preflight_to_document_add_comment(self):
        document = docx.Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run("field result")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        run._r.insert(0, begin)

        with pytest.raises(UnsupportedStructureError, match="field boundary"):
            document.add_comment(run, "review", author="A")

        assert len(document.comments) == 0

    def it_keeps_a_span_live_after_an_unrelated_preceding_insert(self):
        document = docx.Document()
        target = document.add_paragraph("target")
        span = find_one(document, "target")
        target._p.addprevious(parse_xml(f'<w:p {nsdecls("w")}><w:r><w:t>other</w:t></w:r></w:p>'))

        span.replace("changed")

        assert document.paragraphs[-1].text == "changed"

    def it_filters_noops_and_rolls_back_a_stale_replace_batch(self, monkeypatch):
        document = docx.Document()
        document.add_paragraph("same same")
        assert replace_all(document, "same", "same").replaced_count == 0

        original = Span.replace
        calls = 0

        def fail_second(self, *args, **kwargs):
            nonlocal calls
            calls += 1
            if calls == 2:
                raise TargetNotFoundError("forced stale target")
            return original(self, *args, **kwargs)

        monkeypatch.setattr(Span, "replace", fail_second)
        before = document.element.xml
        with pytest.raises(TargetNotFoundError, match="forced stale"):
            replace_all(document, "same", "changed")
        assert document.element.xml == before

    def it_shares_one_freshness_census_across_a_replace_batch(self, monkeypatch):
        document = docx.Document()
        for _ in range(100):
            document.add_paragraph("target")
        original = search_module._story_atoms
        calls = 0

        def counted(*args, **kwargs):
            nonlocal calls
            calls += 1
            return original(*args, **kwargs)

        monkeypatch.setattr(search_module, "_story_atoms", counted)

        result = replace_all(document, "target", "changed")

        assert result.replaced_count == 100
        assert calls == 2

    def it_treats_edit_none_as_unrestricted(self):
        document = docx.Document()
        document.add_paragraph("target")
        protection = OxmlElement("w:documentProtection")
        protection.set(qn("w:edit"), "none")
        protection.set(qn("w:enforcement"), "1")
        document.settings.element.append(protection)

        find_one(document, "target").replace("changed")

        assert document.paragraphs[0].text == "changed"

    def it_matches_comment_para_ids_case_insensitively(self):
        document = docx.Document()
        comment = document.comments.add_comment("review")
        comment.paragraphs[-1]._p.set(qn("w14:paraId"), "abcdef12")
        root = _comments_extended_root(document, create=True)
        root.append(
            parse_xml(
                '<w15:commentEx xmlns:w15="http://schemas.microsoft.com/'
                'office/word/2012/wordml" w15:paraId="ABCDEF12" w15:done="1"/>'
            )
        )

        assert is_resolved(document, comment) is True
        resolve(document, comment, resolved=False)

        assert len(root) == 1
        assert is_resolved(document, comment) is False

    def it_refuses_generic_replace_on_a_typed_control_surface(self):
        document = docx.Document()
        document.element.body.insert(
            0,
            parse_xml(
                f"<w:p {nsdecls('w', 'w14')}><w:sdt><w:sdtPr>"
                "<w14:checkbox/></w:sdtPr><w:sdtContent><w:r><w:t>box</w:t>"
                "</w:r></w:sdtContent></w:sdt></w:p>"
            ),
        )
        with pytest.raises(UnsupportedStructureError, match="checkbox control"):
            find_one(document, "box").replace("changed")

    def it_refuses_generic_replace_on_a_locked_text_control(self):
        document = docx.Document()
        document.element.body.insert(
            0,
            parse_xml(
                f'<w:p {nsdecls("w")}><w:sdt><w:sdtPr>'
                '<w:text/><w:lock w:val="contentLocked"/></w:sdtPr>'
                '<w:sdtContent><w:r><w:t>locked</w:t></w:r></w:sdtContent></w:sdt></w:p>'
            ),
        )

        with pytest.raises(UnsupportedStructureError, match="locked"):
            find_one(document, "locked").replace("changed")

        control = next(iter_controls(document))
        with pytest.raises(UnsupportedStructureError, match="locked"):
            control.set_value("changed")

    def it_applies_document_protection_to_inherited_comment_mutators(self):
        document = docx.Document()
        comments = document.comments
        protection = OxmlElement("w:documentProtection")
        protection.set(qn("w:edit"), "readOnly")
        protection.set(qn("w:enforcement"), "1")
        document.settings.element.append(protection)

        with pytest.raises(DocumentProtectedError):
            comments.add_comment("blocked", author="A")

    def it_rolls_back_inherited_comment_authoring_on_invalid_text(self):
        document = docx.Document()
        comments = document.comments

        with pytest.raises(ValueError, match="XML"):
            comments.add_comment("invalid \x00 text", author="A")

        assert len(comments) == 0

    def it_refuses_a_comment_range_that_contains_a_field_boundary(self):
        document = docx.Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        run._r.append(begin)
        run.add_text("target")

        with pytest.raises(UnsupportedStructureError, match="field boundary"):
            find_one(document, "target").comment("review", author="A")

        assert len(document.comments) == 0

    def it_revalidates_an_existing_comment_range_before_replying(self):
        document = docx.Document()
        paragraph = document.add_paragraph("target")
        comment = find_one(document, "target").comment("review", author="A")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        paragraph.runs[0]._r.insert(0, begin)

        with pytest.raises(UnsupportedStructureError, match="field boundary"):
            reply(document, comment, "late", author="B")

        assert len(document.comments) == 1

    def it_refuses_an_ambiguous_duplicate_comment_anchor(self):
        document = docx.Document()
        paragraph = document.add_paragraph("target")
        comment = find_one(document, "target").comment("review", author="A")
        start = next(paragraph._p.iter(qn("w:commentRangeStart")))
        start.addnext(copy.deepcopy(start))

        with pytest.raises(UnsupportedStructureError, match="duplicate anchor"):
            reply(document, comment, "late", author="B")

        assert len(document.comments) == 1
