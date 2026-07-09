"""Regression tests for confirmed defects — at least one test per fixed
defect."""

from __future__ import annotations

import datetime as dt
import shutil
from pathlib import Path

import pytest

import docx
from docx.errors import (
    BoundaryViolationError,
    UnsupportedStructureError,
)
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.search import find_one, find_text, replace_all
from docx.story import iter_blocks, outline

from .harness.paths import fixture_path

FROZEN = dt.datetime(2026, 7, 8, 9, 0, 0, tzinfo=dt.timezone.utc)
MINIMAL = "generated/minimal-clean/minimal.docx"
TOC = "generated/feature-isolated/toc-field.docx"
PLACEHOLDER = "generated/feature-isolated/placeholder-control.docx"
NOISY = "generated/feature-isolated/noisy-markup.docx"
W = nsdecls("w")


def _doc(relpath: str = MINIMAL):
    return docx.Document(str(fixture_path(relpath)))


def _copy(relpath: str, tmp_path: Path) -> Path:
    target = tmp_path / Path(relpath).name
    shutil.copyfile(fixture_path(relpath), target)
    return target


class DescribeMultiParagraphFields:
    """The critical finding: TOC-shaped fields must stay guarded across
    paragraph boundaries, in search, story, and block ops."""

    def it_flags_continuation_paragraph_spans_as_in_field(self):
        span = find_one(_doc(TOC), "Chapter Two entry")
        assert span.in_field

    def it_flags_continuation_blocks_as_field_content(self):
        blocks = [b for b in iter_blocks(_doc(TOC)) if "Chapter Two" in b.text]
        assert blocks and blocks[0].has_field

    def it_refuses_replacing_a_toc_entry(self):
        span = find_one(_doc(TOC), "Chapter Two entry")
        with pytest.raises(UnsupportedStructureError, match="field result"):
            span.replace("Chapter 2 entry")

    def it_reports_toc_matches_as_refused_in_replace_all(self):
        result = replace_all(_doc(TOC), "Chapter Two entry", "X")
        assert result.replaced_count == 0 and len(result.refused) == 1

    def it_refuses_block_insertion_inside_the_field_result(self):
        from docx.blocks import insert_section_after

        with pytest.raises(UnsupportedStructureError, match="field result"):
            insert_section_after(
                _doc(TOC), "Chapter Two entry", heading="H", paragraphs=[]
            )

    def it_refuses_tracked_deletion_of_a_toc_entry_paragraph(self):
        from docx.blocks import tracked_delete_paragraphs

        with pytest.raises(UnsupportedStructureError, match="field result"):
            tracked_delete_paragraphs(
                _doc(TOC), "Chapter Two entry", count=1,
                author="Carol QA", date=FROZEN,
            )

    def and_it_still_allows_edits_outside_the_field(self):
        document = _doc(TOC)
        find_one(document, "Body paragraph after the TOC field.").replace(
            "Body paragraph following the TOC field."
        )


class DescribeBookmarkHollowingEndBoundary:
    def it_refuses_when_the_span_ends_exactly_at_the_bookmark_end(self):
        """Start before the bookmark, end consuming its last text node —
        previously the end marker escaped the window and the bookmark was
        silently emptied."""
        document = _doc("generated/feature-isolated/bookmarks.docx")
        span = find_one(document, "See the Master Agreement")
        with pytest.raises(UnsupportedStructureError, match="DefinedTerm"):
            span.replace("Refer to the Purchase Contract")


class DescribePlaceholderEdgeCases:
    def it_refuses_tracked_replacement_of_placeholder_text(self, tmp_path: Path):
        """A w:del of prompt text would claim it was real content."""
        document = docx.Document(str(_copy(PLACEHOLDER, tmp_path)))
        span = find_one(document, "Click or tap here to enter text.")
        with pytest.raises(UnsupportedStructureError, match="PLACEHOLDER"):
            span.replace("value", tracked=True, author="Carol QA", date=FROZEN)

    def it_refuses_partial_placeholder_fills(self, tmp_path: Path):
        document = docx.Document(str(_copy(PLACEHOLDER, tmp_path)))
        span = find_one(document, "Click or tap")
        with pytest.raises(UnsupportedStructureError, match="whole prompt"):
            span.replace("Filled")


class DescribeHyperlinkBoundaryBothModes:
    def it_refuses_untracked_replaces_crossing_the_link_boundary(self):
        document = _doc()
        body = document.element.body
        body.insert(len(body) - 1, parse_xml(
            f"<w:p {W}>"
            '<w:r><w:t xml:space="preserve">See </w:t></w:r>'
            '<w:hyperlink w:anchor="X"><w:r><w:t>Section 3</w:t></w:r></w:hyperlink>'
            '<w:r><w:t xml:space="preserve"> above.</w:t></w:r>'
            "</w:p>"
        ))
        span = find_one(document, "See Section 3")
        with pytest.raises(BoundaryViolationError, match="hyperlink boundary"):
            span.replace("Consult Section 4")


class DescribeSpanCommentOrdering:
    def it_never_reorders_run_content_around_the_anchor(self, tmp_path: Path):
        """Boundary runs holding a tab or second w:t must keep their order
        when comment range marks are isolated."""
        from docx.commentops import anchored_text

        path = _copy(MINIMAL, tmp_path)
        document = docx.Document(str(path))
        paragraph = document.add_paragraph()
        run = paragraph.add_run("Hello")
        run.add_tab()
        run.add_text("World")
        comment = find_one(document, "ell").comment(
            "note", author="Carol QA", date=FROZEN
        )
        assert anchored_text(document, comment) == "ell"
        current = next(
            b.text for b in iter_blocks(document) if "World" in b.text
        )
        assert current == "Hello\tWorld"


class DescribeCommentCleanupOnResolution:
    def it_removes_the_whole_comment_when_its_anchor_is_resolved_away(
        self, tmp_path: Path
    ):
        from docx.blocks import tracked_delete_paragraphs

        path = _copy(NOISY, tmp_path)
        document = docx.Document(str(path))
        tracked_delete_paragraphs(
            document, "This clause carries a reviewer comment.", count=1,
            author="Carol QA", date=FROZEN,
        )
        document.revisions.accept_all()
        assert len(document.comments) == 0, "orphan comment left behind"
        assert not document.element.body.xpath("//w:commentRangeStart"), (
            "orphan range markers left behind"
        )

    def it_keeps_the_comment_when_the_deletion_is_rejected(self, tmp_path: Path):
        from docx.blocks import tracked_delete_paragraphs
        from docx.commentops import anchored_text

        path = _copy(NOISY, tmp_path)
        document = docx.Document(str(path))
        tracked_delete_paragraphs(
            document, "This clause carries a reviewer comment.", count=1,
            author="Carol QA", date=FROZEN,
        )
        document.revisions.reject_all()
        (comment,) = document.comments
        assert anchored_text(document, comment) == (
            "This clause carries a reviewer comment."
        )


class DescribeControlEdgeCases:
    def it_reads_value_only_list_items(self):
        from docx.controls import get_control

        document = _doc()
        body = document.element.body
        body.insert(len(body) - 1, parse_xml(
            f'<w:p {W}><w:sdt><w:sdtPr><w:tag w:val="vo"/>'
            '<w:dropDownList><w:listItem w:value="OnlyValue"/></w:dropDownList>'
            "</w:sdtPr><w:sdtContent><w:r><w:t>x</w:t></w:r></w:sdtContent>"
            "</w:sdt></w:p>"
        ))
        assert get_control(document, tag="vo").info().choices == ("OnlyValue",)

    def it_refuses_overwriting_nested_controls(self):
        from docx.controls import set_control_value

        document = _doc()
        body = document.element.body
        body.insert(len(body) - 1, parse_xml(
            f'<w:p {W}><w:sdt><w:sdtPr><w:tag w:val="outer"/></w:sdtPr>'
            "<w:sdtContent>"
            '<w:sdt><w:sdtPr><w:tag w:val="inner"/></w:sdtPr>'
            "<w:sdtContent><w:r><w:t>inner text</w:t></w:r></w:sdtContent></w:sdt>"
            "</w:sdtContent></w:sdt></w:p>"
        ))
        with pytest.raises(UnsupportedStructureError, match="nested"):
            set_control_value(document, "flat", tag="outer")

    def it_drops_a_stale_full_date_when_setting_a_display_string(self):
        from docx.controls import set_control_value

        document = _doc()
        body = document.element.body
        body.insert(len(body) - 1, parse_xml(
            f'<w:p {W}><w:sdt><w:sdtPr><w:tag w:val="d"/>'
            '<w:date w:fullDate="2020-01-01T00:00:00Z"/></w:sdtPr>'
            "<w:sdtContent><w:r><w:t>old</w:t></w:r></w:sdtContent></w:sdt></w:p>"
        ))
        set_control_value(document, "next Tuesday", tag="d")
        (date_pr,) = body.xpath("//w:sdtPr/w:date")
        assert date_pr.get(qn("w:fullDate")) is None


class DescribeCrossParagraphCommentAnchors:
    def it_reads_the_full_anchored_text_across_paragraphs(self, tmp_path: Path):
        from docx.commentops import anchored_text

        path = _copy(MINIMAL, tmp_path)
        document = docx.Document(str(path))
        p1 = document.paragraphs[1]
        p2 = document.paragraphs[2]
        comment = document.add_comment(
            [p1.runs[0], p2.runs[0]], text="both", author="Carol QA"
        )
        assert anchored_text(document, comment) == (
            "First body paragraph with perfectly ordinary text."
            "Second body paragraph, equally unremarkable."
        )


class DescribeMiscEdgeCases:
    def it_validates_control_characters_in_table_values(self):
        from docx.tableops import insert_row_after, update_cell

        document = _doc()
        table = document.add_table(rows=1, cols=1)  # empty cell path
        with pytest.raises(ValueError, match="control character"):
            update_cell(table, 0, 0, "a\nb")
        with pytest.raises(ValueError, match="control character"):
            insert_row_after(table, 0, ["a\tb"])

    def it_counts_charts_as_embedded_objects(self):
        document = _doc()
        body = document.element.body
        body.insert(len(body) - 1, parse_xml(
            "<w:p "
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
            'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
            'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing">'
            '<w:r><w:drawing><wp:inline><wp:extent cx="1" cy="1"/>'
            '<wp:docPr id="99" name="Chart"/>'
            '<a:graphic><a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/chart"/>'
            "</a:graphic></wp:inline></w:drawing></w:r></w:p>"
        ))
        assert outline(document).blind_region_counts["embedded_objects"] == 1

    def it_narrows_cross_paragraph_spans_without_false_staleness(self):
        document = _doc()
        document.add_paragraph("Alpha beta.")
        document.add_paragraph("Gamma delta.")
        span = find_one(document, "Alpha beta. Gamma delta.")
        span.replace("Alpha beta. Gamma omega.")
        texts = [b.text for b in iter_blocks(document)]
        assert "Gamma omega." in texts and "Alpha beta." in texts

    def it_allocates_revision_ids_above_move_ids(self, tmp_path: Path):
        path = _copy("generated/feature-isolated/tracked-moves.docx", tmp_path)
        document = docx.Document(str(path))
        result = find_one(document, "Paragraph before the tracked move.").replace(
            "Paragraph ahead of the tracked move.",
            tracked=True, author="Carol QA", date=FROZEN,
        )
        assert min(result.revision_ids) > 64  # move ids run up to 64

    def it_leaves_no_phantom_insertion_after_layered_consumption(self, tmp_path: Path):
        path = _copy(MINIMAL, tmp_path)
        document = docx.Document(str(path))
        find_one(document, "perfectly ordinary").replace(
            "quite modern", tracked=True, author="Alice Editor", date=FROZEN
        )
        find_one(document, "with quite modern").replace(
            "with thoroughly modern", tracked=True,
            author="Alice Editor", date=FROZEN,
        )
        empty_ins = [
            ins
            for ins in document.element.body.xpath("//w:ins")
            if not "".join(ins.itertext())
        ]
        assert not empty_ins, "phantom empty w:ins left behind"

    def it_diagnoses_rel_resolved_main_parts(self, tmp_path: Path):
        import zipfile

        from docx.package import diagnose

        source = fixture_path(MINIMAL)
        moved = tmp_path / "moved-main.docx"
        with zipfile.ZipFile(source) as zin:
            with zipfile.ZipFile(moved, "w") as zout:
                for name in zin.namelist():
                    blob = zin.read(name)
                    if name in ("[Content_Types].xml", "_rels/.rels"):
                        blob = blob.replace(b"word/document.xml", b"word/doc2.xml")
                    if name == "word/document.xml":
                        name = "word/doc2.xml"
                    zout.writestr(name, blob)
        report = diagnose(moved)
        assert report.kind == "docx" and report.readable

    def it_pads_inserted_tables_against_fusion(self, tmp_path: Path):
        from docx.blocks import TableBlock, insert_blocks_after

        path = _copy(MINIMAL, tmp_path)
        document = docx.Document(str(path))
        document.add_table(rows=1, cols=1)  # existing table right after body end
        insert_blocks_after(
            document, "Second body paragraph",
            blocks=[TableBlock(rows=[["a"]]), TableBlock(rows=[["b"]])],
        )
        body = document.element.body
        children = [c.tag.rsplit("}", 1)[-1] for c in body]
        for index, tag in enumerate(children[:-1]):
            assert not (tag == "tbl" and children[index + 1] == "tbl"), (
                f"adjacent tables at {index}: {children}"
            )
