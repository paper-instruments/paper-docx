"""Rollback regressions for failures raised after compound mutation begins."""

from __future__ import annotations

import struct
from io import BytesIO

import pytest
from lxml import etree

import docx
import docx.revision as revision_module
from docx import commentops, composition
from docx._transaction import rollback_on_error
from docx.enum.style import WD_STYLE_TYPE
from docx.errors import UnsupportedStructureError
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml.text.run import CT_R
from docx.search import find_one
from docx.text.run import Run

from .harness.contract import assert_refusal_atomic


def _late_refusal(*args, **kwargs):  # noqa: ANN002, ANN003
    raise UnsupportedStructureError("forced late refusal")


def _tracked_insert(revision_id: int, text: str):
    return parse_xml(
        f'<w:ins {nsdecls("w")} w:id="{revision_id}" w:author="Audit">'
        f"<w:r><w:t>{text}</w:t></w:r></w:ins>"
    )


def _two_revision_document():
    document = docx.Document()
    paragraph = document.add_paragraph()._p
    paragraph.append(_tracked_insert(1, "first"))
    paragraph.append(_tracked_insert(2, "second"))
    return document


def _bmp(red: int, green: int, blue: int) -> bytes:
    return (
        struct.pack("<2sIHHI", b"BM", 58, 0, 0, 54)
        + struct.pack("<IiiHHIIiiII", 40, 1, 1, 1, 24, 0, 4, 2835, 2835, 0, 0)
        + bytes((blue, green, red, 0))
    )


class DescribeTransactionKernel:
    def it_does_not_reparent_unchanged_namespace_scopes(self):
        document = docx.Document()
        node = etree.fromstring(
            b'<x:child xmlns:x="urn:audit" xmlns:y="urn:audit" y:value="1" ref="y:Thing"/>'
        )
        document.element.body.append(node)
        before = etree.tostring(node)

        refusal = pytest.raises(UnsupportedStructureError)
        with refusal, rollback_on_error(document):  # noqa: PT012
            raise UnsupportedStructureError("forced late refusal")

        assert etree.tostring(node) == before
        assert node.get("ref") == "y:Thing"
        assert node.nsmap["y"] == "urn:audit"

    def it_restores_element_lazy_state(self):
        document = docx.Document()
        paragraph = document.add_paragraph("state")._p
        paragraph.__dict__["audit_cached"] = "before"

        refusal = pytest.raises(UnsupportedStructureError)
        with refusal, rollback_on_error(document):  # noqa: PT012
            paragraph.__dict__["audit_cached"] = "after"
            paragraph.__dict__["audit_new"] = True
            raise UnsupportedStructureError("forced late refusal")

        assert paragraph.__dict__["audit_cached"] == "before"
        assert "audit_new" not in paragraph.__dict__

    def it_restores_cache_only_and_image_list_only_parts(self):
        document = docx.Document()
        package = document.part.package
        assert package is not None
        cached_part = Part(
            PackURI("/word/audit-cache.bin"),
            "application/octet-stream",
            b"cached-before",
            package,
        )
        document.part.rels._target_parts_by_rId["rIdAudit"] = cached_part
        image_part = package.get_or_add_image_part(BytesIO(_bmp(1, 2, 3)))
        image_parts = package.image_parts._image_parts
        image_items = tuple(image_parts)
        image_blob = image_part.blob

        refusal = pytest.raises(UnsupportedStructureError)
        with refusal, rollback_on_error(document):  # noqa: PT012
            cached_part._blob = b"cached-after"
            image_part._blob = b"image-after"
            raise UnsupportedStructureError("forced late refusal")

        assert cached_part.blob == b"cached-before"
        assert image_part.blob == image_blob
        assert package.image_parts._image_parts is image_parts
        assert len(image_parts) == len(image_items)
        assert all(current is original for current, original in zip(image_parts, image_items))

    def it_restores_mutated_relationship_records(self):
        document = docx.Document()
        relationship = next(rel for rel in document.part.rels.values() if not rel.is_external)
        original_target = relationship.target_part
        replacement = Part(
            PackURI("/word/audit-target.bin"),
            "application/octet-stream",
            b"replacement",
            document.part.package,
        )

        refusal = pytest.raises(UnsupportedStructureError)
        with refusal, rollback_on_error(document):  # noqa: PT012
            relationship._target = replacement
            raise UnsupportedStructureError("forced late refusal")

        assert relationship.target_part is original_target
        assert document.part.rels.related_parts[relationship.rId] is original_target
        assert replacement.package is None


class DescribeCompositionRollback:
    def it_removes_a_style_imported_before_a_late_refusal(self, monkeypatch):
        source = docx.Document()
        style = source.styles.add_style("LateFailureStyle", WD_STYLE_TYPE.PARAGRAPH)
        source.add_paragraph("source", style=style)
        destination = docx.Document()
        anchor = destination.add_paragraph("anchor")
        anchor_element = anchor._p
        monkeypatch.setattr(composition, "_remap_numbering", _late_refusal)

        assert_refusal_atomic(
            destination,
            lambda document: composition.insert_blocks_from(
                document, source, "source", anchor="anchor"
            ),
            UnsupportedStructureError,
        )

        assert "LateFailureStyle" not in {item.name for item in destination.styles}
        assert destination.paragraphs[0]._p is anchor_element
        assert anchor.text == "anchor"

    def it_removes_new_media_and_image_bookkeeping(self, monkeypatch):
        source = docx.Document()
        paragraph = source.add_paragraph("source image")
        paragraph.add_run().add_picture(BytesIO(_bmp(255, 0, 0)))
        destination = docx.Document()
        destination.add_paragraph("anchor")
        package = destination.part.package
        assert package is not None
        assert "image_parts" not in package.__dict__
        related_parts = destination.part.rels.related_parts
        relationships_before = dict(related_parts)
        monkeypatch.setattr(composition, "_recreate_hyperlinks", _late_refusal)

        assert_refusal_atomic(
            destination,
            lambda document: composition.insert_blocks_from(
                document, source, "source image", anchor="anchor"
            ),
            UnsupportedStructureError,
        )

        assert "image_parts" not in package.__dict__
        assert destination.part.rels.related_parts is related_parts
        assert destination.part.rels.related_parts == relationships_before
        assert not any("/media/" in str(part.partname) for part in package.iter_parts())

    def it_preserves_an_existing_image_parts_collection(self, monkeypatch):
        source = docx.Document()
        source_paragraph = source.add_paragraph("source image")
        source_paragraph.add_run().add_picture(BytesIO(_bmp(255, 0, 0)))
        destination = docx.Document()
        destination.add_paragraph().add_run().add_picture(BytesIO(_bmp(0, 0, 255)))
        destination.add_paragraph("anchor")
        package = destination.part.package
        assert package is not None
        image_parts = package.image_parts._image_parts
        image_items = tuple(image_parts)
        monkeypatch.setattr(composition, "_recreate_hyperlinks", _late_refusal)

        assert_refusal_atomic(
            destination,
            lambda document: composition.insert_blocks_from(
                document, source, "source image", anchor="anchor"
            ),
            UnsupportedStructureError,
        )

        assert package.image_parts._image_parts is image_parts
        assert len(image_parts) == len(image_items)
        assert all(current is original for current, original in zip(image_parts, image_items))

    def it_rolls_back_append_after_compose_returns(self, monkeypatch):
        source = docx.Document()
        source.add_paragraph("appended content")
        destination = docx.Document()
        anchor = destination.add_paragraph("anchor")
        anchor_element = anchor._p
        original = composition._compose

        def compose_then_refuse(*args, **kwargs):  # noqa: ANN002, ANN003
            original(*args, **kwargs)
            raise UnsupportedStructureError("forced late refusal")

        monkeypatch.setattr(composition, "_compose", compose_then_refuse)

        assert_refusal_atomic(
            destination,
            lambda document: composition.append_document(document, source),
            UnsupportedStructureError,
        )

        assert destination.paragraphs[0]._p is anchor_element
        assert anchor.text == "anchor"


class DescribeRevisionRollback:
    @pytest.mark.parametrize("operation", ["accept_all", "reject_all"])
    def it_restores_an_earlier_revision_when_a_batch_fails(self, monkeypatch, operation: str):
        document = _two_revision_document()
        revisions = document.revisions
        original_elements = tuple(item._element for item in revisions)
        original = revision_module._resolve_one
        calls = 0

        def resolve_then_refuse(node, *, accept, document):
            nonlocal calls
            calls += 1
            if calls == 2:
                raise UnsupportedStructureError("forced late refusal")
            return original(node, accept=accept, document=document)

        monkeypatch.setattr(revision_module, "_resolve_one", resolve_then_refuse)

        assert_refusal_atomic(
            document,
            lambda _document: getattr(revisions, operation)(),
            UnsupportedStructureError,
        )

        current = document.revisions
        assert len(current) == 2
        assert all(
            current_item._element is original_element
            for current_item, original_element in zip(current, original_elements)
        )

    def it_restores_a_single_revision_after_its_mutation(self, monkeypatch):
        document = _two_revision_document()
        revision = document.revisions[0]
        revision_element = revision._element
        original = revision_module._resolve_one

        def resolve_then_refuse(node, *, accept, document):
            original(node, accept=accept, document=document)
            raise UnsupportedStructureError("forced late refusal")

        monkeypatch.setattr(revision_module, "_resolve_one", resolve_then_refuse)

        assert_refusal_atomic(
            document,
            lambda _document: revision.accept(),
            UnsupportedStructureError,
        )

        assert document.revisions[0]._element is revision_element
        assert revision_element.getparent() is not None

    def it_restores_raw_xml_tags_changed_before_a_refusal(self, monkeypatch):
        document = docx.Document()
        paragraph = document.add_paragraph()._p
        deletion = parse_xml(
            f'<w:del {nsdecls("w")} w:id="1" w:author="Audit">'
            "<w:r><w:delText>deleted</w:delText></w:r></w:del>"
        )
        paragraph.append(deletion)
        deleted_text = next(deletion.iter(qn("w:delText")))
        revision = document.revisions[0]
        original = revision_module._resolve_one

        def resolve_then_refuse(node, *, accept, document):
            original(node, accept=accept, document=document)
            raise UnsupportedStructureError("forced late refusal")

        monkeypatch.setattr(revision_module, "_resolve_one", resolve_then_refuse)

        assert_refusal_atomic(
            document,
            lambda _document: revision.reject(),
            UnsupportedStructureError,
        )

        assert deleted_text.tag == qn("w:delText")
        assert deleted_text.getparent() is not None


class DescribeCommentAuthoringRollback:
    def it_refuses_comment_ranges_from_another_document(self):
        document = docx.Document()
        document.add_paragraph("destination")
        foreign = docx.Document()
        foreign_run = foreign.add_paragraph("foreign").runs[0]

        assert_refusal_atomic(
            document,
            lambda candidate: candidate.add_comment(foreign_run, "note", author="Reviewer"),
            UnsupportedStructureError,
        )

        assert foreign.paragraphs[0].text == "foreign"
        assert not tuple(foreign.element.iter(qn("w:commentRangeStart")))

    def it_refuses_a_direct_range_across_documents(self):
        document = docx.Document()
        run = document.add_paragraph("first").runs[0]
        foreign = docx.Document()
        foreign_run = foreign.add_paragraph("last").runs[0]
        foreign_xml = foreign.element.xml

        assert_refusal_atomic(
            document,
            lambda _document: run.mark_comment_range(foreign_run, 7),
            UnsupportedStructureError,
        )

        assert foreign.element.xml == foreign_xml

    def it_refuses_a_comment_range_that_ends_before_it_starts(self):
        document = docx.Document()
        paragraph = document.add_paragraph()
        first_run = paragraph.add_run("first")
        last_run = paragraph.add_run("last")

        assert_refusal_atomic(
            document,
            lambda candidate: candidate.add_comment(
                [last_run, first_run], "note", author="Reviewer"
            ),
            UnsupportedStructureError,
        )

        assert paragraph.text == "firstlast"
        assert not any(
            relationship.reltype == RT.COMMENTS for relationship in document.part.rels.values()
        )

    def it_restores_a_span_and_its_split_runs_after_a_late_refusal(self, monkeypatch):
        document = docx.Document()
        document.add_paragraph("prefix target suffix")
        span = find_one(document, "target")
        offsets = (span._start_offset, span._end_offset)
        monkeypatch.setattr(CT_R, "insert_comment_range_end_and_reference_below", _late_refusal)

        assert_refusal_atomic(
            document,
            lambda _document: span.comment("note", author="Reviewer"),
            UnsupportedStructureError,
        )

        assert (span._start_offset, span._end_offset) == offsets
        assert span.text == "target"
        span._validate_fresh()

    def it_removes_a_new_comment_when_range_authoring_refuses(self, monkeypatch):
        document = docx.Document()
        paragraph = document.add_paragraph("comment target")
        run = paragraph.runs[0]

        def mark_then_refuse(candidate: Run, last_run: Run, comment_id: int) -> None:
            candidate._r.insert_comment_range_start_above(comment_id)
            raise UnsupportedStructureError("forced late refusal")

        monkeypatch.setattr(Run, "mark_comment_range", mark_then_refuse)

        assert_refusal_atomic(
            document,
            lambda candidate: candidate.add_comment(run, "note", author="Reviewer"),
            UnsupportedStructureError,
        )

        assert paragraph.runs[0]._r is run._r
        assert not tuple(document.element.iter(qn("w:commentRangeStart")))
        assert not any(
            relationship.reltype == RT.COMMENTS for relationship in document.part.rels.values()
        )

    def it_removes_a_range_start_when_range_authoring_refuses(self, monkeypatch):
        document = docx.Document()
        paragraph = document.add_paragraph()
        first_run = paragraph.add_run("first")
        last_run = paragraph.add_run("last")
        monkeypatch.setattr(CT_R, "insert_comment_range_end_and_reference_below", _late_refusal)

        assert_refusal_atomic(
            document,
            lambda _document: first_run.mark_comment_range(last_run, 7),
            UnsupportedStructureError,
        )

        assert first_run._r.getparent() is paragraph._p
        assert last_run._r.getparent() is paragraph._p
        assert not tuple(document.element.iter(qn("w:commentRangeStart")))


class DescribeCommentThreadRollback:
    @pytest.mark.parametrize("operation", ["reply", "resolve"])
    def it_restores_all_thread_parts_after_late_refusal(self, monkeypatch, operation: str):
        document = docx.Document()
        document.add_paragraph("target")
        comment = find_one(document, "target").comment("parent", author="Reviewer")
        comment_element = comment._comment_elm
        monkeypatch.setattr(commentops, "_entry_for", _late_refusal)

        def mutate(candidate):
            if operation == "reply":
                commentops.reply(candidate, comment, "reply", author="Second Reviewer")
            else:
                commentops.resolve(candidate, comment)

        assert_refusal_atomic(
            document,
            mutate,
            UnsupportedStructureError,
        )

        restored = document.comments.get(comment.comment_id)
        assert restored is not None
        assert restored._comment_elm is comment_element
        assert len(document.comments) == 1
        assert comment.text == "parent"
