"""Revision-resolution completion.

First we pin that detection fires on every revision shape the resolution
pipeline builds on (the ground we stand on); then we pin the resolution
semantics themselves against the hand-computed accepted ground truth.
"""

from __future__ import annotations

from pathlib import Path

import pytest

import docx
from docx.errors import UnsupportedStructureError

from .harness.contract import save_and_reopen
from .harness.paths import fixture_path

MULTIROUND = "generated/redline/multiround.docx"
MULTIROUND_ACCEPTED = "generated/redline/multiround-accepted.docx"
ROW_REVISIONS = "generated/feature-isolated/row-revisions.docx"
FORMAT_RICH = "generated/feature-isolated/format-changes-rich.docx"
PARAGRAPH_MERGE = "generated/feature-isolated/paragraph-merge.docx"
MINIMAL_TABLE = "generated/feature-isolated/table-merged-nested.docx"
GAUNTLET = "generated/gauntlet/gauntlet.docx"


def _doc(relpath: str):
    return docx.Document(str(fixture_path(relpath)))


class DescribePhase0Detection:
    """Detection fires on every shape the resolver later resolves."""

    def it_enumerates_every_revision_in_the_multiround_redline(self):
        revisions = _doc(MULTIROUND).revisions
        assert len(revisions) == 20
        assert {r.author for r in revisions} == {"Alice Editor", "Bob Reviewer"}

    def it_enumerates_the_move_pair_with_mark_stamps(self):
        revisions = _doc(MULTIROUND).revisions
        move_from = [r for r in revisions if r.revision_type == "move_from"]
        move_to = [r for r in revisions if r.revision_type == "move_to"]
        assert len(move_from) == 2 and len(move_to) == 2  # wrapper + mark each
        assert any(r.is_paragraph_mark for r in move_from)
        assert any(r.is_paragraph_mark for r in move_to)

    def it_enumerates_rich_format_changes_including_the_mark_change(self):
        revisions = _doc(FORMAT_RICH).revisions
        format_changes = [r for r in revisions if r.revision_type == "format_change"]
        assert len(format_changes) == 3
        assert sum(1 for r in format_changes if r.is_paragraph_mark) == 1

    def it_enumerates_both_paragraph_mark_revisions(self):
        revisions = _doc(PARAGRAPH_MERGE).revisions
        marks = [(r.revision_type, r.author) for r in revisions if r.is_paragraph_mark]
        assert sorted(marks) == [
            ("deletion", "Bob Reviewer"),
            ("insertion", "Alice Editor"),
        ]

    def it_enumerates_all_ten_row_revision_nodes(self):
        revisions = _doc(ROW_REVISIONS).revisions
        assert len(revisions) == 10  # 2 trPr markers + 4 content + 4 mark stamps
        assert {r.author for r in revisions} == {"Alice Editor", "Bob Reviewer"}

    def it_finds_zero_revisions_in_the_accepted_ground_truth(self):
        assert len(_doc(MULTIROUND_ACCEPTED).revisions) == 0


def _rescan(document) -> dict:
    from docx.revision import _remaining_markup

    return _remaining_markup(document)


def _body_texts(document) -> list:
    from docx.story import iter_blocks

    return [b.text for b in iter_blocks(document) if b.story == "word/document.xml"]


class DescribeFormatChangeResolution:
    """w:rPrChange/w:pPrChange accept and reject."""

    def it_accepts_a_run_format_change_keeping_current_properties(
        self, tmp_path: Path
    ):
        document = _doc(FORMAT_RICH)
        for revision in document.revisions:
            if not revision.is_paragraph_mark and "Delivery" in revision.text:
                revision.accept()
                break
        else:
            pytest.fail("run format change not found")
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        paragraph = next(
            p for p in reopened.paragraphs if "Delivery" in p.text
        )
        assert paragraph.runs[0].font.bold is True
        assert all(
            r.revision_type != "format_change" or "Delivery" not in r.text
            for r in reopened.revisions
        )

    def it_rejects_a_run_format_change_restoring_stored_properties(
        self, tmp_path: Path
    ):
        document = _doc(FORMAT_RICH)
        for revision in document.revisions:
            if not revision.is_paragraph_mark and "Delivery" in revision.text:
                revision.reject()
                break
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        run = next(p for p in reopened.paragraphs if "Delivery" in p.text).runs[0]
        assert run.font.bold is None  # bold dropped
        assert run.font.italic is True  # stored previous restored
        assert run.font.size is not None and run.font.size.pt == 14  # sz 28 half-points

    def it_accepts_a_paragraph_format_change(self, tmp_path: Path):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        document = _doc(FORMAT_RICH)
        for revision in document.revisions:
            if "right-aligned" in revision.text:
                revision.accept()
                break
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        paragraph = next(p for p in reopened.paragraphs if "right-aligned" in p.text)
        assert paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def it_rejects_a_paragraph_format_change_restoring_stored_properties(
        self, tmp_path: Path
    ):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        document = _doc(FORMAT_RICH)
        for revision in document.revisions:
            if "right-aligned" in revision.text:
                revision.reject()
                break
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        paragraph = next(p for p in reopened.paragraphs if "right-aligned" in p.text)
        assert paragraph.alignment == WD_ALIGN_PARAGRAPH.RIGHT

    def it_resolves_the_paragraph_mark_format_change_both_ways(
        self, tmp_path: Path
    ):
        for accept in (True, False):
            document = _doc(FORMAT_RICH)
            mark = next(r for r in document.revisions if r.is_paragraph_mark)
            assert mark.revision_type == "format_change"
            (mark.accept if accept else mark.reject)()
            reopened = save_and_reopen(document, tmp_path / f"out-{accept}.docx")
            # only the mark change resolved; the run/paragraph changes remain
            assert _rescan(reopened) == {"pPrChange": 1, "rPrChange": 1}
            texts = _body_texts(reopened)
            assert "The paragraph mark itself was re-formatted." in texts

    def it_accepts_all_and_a_rescan_finds_zero_markup(self, tmp_path: Path):
        document = _doc(FORMAT_RICH)
        assert document.revisions.accept_all() == 3
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _rescan(reopened) == {}
        assert len(reopened.revisions) == 0

    def it_rejects_all_and_a_rescan_finds_zero_markup(self, tmp_path: Path):
        document = _doc(FORMAT_RICH)
        assert document.revisions.reject_all() == 3
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _rescan(reopened) == {}


class DescribeRowRevisionResolution:
    """w:trPr row markers — the ghost-row defect closed."""

    def _grid(self, document) -> list:
        return [
            [cell.text for cell in row.cells] for row in document.tables[0].rows
        ]

    def it_accepts_all_keeping_inserted_and_dropping_deleted_rows(
        self, tmp_path: Path
    ):
        document = _doc(ROW_REVISIONS)
        document.revisions.accept_all()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert self._grid(reopened) == [["Item", "Amount"], ["Filing fee", "$100"]]
        assert _rescan(reopened) == {}

    def it_rejects_all_restoring_the_original_table(self, tmp_path: Path):
        document = _doc(ROW_REVISIONS)
        document.revisions.reject_all()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert self._grid(reopened) == [["Item", "Amount"], ["Old charge", "$50"]]
        assert _rescan(reopened) == {}

    def it_classifies_row_markers_distinctly(self):
        revisions = _doc(ROW_REVISIONS).revisions
        types = sorted(
            r.revision_type
            for r in revisions
            if r.revision_type.startswith("row_")
        )
        assert types == ["row_deletion", "row_insertion"]

    def it_accepts_a_single_row_insertion_removing_only_the_marker(
        self, tmp_path: Path
    ):
        document = _doc(ROW_REVISIONS)
        row_ins = next(
            r for r in document.revisions if r.revision_type == "row_insertion"
        )
        row_ins.accept()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert len(reopened.tables[0].rows) == 3  # row kept, marker gone
        assert "row_insertion" not in {r.revision_type for r in reopened.revisions}

    def it_refuses_a_single_row_insertion_that_would_consume_cell_revisions(
        self, tmp_path: Path
    ):
        document = _doc(ROW_REVISIONS)
        row_ins = next(
            r for r in document.revisions if r.revision_type == "row_insertion"
        )
        before = document.element.xml
        with pytest.raises(UnsupportedStructureError, match="nested unselected"):
            row_ins.reject()
        assert document.element.xml == before

        assert document.revisions.reject_all(author="Alice Editor") == 5
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        # inserted row gone; the deleted-marked row remains (its delText
        # content is invisible to upstream cell.text — still pending)
        assert len(reopened.tables[0].rows) == 2
        assert "Filing fee" not in reopened.element.xml
        assert "row_deletion" in {r.revision_type for r in reopened.revisions}

    def it_removes_the_whole_table_when_its_last_row_resolves_away(
        self, tmp_path: Path
    ):
        """Word's semantic for a fully tracked-deleted table: accepting the
        last row's deletion removes the table itself (never an invalid
        zero-row w:tbl)."""
        document = _doc(ROW_REVISIONS)
        table = document.tables[0]._tbl
        # strip the table down to ONLY the deleted-marked row
        for row in list(table.tr_lst):
            if "Old charge" not in "".join(row.itertext()):
                table.remove(row)
        document.revisions.accept_all()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert len(reopened.tables) == 0
        assert _rescan(reopened) == {}


class DescribeParagraphMarkResolution:
    """The pilcrow itself, both directions explicitly."""

    def it_accepts_a_deleted_mark_merging_with_the_next_paragraph(
        self, tmp_path: Path
    ):
        document = _doc(PARAGRAPH_MERGE)
        mark = next(
            r
            for r in document.revisions
            if r.is_paragraph_mark and r.revision_type == "deletion"
        )
        mark.accept()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        texts = _body_texts(reopened)
        assert "This sentence continues onto the following line." in texts
        assert "onto the following line." not in texts

    def it_rejects_a_deleted_mark_keeping_both_paragraphs(self, tmp_path: Path):
        document = _doc(PARAGRAPH_MERGE)
        mark = next(
            r
            for r in document.revisions
            if r.is_paragraph_mark and r.revision_type == "deletion"
        )
        mark.reject()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        texts = _body_texts(reopened)
        assert "This sentence continues " in texts
        assert "onto the following line." in texts

    def it_accepts_an_inserted_mark_keeping_the_split(self, tmp_path: Path):
        document = _doc(PARAGRAPH_MERGE)
        mark = next(
            r
            for r in document.revisions
            if r.is_paragraph_mark and r.revision_type == "insertion"
        )
        mark.accept()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        texts = _body_texts(reopened)
        assert "A tracked split divides " in texts
        assert "this once-single sentence." in texts

    def it_rejects_an_inserted_mark_merging_the_split_back(self, tmp_path: Path):
        document = _doc(PARAGRAPH_MERGE)
        mark = next(
            r
            for r in document.revisions
            if r.is_paragraph_mark and r.revision_type == "insertion"
        )
        mark.reject()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        texts = _body_texts(reopened)
        assert "A tracked split divides this once-single sentence." in texts

    def it_resolves_the_whole_fixture_both_ways_to_zero_markup(
        self, tmp_path: Path
    ):
        for accept in (True, False):
            document = _doc(PARAGRAPH_MERGE)
            resolver = (
                document.revisions.accept_all if accept else document.revisions.reject_all
            )
            assert resolver() == 2
            reopened = save_and_reopen(document, tmp_path / f"out-{accept}.docx")
            assert _rescan(reopened) == {}


def _story_texts(document) -> dict:
    from docx.story import iter_blocks

    texts: dict = {}
    for block in iter_blocks(document):
        texts.setdefault(block.story, []).append(block.text)
    return texts


TRACKED_MOVES = "generated/feature-isolated/tracked-moves.docx"
MOVED_TEXT_V01 = "The indemnity clause relocated by tracked move."


class DescribeMoveResolution:
    """Moves resolve as paired units, never one site alone."""

    def it_accepts_the_multiround_redline_to_match_the_ground_truth(
        self, tmp_path: Path
    ):
        document = _doc(MULTIROUND)
        assert document.revisions.accept_all() == 20
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _rescan(reopened) == {}
        assert len(reopened.revisions) == 0
        assert _story_texts(reopened) == _story_texts(_doc(MULTIROUND_ACCEPTED))

    def it_rejects_the_multiround_redline_back_to_the_original(
        self, tmp_path: Path
    ):
        document = _doc(MULTIROUND)
        assert document.revisions.reject_all() == 20
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _rescan(reopened) == {}
        moved = "The indemnity clause shall survive termination of this agreement."
        assert _body_texts(reopened) == [
            "Engagement letter, revised across two rounds.",
            "This agreement is made between the parties.",
            moved,  # restored at the SOURCE
            "Middle paragraph between the move sites.",
            # destination paragraph is gone entirely
            "Payment is due within thirty days of invoice.",
            "Delivery follows the schedule in Exhibit A.",
            "Item\nAmount\nOld charge\n$50",
            "This sentence continues ",  # mark deletion rejected: no merge
            "onto the following line.",
            "A tracked split divides this once-single sentence.",  # split undone
            "Signed at .",  # Alice's insertion rejected; comment went with it
            "Closing paragraph after all tracked activity.",
        ]

    def it_removes_the_rejected_insertions_comment_entirely(self, tmp_path: Path):
        document = _doc(MULTIROUND)
        document.revisions.reject_all()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        xml = reopened.element.xml
        assert "commentRangeStart" not in xml and "commentReference" not in xml

    def it_resolves_a_move_pair_from_the_source_site(self, tmp_path: Path):
        document = _doc(TRACKED_MOVES)
        move_from = next(
            r for r in document.revisions if r.revision_type == "move_from"
        )
        move_from.accept()  # resolving EITHER site resolves the pair
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _rescan(reopened) == {}
        body = "\n".join(_body_texts(reopened))
        assert body.count(MOVED_TEXT_V01) == 1
        # destination follows the middle paragraph; source paragraph emptied
        # (no mark stamps in this fixture, so the empty shell remains)
        texts = _body_texts(reopened)
        assert texts.index("Paragraph between the move ends.") < texts.index(
            MOVED_TEXT_V01
        )

    def it_rejects_a_move_pair_from_the_destination_site(self, tmp_path: Path):
        document = _doc(TRACKED_MOVES)
        move_to = next(
            r for r in document.revisions if r.revision_type == "move_to"
        )
        move_to.reject()
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _rescan(reopened) == {}
        texts = _body_texts(reopened)
        assert "\n".join(texts).count(MOVED_TEXT_V01) == 1
        assert texts.index(MOVED_TEXT_V01) < texts.index(
            "Paragraph between the move ends."
        )

    def it_reports_moves_as_resolvable_in_the_census(self):
        assert _doc(TRACKED_MOVES).revisions.remaining_unsupported() == {}

    def it_refuses_an_orphaned_move_atomically(self):
        document = _doc(TRACKED_MOVES)
        body = document.element.body
        destination = next(
            p
            for p in body.findall(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p"
            )
            if p.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}moveTo"
            )
            is not None
        )
        body.remove(destination)
        before = document.element.xml
        with pytest.raises(UnsupportedStructureError, match="move"):
            document.revisions.accept_all()
        assert document.element.xml == before

    def it_refuses_a_cross_story_move(self):
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        w = nsdecls("w")
        document = _doc("generated/feature-isolated/header-footer-sections.docx")
        body_p = document.paragraphs[0]._p
        body_p.addprevious(
            parse_xml(
                f'<w:moveFromRangeStart {w} w:id="801" w:author="A"'
                ' w:date="2026-06-01T09:30:00Z" w:name="crossMove"/>'
            )
        )
        body_p.append(
            parse_xml(
                f'<w:moveFrom {w} w:id="802" w:author="A"'
                ' w:date="2026-06-01T09:30:00Z">'
                "<w:r><w:t>wandering text</w:t></w:r></w:moveFrom>"
            )
        )
        body_p.addnext(parse_xml(f'<w:moveFromRangeEnd {w} w:id="801"/>'))
        header_p = document.sections[0].header.paragraphs[0]._p
        header_p.addprevious(
            parse_xml(
                f'<w:moveToRangeStart {w} w:id="803" w:author="A"'
                ' w:date="2026-06-01T09:30:00Z" w:name="crossMove"/>'
            )
        )
        header_p.append(
            parse_xml(
                f'<w:moveTo {w} w:id="804" w:author="A"'
                ' w:date="2026-06-01T09:30:00Z">'
                "<w:r><w:t>wandering text</w:t></w:r></w:moveTo>"
            )
        )
        header_p.addnext(parse_xml(f'<w:moveToRangeEnd {w} w:id="803"/>'))
        with pytest.raises(UnsupportedStructureError, match="crosses stories"):
            document.revisions.accept_all()


class DescribeGauntletResolution:
    """The everything-document resolves completely, both ways."""

    @pytest.mark.parametrize("accept", [True, False], ids=["accept", "reject"])
    def it_resolves_the_entire_gauntlet_to_zero_markup(
        self, accept: bool, tmp_path: Path
    ):
        document = _doc(GAUNTLET)
        resolver = (
            document.revisions.accept_all if accept else document.revisions.reject_all
        )
        assert resolver() > 0
        reopened = save_and_reopen(document, tmp_path / "out.docx")
        assert _rescan(reopened) == {}
        assert len(reopened.revisions) == 0


class DescribeExoticTypesStayRefused:
    """Everything not resolved is enumerated and refused BY NAME."""

    @pytest.mark.parametrize(
        ("markup", "expected_type"),
        [
            (
                '<w:tblPrChange w:id="900" w:author="A" w:date="2026-06-01T09:30:00Z">'
                "<w:tblPr/></w:tblPrChange>",
                "table_property_change",
            ),
            (
                '<w:tcPrChange w:id="901" w:author="A" w:date="2026-06-01T09:30:00Z">'
                "<w:tcPr/></w:tcPrChange>",
                "table_property_change",
            ),
            (
                '<w:cellMerge w:id="902" w:author="A" w:date="2026-06-01T09:30:00Z"/>',
                "cell_revision",
            ),
        ],
    )
    def it_enumerates_and_refuses_each_exotic_type_by_name(
        self, markup: str, expected_type: str
    ):
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls, qn

        document = _doc(MINIMAL_TABLE)
        table = document.tables[0]._tbl
        first_row = table.tr_lst[0]
        cell = first_row.tc_lst[0]
        if "tblPrChange" in markup:
            table.tblPr.append(parse_xml(f"<w:x {nsdecls('w')}>{markup}</w:x>")[0])
        else:
            tc_pr = cell.get_or_add_tcPr()
            tc_pr.append(parse_xml(f"<w:x {nsdecls('w')}>{markup}</w:x>")[0])
        revisions = document.revisions
        exotic = [r for r in revisions if r.revision_type == expected_type]
        assert exotic, f"{expected_type} not enumerated"
        with pytest.raises(UnsupportedStructureError, match=expected_type):
            exotic[0].accept()
        with pytest.raises(UnsupportedStructureError, match="not resolve"):
            revisions.accept_all()
        assert revisions.remaining_unsupported() == {expected_type: 1}
