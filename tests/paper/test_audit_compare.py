"""Regression coverage for strict compare safety guarantees."""

from __future__ import annotations

import datetime as dt
import io
import zipfile
from pathlib import Path

import pytest

import docx
from docx.errors import UnsupportedStructureError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.package import compare

FROZEN = dt.datetime(2026, 7, 10, 12, 0, tzinfo=dt.timezone.utc)


def _save_pair(tmp_path: Path, build_original, build_revised):
    paths = []
    for name, build in (("original", build_original), ("revised", build_revised)):
        document = docx.Document()
        build(document)
        path = tmp_path / f"{name}.docx"
        document.save(path)
        paths.append(path)
    return paths


class DescribeStrictComparePreflight:
    def it_refuses_style_only_changes_without_touching_inputs(self, tmp_path: Path):
        def original(document):
            document.add_paragraph("Same text")

        def revised(document):
            document.add_paragraph("Same text")
            document.styles["Normal"].font.name = "Courier New"

        a, b = _save_pair(tmp_path, original, revised)
        before = (a.read_bytes(), b.read_bytes())
        with pytest.raises(UnsupportedStructureError, match="package part 'word/styles.xml'"):
            compare(a, b, author="Reviewer", date=FROZEN)
        assert (a.read_bytes(), b.read_bytes()) == before

    def it_refuses_text_and_formatting_changes(self, tmp_path: Path):
        def original(document):
            document.add_paragraph("Alpha content")

        def revised(document):
            run = document.add_paragraph().add_run("Alphb content")
            run.bold = True

        a, b = _save_pair(tmp_path, original, revised)
        with pytest.raises(UnsupportedStructureError, match="formatting or structural"):
            compare(a, b, author="Reviewer", date=FROZEN)

    def it_refuses_unmodeled_custom_xml_content(self, tmp_path: Path):
        def build(value):
            def _build(document):
                wrapper = OxmlElement("w:customXml")
                paragraph = OxmlElement("w:p")
                run = OxmlElement("w:r")
                text = OxmlElement("w:t")
                text.text = value
                run.append(text)
                paragraph.append(run)
                wrapper.append(paragraph)
                document.element.body.insert(0, wrapper)

            return _build

        a, b = _save_pair(tmp_path, build("ORIGINAL SECRET"), build("REVISED SECRET"))
        with pytest.raises(UnsupportedStructureError, match="does not match"):
            compare(a, b, author="Reviewer", date=FROZEN)

    def it_refuses_package_relationship_changes(self, tmp_path: Path):
        png = (
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
            b"\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89"
            b"\x00\x00\x00\rIDAT\x08\xd7c\xf8\xcf\xc0\xf0\x1f\x00\x05"
            b"\x00\x01\xff\x89\x99=\x1d\x00\x00\x00\x00IEND\xaeB`\x82"
        )
        image = tmp_path / "pixel.png"
        image.write_bytes(png)

        def original(document):
            document.add_paragraph("Alpha")

        def revised(document):
            document.add_paragraph("Alpha").add_run().add_picture(str(image))

        a, b = _save_pair(tmp_path, original, revised)
        with pytest.raises(UnsupportedStructureError, match="package-part addition"):
            compare(a, b, author="Reviewer", date=FROZEN)

    def it_refuses_malformed_numbering_as_a_typed_error(self, tmp_path: Path):
        def original(document):
            document.add_paragraph("Anchor")

        def revised(document):
            document.add_paragraph("Anchor")
            paragraph = document.add_paragraph("Malformed list item")
            num_pr = OxmlElement("w:numPr")
            num_id = OxmlElement("w:numId")
            num_id.set(qn("w:val"), "not-an-integer")
            num_pr.append(num_id)
            paragraph._p.get_or_add_pPr().append(num_pr)

        a, b = _save_pair(tmp_path, original, revised)
        before = (a.read_bytes(), b.read_bytes())
        with pytest.raises(UnsupportedStructureError, match="numbering id"):
            compare(a, b, author="Reviewer", date=FROZEN)
        assert (a.read_bytes(), b.read_bytes()) == before


class DescribeComparePairingBudget:
    def it_refuses_an_oversized_word_pairing_before_allocating(self, tmp_path: Path):
        def build(prefix):
            def _build(document):
                for index in range(101):
                    document.add_paragraph(f"{prefix} paragraph {index:03d}")

            return _build

        a, b = _save_pair(tmp_path, build("Old"), build("New"))
        with pytest.raises(UnsupportedStructureError, match="pairing budget"):
            compare(a, b, author="Reviewer", date=FROZEN)

    def it_uses_linear_whole_block_edits_in_block_mode(self, tmp_path: Path):
        def build(prefix):
            def _build(document):
                for index in range(101):
                    document.add_paragraph(f"{prefix} paragraph {index:03d}")

            return _build

        a, b = _save_pair(tmp_path, build("Old"), build("New"))
        result = compare(
            a,
            b,
            author="Reviewer",
            date=FROZEN,
            granularity="block",
        )
        result.document.revisions.accept_all()
        assert [p.text for p in result.document.paragraphs] == [
            p.text for p in docx.Document(b).paragraphs
        ]


def it_saves_compare_output_with_deterministic_zip_metadata(tmp_path: Path):
    a, b = _save_pair(
        tmp_path,
        lambda document: document.add_paragraph("Original text"),
        lambda document: document.add_paragraph("Revised text"),
    )
    result = compare(a, b, author="Reviewer", date=FROZEN)
    stream = io.BytesIO()
    result.document.save(stream)

    with zipfile.ZipFile(stream) as package:
        assert {info.date_time for info in package.infolist()} == {
            (1980, 1, 1, 0, 0, 0)
        }


class DescribeCompareAlignmentTrim:
    """The sequence budget measures the CHANGE, not the document length."""

    def it_self_compares_a_long_document_with_zero_revisions(self, tmp_path: Path):
        def build(document):
            for index in range(1_050):
                document.add_paragraph(f"Clause {index}: unique wording {index}.")

        a, _b = _save_pair(tmp_path, build, build)
        result = compare(a, a, author="Reviewer")
        assert len(result.document.revisions) == 0

    def and_it_redlines_one_edit_in_a_long_document(self, tmp_path: Path):
        def original(document):
            for index in range(1_050):
                document.add_paragraph(f"Clause {index} body text.")

        def revised(document):
            for index in range(1_050):
                text = f"Clause {index} body text."
                if index == 525:
                    text = "Clause 525 amended body text."
                document.add_paragraph(text)

        a, b = _save_pair(tmp_path, original, revised)
        result = compare(a, b, author="Reviewer")
        assert len(result.document.revisions) > 0
        result.document.revisions.accept_all()
        assert (
            result.document.paragraphs[525].text == "Clause 525 amended body text."
        )


class DescribeProofErrTransparency:
    """w:proofErr is the proofing tool's cache, not content."""

    def _flagged_paragraph(self, document, text: str) -> None:
        paragraph = document.add_paragraph(text)
        run = paragraph.runs[0]._r
        start = OxmlElement("w:proofErr")
        start.set(qn("w:type"), "spellStart")
        end = OxmlElement("w:proofErr")
        end.set(qn("w:type"), "spellEnd")
        run.addprevious(start)
        run.addnext(end)

    def it_redlines_a_deletion_beside_proofing_markers(self, tmp_path: Path):
        def original(document):
            self._flagged_paragraph(document, "Klause with a flagged typo.")
            document.add_paragraph("A second paragraph to be deleted.")

        def revised(document):
            self._flagged_paragraph(document, "Klause with a flagged typo.")

        a, b = _save_pair(tmp_path, original, revised)
        result = compare(a, b, author="Reviewer")
        result.document.revisions.accept_all()
        texts = [p.text for p in result.document.paragraphs]
        assert "A second paragraph to be deleted." not in texts

    def and_a_proofing_only_difference_is_no_difference(self, tmp_path: Path):
        def original(document):
            self._flagged_paragraph(document, "Klause with a flagged typo.")

        def revised(document):
            document.add_paragraph("Klause with a flagged typo.")

        a, b = _save_pair(tmp_path, original, revised)
        result = compare(a, b, author="Reviewer")
        assert len(result.document.revisions) == 0
