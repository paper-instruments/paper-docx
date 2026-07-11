"""Audit regressions for inspection and guarded editing edge cases."""

from __future__ import annotations

import datetime as dt

import pytest
from lxml import etree

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.errors import (
    BoundaryViolationError,
    DocumentProtectedError,
    TargetNotFoundError,
    UnsupportedStructureError,
)
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import OxmlElement, parse_xml

W = nsdecls("w")
W14 = nsdecls("w", "w14")


def _append_body_block(document, element) -> None:
    sect_pr = document.element.body.find(qn("w:sectPr"))
    if sect_pr is None:
        document.element.body.append(element)
    else:
        sect_pr.addprevious(element)


class DescribeRunTextProjection:
    def it_keeps_field_instructions_out_of_visible_story_text(self):
        from docx.story import iter_blocks

        document = docx.Document()
        _append_body_block(
            document,
            parse_xml(
                f'<w:p {W}><w:r><w:fldChar w:fldCharType="begin"/></w:r>'
                '<w:r><w:instrText> REF InternalTarget </w:instrText></w:r>'
                '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
                '<w:r><w:t>Visible result</w:t></w:r>'
                '<w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>'
            ),
        )
        assert [block.text for block in iter_blocks(document)] == ["Visible result"]

    def it_models_no_break_hyphens_for_inspection_search_and_safe_narrowing(self):
        from docx.search import find_one
        from docx.story import iter_blocks

        document = docx.Document()
        _append_body_block(
            document,
            parse_xml(
                f"<w:p {W}><w:r><w:t>non</w:t><w:noBreakHyphen/>"
                "<w:t>breaking</w:t></w:r></w:p>"
            ),
        )
        assert [block.text for block in iter_blocks(document)] == ["non-breaking"]
        find_one(document, "non-breaking").replace("non-binding")
        assert [block.text for block in iter_blocks(document)] == ["non-binding"]
        assert document.element.body.xpath("//w:noBreakHyphen")

    def it_does_not_match_through_unmodeled_visible_run_content(self):
        from docx.search import find_one, find_text

        document = docx.Document()
        _append_body_block(
            document,
            parse_xml(
                f"<w:p {W}><w:r><w:t>alpha</w:t>"
                '<w:sym w:font="Wingdings" w:char="F0FC"/>'
                "<w:t>omega</w:t></w:r></w:p>"
            ),
        )
        assert find_text(document, "alphaomega") == []
        assert find_one(document, "alpha").text == "alpha"
        assert find_one(document, "omega").text == "omega"


class DescribeAlternateContentSelection:
    def it_selects_the_first_supported_choice_then_falls_back(self):
        from docx.story import iter_blocks

        document = docx.Document()
        namespace_attrs = (
            f'{W} xmlns:mc="http://schemas.openxmlformats.org/'
            'markup-compatibility/2006" '
            'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
            'xmlns:x="urn:paper:unsupported"'
        )
        _append_body_block(
            document,
            parse_xml(
                f"<w:p {namespace_attrs}><mc:AlternateContent>"
                '<mc:Choice Requires="x"><w:r><w:t>wrong</w:t></w:r></mc:Choice>'
                '<mc:Choice Requires="w14"><w:r><w:t>supported</w:t></w:r>'
                "</mc:Choice><mc:Fallback><w:r><w:t>fallback one</w:t></w:r>"
                "</mc:Fallback></mc:AlternateContent></w:p>"
            ),
        )
        _append_body_block(
            document,
            parse_xml(
                f"<w:p {namespace_attrs}><mc:AlternateContent>"
                '<mc:Choice Requires="x"><w:r><w:t>wrong again</w:t></w:r>'
                "</mc:Choice><mc:Fallback><w:r><w:t>fallback two</w:t></w:r>"
                "</mc:Fallback></mc:AlternateContent></w:p>"
            ),
        )
        assert [block.text for block in iter_blocks(document)] == [
            "supported",
            "fallback two",
        ]


class DescribeCellReplacementGuards:
    @pytest.mark.parametrize("kind", ["tab", "break"])
    def it_refuses_non_text_run_content_without_leaving_it_behind(self, kind: str):
        from docx.tableops import update_cell

        document = docx.Document()
        table = document.add_table(rows=1, cols=1)
        paragraph = table.cell(0, 0).paragraphs[0]
        paragraph.clear()
        run = paragraph.add_run("left")
        run.add_tab() if kind == "tab" else run.add_break()
        run.add_text("right")
        before = table._tbl.xml
        with pytest.raises(UnsupportedStructureError, match="tab, break"):
            update_cell(table, 0, 0, "replacement")
        assert table._tbl.xml == before


def _change_span_context(run, kind: str) -> None:
    if kind == "field":
        begin = OxmlElement("w:r")
        begin_char = OxmlElement("w:fldChar")
        begin_char.set(qn("w:fldCharType"), "begin")
        begin.append(begin_char)
        end = OxmlElement("w:r")
        end_char = OxmlElement("w:fldChar")
        end_char.set(qn("w:fldCharType"), "end")
        end.append(end_char)
        run.addprevious(begin)
        run.addnext(end)
        return

    parent = run.getparent()
    position = parent.index(run)
    parent.remove(run)
    if kind == "hyperlink":
        wrapper = OxmlElement("w:hyperlink")
        wrapper.set(qn("w:anchor"), "ChangedTarget")
        wrapper.append(run)
    elif kind == "sdt":
        wrapper = OxmlElement("w:sdt")
        wrapper.append(OxmlElement("w:sdtPr"))
        content = OxmlElement("w:sdtContent")
        content.append(run)
        wrapper.append(content)
    else:
        wrapper = OxmlElement("w:ins")
        wrapper.set(qn("w:id"), "91")
        wrapper.set(qn("w:author"), "Reviewer")
        wrapper.append(run)
    parent.insert(position, wrapper)


class DescribeSpanContextFreshness:
    @pytest.mark.parametrize("kind", ["field", "hyperlink", "sdt", "revision"])
    def it_rejects_same_text_after_its_edit_context_changes(self, kind: str):
        from docx.search import find_one

        document = docx.Document()
        run = document.add_paragraph("context-sensitive target").runs[0]._r
        span = find_one(document, "context-sensitive target")
        _change_span_context(run, kind)
        with pytest.raises(TargetNotFoundError, match="stale"):
            span.replace("changed")
        assert "context-sensitive target" in document.element.body.xml

    def it_refuses_when_a_text_atom_is_inserted_inside_the_captured_interval(self):
        from docx.search import find_one

        document = docx.Document()
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_text("adjacent ")
        run.add_text("atoms")
        span = find_one(document, "adjacent atoms")
        inserted = OxmlElement("w:t")
        inserted.text = "unexpected "
        run._r.insert(1, inserted)
        before = etree.tostring(document.element)

        with pytest.raises(TargetNotFoundError, match="stale"):
            span.replace("replacement")

        assert etree.tostring(document.element) == before

    def it_allows_text_atoms_inserted_outside_the_captured_interval(self):
        from docx.search import find_one
        from docx.story import iter_blocks

        document = docx.Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("adjacent ")
        paragraph.add_run("atoms")
        span = find_one(document, "adjacent atoms")
        prefix = OxmlElement("w:r")
        prefix.add_t("before ")
        paragraph._p.insert(0, prefix)
        suffix = OxmlElement("w:r")
        suffix.add_t(" after")
        paragraph._p.append(suffix)

        span.replace("replacement")

        assert [block.text for block in iter_blocks(document)] == [
            "before replacement after"
        ]


class DescribeControlSetterHardening:
    @staticmethod
    def _text_control(document, tag: str):
        from docx.controls import get_control

        paragraph = document.add_paragraph()
        paragraph._p.append(
            parse_xml(
                f'<w:sdt {W}><w:sdtPr><w:tag w:val="{tag}"/><w:text/>'
                "</w:sdtPr><w:sdtContent><w:r><w:t>original</w:t></w:r>"
                "</w:sdtContent></w:sdt>"
            )
        )
        return get_control(document, tag=tag)

    @staticmethod
    def _protect(document) -> None:
        protection = OxmlElement("w:documentProtection")
        protection.set(qn("w:edit"), "readOnly")
        protection.set(qn("w:enforcement"), "1")
        document.settings._element.append(protection)

    def it_refuses_a_detached_control_before_protection_or_mutation(self):
        document = docx.Document()
        control = self._text_control(document, "detached")
        control._sdt.getparent().remove(control._sdt)
        self._protect(document)
        before_document = etree.tostring(document.element)
        before_control = etree.tostring(control._sdt)

        with pytest.raises(TargetNotFoundError, match="stale"):
            control.set_value("replacement")

        assert etree.tostring(document.element) == before_document
        assert etree.tostring(control._sdt) == before_control

    def it_refuses_a_control_moved_to_another_document_atomically(self):
        source = docx.Document()
        control = self._text_control(source, "foreign")
        destination = docx.Document()
        destination.add_paragraph()._p.append(control._sdt)
        self._protect(source)
        before_source = etree.tostring(source.element)
        before_destination = etree.tostring(destination.element)

        with pytest.raises(BoundaryViolationError, match="different document"):
            control.set_value("replacement")

        assert etree.tostring(source.element) == before_source
        assert etree.tostring(destination.element) == before_destination

    def it_honors_declared_checkbox_glyphs(self):
        from docx.controls import get_control, set_control_value

        document = docx.Document()
        _append_body_block(
            document,
            parse_xml(
                f'<w:p {W14}><w:sdt><w:sdtPr><w:tag w:val="choice"/>'
                '<w14:checkbox><w14:checked w14:val="0"/>'
                '<w14:checkedState w14:val="2611" w14:font="Arial Unicode MS"/>'
                '<w14:uncheckedState w14:val="2610" w14:font="Arial Unicode MS"/>'
                "</w14:checkbox></w:sdtPr><w:sdtContent><w:r><w:t>☐</w:t>"
                "</w:r></w:sdtContent></w:sdt></w:p>"
            ),
        )
        set_control_value(document, True, tag="choice")
        assert get_control(document, tag="choice").value is True
        assert "☑" in document.element.body.xml
        (fonts,) = document.element.body.xpath("//w:sdtContent//w:rFonts")
        assert fonts.get(qn("w:ascii")) == "Arial Unicode MS"
        assert fonts.get(qn("w:hAnsi")) == "Arial Unicode MS"

    def it_validates_checkbox_content_before_changing_its_state(self):
        from docx.controls import get_control, set_control_value

        document = docx.Document()
        _append_body_block(
            document,
            parse_xml(
                f'<w:p {W14}><w:sdt><w:sdtPr><w:tag w:val="complex"/>'
                '<w14:checkbox><w14:checked w14:val="0"/></w14:checkbox>'
                "</w:sdtPr><w:sdtContent><w:tbl><w:tr><w:tc><w:p/>"
                "</w:tc></w:tr></w:tbl></w:sdtContent></w:sdt></w:p>"
            ),
        )
        control = get_control(document, tag="complex")
        before = etree.tostring(control._sdt)
        with pytest.raises(UnsupportedStructureError, match="table"):
            set_control_value(document, True, tag="complex")
        assert etree.tostring(control._sdt) == before

    def it_formats_aware_dates_from_utc_using_the_declared_format(self):
        from docx.controls import get_control, set_control_value

        document = docx.Document()
        _append_body_block(
            document,
            parse_xml(
                f'<w:p {W}><w:sdt><w:sdtPr><w:tag w:val="signed"/>'
                '<w:date w:fullDate="2020-01-01T00:00:00Z">'
                '<w:dateFormat w:val="MMMM d, yyyy"/><w:lid w:val="en-US"/>'
                "</w:date></w:sdtPr><w:sdtContent><w:r><w:t>old</w:t></w:r>"
                "</w:sdtContent></w:sdt></w:p>"
            ),
        )
        value = dt.datetime(
            2026, 7, 10, 1, 30, tzinfo=dt.timezone(dt.timedelta(hours=2))
        )
        set_control_value(document, value, tag="signed")
        assert get_control(document, tag="signed").value == "July 9, 2026"
        (date_pr,) = document.element.body.xpath("//w:sdtPr/w:date")
        assert date_pr.get(qn("w:fullDate")) == "2026-07-09T23:30:00Z"

    def it_refuses_a_date_without_a_declared_format_before_mutation(self):
        from docx.controls import get_control, set_control_value

        document = docx.Document()
        _append_body_block(
            document,
            parse_xml(
                f'<w:p {W}><w:sdt><w:sdtPr><w:tag w:val="unformatted"/>'
                '<w:date w:fullDate="2020-01-01T00:00:00Z"/></w:sdtPr>'
                "<w:sdtContent><w:r><w:t>old</w:t></w:r></w:sdtContent>"
                "</w:sdt></w:p>"
            ),
        )
        control = get_control(document, tag="unformatted")
        before = etree.tostring(control._sdt)
        with pytest.raises(UnsupportedStructureError, match="w:dateFormat"):
            set_control_value(document, dt.date(2026, 7, 10), tag="unformatted")
        assert etree.tostring(control._sdt) == before

    def it_refuses_named_date_tokens_without_a_declared_locale_before_mutation(self):
        from docx.controls import get_control, set_control_value

        document = docx.Document()
        _append_body_block(
            document,
            parse_xml(
                f'<w:p {W}><w:sdt><w:sdtPr><w:tag w:val="unlocalized"/>'
                '<w:date w:fullDate="2020-01-01T00:00:00Z">'
                '<w:dateFormat w:val="MMMM d, yyyy"/></w:date></w:sdtPr>'
                '<w:sdtContent><w:r><w:rPr><w:lang w:val="fr-FR"/></w:rPr>'
                "<w:t>ancien</w:t></w:r></w:sdtContent></w:sdt></w:p>"
            ),
        )
        control = get_control(document, tag="unlocalized")
        before = etree.tostring(control._sdt)
        with pytest.raises(UnsupportedStructureError, match="w:lid"):
            set_control_value(document, dt.date(2026, 7, 10), tag="unlocalized")
        assert etree.tostring(control._sdt) == before

    @pytest.mark.parametrize(
        "declaration",
        [
            '<w:dateFormat w:val="Q yyyy"/>',
            '<w:calendar w:val="hijri"/>',
        ],
    )
    def it_refuses_unsupported_date_declarations_before_mutation(
        self, declaration: str
    ):
        from docx.controls import get_control, set_control_value

        document = docx.Document()
        _append_body_block(
            document,
            parse_xml(
                f'<w:p {W}><w:sdt><w:sdtPr><w:tag w:val="unsupported"/>'
                '<w:date w:fullDate="2020-01-01T00:00:00Z">'
                f"{declaration}</w:date></w:sdtPr>"
                "<w:sdtContent><w:r><w:t>old</w:t></w:r></w:sdtContent>"
                "</w:sdt></w:p>"
            ),
        )
        control = get_control(document, tag="unsupported")
        before = etree.tostring(control._sdt)
        with pytest.raises(UnsupportedStructureError, match="unsupported"):
            set_control_value(document, dt.date(2026, 7, 10), tag="unsupported")
        assert etree.tostring(control._sdt) == before


class DescribeNumberingAudit:
    def it_reports_numbering_inherited_through_based_on_and_honors_num_id_zero(self):
        from docx.numbering import ensure_decimal_definition, list_numbering

        document = docx.Document()
        num_id = ensure_decimal_definition(document)
        base = document.styles.add_style("AuditListBase", WD_STYLE_TYPE.PARAGRAPH)
        base.element.append(
            parse_xml(
                f'<w:pPr {W}><w:numPr><w:ilvl w:val="1"/>'
                f'<w:numId w:val="{num_id}"/></w:numPr></w:pPr>'
            )
        )
        leaf = document.styles.add_style("AuditListLeaf", WD_STYLE_TYPE.PARAGRAPH)
        leaf.base_style = base
        document.add_paragraph("inherited item", style=leaf)
        suppressed = document.add_paragraph("not numbered", style=leaf)
        suppressed_num_pr = suppressed._p.get_or_add_pPr().get_or_add_numPr()
        suppressed_num_pr.get_or_add_numId().val = 0

        report = list_numbering(document)
        assert [(item.text, item.num_id, item.level) for item in report.numbered_paragraphs] == [
            ("inherited item", num_id, 1)
        ]

    def it_reuses_only_structurally_canonical_definitions(self):
        from docx.numbering import ensure_bullet_definition
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        document = docx.Document()
        first_num_id = ensure_bullet_definition(document)
        numbering = document.part.part_related_by(RT.NUMBERING)._element
        first_num = next(
            num
            for num in numbering.findall(qn("w:num"))
            if int(num.get(qn("w:numId"))) == first_num_id
        )
        abstract_id = first_num.find(qn("w:abstractNumId")).get(qn("w:val"))
        abstract = next(
            item
            for item in numbering.findall(qn("w:abstractNum"))
            if item.get(qn("w:abstractNumId")) == abstract_id
        )
        abstract.find(f".//{qn('w:lvlText')}").set(qn("w:val"), "not canonical")

        replacement_num_id = ensure_bullet_definition(document)
        assert replacement_num_id != first_num_id
        assert ensure_bullet_definition(document) == replacement_num_id


class DescribeHiddenTextScrubbing:
    def it_removes_style_inherited_vanish_but_honors_direct_false(self):
        document = docx.Document()
        style = document.styles.add_style("HiddenAudit", WD_STYLE_TYPE.PARAGRAPH)
        style.font.hidden = True
        paragraph = document.add_paragraph(style=style)
        paragraph.add_run("remove me")
        visible = paragraph.add_run("keep me")
        visible.font.hidden = False

        report = document.scrub(
            comments=False,
            metadata=False,
            track_changes_setting=False,
            hidden_text=True,
        )
        assert report.hidden_runs_removed == 1
        assert "remove me" not in paragraph._p.xml
        assert "keep me" in paragraph._p.xml


class DescribeMetadataScrubbing:
    def it_clears_core_timestamps_and_revision(self):
        document = docx.Document()
        core = document.core_properties
        core.created = dt.datetime(2020, 1, 2, 3, 4)
        core.modified = dt.datetime(2021, 2, 3, 4, 5)
        core.last_printed = dt.datetime(2022, 3, 4, 5, 6)
        core.revision = 42

        report = document.scrub(
            comments=False,
            track_changes_setting=False,
        )

        assert core.created is None
        assert core.modified is None
        assert core.last_printed is None
        assert core.revision == 0
        for name in ("created", "modified", "last_printed", "revision"):
            assert f"core:{name}" in report.metadata_fields_cleared

    def it_refuses_malformed_app_metadata_before_removing_comments(self):
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        document = docx.Document()
        paragraph = document.add_paragraph("commented")
        document.add_comment(paragraph.runs, text="keep", author="Audit")
        app_part = document.part.package.part_related_by(RT.EXTENDED_PROPERTIES)
        app_part._blob = b"<Properties>"  # noqa: SLF001 - malformed fixture
        before = document.element.xml

        with pytest.raises(UnsupportedStructureError, match="app.xml is malformed"):
            document.scrub()

        assert document.element.xml == before
        assert len(document.comments) == 1


class DescribeFormatOnlyProtection:
    def it_reports_and_enforces_formatting_only_protection(self):
        from docx.protection import acknowledge_protection, protection_status
        from docx.search import find_one

        document = docx.Document()
        document.add_paragraph("protected formatting")
        protection = OxmlElement("w:documentProtection")
        protection.set(qn("w:formatting"), "1")
        protection.set(qn("w:enforcement"), "1")
        document.settings._element.append(protection)

        status = protection_status(document)
        assert status.edit is None
        assert status.formatting
        assert status.enforced
        assert status.blocks_paper_edits
        assert status.to_dict()["formatting"] is True
        with pytest.raises(DocumentProtectedError, match="formatting-only"):
            find_one(document, "protected formatting").replace("changed")

        acknowledge_protection(document)
        report = document.scrub(
            comments=False, metadata=False, track_changes_setting=False
        )
        assert report.document_protection == {
            "edit": None,
            "enforced": True,
            "note": "reported, never removed (docx.protection)",
            "formatting": True,
        }


class DescribeFormattingDisclosure:
    def it_declares_uncomputed_run_and_paragraph_categories_unresolved(self):
        from docx.formatting import format_of

        document = docx.Document()
        paragraph = document.add_paragraph()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "120")
        paragraph._p.get_or_add_pPr().append(spacing)
        run = paragraph.add_run("formatted")
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "FFFF00")
        run._r.get_or_add_rPr().append(shading)

        unresolved = format_of(run).unresolved
        assert "run_shading" in unresolved
        assert "paragraph_spacing" in unresolved
        assert "theme_color_resolution" in unresolved
