"""Standing job evals (v0.1 Phase E).

The v0 gap review's panel found what 326 organ tests didn't, because organ
tests test organs and the panel tested JOBS. These scenarios drive the public
surface the way agents actually do. The bar for every step, forever:
**green or explicitly-refusing — never silently wrong.**

Steps that v0 gets wrong carry `@pytest.mark.xfail(strict=True, reason=...)`
naming the phase that fixes them; landing that phase REMOVES the marker in
the same commit (strict xfail fails the suite the moment the fix works, so a
marker can never go stale silently).
"""

from __future__ import annotations

import datetime as dt
import shutil
from pathlib import Path

import pytest

import docx
from docx.errors import PaperRefusal, UnsupportedStructureError
from docx.search import find_one, find_text
from docx.story import iter_blocks, outline

from .harness.contract import assert_changed_parts
from .harness.paths import fixture_path

FROZEN = dt.datetime(2026, 7, 8, 9, 0, 0, tzinfo=dt.timezone.utc)
GAUNTLET = "generated/gauntlet/gauntlet.docx"
MOVES = "generated/feature-isolated/tracked-moves.docx"
FORMAT_CHANGES = "generated/feature-isolated/format-changes.docx"
FIELDS = "generated/feature-isolated/fields.docx"
PLACEHOLDER = "generated/feature-isolated/placeholder-control.docx"
MERGED_TABLE = "generated/feature-isolated/table-merged-nested.docx"
MINIMAL = "generated/minimal-clean/minimal.docx"
MOVED_TEXT = "The indemnity clause relocated by tracked move."


def _copy(relpath: str, tmp_path: Path) -> Path:
    target = tmp_path / Path(relpath).name
    shutil.copyfile(fixture_path(relpath), target)
    return target


class DescribeMultiRoundRedlineJob:
    """Contract redlining: edit tracked, enumerate, layer, resolve, report."""

    def it_redlines_and_resolves_a_clean_document(self, tmp_path: Path):
        path = _copy(MINIMAL, tmp_path)
        doc = docx.Document(str(path))
        find_one(doc, "perfectly ordinary").replace(
            "entirely unremarkable", tracked=True, author="Alice Editor", date=FROZEN
        )
        authors = {r.author for r in doc.revisions}
        assert authors == {"Alice Editor"}
        assert doc.revisions.reject_all() > 0
        texts = [b.text for b in iter_blocks(doc)]
        assert "First body paragraph with perfectly ordinary text." in texts

    def it_enumerates_moves_instead_of_reporting_a_clean_document(self):
        """A doc with pending tracked MOVES must never read as revision-free."""
        doc = docx.Document(str(fixture_path(MOVES)))
        move_types = {r.revision_type for r in doc.revisions}
        assert "move_from" in move_types and "move_to" in move_types

    def it_shows_moved_text_exactly_once_per_view(self):
        doc = docx.Document(str(fixture_path(MOVES)))
        current = "\n".join(b.text for b in iter_blocks(doc, view="current"))
        original = "\n".join(b.text for b in iter_blocks(doc, view="original"))
        assert current.count(MOVED_TEXT) == 1
        assert original.count(MOVED_TEXT) == 1

    def it_enumerates_format_change_revisions(self):
        """Bolding a word with tracking on must be visible to revisions."""
        doc = docx.Document(str(fixture_path(FORMAT_CHANGES)))
        kinds = {r.revision_type for r in doc.revisions}
        assert "format_change" in kinds

    @pytest.mark.parametrize("relpath", [MOVES, FORMAT_CHANGES])
    def it_refuses_to_claim_resolution_of_unsupported_revisions(self, relpath: str):
        doc = docx.Document(str(fixture_path(relpath)))
        with pytest.raises(UnsupportedStructureError):
            doc.revisions.accept_all()

    @pytest.mark.xfail(
        strict=True, reason="same-author layering lands in Phase 1 (S4)"
    )
    def it_lets_the_same_author_extend_their_own_pending_edit(self, tmp_path: Path):
        path = _copy(MINIMAL, tmp_path)
        doc = docx.Document(str(path))
        find_one(doc, "perfectly ordinary").replace(
            "quite modern", tracked=True, author="Alice Editor", date=FROZEN
        )
        # straddles her own insertion boundary: base "with " + inserted "quite"
        span = find_one(doc, "with quite modern")
        span.replace("with thoroughly modern", tracked=True,
                     author="Alice Editor", date=FROZEN)
        doc.revisions.reject_all()
        texts = [b.text for b in iter_blocks(doc)]
        assert "First body paragraph with perfectly ordinary text." in texts

    def it_still_refuses_cross_author_straddles(self, tmp_path: Path):
        path = _copy(MINIMAL, tmp_path)
        doc = docx.Document(str(path))
        find_one(doc, "perfectly ordinary").replace(
            "quite modern", tracked=True, author="Alice Editor", date=FROZEN
        )
        span = find_one(doc, "with quite modern")
        with pytest.raises(UnsupportedStructureError):
            span.replace("with thoroughly modern", tracked=True,
                         author="Bob Reviewer", date=FROZEN)


class DescribeFormFillJob:
    """Template filling: placeholder controls, tokens, boxed and tabled text."""

    def it_fills_a_placeholder_control_so_word_sees_it_filled(self, tmp_path: Path):
        path = _copy(PLACEHOLDER, tmp_path)
        doc = docx.Document(str(path))
        find_one(doc, "Click or tap here to enter text.").replace("Paper Instruments LLC")
        body = doc.element.body
        assert not body.xpath("//w:sdtPr/w:showingPlcHdr"), (
            "control still flagged as showing placeholder text"
        )
        assert not body.xpath('//w:rStyle[@w:val="PlaceholderText"]'), (
            "filled value still styled as placeholder"
        )
        texts = [b.text for b in iter_blocks(doc)]
        assert "Client name: Paper Instruments LLC" in texts

    def it_replaces_a_fragmented_token(self, tmp_path: Path):
        path = _copy(MINIMAL, tmp_path)
        doc = docx.Document(str(path))
        p = doc.add_paragraph()
        p.add_run("Prepared for {{cl")
        p.add_run("ient_na")
        p.add_run("me}} today.")
        find_one(doc, "{{client_name}}").replace("Paper Instruments LLC")
        assert p.text == "Prepared for Paper Instruments LLC today."

    def it_replaces_text_inside_a_text_box(self, tmp_path: Path):
        path = _copy("generated/feature-isolated/textbox.docx", tmp_path)
        doc = docx.Document(str(path))
        find_one(doc, "living inside the text box").replace("living happily in the box")
        boxed = [b.text for b in iter_blocks(doc) if b.in_text_box]
        assert boxed == ["Text living happily in the box."]

    def it_refuses_to_fill_a_field_result(self):
        """Field results are Word's to regenerate; edits there vanish."""
        doc = docx.Document(str(fixture_path(FIELDS)))
        span = find_one(doc, "June 1, 2026")
        assert span.in_field
        with pytest.raises(UnsupportedStructureError):
            span.replace("July 9, 2026")

    @pytest.mark.xfail(
        strict=True, reason="cell-wise table guards land in Phase 1 (S1)"
    )
    def it_updates_a_plain_cell_in_a_table_with_merged_headers(self, tmp_path: Path):
        from docx.tableops import update_cell

        path = _copy(MERGED_TABLE, tmp_path)
        doc = docx.Document(str(path))
        table = doc.tables[0]
        update_cell(table, 1, 0, "updated plain cell")  # unmerged target cell
        assert "updated plain cell" in table.cell(1, 0).text

    @pytest.mark.xfail(strict=True, reason="replace_all lands in Phase 1 (S3)")
    def it_fills_ten_placeholders_in_one_call(self, tmp_path: Path):
        from docx.search import replace_all

        path = _copy(MINIMAL, tmp_path)
        doc = docx.Document(str(path))
        for i in range(10):
            doc.add_paragraph(f"Slot {i}: {{{{token}}}} end.")
        result = replace_all(doc, "{{token}}", "VALUE")
        assert result.replaced_count == 10
        assert len(find_text(doc, "{{token}}")) == 0


class DescribeReportAssemblyJob:
    """Build a report section after an anchor; verify structure and budget."""

    def it_inserts_a_plain_section_with_a_narrow_budget(self, tmp_path: Path):
        from docx.blocks import insert_section_after
        from docx.package import patch_save

        source = _copy(MINIMAL, tmp_path)
        doc = docx.Document(str(source))
        insert_section_after(
            doc, "First body paragraph",
            heading="Findings", paragraphs=["First finding.", "Second finding."],
        )
        out = tmp_path / "report.docx"
        patch_save(source, doc, out)
        assert_changed_parts(source, out, {"word/document.xml"})
        reopened = docx.Document(str(out))
        texts = [b.text for b in iter_blocks(reopened)]
        assert "Findings" in texts and "Second finding." in texts

    @pytest.mark.xfail(
        strict=True, reason="rich block insertion lands in Phase 2 (V3)"
    )
    def it_inserts_a_rich_section_with_list_and_table(self, tmp_path: Path):
        from docx.blocks import (
            ListBlock,
            RichParagraph,
            TableBlock,
            TextRun,
            insert_blocks_after,
        )

        path = _copy(MINIMAL, tmp_path)
        doc = docx.Document(str(path))
        insert_blocks_after(
            doc, "First body paragraph",
            blocks=[
                RichParagraph(runs=[TextRun("Key findings", bold=True)]),
                ListBlock(items=["Alpha", "Beta"], kind="bullet"),
                TableBlock(rows=[["Metric", "Value"], ["Uptime", "99.9%"]]),
            ],
        )
        blocks = list(iter_blocks(doc))
        assert any(b.text == "Key findings" for b in blocks)
        assert any(b.kind == "table" and "Uptime" in b.text for b in blocks)

    @pytest.mark.xfail(
        strict=True, reason="numbering authoring lands in Phase 2 (V2)"
    )
    def it_creates_a_real_bullet_list_in_a_definitionless_document(
        self, tmp_path: Path
    ):
        """THE 'real bullet' gap: a doc that never had a list gets one."""
        import zipfile

        from docx.numbering import ensure_bullet_definition, apply_numbering

        source = fixture_path(MINIMAL)
        stripped = tmp_path / "no-numbering.docx"
        with zipfile.ZipFile(source) as zin:
            with zipfile.ZipFile(stripped, "w") as zout:
                for name in zin.namelist():
                    if "numbering" not in name:
                        blob = zin.read(name)
                        if name.endswith(".rels") or name.endswith(".xml"):
                            blob = blob.replace(b"numbering.xml", b"gone.xml")
                        zout.writestr(name, blob)
        doc = docx.Document(str(stripped))
        num_id = ensure_bullet_definition(doc)
        p = doc.add_paragraph("A real bullet at last")
        apply_numbering(p, num_id=num_id, level=0)

    def it_summarizes_what_changed_for_a_human(self, tmp_path: Path):
        from docx.package import diff_package

        source = _copy(MINIMAL, tmp_path)
        doc = docx.Document(str(source))
        find_one(doc, "equally unremarkable").replace("quite peculiar")
        out = tmp_path / "edited.docx"
        from docx.package import patch_save

        patch_save(source, doc, out)
        assert diff_package(source, out).semantic_changed_parts() == (
            "word/document.xml",
        )


class DescribeJobSafetyNet:
    """Cross-job honesty: outline never hides what it cannot read."""

    def it_confesses_unreadable_and_unresolved_regions(self):
        counts = outline(docx.Document(str(fixture_path(GAUNTLET)))).blind_region_counts
        for key in ("moves", "format_changes", "fields"):
            assert key in counts, f"blind_region_counts missing {key!r}"
        assert counts["moves"] == 2
        assert counts["format_changes"] == 2
        assert counts["fields"] == 2

    def it_never_lets_a_refusal_mutate_the_document(self, tmp_path: Path):
        from .harness.contract import assert_refusal_atomic

        path = _copy(FIELDS, tmp_path)
        doc = docx.Document(str(path))

        def edit_field_result(document):
            find_one(document, "June 1, 2026").replace("never")

        assert_refusal_atomic(doc, edit_field_result, PaperRefusal, on_disk=(path,))
