"""Rollback at the wrap sites added around already-mutating paper APIs."""

from __future__ import annotations

import datetime as dt

import pytest
from lxml import etree

import docx
from docx import blocks, fields
from docx.blocks import (
    RichParagraph,
    TextRun,
    insert_blocks_after,
    insert_section_after,
    tracked_delete_paragraphs,
    tracked_replace_paragraphs,
)
from docx.bookmarks import create_bookmark, delete_bookmark, list_bookmarks
from docx.fields import (
    add_date_field,
    add_page_count_field,
    add_page_number_field,
    add_reference_field,
    insert_toc_after,
)
from docx.oxml.comments import CT_Comment
from docx.oxml.ns import qn
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.search import find_one

from .harness.contract import assert_refusal_atomic
from .harness.paths import fixture_path

MINIMAL = "generated/minimal-clean/minimal.docx"
BOOKMARKS = "generated/feature-isolated/bookmarks.docx"
FROZEN = dt.datetime(2026, 7, 7, 12, 0, 0, tzinfo=dt.timezone.utc)


def _doc(relpath: str = MINIMAL):
    return docx.Document(str(fixture_path(relpath)))


def _fail_after(original):
    def wrapper(*args, **kwargs):
        original(*args, **kwargs)
        raise RuntimeError("forced after mutation")

    return wrapper


class DescribeBlockRollbackWrappers:
    def it_restores_insert_section_after_a_late_failure(self, monkeypatch):
        document = _doc()
        monkeypatch.setattr(blocks, "_insert_after", _fail_after(blocks._insert_after))

        assert_refusal_atomic(
            document,
            lambda candidate: insert_section_after(
                candidate,
                "First body paragraph",
                heading="Rolled Back",
                paragraphs=["Should not stay."],
            ),
            RuntimeError,
        )

        assert "Rolled Back" not in [paragraph.text for paragraph in document.paragraphs]

    def it_restores_tracked_delete_after_markup_is_written(self, monkeypatch):
        document = _doc()
        target = "First body paragraph with perfectly ordinary text."
        monkeypatch.setattr(
            blocks, "_mark_paragraph_deleted", _fail_after(blocks._mark_paragraph_deleted)
        )

        assert_refusal_atomic(
            document,
            lambda candidate: tracked_delete_paragraphs(
                candidate, "First body paragraph", count=1, author="Carol QA", date=FROZEN
            ),
            RuntimeError,
        )

        assert target in [paragraph.text for paragraph in document.paragraphs]
        assert not document.element.body.xpath("//w:del")

    def it_restores_tracked_replace_after_insert_starts(self, monkeypatch):
        document = _doc()
        target = "First body paragraph with perfectly ordinary text."
        monkeypatch.setattr(blocks, "_insert_after", _fail_after(blocks._insert_after))

        assert_refusal_atomic(
            document,
            lambda candidate: tracked_replace_paragraphs(
                candidate,
                "First body paragraph",
                ["Replacement that must not stay."],
                author="Carol QA",
                date=FROZEN,
            ),
            RuntimeError,
        )

        texts = [paragraph.text for paragraph in document.paragraphs]
        assert target in texts
        assert "Replacement that must not stay." not in texts
        assert not document.element.body.xpath("//w:del")
        assert not document.element.body.xpath("//w:ins")

    def it_restores_insert_blocks_after_a_late_failure(self, monkeypatch):
        document = _doc()
        monkeypatch.setattr(blocks, "_insert_after", _fail_after(blocks._insert_after))

        assert_refusal_atomic(
            document,
            lambda candidate: insert_blocks_after(
                candidate,
                "First body paragraph",
                blocks=[RichParagraph(runs=[TextRun("Must not stay")])],
            ),
            RuntimeError,
        )

        assert "Must not stay" not in [paragraph.text for paragraph in document.paragraphs]


class DescribeFieldRollbackWrappers:
    @pytest.mark.parametrize(
        "operation",
        [
            pytest.param(
                lambda document: add_page_number_field(document.paragraphs[0]),
                id="page_number",
            ),
            pytest.param(
                lambda document: add_page_count_field(document.paragraphs[0]),
                id="page_count",
            ),
            pytest.param(
                lambda document: add_date_field(document.paragraphs[0]),
                id="date",
            ),
            pytest.param(
                lambda document: add_reference_field(
                    document.paragraphs[0], bookmark="KeyPhrase"
                ),
                id="reference",
            ),
            pytest.param(
                lambda document: insert_toc_after(document, "First body paragraph"),
                id="toc",
            ),
        ],
    )
    def it_restores_a_field_insert_after_settings_are_armed(self, monkeypatch, operation):
        document = _doc()
        create_bookmark(document, find_one(document, "perfectly ordinary"), "KeyPhrase")
        monkeypatch.setattr(
            fields, "_set_update_fields_on_open", _fail_after(fields._set_update_fields_on_open)
        )

        assert_refusal_atomic(document, operation, RuntimeError)

        assert document.settings.element.find(qn("w:updateFields")) is None
        assert not document.element.body.xpath("//w:fldSimple")
        assert not document.element.body.xpath("//w:fldChar")


class DescribeBookmarkDeleteRollback:
    def it_restores_markers_if_delete_fails_after_the_first_remove(self, monkeypatch):
        document = _doc(BOOKMARKS)
        names_before = {bookmark.name for bookmark in list_bookmarks(document)}
        assert "DefinedTerm" in names_before
        removed = {"count": 0}

        def remove_then_fail(self, child):
            etree._Element.remove(self, child)
            removed["count"] += 1
            if removed["count"] == 1:
                raise RuntimeError("forced after mutation")

        monkeypatch.setattr(BaseOxmlElement, "remove", remove_then_fail)

        assert_refusal_atomic(
            document,
            lambda candidate: delete_bookmark(candidate, "DefinedTerm"),
            RuntimeError,
        )

        assert {bookmark.name for bookmark in list_bookmarks(document)} == names_before
        assert "See the Master Agreement for definitions." in [
            paragraph.text for paragraph in document.paragraphs
        ]


class DescribeCommentAttributeRollback:
    def it_restores_author_after_a_late_write_failure(self, monkeypatch):
        document = docx.Document()
        comment = document.comments.add_comment("note", author="Original")
        original_set = CT_Comment.author.fset

        def set_then_fail(element, value):
            original_set(element, value)
            raise RuntimeError("forced after mutation")

        monkeypatch.setattr(CT_Comment, "author", property(CT_Comment.author.fget, set_then_fail))

        assert_refusal_atomic(
            document,
            lambda _candidate: setattr(comment, "author", "Changed"),
            RuntimeError,
        )
        assert comment.author == "Original"

    def it_restores_initials_after_a_late_write_failure(self, monkeypatch):
        document = docx.Document()
        comment = document.comments.add_comment("note", author="Original", initials="OR")
        original_set = CT_Comment.initials.fset

        def set_then_fail(element, value):
            original_set(element, value)
            raise RuntimeError("forced after mutation")

        monkeypatch.setattr(
            CT_Comment, "initials", property(CT_Comment.initials.fget, set_then_fail)
        )

        assert_refusal_atomic(
            document,
            lambda _candidate: setattr(comment, "initials", "CH"),
            RuntimeError,
        )
        assert comment.initials == "OR"
